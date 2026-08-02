#!/usr/bin/env python3
"""Build CKI NAR Submission Package v26.

v26 incorporates P0+P1 fixes from v25 expert panel review (Score: 7.50/10).
All 10 new issues (N1-N10) addressed, except N5 (reference renumbering, deferred).

P0 Fixes (v25 expert panel, ~1h):
  N1 — Supplementary title "Selective" → "Baseline-Normalized"
  N2 — Global cleanup "Cohen's d" → "SES" (7 residues across 3 docs)
  N3 — Table 1 values: 99→102 cell types, 4,851→5,151 pairs (regenerated)
  N4 — "neutral" terminology cleanup in Figure legends (→"constrained"/"baseline")

P1 Fixes (v25 expert panel, ~1.5h):
  N6 — Repro Guide Section 3.2: brain k_n clarified (per-pair, not common)
  N7 — Supplementary: EVT GPD fit diagnostics reference added
  N8 — Limitations: duplicate "Seventh" → "Eleventh"
  N9 — Brain PMI discussion expanded (regional heterogeneity, cell-type RNA stability)

Deferred:
  N5 — Reference numbering alignment (requires systematic renumbering)
  N10 — Python version consistency (minor)

Inherits from v25:
  - C1-C7 Critical fixes (v22+v23+v24)
  - M1-M20 Major Issue fixes (v25)
  - Phase A-E expert panel fixes (v20)
  - Figures (1-6, S1-S7)
  - Graphical Abstract

Files regenerated:
  - CKI_NAR_Manuscript.docx (fresh, N1/N2/N4/N8/N9 fixes)
  - CKI_NAR_Supplementary.docx (fresh, N1/N2/N7 fixes)
  - CKI_NAR_Cover_Letter.docx (fresh, N1 fix)
  - CKI_NAR_Reproducibility_Guide.docx (fresh, N2/N6 fixes)
  - Table1-2.docx (fresh, extracted from regenerated manuscript)
"""

import os
import sys
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

BASE_DIR = Path(__file__).parent.resolve()
VERSION3_DIR = BASE_DIR / "version3"
RESULTS_DIR = BASE_DIR / "results"
FIGURES_DIR = RESULTS_DIR / "figures_final"
PYTHON = r"C:\Users\KnightZ\.workbuddy\binaries\python\envs\default\Scripts\python.exe"
NODE = r"C:\Users\KnightZ\.workbuddy\binaries\node\versions\22.22.2\node.exe"
NODE_PATH = r"C:\Users\KnightZ\.workbuddy\binaries\node\workspace\node_modules"

V26_ZIP = VERSION3_DIR / "CKI_NAR_Submission_v26.zip"
WORK_DIR = VERSION3_DIR / "CKI_NAR_Submission_v26"
V25_DIR = VERSION3_DIR / "CKI_NAR_Submission_v25"


def run_script(cmd, label, env=None):
    """Run a subprocess and check return code."""
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"{'='*60}")
    result = subprocess.run(cmd, shell=True, capture_output=True, text=True,
                            cwd=str(BASE_DIR), env=env)
    if result.stdout:
        print(result.stdout[-800:] if len(result.stdout) > 800 else result.stdout)
    if result.returncode != 0:
        print(f"  WARNING: Return code {result.returncode}")
        if result.stderr:
            print(f"  STDERR: {result.stderr[-500:]}")
    else:
        print(f"  OK")
    return result.returncode == 0


def build_v26():
    print("=" * 60)
    print("  CKI NAR Submission Package v26 Builder")
    print("  v25 Expert Panel P0+P1 Fixes (N1-N10 except N5)")
    print("=" * 60)

    # 0. Prepare work directory
    print(f"\n[0] Preparing {WORK_DIR.name}...")
    if WORK_DIR.exists():
        shutil.rmtree(WORK_DIR)
    WORK_DIR.mkdir(parents=True)

    # Copy base files from v25 (figures, graphical abstract)
    # NOTE: _fulltext.txt files are regenerated from fresh DOCX below
    if V25_DIR.exists():
        for f in V25_DIR.iterdir():
            if f.name.startswith("MANIFEST"):
                continue
            # DOCX and fulltext files will be regenerated fresh
            if f.name.endswith(".docx") or f.name.endswith("_fulltext.txt"):
                continue
            shutil.copy2(f, WORK_DIR / f.name)
        print(f"  Copied figures/GA from v25")
    else:
        print(f"  ERROR: v25 directory not found at {V25_DIR}")
        return False

    # 1. Regenerate DOCX files fresh
    print(f"\n[1] Regenerating all DOCX files with v26 fixes...")

    # 1a. Main manuscript (N1/N2/N4/N8/N9 fixes)
    run_script(
        f'"{PYTHON}" -u generate_manuscript_nar.py',
        "Generating CKI_NAR_Manuscript.docx (N1/N2/N4/N8/N9 fixes)"
    )
    ms_src = RESULTS_DIR / "CKI_NAR_Manuscript.docx"
    if ms_src.exists():
        shutil.copy2(ms_src, WORK_DIR / "CKI_NAR_Manuscript.docx")
        print(f"  Copied: {ms_src.name} ({ms_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Manuscript not generated!")
        return False

    # 1b. Supplementary materials (N1/N2/N7 fixes)
    run_script(
        f'"{PYTHON}" -u notebooks/68_gen_supplementary_en.py',
        "Generating CKI_NAR_Supplementary.docx (N1/N2/N7 fixes)"
    )
    sm_src = RESULTS_DIR / "CKI_NAR_Supplementary.docx"
    if sm_src.exists():
        shutil.copy2(sm_src, WORK_DIR / "CKI_NAR_Supplementary.docx")
        print(f"  Copied: {sm_src.name} ({sm_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Supplementary not generated!")
        return False

    # 1c. Cover letter (N1 fix)
    run_script(
        f'"{PYTHON}" -u generate_cover_letter_nar.py',
        "Generating CKI_NAR_Cover_Letter.docx (N1 fix)"
    )
    cl_src = RESULTS_DIR / "CKI_NAR_Cover_Letter.docx"
    if cl_src.exists():
        shutil.copy2(cl_src, WORK_DIR / "CKI_NAR_Cover_Letter.docx")
        print(f"  Copied: {cl_src.name} ({cl_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Cover letter not generated!")
        return False

    # 1d. Reproducibility guide (N2/N6 fixes)
    node_env = os.environ.copy()
    node_env["NODE_PATH"] = NODE_PATH
    run_script(
        f'"{NODE}" notebooks/100_gen_reproducibility_docx.js',
        "Generating CKI_NAR_Reproducibility_Guide.docx (N2/N6 fixes)"
    )
    rg_src = RESULTS_DIR / "CKI_Reproducibility_Guide.docx"
    if rg_src.exists():
        shutil.copy2(rg_src, WORK_DIR / "CKI_NAR_Reproducibility_Guide.docx")
        print(f"  Copied: {rg_src.name} ({rg_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Reproducibility guide not generated!")
        return False

    # 1e. Table1-2 (N3 fix: regenerated from fresh manuscript)
    run_script(
        f'"{PYTHON}" -u notebooks/_extract_table1_2.py',
        "Generating Table1-2.docx (N3: extracted from fresh manuscript)"
    )
    tb_src = RESULTS_DIR / "Table1-2.docx"
    if tb_src.exists():
        shutil.copy2(tb_src, WORK_DIR / "Table1-2.docx")
        print(f"  Copied: {tb_src.name} ({tb_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Table1-2 not generated!")
        return False

    # 1f. Regenerate fulltext files from fresh DOCX
    print(f"\n[1f] Regenerating fulltext files from fresh DOCX...")
    import docx as _docx
    fulltext_map = {
        "CKI_NAR_Manuscript.docx": "CKI_NAR_Manuscript_fulltext.txt",
        "CKI_NAR_Supplementary.docx": "CKI_NAR_Supplementary_fulltext.txt",
        "CKI_NAR_Cover_Letter.docx": "CKI_NAR_Cover_Letter_fulltext.txt",
        "CKI_NAR_Reproducibility_Guide.docx": "CKI_NAR_Reproducibility_Guide_fulltext.txt",
        "Table1-2.docx": "Table1-2_fulltext.txt",
    }
    for docx_name, txt_name in fulltext_map.items():
        docx_path = WORK_DIR / docx_name
        txt_path = WORK_DIR / txt_name
        d = _docx.Document(str(docx_path))
        lines = [p.text for p in d.paragraphs]
        for table in d.tables:
            lines.append("")
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"  {txt_name}: {txt_path.stat().st_size:,} bytes")

    # 2. Write manifest
    print(f"\n[2] Writing MANIFEST_v26.txt...")
    manifest = f"""CKI NAR Submission Package v26
Built: {datetime.now().strftime('%Y-%m-%d %H:%M')}

v25 Expert Panel Review P0+P1 Fixes (Score: 7.50/10 → target ~7.8+):

P0 Fixes (v25 expert panel N1-N4):
  N1 — Supplementary title: "Selective" → "Baseline-Normalized" (68_gen_supplementary_en.py:67)
  N2 — Global cleanup: "Cohen's d" → "SES" (7 residues: 3 Supp + 2 Repro Guide + 1 Manuscript Limitations + 1 Cover Letter)
  N3 — Table 1 values aligned: 99→102 cell types, 4,851→5,151 pairs (regenerated via _extract_table1_2.py)
  N4 — Figure legends: "neutral" → "constrained"/"baseline" (Fig 1-3 legends, 5 occurrences)
  N8 — Limitations: duplicate "Seventh" → "Eleventh" (calibration factor paragraph)

P1 Fixes (v25 expert panel N6-N9):
  N6 — Repro Guide Section 3.2: brain k_n clarified as per-pair (not common k_n scale)
  N7 — Supplementary SN 3.3: EVT GPD fit diagnostics reference added
  N9 — Brain PMI discussion expanded (regional heterogeneity, cell-type RNA stability)

Deferred:
  N5 — Reference citation-number ordering (requires systematic renumbering)
  N10 — Python version consistency (minor, 3 places)

All Critical Issues (C1-C7) Resolved (inherited from v24):
  C1: BH-FDR q-value description (EVT, m=31,764)
  C2: k_n floor parameter table (1e-4)
  C3: Mouse k_f scheme (hybrid vs global HVG)
  C4: HK constrained baseline terminology
  C5: OPC internal consistency check
  C6: Title wording
  C7: NAR formatting (Keywords/Running Title/OS)

20 Major Issues (M1-M20) Resolved (inherited from v25):
  M5: "Cohen's d" → "SES" — now complete (N2 fix)
  All other M1-M20 fixes intact

Bootstrap Status (all 4 datasets):
  Mouse (Tabula Muris): 8/15 significant, B=1000, one-sided + BH FDR
  Human (Tabula Sapiens): 15/16 significant, P=9.99e-04, B=1000
  TCGA (BRCA/LIHC/LUAD): descriptive + SES, B=1000
  Brain (Siletti Atlas): 10/10 significant, P<0.01, FDR<0.05, B=1000

Residual Model (Brain):
  30 Strong candidates, BH-FDR across 31,764 EVT-extrapolated P-values
  16/30 significant (FDR<0.05): 6 astrocyte + 10 oligodendrocyte

Contents:
1. CKI_NAR_Manuscript.docx - Main manuscript (English, fresh, N1/N2/N4/N8/N9)
2. CKI_NAR_Supplementary.docx - Supplementary materials (English, fresh, N1/N2/N7)
3. CKI_NAR_Cover_Letter.docx - Cover letter (fresh, N1)
4. CKI_NAR_Reproducibility_Guide.docx - Computational reproducibility guide (fresh, N2/N6)
5. Table1-2.docx - Standalone tables (fresh, N3: regenerated from manuscript)
6. figure1.pdf through figure6.pdf - Main figures (unchanged)
7. Supplementary_Figure_S1.pdf through Supplementary_Figure_S7.pdf
8. CKI_graphical_abstract.png/pdf/svg - Graphical Abstract

Inherited from v20-v25:
  - Phase D Critical Fixes (4 items)
  - Phase D Major Fixes (19 items)
  - Phase B Statistical Upgrades
  - Phase C Methodological Reinforcement
  - Phase A Text Polish
  - v20 Review Fixes (7 items)
"""
    manifest_path = WORK_DIR / "MANIFEST_v26.txt"
    with open(manifest_path, "w", encoding="utf-8") as f:
        f.write(manifest)
    print(f"  Wrote {manifest_path.name}")

    # 3. Create ZIP
    print(f"\n[3] Creating ZIP...")
    with zipfile.ZipFile(V26_ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(WORK_DIR):
            for fname in sorted(files):
                fpath = Path(root) / fname
                arcname = f"CKI_NAR_Submission_v26/{fname}"
                zf.write(fpath, arcname)

    # 4. Verify
    zip_size_mb = V26_ZIP.stat().st_size / (1024 * 1024)
    print(f"\n{'='*60}")
    print(f"  v26 Package Built")
    print(f"{'='*60}")
    print(f"ZIP: {V26_ZIP}")
    print(f"Size: {zip_size_mb:.1f} MB")

    with zipfile.ZipFile(V26_ZIP, "r") as zf:
        infos = sorted(zf.infolist(), key=lambda x: x.filename)
        print(f"Files: {len(infos)}")
        for info in infos:
            print(f"  {info.file_size:>10,}  {info.filename}")

    # 5. Final consistency check
    print(f"\n{'='*60}")
    print(f"  Final Consistency Check")
    print(f"{'='*60}")
    checks_passed = 0
    checks_failed = 0

    docx_files = [
        ("CKI_NAR_Manuscript.docx", 50),
        ("CKI_NAR_Supplementary.docx", 35),
        ("CKI_NAR_Cover_Letter.docx", 30),
        ("CKI_NAR_Reproducibility_Guide.docx", 15),
        ("Table1-2.docx", 5),
    ]
    for fname, min_kb in docx_files:
        fpath = WORK_DIR / fname
        if fpath.exists():
            size_kb = fpath.stat().st_size / 1024
            status = "OK" if size_kb >= min_kb else "WARNING: small"
            print(f"  {fname}: {size_kb:.1f} KB [{status}]")
            if size_kb >= min_kb:
                checks_passed += 1
            else:
                checks_failed += 1
        else:
            print(f"  {fname}: MISSING!")
            checks_failed += 1

    for i in range(1, 7):
        fpath = WORK_DIR / f"figure{i}.pdf"
        if fpath.exists():
            checks_passed += 1
        else:
            print(f"  figure{i}.pdf: MISSING!")
            checks_failed += 1

    for i in range(1, 8):
        fpath = WORK_DIR / f"Supplementary_Figure_S{i}.pdf"
        if fpath.exists():
            checks_passed += 1
        else:
            print(f"  Supplementary_Figure_S{i}.pdf: MISSING!")
            checks_failed += 1

    for ext in ["png", "pdf", "svg"]:
        fpath = WORK_DIR / f"CKI_graphical_abstract.{ext}"
        if fpath.exists():
            checks_passed += 1
        else:
            print(f"  CKI_graphical_abstract.{ext}: MISSING!")
            checks_failed += 1

    print(f"\n  Checks passed: {checks_passed}")
    print(f"  Checks failed: {checks_failed}")
    if checks_failed == 0:
        print(f"  ALL CHECKS PASSED")
    else:
        print(f"  SOME CHECKS FAILED - review above")

    return checks_failed == 0


if __name__ == "__main__":
    success = build_v26()
    sys.exit(0 if success else 1)
