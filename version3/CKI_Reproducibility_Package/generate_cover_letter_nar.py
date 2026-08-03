"""
Generate Nucleic Acids Research Cover Letter
CKI Project — NAR submission
"""
from pathlib import Path
from datetime import date
from docx.shared import Pt, Inches

PROJECT_ROOT = Path(__file__).parent
OUTPUT_DIR = PROJECT_ROOT / "results"
OUTPUT_DIR.mkdir(exist_ok=True)

# Sender address grouped by institution (no empty-string spacers)
ADDRESS_GROUPS = [
    [
        "Li Zhang",
        "Institute of Blood Transfusion",
        "Chinese Academy of Medical Sciences & Peking Union Medical College",
        "Chengdu, China",
    ],
    [
        "Chinese Institute for Brain Research, Beijing",
        "Beijing, China",
    ],
    [
        "Email: knightz@pumc.edu.cn",
        "ORCID: 0000-0002-0698-0754",
    ],
]

RECIPIENT = [
    "The Editors",
    "Nucleic Acids Research",
    "Oxford University Press",
]


def add_para(text, doc, space_after=6, align=None, bold=False, size=None):
    """Add a paragraph with controlled spacing — no empty-paragraph spacers."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if align is not None:
        para.alignment = align
    return para


def run():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Sender address (right-aligned, grouped by institution) ──
    for gi, group in enumerate(ADDRESS_GROUPS):
        for li, line in enumerate(group):
            is_last_of_group = (li == len(group) - 1)
            is_last_group = (gi == len(ADDRESS_GROUPS) - 1)
            # Tight within group, 6pt between groups, 12pt after entire address
            sa = 0 if not is_last_of_group else (12 if is_last_group else 6)
            add_para(line, doc, space_after=sa,
                     align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Date ──
    add_para(date.today().strftime("%B %d, %Y"), doc, space_after=12)

    # ── Recipient ──
    for line in RECIPIENT:
        add_para(line, doc, space_after=0)
    # Add spacing after last recipient line
    doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

    # ── Re: line ──
    add_para("Re: Submission of Original Research Article", doc, space_after=12)

    # ── Title (centered, bold) ──
    add_para(
        "CKI: A Cell-state Kinetic Index for Quantifying Baseline-Normalized "
        "Transcriptomic Remodeling",
        doc, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True, size=13,
    )

    # ── Salutation ──
    add_para("Dear Editors,", doc, space_after=6)

    # ── Body ──
    add_para(
        "On behalf of my co-author, Dr. Xianming Wu (Chinese Institute for Brain "
        "Research, Beijing), I am pleased to submit our original research article, "
        "\u201cCKI: A Cell-state Kinetic Index for Quantifying Baseline-Normalized "
        "Transcriptomic Remodeling\u201d, for consideration for publication "
        "in Nucleic Acids Research.",
        doc,
    )

    add_para(
        "Cell identity is fundamentally defined by its transcriptome. However, a "
        "standardized, assumption-free metric to quantify transcriptomic divergence "
        "between cell states has been lacking. Drawing inspiration from evolutionary "
        "biology\u2019s Ka/Ks ratio, we propose CKI (Cell-state Kinetic Index), "
        "which decomposes Jensen\u2013Shannon divergence into two complementary components: "
        "k_n (baseline divergence rate, from housekeeping gene expression) and k_f "
        "(functional divergence rate, from identity gene expression). The ratio "
        "\u03c9 = k_f/k_n provides a robust, interpretable measure of baseline-normalized "
        "transcriptomic remodeling. As a new computational method with rigorous "
        "validation across multiple datasets, CKI aligns well with "
        "Nucleic Acids Research\u2019s scope in methods development and genomic analysis.",
        doc,
    )

    add_para(
        "We validated CKI \u03c9 across three dimensions: (1) an independent information dimension\u2014"
        "CKI \u03c9 captures information that standard metrics miss, showing negative "
        "correlation with all four standard distance metrics (Spearman r = \u22120.57 "
        "to \u22120.38, all P < 0.001), demonstrating that it measures something fundamentally "
        "different from Cosine similarity, JS divergence, and other existing approaches; "
        "(2) cross-dataset consistency\u2014CKI \u03c9 rankings are reproducible across independent "
        "datasets and parameter settings, with calibration establishing an empirical baseline for equivalent populations; and (3) two biological "
        "applications\u2014pan-cancer analysis revealing that tumors are more transcriptionally "
        "homogeneous than normal tissues (median NN/TT ratio 1.40\u20132.83), and "
        "brain regional analysis identifying 30 threshold-passing candidates (16 statistically significant) "
        "among 31,764 cross-region comparisons through a multiplicative residual model.",
        doc,
    )

    add_para(
        "Both authors have approved the manuscript and declare no competing interests. "
        "This work has not been published elsewhere, is not under consideration "
        "by any other journal, and has not been previously submitted to "
        "Nucleic Acids Research. AI tools (LLMs) were used for writing assistance; "
        "all AI-generated text was reviewed and revised by the authors, who take full "
        "responsibility. The CKI Python package (v0.3.1, MIT License) and all analysis "
        "code are publicly available at https://github.com/zhanglknt/CKI-cell-type-identification "
        "(permanent archive: Zenodo DOI: 10.5281/zenodo.15670808). "
        "The CKI package is implemented in Python (≥3.10) and runs on Linux, macOS, and Windows. "
        "Dependencies include scanpy, scipy, numpy, and pandas (see repository for full requirements).",
        doc,
    )

    add_para(
        "We suggest the following potential reviewers with relevant expertise: "
        "(1) Prof. Fabian Theis, Helmholtz Munich "
        "(fabian.theis@helmholtz-munich.de); "
        "(2) Prof. Sarah Teichmann, Wellcome Sanger Institute "
        "(st1@sanger.ac.uk); "
        "(3) Prof. Joshua Welch, University of Michigan "
        "(jdwlch@umich.edu); "
        "(4) Prof. Itai Yanai, NYU Grossman School of Medicine "
        "(itai.yanai@nyulangone.org); "
        "(5) Prof. Zemin Zhang, Peking University "
        "(zeminzhang@pku.edu.cn); "
        "(6) Prof. Xiaowo Wang, Tsinghua University "
        "(xww@tsinghua.edu.cn). "
        "None of the suggested reviewers have recent collaborations or conflicts "
        "of interest with the authors.",
        doc,
    )

    add_para(
        "Thank you for considering our work for publication in Nucleic Acids Research. "
        "We look forward to hearing from you.",
        doc,
    )

    # ── Closing ──
    add_para("Sincerely,", doc, space_after=12)
    add_para("Li Zhang (Corresponding Author)", doc, space_after=0)
    add_para("Xianming Wu (First Author)", doc, space_after=0)

    # Save
    out = str(OUTPUT_DIR / "CKI_NAR_Cover_Letter.docx")
    doc.save(out)
    print(f"Saved: {out}")
    return out


if __name__ == "__main__":
    run()
