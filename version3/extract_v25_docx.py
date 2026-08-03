"""Extract text from all DOCX files in v25 submission package."""
import docx
import os
import sys

def extract_docx(path):
    """Extract all text from a docx file."""
    doc = docx.Document(path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # Also extract tables
    for table in doc.tables:
        lines.append("\n--- TABLE ---")
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("--- END TABLE ---\n")
    return "\n".join(lines)

base = r"C:\Users\KnightZ\Desktop\细胞受选择\version3\CKI_NAR_Submission_v25"
out_dir = r"C:\Users\KnightZ\Desktop\细胞受选择\version3\v25_extracted"
os.makedirs(out_dir, exist_ok=True)

docx_files = [
    "CKI_NAR_Manuscript.docx",
    "CKI_NAR_Supplementary.docx",
    "CKI_NAR_Cover_Letter.docx",
    "CKI_NAR_Reproducibility_Guide.docx",
    "Table1-2.docx",
]

for f in docx_files:
    path = os.path.join(base, f)
    if os.path.exists(path):
        text = extract_docx(path)
        out_path = os.path.join(out_dir, f.replace(".docx", ".txt"))
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write(text)
        print(f"Extracted {f}: {len(text)} chars -> {out_path}")
    else:
        print(f"NOT FOUND: {f}")

print("Done.")
