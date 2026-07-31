"""
Simplified rebuild: Diff V1 clean baseline vs v4 final manuscript → Track Changes.

Phase 1: Accept V1 Track Changes → clean baseline
Phase 2: Load v4 (already has ALL fixes applied) as target
Phase 3: Paragraph-align + word-diff → generate Track Changes
Phase 4: Save final manuscript with TC markup

Output: version2/NAR_Submission/CKI_NAR_Manuscript.docx
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import difflib
import os

V1_PATH = "version2/CKI_manuscript_submit_V1/CKI_NAR_Manuscript.docx"
TARGET_PATH = "version2/CKI_NAR_Manuscript_v4.docx"
CLEAN_BASELINE = "version2/_clean_baseline.docx"
OUTPUT = "version2/NAR_Submission/CKI_NAR_Manuscript.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ============================================================
# Phase 1: Accept V1 Track Changes
# ============================================================
print("Phase 1: Accept V1 Track Changes")
doc = Document(V1_PATH)
body = doc.element.body

ins_count = 0
while True:
    ins_elements = body.findall('.//{%s}ins' % W)
    if not ins_elements: break
    for ins in ins_elements:
        parent = ins.getparent()
        if parent is None: continue
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, deepcopy(child))
            idx += 1
        parent.remove(ins)
        ins_count += 1

del_count = 0
while True:
    del_elements = body.findall('.//{%s}del' % W)
    if not del_elements: break
    for de in del_elements:
        parent = de.getparent()
        if parent is not None:
            parent.remove(de)
            del_count += 1

for rsid in body.findall('.//{%s}rsidRPr' % W):
    rsid.getparent().remove(rsid)

# Fix acceptance artifacts
paras = doc.paragraphs
for p in paras:
    if "phenomenon.While" in p.text and "phenomenon. While" not in p.text:
        for run in p.runs:
            if run.text.strip().endswith("phenomenon."):
                run.text = run.text + " "
                break

for p in paras:
    txt = p.text
    idx_end = txt.rfind("omega-based inference.")
    if idx_end >= 0:
        after = txt[idx_end + len("omega-based inference."):]
        if ", which effectively" in after:
            for run in reversed(p.runs):
                if ", which effectively" in run.text:
                    run.text = run.text.split(", which effectively")[0]
                    break

doc.save(CLEAN_BASELINE)
print(f"  Accepted {ins_count} insertions, {del_count} deletions")
print(f"  Clean baseline: {CLEAN_BASELINE}")

# ============================================================
# Phase 2: Load v4 target
# ============================================================
print("\nPhase 2: Load v4 target manuscript")
target_doc = Document(TARGET_PATH)
target_paras = [p.text for p in target_doc.paragraphs]
baseline_doc = Document(CLEAN_BASELINE)
baseline_paras = [p.text for p in baseline_doc.paragraphs]

print(f"  Baseline: {len(baseline_paras)} paragraphs")
print(f"  Target (v4): {len(target_paras)} paragraphs")

# ============================================================
# Phase 3: Paragraph-align + diff → Track Changes
# ============================================================
print("\nPhase 3: Generate Track Changes markup")

# Align paragraphs
sm = difflib.SequenceMatcher(None, baseline_paras, target_paras)
alignments = []  # (baseline_idx, target_idx) pairs
target_new = []  # target indices that are insertions

for tag, i1, i2, j1, j2 in sm.get_opcodes():
    if tag == 'equal':
        for bi, ti in zip(range(i1, i2), range(j1, j2)):
            alignments.append((bi, ti))
    elif tag == 'replace':
        n = min(i2 - i1, j2 - j1)
        for k in range(n):
            alignments.append((i1 + k, j1 + k))
        # Extra target paragraphs = insertions
        for k in range(n, j2 - j1):
            target_new.append(j1 + k)
    elif tag == 'insert':
        for ti in range(j1, j2):
            target_new.append(ti)

print(f"  Matched pairs: {len(alignments)}")
print(f"  New paragraphs: {len(target_new)}")

# Get paragraph elements
target_p_elems = [p._element for p in target_doc.paragraphs]
baseline_p_elems = [p._element for p in baseline_doc.paragraphs]

tc_counter = [0]

def create_equal_run(text):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def create_del_element(text):
    tc_counter[0] += 1
    de = OxmlElement('w:del')
    de.set(qn('w:id'), str(tc_counter[0]))
    de.set(qn('w:author'), 'WorkBuddy')
    de.set(qn('w:date'), '2026-07-22T00:00:00Z')
    dr = OxmlElement('w:r')
    drPr = OxmlElement('w:rPr')
    dr.append(drPr)
    dt = OxmlElement('w:delText')
    dt.set(qn('xml:space'), 'preserve')
    dt.text = text
    dr.append(dt)
    de.append(dr)
    return de

def create_ins_element(text):
    tc_counter[0] += 1
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(tc_counter[0]))
    ins.set(qn('w:author'), 'WorkBuddy')
    ins.set(qn('w:date'), '2026-07-22T00:00:00Z')
    ir = OxmlElement('w:r')
    irPr = OxmlElement('w:rPr')
    ir.append(irPr)
    it = OxmlElement('w:t')
    it.set(qn('xml:space'), 'preserve')
    it.text = text
    ir.append(it)
    ins.append(ir)
    return ins

def diff_and_replace_para(target_elem, old_text, new_text):
    """Replace paragraph content with Track Changes markup based on text diff."""
    if old_text == new_text:
        return

    sm = difflib.SequenceMatcher(None, old_text, new_text)
    opcodes = sm.get_opcodes()

    new_children = []
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal':
            new_children.append(create_equal_run(old_text[i1:i2]))
        elif tag == 'replace':
            if i2 > i1:
                new_children.append(create_del_element(old_text[i1:i2]))
            if j2 > j1:
                new_children.append(create_ins_element(new_text[j1:j2]))
        elif tag == 'delete':
            new_children.append(create_del_element(old_text[i1:i2]))
        elif tag == 'insert':
            new_children.append(create_ins_element(new_text[j1:j2]))

    # Clear existing runs/ins/del from paragraph, keep pPr
    pPr = target_elem.find('{%s}pPr' % W)
    for child in list(target_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'ins', 'del', 'bookmarkStart', 'bookmarkEnd', 'proofErr'):
            if child is not pPr:
                target_elem.remove(child)

    insert_pos = 0
    if pPr is not None:
        insert_pos = list(target_elem).index(pPr) + 1
    for elem in new_children:
        target_elem.insert(insert_pos, elem)
        insert_pos += 1

# Process matched paragraphs
for bi, ti in alignments:
    if bi >= len(baseline_paras) or ti >= len(target_paras):
        continue
    b_text = baseline_paras[bi]
    t_text = target_paras[ti]
    if b_text != t_text:
        diff_and_replace_para(target_p_elems[ti], b_text, t_text)

# Process new paragraphs (mark as insertion)
for ti in target_new:
    if ti >= len(target_paras):
        continue
    t_text = target_paras[ti]
    if not t_text.strip():
        continue
    diff_and_replace_para(target_p_elems[ti], "", t_text)

print(f"  Track Changes operations: {tc_counter[0]}")

# ============================================================
# Phase 4: Save & verify
# ============================================================
print("\nPhase 4: Save & verify")
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
target_doc.save(OUTPUT)

final = Document(OUTPUT)
final_paras = final.paragraphs
word_count = sum(len(p.text.split()) for p in final_paras if p.text.strip())
para_count = sum(1 for p in final_paras if p.text.strip())
ins_tags = final.element.body.findall('.//{%s}ins' % W)
del_tags = final.element.body.findall('.//{%s}del' % W)

print(f"  Output: {OUTPUT}")
print(f"  Words: ~{word_count}")
print(f"  Paragraphs: {para_count}")
print(f"  Track Changes: {len(ins_tags)} insertions, {len(del_tags)} deletions")

# Content checks
checks = [
    ("FDR/Benjamini", "Benjamini-Hochberg"),
    ("HK sensitivity", "sensitivity analysis by varying"),
    ("Ka/Ks \u03c9=1 neutral point", "\u03c9 = 1 in CKI carries no population-genetic"),
    ("CellTypist [35]", "CellTypist [35]"),
    ("TCGA deconvolution", "Computational deconvolution"),
    ("np.log2 declaration", "np.log2 in Python"),
    ("OPC circularity", "a potential circularity"),
    ("Ontogenetic time scale", "ontogenetic time scale"),
    ("Eleventh limitation", "Eleventh"),
    ("requirements.txt", "requirements.txt"),
    ("v0.3.2", "v0.3.2"),
]
all_ok = True
for label, keyword in checks:
    found = any(keyword in p.text for p in final_paras)
    status = "OK" if found else "MISS"
    if not found: all_ok = False
    print(f"  {status}: {label}")

if all_ok:
    print("\nAll content checks passed!")
else:
    print("\nWARNING: Some content checks failed")

print("\nDone!")
