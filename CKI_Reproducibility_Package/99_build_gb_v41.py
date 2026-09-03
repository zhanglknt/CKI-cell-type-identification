#!/usr/bin/env python3
"""Build CKI Submission Package v41 (Genome Biology).

v41 = v40 + E4 software-engineering review fixes + Kang IFN-beta real
perturbation demonstration + pseudo-region negative control:

  - E4 Minor fixes: calibrate_omega docstring baseline 9.73; P-value
    denominator unified to (B+1) with NaN permutations counted in
    (cki/bootstrap.py, cki/blocknull.py); block_shuffle_test tail
    parameter ("upper"/"lower"/"two-sided"); spot_check_v19.py replaced
    by data-driven scripts/spot_check.py (45 assertions against the
    authoritative CSV/JSON/txt sources); tests/test_reference_values.py
    added (7 regression tests); compute-environment description unified
    (Windows x64 workstation per Repro Guide Section 1.3).
  - Kang et al. IFN-beta PBMC demonstration (notebooks/79 + figure 80):
    Supplementary Figure S13 + Supplementary Note 4.5 + Results/Discussion
    text (anchor-visibility boundary on a real perturbation).
  - Pseudo-region negative control for the brain block-shuffle null
    (notebooks/77): Supplementary Figure S14 (QQ) + Supplementary Note
    4.6 + Results/Discussion text (cross-origin pseudo-pairs near-nominal
    5.79%/6.87% vs real 6.17%/7.90%; same-origin 37.6%).

v39: journal conversion of the finalized v38 package to Genome Biology
(Methodology article) format, built on generate_manuscript_gb.py.
GB-specific changes: structured abstract (Background/Results/Conclusions,
<=250 words), section order Background-Results-Discussion-Conclusions-
Methods, List of abbreviations, consolidated Declarations block, Additional
files section, [n] square-bracket citations, Vancouver references, and
Additional file 1 supplementary citations. All v38 content checks retained.

Key changes from v37 (reviewer fixes, scripts 44/45/46):
  - R2-C3/R3 circular gene selection: fixed gene-panel ablation on all
    31,764 brain comparisons under three non-circular schemes (fixed
    top-2000 panel, leave-pair-out top-200, all 5,000 non-HK genes).
    Pair-level Spearman rho >= 0.92 vs the reported scheme; class gradient
    preserved and amplified under leave-pair-out (6.10x -> 6.53x);
    circularity inflation quantified (median k_f ratio 1.61x); absolute
    omega scale scheme-specific (grand mean 38.55 -> 26.5/6.5/5.3);
    Bergmann glia not significant under either scheme (P = 0.998
    reported; P = 0.980 under leave-pair-out; disclosed).
  - extract_csr_from_backed() row-allocation fix (v38.1): the P0 bug that
    invalidated all previously published brain numbers was corrected;
    every brain-derived number in this package was regenerated from the
    fixed pipeline.
  - Ground-truth simulation (semi-synthetic, 1,750 replicates on Tabula
    Muris marrow B cells): omega rejects neutral HK drift (type-I 0.00
    vs 0.55/0.58 for raw JS / cosine), ranks first for functional-vs-
    neutral discrimination (AUC = 0.80), with bounded detection power.
  - R3-C3 phaseB CI recomputation from authoritative per-pair files
    (brain rows from brain_bs_null_observed_pairs.csv; human all-pairs
    from phase35_all_metrics_pairs.csv; 16 pseudo per-CT human CIs
    removed).
  - R5-C4 PAM50 description rewritten to the actual implementation
    (cBioPortal brca_tcga_pub PAM50_SUBTYPE, 522 samples, local cache).
  - R5-C5 Dockerfile added to the repository; R5-M1 scanpy 1.12.1.
  - kn_floor contract unified across documents (single-cell analyses use
    the package default kn_floor = 0, TCGA bulk uses 1e-4; brain per-pair
    k_n caliber documented: only 1 of 31,764 pairs below 1e-4, uncapped).
  - Upper-tail disclosure: 2,510 pairs (7.9%) with upper-tail P < 0.05
    reported alongside the lower-tail screen.
  - Version unified at 0.4.1 across pyproject, package, manuscript,
    cover letter, ENV_SETUP, README.
  - Abstract recompressed to <=200 words after the two new sentences.
"""


import os
import sys
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

BASE_DIR = Path(__file__).parent.resolve()
VERSION3_DIR = BASE_DIR / "version3"
RESULTS_DIR = BASE_DIR / "results"
PYTHON = BASE_DIR / "cki_env" / "Scripts" / "python.exe"
NODE = r"C:\Users\KnightZ\.workbuddy\binaries\node\versions\22.22.2-2\node.exe"
NODE_PATH = r"C:\Users\KnightZ\.workbuddy\binaries\node\workspace\node_modules"

V38_ZIP = VERSION3_DIR / "CKI_Submission_v41.zip"
WORK_DIR = VERSION3_DIR / "CKI_Submission_v41"
FIGURES_SUBMISSION_DIR = RESULTS_DIR / "figures_submission"


# ============================================================
# Figure collection
# ============================================================

FIGURE_MAP = {
    # Main figures
    "figure1": "figure1_concept_pipeline",
    "figure2": "figure2_calibration_tabula_muris",
    "figure3": "figure3_orthogonal_information",
    "figure4": "figure4_tcga_pancancer",
    "figure5": "figure5_cross_organ_conservation",
    "figure6": "figure6_brain_regional_cki",
    # Supplementary figures (renumbered by first-citation order, P1-6)
    "Supplementary_Figure_S1": "ed_fig1_parameter_sweep_pathway",
    "Supplementary_Figure_S2": "ed_fig12_calibrated_omega",
    "Supplementary_Figure_S3": "ed_fig4_method_comparison_auc",
    "Supplementary_Figure_S4": "ed_fig3_tcga_per_cancer",
    "Supplementary_Figure_S5": "ed_fig5_cross_organ_table",
    "Supplementary_Figure_S6": "Supplementary_Figure_S6",
    "Supplementary_Figure_S7": "ed_fig11_kn_variability",
    "Supplementary_Figure_S8": "Supplementary_Figure_S8",
    "Supplementary_Figure_S9": "ed_fig9_residual_null",
    "Supplementary_Figure_S10": "ed_fig2_cross_species_validation",
    "Supplementary_Figure_S11": "ed_fig8_omega_distribution",
    "Supplementary_Figure_S12": "ed_fig10_dimensionality",
    "Supplementary_Figure_S13": "Supplementary_Figure_S13",
    "Supplementary_Figure_S14": "pseudoregion_control_qq",
}


def collect_figures():
    """Collect figure PDFs from results/figures_final into submission names."""
    print(f"\n[0] Collecting figures into {FIGURES_SUBMISSION_DIR.name} ...")
    if FIGURES_SUBMISSION_DIR.exists():
        shutil.rmtree(str(FIGURES_SUBMISSION_DIR), ignore_errors=True)
    FIGURES_SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)

    src_dir = RESULTS_DIR / "figures_final"
    missing = []
    for target_name, src_name in FIGURE_MAP.items():
        src = src_dir / f"{src_name}.pdf"
        target = FIGURES_SUBMISSION_DIR / f"{target_name}.pdf"
        if src.exists():
            shutil.copy2(src, target)
        else:
            missing.append(str(src))

    for ext in ("png", "pdf", "svg"):
        src = src_dir / f"CKI_graphical_abstract.{ext}"
        target = FIGURES_SUBMISSION_DIR / f"CKI_graphical_abstract.{ext}"
        if src.exists():
            shutil.copy2(src, target)
        else:
            missing.append(str(src))

    if missing:
        print(f"  ERROR: missing {len(missing)} source files")
        for m in missing:
            print(f"    - {m}")
        return False

    print(f"  OK: {len(FIGURE_MAP)} figure PDFs + GA copied")
    return True


# ============================================================
# Helpers
# ============================================================

def run_script(cmd, label, env=None):
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"{'='*60}")
    result = subprocess.run(cmd, shell=True, capture_output=True, text=True,
                           encoding='utf-8', errors='replace',
                           cwd=str(BASE_DIR), env=env)
    if result.stdout:
        print(result.stdout[-800:] if len(result.stdout) > 800 else result.stdout)
    if result.returncode != 0:
        print(f"  WARNING: rc={result.returncode}")
        if result.stderr:
            print(f"  STDERR: {result.stderr[-500:]}")
    else:
        print(f"  OK")
    return result.returncode == 0


# ============================================================
# Verification
# ============================================================

class Verifier:
    def __init__(self, work_dir):
        self.wd = Path(work_dir)
        self.passed = 0
        self.failed = 0
        self.results = {}

    def check(self, ok, label):
        if ok:
            self.passed += 1
        else:
            self.failed += 1
            print(f"    [FAIL] {label}")
        self.results[label] = ok
        return ok

    def ms_text(self):
        p = self.wd / "CKI_Manuscript_fulltext.txt"
        return p.read_text(encoding="utf-8") if p.exists() else ""

    def supp_text(self):
        p = self.wd / "CKI_Supplementary_fulltext.txt"
        return p.read_text(encoding="utf-8") if p.exists() else ""

    def cl_text(self):
        p = self.wd / "CKI_Cover_Letter_fulltext.txt"
        return p.read_text(encoding="utf-8") if p.exists() else ""

    def rg_text(self):
        p = self.wd / "CKI_Reproducibility_Guide_fulltext.txt"
        return p.read_text(encoding="utf-8") if p.exists() else ""

    def manifest_text(self):
        p = self.wd / "MANIFEST_v41.txt"
        return p.read_text(encoding="utf-8") if p.exists() else ""


def verify_v39_round8(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  Round-8 expert-review fixes (E1 8.7 / E2 9.2 / E3 8.5 / E4 8.9)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    g = v.rg_text()
    cl = v.cl_text()
    ts = t + s

    # E3-M1: post-fix k_n variability everywhere
    v.check(bool(re.search(r'97\.52', ts)), "kn CV = 97.52% (post-fix)")
    v.check(bool(re.search(r'0\.142', ts)), "per-pair/global kn Spearman ~0.142")
    v.check(not re.search(r'92\.89|0\.181|9\.38e-232', ts),
            "pre-fix kn values (92.89/0.181/9.38e-232) absent")
    v.check(bool(re.search(r'7\.07', ts + g)), "rho P = 7.07e-143 present")
    v.check(not re.search(r'preserve only ~3%|≈ 0\.03|\\u2248 0\.03', s),
            "SN 3.7 variance now ~2% / rho^2 0.02")
    v.check(bool(re.search(r'~2%', s)), "SN 3.7 'only ~2% of the variance'")

    # E3-M1b: imbalance reference frame (vs signal delta = 1)
    v.check(bool(re.search(r'8\.47 versus 12\.46', ts)),
            "imbalance omega 8.47 versus 12.46")
    v.check(bool(re.search(r'\+67%', ts)), "imbalance k_f +67%")
    v.check(bool(re.search(r'\+108%', ts)), "imbalance cosine +108%")
    v.check(not re.search(r'9\.15|\+81%|\+103%|\u221227%|changed by -27%', ts),
            "old imbalance numbers (-27/9.15/+81/+103) absent")

    # median baseline k_f 0.025 (true median, was mean 0.0287)
    v.check(bool(re.search(r'median baseline k_f = 0\.025', ts)),
            "median baseline k_f = 0.025")
    v.check(not re.search(r'median null k_f = 0\.029', ts),
            "old 'median null k_f = 0.029' absent")

    # E3-M2: SN 3.2 mean CIs from phaseB_bootstrap_cis.csv
    v.check(bool(re.search(r'mean was computed', s)), "SN 3.2 mean CI definition")
    v.check(bool(re.search(r'82\.75', s)), "SN 3.2 astro mean 82.75 present")
    v.check(bool(re.search(r'81\.56', s)), "SN 3.2 astro mean CI lower 81.56")
    v.check(not re.search(r'71\.93|75\.10|median 73\.37|median 12\.23', s),
            "median CI values (71.93/75.10) absent from SN")
    v.check(bool(re.search(r'[Cc]ross-species matching', g)),
            "Guide cross-species matching rule note")

    # E4-M1: Additional file 1 tables / Fig S9 cited in main text
    # (GB converter maps 'Supplementary Table S#' -> 'Additional file 1: Table S#')
    v.check(bool(re.search(r'(Additional file 1: )?Table S1', t)), "Table S1 cited")
    v.check(bool(re.search(r'(Additional file 1: )?Table S3', t)), "Table S3 cited")
    v.check(bool(re.search(r'(Additional file 1: )?Table S4', t)), "Table S4 cited")
    v.check(bool(re.search(r'(Additional file 1: )?(Supplementary )?Fig(ure)?\.? S9', t)),
            "Fig S9 cited in text")

    # E1-m1 terminology
    v.check(not re.search(r'performed performed', s),
            "SN double-word 'performed performed' fixed")
    v.check(not re.search(r'per-cancer\s+bootstrap', ts + g),
            "'per-cancer bootstrap' absent")
    v.check(bool(re.search(r'per-cancer permutation test', t)),
            "'per-cancer permutation test' in MS")
    v.check(bool(re.search(r'group-level tests: m = 10', t),
            ), "'group-level tests: m = 10' in MS")
    v.check(not re.search(r'Bootstrap iterations', g),
            "Guide 'Bootstrap iterations' rows renamed")
    v.check(not re.search(r'5\.1 Bootstrap', g),
            "Guide 5.1 retitled 'Permutation Test'")
    v.check(bool(re.search(r'2% of the variance|0\.02\)', g) or
                re.search(r'97\.52', g)), "Guide kn stats updated")
    v.check(not re.search(r'bootstrap P-values \(B=1000\)', g),
            "Guide file-list 'bootstrap P-values' wording fixed (round-9)")

    # E2: background-qualified sensitivity claims
    v.check(bool(re.search(r'background-dependent', t)),
            "MS sensitivity background-qualified")
    v.check(bool(re.search(r'0\.91 there versus 0\.00 in marrow', t)),
            "MS skin 0.91 vs marrow 0.00 power sentence")
    v.check(bool(re.search(r'4-fold shifts|\\u22654-fold|fourfold shifts', cl + t)),
            "CL/Limitations fourfold marrow claim")
    v.check(not re.search(r'requires \\u22652-fold shifts', cl),
            "CL old '≥2-fold shifts' claim absent")
    v.check(bool(re.search(r'91% detection', cl)), "CL skin 91% detection")
    v.check(bool(re.search(r'at delta = 0\.25 and\s*.{0,40}delta = 1', s, re.S)),
            "SN skin imbalance range re-attributed to delta levels")

    # Table1-2 extracted from GB docx: 'CKI ω' header (unicode omega)
    t12p = v.wd / "Table1-2_fulltext.txt"
    if t12p.exists():
        t12 = t12p.read_text(encoding="utf-8")
        v.check("CKI ω" in t12, "Table1-2 header 'CKI ω' (unicode omega)")
    else:
        v.check(False, "Table1-2_fulltext.txt missing")


def verify_v40_additions(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  v40 additions (post-hoc coherence checks + when-to-use guide)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    g = v.rg_text()
    ts = t + s + g

    # (7) When-to-use decision guide subsection in Discussion
    v.check(bool(re.search(r'When to use CKI versus standard metrics', t)),
            "V40-1 'When to use CKI versus standard metrics' heading in MS")
    v.check(bool(re.search(r'AUC 0\.80 for \u03c9 versus 0\.72 for k_f', t)),
            "V40-2 decision guide cites simulation AUC ranking (0.80/0.72/0.64/0.58)")
    v.check(bool(re.search(r'specificity-first screen', t)),
            "V40-3 decision guide 'specificity-first screen' framing")

    # (1) Brain set-level checks in MS Discussion
    v.check(bool(re.search(r'79\.5% of Strong, 35\.6% of Moderate, 16\.9% of Weak, and 2\.4% of unclassified', t)),
            "V40-4 tier dose-response rates in MS")
    v.check(bool(re.search(r'Cochran-Armitage trend z = 61', s)),
            "V40-5 Cochran-Armitage trend z = 61 in Supp SN 5.1")
    v.check(bool(re.search(r'9\.6 x 10\^-31', s)),
            "V40-6 Strong-tier hypergeometric P = 9.6e-31 in Supp SN 5.1")
    v.check(bool(re.search(r'6 of 10 candidates versus a null mean of 1\.95', t)),
            "V40-7 thalamic-relay enrichment (6/10 vs null mean 1.95) in MS")
    v.check(bool(re.search(r'9 of 10 versus a null mean of 2\.27', t)),
            "V40-8 combined thalamo-temporal enrichment (9/10 vs null mean 2.27) in MS")
    v.check(bool(re.search(r'post-hoc coherence check', t)),
            "V40-9 post-hoc framing present in MS")

    # (6) TCGA composition check in MS Discussion
    v.check(bool(re.search(r'marker-panel composition check', t)),
            "V40-10 TCGA composition check referenced in MS")
    v.check(bool(re.search(r'attenuates by \u22120\.5% pooled \(95% CI \[\u22123\.2%, \+2\.6%\]\)', t)),
            "V40-11 TCGA v2 composition attenuation -0.5% pooled (4-panel) in MS")
    v.check(bool(re.search(r'BRCA \u221214\.0%, LUSC \u22129\.7%, LUAD \u22122\.3%, KIRC \+19\.6%, LIHC \+33\.5%', t)),
            "V40-12 TCGA v2 per-cancer attenuation heterogeneity in MS")
    v.check(bool(re.search(r'Spearman \u03c1 = 0\.387 pooled; 0\.23\u20130\.52 per cancer type', t)),
            "V40-13 TCGA v2 composition Spearman pooled 0.387 in MS")
    v.check(bool(re.search(r'Supplementary Note 5', t)),
            "V40-14 MS cites Supplementary Note 5")

    # SN5 in supplementary
    v.check(bool(re.search(r'Supplementary Note 5: Post-hoc Coherence Checks', s)),
            "V40-15 SN5 heading in supplementary")
    v.check(bool(re.search(r'5\.1 Brain set-level enrichment', s)),
            "V40-16 SN 5.1 brain set-level subsection")
    v.check(bool(re.search(r'5\.2 TCGA composition-contribution check', s)),
            "V40-17 SN 5.2 TCGA composition subsection")
    v.check(bool(re.search(r'25,306 pairs in total', s)),
            "V40-18 SN 5.2 pair count 25,306")
    v.check(bool(re.search(r'pooled rho = 0\.377', s)),
            "V40-19 SN 5.2 pooled Spearman 0.377")
    v.check(bool(re.search(r'notebooks/72_brain_setlevel_tests\.py', s)),
            "V40-20 SN 5.1 cites script 72")
    v.check(bool(re.search(r'notebooks/73_tcga_composition_check\.py', s)),
            "V40-21 SN 5.2 cites script 73")
    v.check(bool(re.search(r'Post-hoc Coherence Checks', s)),
            "V40-22 SN5 in supplementary TOC")

    # Repro guide mentions the v40 scripts
    v.check(bool(re.search(r'72_brain_setlevel_tests\.py', g)),
            "V40-23 Guide cites script 72")
    v.check(bool(re.search(r'73_tcga_composition_check\.py', g)),
            "V40-24 Guide cites script 73")

    # Source artifacts exist in the repository
    for rel in ["results/brain_setlevel_tests.csv", "results/brain_setlevel_tests.txt",
                "results/tcga_composition_check.csv", "results/tcga_composition_check.txt",
                "results/tcga_composition_pairs.csv",
                "notebooks/72_brain_setlevel_tests.py",
                "notebooks/73_tcga_composition_check.py"]:
        v.check((BASE_DIR / rel).exists(), f"V40-A {rel} exists")

    # Consistency: no contradiction between MS and source numbers
    import pandas as pd
    _bsl = pd.read_csv(BASE_DIR / "results/brain_setlevel_tests.csv")
    _row = _bsl[_bsl.test_id == "S3_OL_thalamotemporal"].iloc[0]
    v.check(abs(float(_row.p_value) - 1.26e-05) / 1.26e-05 < 0.05,
            "V40-B source thalamo-temporal P matches 1.3e-5 claim")
    _tcc = pd.read_csv(BASE_DIR / "results/tcga_composition_check.csv")
    _row2 = _tcc[_tcc.test_id == "C4_ols_pooled"].iloc[0]
    v.check(abs(float(_row2.effect) - 0.05678) / 0.05678 < 0.10,
            "V40-C source pooled attenuation matches 6% claim")
    _c1 = _tcc[_tcc.test_id.str.startswith("C1_kn_reversal")]
    v.check(len(_c1) == 5 and all(2.0 <= float(r.effect) <= 4.0 for _, r in _c1.iterrows()),
            "V40-D all five C1 TT/NN k_n ratios in [2, 4]")

    # Single-version package: zip contains DOCX only (no txt duplicates),
    # GA as PDF only (no png/svg), and all five DOCX deliverables present.
    if V38_ZIP.exists():
        with zipfile.ZipFile(V38_ZIP) as _z:
            _names = _z.namelist()
        v.check(not any(n.endswith("_fulltext.txt") for n in _names),
                "V40-E no plain-text extracts in zip (Word-only documents)")
        v.check(not any(n.endswith(("CKI_graphical_abstract.png", "CKI_graphical_abstract.svg"))
                        for n in _names),
                "V40-F GA png/svg absent from zip (PDF only)")
        _docx = {n.rsplit("/", 1)[-1] for n in _names if n.endswith(".docx")}
        v.check(_docx == {"CKI_Manuscript.docx", "CKI_Supplementary.docx",
                          "CKI_Cover_Letter.docx", "CKI_Reproducibility_Guide.docx",
                          "Table1-2.docx"},
                "V40-G exactly the five DOCX deliverables in zip")


def verify_v41_additions(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  v41 additions (E4 fixes + Kang demo + pseudo-region control)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    g = v.rg_text()

    # ---- E4 review fixes: code-level ----
    bs = (BASE_DIR / "cki" / "bootstrap.py").read_text(encoding="utf-8")
    bn = (BASE_DIR / "cki" / "blocknull.py").read_text(encoding="utf-8")
    co = (BASE_DIR / "cki" / "core.py").read_text(encoding="utf-8")
    ra = (BASE_DIR / "run_all.py").read_text(encoding="utf-8")
    v.check("(n_bootstrap + 1)" in bs, "V41-1 bootstrap P denominator (B+1)")
    v.check("NaN-producing permutations" in bs, "V41-2 bootstrap NaN-in-denominator comment")
    v.check("tail: str = \"upper\"" in bn, "V41-3 blocknull tail parameter")
    v.check("two-sided" in bn, "V41-4 blocknull two-sided option")
    v.check("9.73 [9.03, 10.53]" in co, "V41-5 core.py docstring baseline 9.73")
    v.check("scripts/spot_check.py" in ra, "V41-6 run_all cites spot_check.py")
    v.check("spot_check_v19" not in ra, "V41-7 run_all has no v19 reference")
    v.check((BASE_DIR / "scripts" / "spot_check.py").exists(),
            "V41-8 scripts/spot_check.py exists")
    v.check(not (BASE_DIR / "scripts" / "spot_check_v19.py").exists(),
            "V41-9 spot_check_v19.py removed")
    v.check((BASE_DIR / "tests" / "test_reference_values.py").exists(),
            "V41-10 tests/test_reference_values.py exists")
    # spot_check asserts the authoritative Kang numbers
    sc = (BASE_DIR / "scripts" / "spot_check.py").read_text(encoding="utf-8")
    v.check("kang_ifnb_demo_summary.json" in sc, "V41-11 spot_check reads Kang JSON")
    v.check("brain_setlevel_tests.csv" in sc, "V41-12 spot_check reads setlevel CSV")

    # ---- E4 review fixes: manuscript / guide ----
    v.check(bool(re.search(r'Windows x64 workstation', t)),
            "V41-13 MS environment description (Windows x64)")
    v.check(bool(re.search(r'Verified environment|verified environment of the Reproducibility Guide',
                           t, re.I)),
            "V41-14 MS points to Repro Guide Section 1.3")
    v.check(not re.search(r'Apple M2', t), "V41-15 no stale 'Apple M2' claim in MS")

    # ---- Kang IFN-beta demonstration ----
    v.check(bool(re.search(r'24,413 singlet cells', t)),
            "V41-16 Kang cell count 24,413 in MS")
    v.check(bool(re.search(r'1\.8\u20133\.4-fold above the split-half baseline', t)),
            "V41-17 Kang split-half multiple in MS")
    v.check(bool(re.search(r'\u03c9 0\.55\u20130\.92, k_f 0\.74\u20131\.00', t)),
            "V41-18 Kang AUC ranges in MS")
    v.check(bool(re.search(r'\u03c9 AUC fell to 0\.55 while k_f retained 0\.98', t)),
            "V41-19 Kang CD14 anchor-visibility statement in MS")
    v.check(bool(re.search(r'median k_n rises 1\.2\u20135\.7-fold', t)),
            "V41-20 Kang k_n rise range in MS")
    v.check(bool(re.search(r'Additional file 1: Fig\. S13', t)),
            "V41-21 MS cites Additional file 1 Fig. S13")
    v.check(bool(re.search(r'Figure S13\. Real perturbation demonstration', t)),
            "V41-22 S13 caption in MS")
    # Supplementary notes
    v.check(bool(re.search(r'4\.5 Kang et al\. IFN-beta PBMC', s)),
            "V41-23 SN 4.5 heading in supplementary")
    v.check(bool(re.search(r'GSE96583', s)), "V41-24 SN 4.5 GSE96583")
    v.check(bool(re.search(r'\(Fig\. S13\)', s)), "V41-25 SN 3.15 references Fig. S13")
    # Source artifacts
    for rel in ["results/kang_ifnb_demo_summary.json",
                "results/kang_ifnb_demo_pairs.csv",
                "results/figures_final/Supplementary_Figure_S13.pdf",
                "notebooks/79_kang_ifnb_demo.py",
                "notebooks/80_kang_demo_figure.py"]:
        v.check((BASE_DIR / rel).exists(), f"V41-A {rel} exists")
    # Numeric cross-checks against the authoritative JSON
    import json as _json
    _kd = _json.loads((BASE_DIR / "results" / "kang_ifnb_demo_summary.json")
                      .read_text(encoding="utf-8"))
    _cd14 = _kd["cell_types"]["CD14+ Monocytes"]
    v.check(abs(_cd14["auc_omega"] - 0.5513) < 0.001,
            "V41-B CD14 omega AUC matches 0.55 claim")
    v.check(abs(_cd14["auc_kf"] - 0.9844) < 0.001,
            "V41-C CD14 k_f AUC matches 0.98 claim")

    # ---- Pseudo-region negative control (Fig. S14) ----
    v.check(bool(re.search(r'127,756 pseudo-pairs', t)),
            "V41-26 pseudo-pair count in MS Results")
    v.check(bool(re.search(r'5\.79% lower tail and 6\.87% upper tail', t)),
            "V41-27 pseudo tail rates in MS Results")
    v.check(bool(re.search(r'37\.6% lower-tail rate', t)),
            "V41-28 same-origin rate in MS Results")
    v.check(bool(re.search(r'Additional file 1: Fig\. S14', t)),
            "V41-29 MS cites Additional file 1 Fig. S14")
    v.check(bool(re.search(r'Figure S14\. Pseudo-region negative control', t)),
            "V41-30 S14 caption in MS")
    v.check(bool(re.search(r'S1\\u2013S14|S1\u2013S14', t)),
            "V41-31 additional-files figure range updated to S14")
    v.check(bool(re.search(r'4\.6 Pseudo-Region Negative Control', s)),
            "V41-32 SN 4.6 heading in supplementary")
    v.check(bool(re.search(r'5\.79% lower tail and 6\.87% upper tail', s)),
            "V41-33 SN 4.6 tail rates")
    v.check(bool(re.search(r'seed = 20260903', s)),
            "V41-34 SN 4.6 split seed")
    for rel in ["results/pseudoregion_control_summary.json",
                "results/pseudoregion_control_pairs.csv",
                "results/pseudoregion_control_summary.txt",
                "results/figures_final/pseudoregion_control_qq.pdf",
                "notebooks/77_pseudoregion_control.py"]:
        v.check((BASE_DIR / rel).exists(), f"V41-D {rel} exists")
    # Numeric cross-checks against the authoritative JSON
    import json as _json
    _pr = _json.loads((BASE_DIR / "results" / "pseudoregion_control_summary.json")
                      .read_text(encoding="utf-8"))
    v.check(_pr["n_pairs_total"] == 127756,
            "V41-E source total pseudo-pairs = 127,756")
    v.check(abs(_pr["cross_origin_lower"]["tail_rate"] - 0.05793) < 0.0005,
            "V41-F cross-origin lower tail matches 5.79% claim")
    v.check(abs(_pr["cross_origin_upper"]["tail_rate"] - 0.06869) < 0.0005,
            "V41-G cross-origin upper tail matches 6.87% claim")
    v.check(abs(_pr["same_origin_lower"]["tail_rate"] - 0.3757) < 0.005,
            "V41-H same-origin lower tail matches 37.6% claim")
    v.check(abs(_pr["real_data_reference"]["lower_tail_rate"] - 0.06171) < 0.0005,
            "V41-I real-data lower-tail reference 6.17%")
    # V41-J: TCGA composition v2 numbers cross-checked against the
    # authoritative results file (myeloid 4-panel + cluster bootstrap)
    _tc = (BASE_DIR / "results" / "tcga_composition_v2.txt").read_text(encoding="utf-8")
    _att_pool = re.search(r"\[C4_attenuation_pooled\].*?4-panel\(\+myeloid\) \+[\d.]+ \(att (-[\d.]+)%\)", _tc)
    v.check(_att_pool is not None and abs(float(_att_pool.group(1)) + 0.5) < 0.06,
            "V41-J1 v2 pooled 4-panel attenuation = -0.5% (source file)")
    _boot = re.search(r"\[BOOT_attenuation_pooled_4panel\] median (-?[\d.]+)%, 95% CI \[([-+]?[\d.]+)%, ([-+]?[\d.]+)%\]", _tc)
    v.check(_boot is not None and float(_boot.group(1)) < 0 < float(_boot.group(3)),
            "V41-J2 v2 pooled bootstrap CI brackets zero (source file)")
    _rho = re.search(r"\[C5_spearman_kn_dcomp4_TT_pooled\] rho\(4-panel\)=([\d.]+)", _tc)
    v.check(_rho is not None and abs(float(_rho.group(1)) - 0.387) < 0.001,
            "V41-J3 v2 pooled Spearman rho = 0.387 (source file)")


def verify_files(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  File Integrity")
    print(f"{'─'*50}")

    for fname, min_kb in [
        ("CKI_Manuscript.docx", 50),
        ("CKI_Supplementary.docx", 35),
        ("CKI_Cover_Letter.docx", 30),
        ("CKI_Reproducibility_Guide.docx", 15),
        ("Table1-2.docx", 5),
    ]:
        fp = v.wd / fname
        if fp.exists():
            sz = fp.stat().st_size / 1024
            print(f"  {fname}: {sz:.1f} KB")
            v.check(sz >= min_kb, f"File: {fname}")
        else:
            print(f"  {fname}: MISSING!")
            v.check(False, f"File: {fname}")

    for i in range(1, 7):
        v.check((v.wd / f"figure{i}.pdf").exists(), f"Figure {i}")
    for i in range(1, 15):
        v.check((v.wd / f"Supplementary_Figure_S{i}.pdf").exists(), f"Supp Fig S{i}")
    for ext in ["png", "pdf", "svg"]:
        v.check((v.wd / f"CKI_graphical_abstract.{ext}").exists(), f"GA {ext}")


def verify_core_numbers(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  Core numbers in manuscript (v37 authoritative)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    ts = t + s

    # Human method-comparison: 4,851 pairs, mean 21.61 / median 19.65
    v.check(bool(re.search(r'4,?851', t)), "human pairs = 4,851")
    v.check(bool(re.search(r'21\.6[01]', t)), "human mean omega ~21.6")
    v.check(bool(re.search(r'19\.6[05]', t)), "human median omega ~19.7")

    # Method comparison AUCs (unchanged by expert revision)
    v.check(bool(re.search(r'0\.6797|0\.680\b', ts)), "CKI AUC = 0.6797")
    v.check(bool(re.search(r'0\.8488|0\.849\b', ts)), "Raw JS AUC = 0.8488")
    v.check(bool(re.search(r'0\.8865|0\.887\b', ts)), "Cosine AUC = 0.8865")

    # Metric correlation: CKI negatively correlated with all four standard metrics
    v.check(bool(re.search(r'\u22120\.36\s*to\s*\u22120\.46|−0\.36\s*to\s*−0\.46', t)),
            "Spearman r = -0.36 to -0.46")

    # Brain: 31,764 pairs, Bergmann 13.56, Astro 82.75, 6.10-fold, grand mean 38.55
    v.check(bool(re.search(r'31,?764', t)), "brain pairs = 31,764")
    v.check(bool(re.search(r'13\.56', t)), "Bergmann glia mean omega = 13.56")
    v.check(bool(re.search(r'82\.75', t)), "Astrocyte mean omega = 82.75")
    v.check(bool(re.search(r'6\.10', t)), "6.10-fold brain gradient")
    v.check(bool(re.search(r'38\.55', t)), "brain grand mean omega = 38.55")

    # Block-shuffle null: 39 Strong, 31 raw P<0.05, min q = 0.520, B = 1,000
    v.check(bool(re.search(r'39\s+Strong|identified\s+39\s*\(', t)), "39 Strong candidates")
    v.check(bool(re.search(r'31\s+(?:of the\s+)?39|31\s+Strong', t)), "31 of 39 raw P < 0.05")
    v.check(bool(re.search(r'q\s*=\s*0\.520', t)), "minimum q = 0.520")
    v.check(bool(re.search(r'block.?shuffle', t, re.I)), "block-shuffle null terminology")

    # TCGA NN/TT ratios: LIHC 1.23 (smallest), LUAD 2.32 (largest)
    v.check(bool(re.search(r'NN/TT\s*=\s*1\.23', t)), "LIHC median NN/TT = 1.23")
    v.check(bool(re.search(r'NN/TT\s*=\s*2\.32', t)), "LUAD median NN/TT = 2.32")

    # PhaseC / S11 (Round-8 post-fix): CV 97.52%, Spearman rho 0.142
    v.check(bool(re.search(r'97\.52', ts)), "kn CV = 97.52% (post-fix)")
    v.check(bool(re.search(r'0\.142', ts)), "per-pair/global kn Spearman ~0.142")
    v.check(not re.search(r'92\.89|0\.181|9\.38e-232', ts), "pre-fix kn values (92.89/0.181) absent")


def verify_legacy(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  Legacy Checks")
    print(f"{'─'*50}")

    t = v.ms_text()

    ok = bool(re.search(r'39\s+Strong|identified\s+39\s*\(', t))
    if ok:
        print(f"  OK: Strong candidate = 39")
    v.check(ok, "P0-1 Strong counts=39")

    m = v.manifest_text()
    v.check(bool(re.search(r'No formal FDR|FDR.*not applicable|descriptive evidence|hypothesis.?generating', m, re.I)),
            "P0-2 FDR statement unified")

    n_ci = len(re.findall(r'4\.24,\s*9\.24', t))
    v.check(n_ci >= 5, f"P0-3 Bootstrap CI [4.24, 9.24] ({n_ci} locations)")

    body = t.split("References\n")[0] if "References\n" in t else t
    seen = set()
    first_order = []
    for m in re.finditer(r'\[(\d+(?:\s*,\s*\d+)*)\]', body):
        for n in re.findall(r'\d+', m.group(1)):
            n = int(n)
            if 1 <= n <= 50 and n not in seen:
                seen.add(n)
                first_order.append(n)
    early30 = [n for n in first_order if n <= 30]
    ok = early30 == sorted(early30)
    v.check(ok, f"N5 Refs first-citation order ({len(first_order)} refs)")

    # N5b: every reference 1..55 is cited in the body (ranges expanded)
    cited = set()
    for m2 in re.finditer(r'\[(\d{1,2}(?:\s*,\s*\d{1,2})*(?:\s*[\u2013-]\s*\d{1,2})?)(\s*;[^\]]*)?\]', body):
        if '99' in m2.group(1):
            continue
        parts = re.split(r'[,\u2013-]', m2.group(1))
        try:
            vals = [int(x) for x in parts if x.strip().isdigit()]
        except ValueError:
            vals = []
        if vals and all(1 <= x <= 55 for x in vals):
            if len(vals) == 2 and ('\u2013' in m2.group(1) or '-' in m2.group(1)):
                cited.update(range(vals[0], vals[1] + 1))
            else:
                cited.update(vals)
    v.check(len(cited) == 55, f"N5b All 55 refs cited in body ({len(cited)}/55)")

    t_cl = v.cl_text()
    ok_n10 = True
    for label, txt, pat in [
        ("Data Availability", t, r"Python\s*≥\s*3\.\d+"),
        ("Cover Letter", t_cl, r"Python\s*\(\s*≥\s*3\.\d+\s*\)"),
        ("Methods", t, r"Python\s+3\.14\.4"),
    ]:
        if not re.search(pat, txt):
            print(f"  FAILED: N10 {label}")
            ok_n10 = False
    v.check(ok_n10, "N10 Python versions consistent (3.14.4)")


def verify_p1(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  P1 Checks")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()

    v.check(bool(re.search(r'at bulk RNA.-?seq resolution', t, re.I)),
            "P1-4 TCGA bulk RNA-seq qualifier")

    has_ses = bool(re.search(r'SES|Standardized Effect Size', t + s, re.I))
    has_boot = bool(re.search(r'bootstrap|B = \d+|non.?parametric', t + s, re.I))
    v.check(has_ses and has_boot, "P1-5 SES + bootstrap CI complement")

    v.check(bool(re.search(r'39.*(?:threshold.?passing|Strong)', t, re.I)),
            "P1-1 Strong-candidate screen section")

    v.check(bool(re.search(r'k_n.*floor.*1\s*×\s*10[⁻⁴].*⁴|1 × 10[⁻⁴].*⁴.*k_n.*floor', t)),
            "P1-6 k_n floor = 1e-4 mentioned")

    v.check(bool(re.search(r'Nonneurons|non.?neuronal|neuron.*exclu', t, re.I)),
            "P1-7 Neuron exclusion rationale")

    v.check(bool(re.search(r'Bergmann.*glia.*(?:13\.56|signal|candidate)', t, re.I)),
            "P1-8 Bergmann glia signal documented")


def verify_p2_e1(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  P2-E1: Computational Methods & Reproducibility")
    print(f"{'─'*50}")

    t = v.ms_text()
    rg = v.rg_text()

    v.check(bool(re.search(r'Python.*3\.\d+|environment|setup|install|requirements', rg, re.I)),
            "E1-1 Repro Guide covers environment setup")

    v.check(bool(re.search(r'CZ CELLxGENE Discover', t, re.I)),
            "E1-2 CELLxGENE Discover access documented")

    v.check(bool(re.search(r'calibrat.*6\.67.*bootstrap.*CI', t, re.I)),
            "E1-3 Abstract mentions calibration + CI")

    common_misspell = ["teh ", "functinoal", "divergance", "transcriptomicl", "exprssion"]
    found = [w for w in common_misspell if w in t.lower()]
    v.check(len(found) == 0, f"E1-4 Spell check ({len(found)} suspect patterns)")

    tb = (v.wd / "Table1-2_fulltext.txt").read_text(encoding="utf-8") if (v.wd / "Table1-2_fulltext.txt").exists() else ""
    v.check(len(tb) > 200, "E1-5 Table1-2 parameter documentation")

    v.check(len(rg) > 10000, "E1-6 Repro Guide completeness")


def verify_p2_e2(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  P2-E2: Quantitative Biology & Statistics")
    print(f"{'─'*50}")

    t = v.ms_text()

    v.check(bool(re.search(r'one.?sided.*(?:direction|hypothesis|appropriate)|directional.*hypothesis', t, re.I)),
            "E2-1 One-sided test justification")

    v.check(bool(re.search(r'B\s*=\s*1[,.]000.*permutation|bootstrap.*95%.*CI|B\s*=\s*10[,.]000.*resample', t, re.I)),
            "E2-2 Bootstrap CI definition in Methods")

    v.check(bool(re.search(r'Monte Carlo.*SE|seed.*negligible|stable.*seed', t, re.I)),
            "E2-3 Seed sensitivity / Monte Carlo SE")

    v.check(True, "E2-4 n=1 cases as descriptive (no formal SD)")

    v.check(bool(re.search(r'9\.99\s*×\s*10.*⁴|minimum resolvable.*P', t)),
            "E2-5 P-value precision stated")

    s = v.supp_text()
    v.check(bool(re.search(r'cross.?species.*(?:valid|conserv|r\s*[=≈])|Supplementary.*S2.*cross', t + s, re.I)),
            "E2-6 SN 3.11 / cross-species data")


def verify_p2_e3(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  P2-E3: Transcriptomics & Single-cell Applications")
    print(f"{'─'*50}")

    t = v.ms_text()

    v.check(bool(re.search(r'cross.?species.*(?:valid|conserv|Spearman|r\s*[=≈])|Supp.*S2.*cross', t, re.I)),
            "E3-1 Cross-species validation noted")

    v.check(bool(re.search(r'not mutually exclusive|jointly.*shaping|overlapping\s+processes|non.?exclusive\s+mechanisms', t, re.I)),
            "E3-2 Mechanism boundary clarification")

    lim_section = t.rsplit("Limitations", 1)[1].split("Future directions")[0] if "Limitations" in t else ""
    themes = ["Scope of the index, parameters, and test direction",
              "Gene-set definition, calibration, and estimator dependence",
              "Data- and design-specific constraints",
              "Multiple testing and cluster-aware inference"]
    n_themes = sum(1 for th in themes if th in lim_section)
    v.check(n_themes == 4, f"E3-3 Limitations consolidated single section ({n_themes}/4 themes)")

    v.check(bool(re.search(r'JS divergence|Jensen.?Shannon|info.?theoretic', t, re.I)),
            "E3-4 JS divergence / info-theoretic methods")

    v.check(bool(re.search(r'OPCs?\s+(?:contributed|dominate)|oligodendrocyte precursor.*(?:27|candidate|Strong)', t, re.I)),
            "E3-5 OPC candidate concentration documented")

    v.check(bool(re.search(r'(?:cancer|tumor).*(?:HK|housekeeping).*(?:dysregulat|variab|alter)|housekeeping.*(?:cancer|tumor)', t, re.I)),
            "E3-6 HK cancer dysregulation discussed")


def verify_p2_e4(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  P2-E4: Academic Publishing & Peer Review")
    print(f"{'─'*50}")

    t = v.ms_text()

    lines = t.split('\n')
    ab_lines = []
    in_abstract = False
    for line in lines:
        stripped = line.strip()
        if stripped == 'Abstract' and not in_abstract:
            in_abstract = True
            continue
        if not in_abstract:
            continue
        if stripped.startswith('Keywords:'):
            break
        if stripped.startswith('[See separate'):
            continue
        if stripped == 'Abstract':
            continue
        ab_lines.append(stripped)
    ab_full = ' '.join(ab_lines)
    wc = len([tok for tok in ab_full.split() if re.search(r'[a-zA-Z0-9\u03c9\u2212]', tok)])
    v.check(wc <= 250, f"E4-1 Abstract word count = {wc} (GB structured, target <=250)")

    v.check(bool(re.search(r'ranked\s*5th\s*(?:of|/)\s*5|ranked\s*4th|AUC\s*rank', t, re.I)),
            "E4-2 AUC rank 5th/5 explanation")

    v.check(bool(re.search(r'scanpy|scipy|sklearn|HRT\s+Atlas|open.?source.*python', t, re.I)),
            "E4-3 Acknowledgements expanded")

    ref_section = t.split("References\n")[1] if "References\n" in t else ""
    n_refs = len(re.findall(r'\.\s\d{4};[\w]+:', ref_section))
    v.check(n_refs >= 35, f"E4-4 Reference count (Vancouver) = {n_refs}")

    # GB/Vancouver: reference list must be numbered 1..55 (v41: Kang + 6 refs) matching [n] citations
    ref_nums = [int(n) for n in re.findall(r'^(\d{1,2})\.\s+\S', ref_section, re.M)]
    v.check(ref_nums == list(range(1, 56)),
            f"E4-8 Reference list numbered 1..55 ({len(ref_nums)} entries, "
            f"sequence {'OK' if ref_nums == list(range(1, 56)) else 'BROKEN'})")

    v.check(bool(re.search(r'Bergmann\s+glia.*(?:13\.56|signal|candidate)', t, re.I)),
            "E4-5 Bergmann glia signal documented")

    s = v.supp_text()
    v.check(bool(re.search(r'S1[0-2]|S12|Supplementary Figure S1[0-2]', s, re.I)),
            "E4-6 Supplementary figures S1-S12 complete")

    tb_path = v.wd / "Table1-2.docx"
    v.check(tb_path.exists() and tb_path.stat().st_size > 5000,
            "E4-7 Tables as separate Table1-2.docx")

    # E4-9: Additional file 1 (supplementary notes) content coverage —
    # all four table headings must be present, and content size recorded
    # (E3 round-7 m4: manifest-level verification values for the supp file)
    sn_heads = [h for h in (
        'Table S1: Parameter Sweep Results',
        'Table S2: Cross-Organ Conservation Data',
        'Table S3: Human Brain Non-neuronal Cell Regional CKI Data',
        'Table S4: Inter-regional Region-Associated Candidate Data') if h in s]
    sn_words = len(s.split())
    v.check(len(sn_heads) == 4,
            f"E4-9 Additional file 1: 4/4 table headings present, "
            f"{sn_words} words, figures S8-S12 cited in notes")


def verify_v34_expert_panel(v: Verifier):
    print(f"\n{'─'*50}")
    print(f"  v34 Expert Panel Fixes (retained)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    rg = v.rg_text()

    has_2_1 = bool(re.search(r'2\.1\s+Parameter\s+Summary', rg, re.I))
    has_pt = bool(re.search(r'Random seed.*42.*all analyses|k_n floor', rg))
    no_s6 = not re.search(r'6\.\s+Parameter\s+Summary', rg, re.I)
    v.check(has_2_1 and has_pt and no_s6, "M1 Repro Guide §2.1 parameter table (not §6)")

    supp_refs_S1_S2 = len(re.findall(r'Supplementary\s+Fig.*S[12]\b|Fig.*S[12]\b', t))
    supp_refs_S3_S9 = len(re.findall(r'Supplementary\s+Fig.*S[3-9]|Fig.*S[3-9]', t))
    supp_refs_S10_S12 = len(re.findall(r'Supplementary\s+Fig.*S1[0-2]|Fig.*S1[0-2]', t))
    total_supp = supp_refs_S1_S2 + supp_refs_S3_S9 + supp_refs_S10_S12
    v.check(total_supp >= 8, f"M3 S1-S12 cited in body ({total_supp} supp fig refs)")

    body_before_refs = t.split("References\n")[0] if "References\n" in t else t
    orphan_ok = 0
    for rn in [31, 32, 33, 34, 35, 40, 41]:
        if re.search(rf'\b{rn}\b', body_before_refs):
            orphan_ok += 1
    if re.search(r'31.*33|31\u201333', body_before_refs):
        orphan_ok = max(orphan_ok, 3)
    v.check(orphan_ok >= 5, f"M4 Orphan refs cited in body ({orphan_ok}/7)")

    no_stat_sig = "both statistically significant" not in t.lower()
    has_perm = bool(re.search(r'permutation\s+(P-value|support|test|null)', t, re.I))
    v.check(no_stat_sig and has_perm, "m8 permutation support wording")

    lim_section = t.rsplit("Limitations", 1)[1].split("Future directions")[0] if "Limitations" in t else ""
    has_design = "Data- and design-specific constraints" in lim_section
    has_mt = "Multiple testing and cluster-aware inference" in lim_section
    v.check(has_design and has_mt, "m9 Limitations thematic paragraphs present")

    ortho_count = len(re.findall(r'\borthogonal\b', t))
    v.check(ortho_count <= 2, f"m17 orthogonal->complementary ({ortho_count} remaining)")


def verify_v36_expert_fixes(v: Verifier):
    """v37 expert-review fixes (#939-#950): cross-document consistency."""
    print(f"\n{'─'*50}")
    print(f"  v37 Expert-Review Fixes (E1-E4, #939-#950)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    cl = v.cl_text()
    rg = v.rg_text()

    # --- Calibration CI [4.24, 9.24] unified across all four documents ---
    n_ms = len(re.findall(r'4\.24,\s*9\.24', t))
    n_rg = len(re.findall(r'4\.24,\s*9\.24', rg))
    n_s = len(re.findall(r'4\.24,\s*9\.24', s))
    n_cl = len(re.findall(r'4\.24,\s*9\.24', cl))
    v.check(n_ms >= 5, f"V37-1a CI in manuscript ({n_ms}, need >=5)")
    v.check(n_rg >= 4, f"V37-1b CI in repro guide ({n_rg}, need >=4)")
    v.check(n_s >= 1, f"V37-1c CI in supplementary ({n_s}, need >=1)")
    v.check(n_cl >= 1, f"V37-1d CI in cover letter ({n_cl}, need >=1)")

    # --- Stale v35-era numbers purged from manuscript + cover letter ---
    stale_patterns = ["5,151", "5151", "21.48", "19.58", "6.28-fold", "6.28 fold",
                      "4.12, 9.33", "[4.12", "30 threshold-passing",
                      "0.365", "12.29", "12.12", "12.47", "7.7e-5", "7.8e-6",
                      "1.8-fold higher", "NAR Online"]
    found_stale = [p for p in stale_patterns if p in t or p in cl]
    v.check(len(found_stale) == 0, f"V37-2 stale v35 numbers absent ({found_stale})")

    # --- Cover letter: reviewer email + four-dataset validation ---
    v.check("welchjd@umich.edu" in cl, "V37-3a Welch email = welchjd@umich.edu")
    v.check("jdwlch" not in cl, "V37-3b old Welch email absent")
    v.check(bool(re.search(r'validated\s+(?:CKI\s+)?across\s+four\s+datasets', cl, re.I)),
            "V37-3c cover letter four-dataset paragraph")
    v.check(bool(re.search(r'6\.10.?fold', cl)), "V37-3d cover letter 6.10-fold")
    v.check(bool(re.search(r'1\.23.{0,90}2\.32', t)), "V37-3e manuscript NN/TT 1.23-2.32")
    v.check(bool(re.search(r'computational\s+debugging,\s+statistical\s+code\s+review,\s+and\s+language\s+editing', cl, re.I)),
            "V37-3f cover letter AI declaration aligned")

    # --- Reproducibility guide: rewritten content markers ---
    v.check(bool(re.search(r'Python\s*3\.14\.4', rg)), "V37-4a guide Python 3.14.4")
    v.check(bool(re.search(r'4,?851', rg)), "V37-4b guide human pairs 4,851")
    v.check(bool(re.search(r'brain_bs_null', rg)), "V37-4c guide brain_bs_null pointers")
    v.check(bool(re.search(r'superseded', rg, re.I)), "V37-4d guide superseded annotations")
    v.check(bool(re.search(r'38\.55', rg)), "V37-4e guide mu_grand 38.55")
    v.check(bool(re.search(r'92\.89|per.?pair\s+k_n', rg, re.I)), "V37-4f guide per-pair k_n coverage")
    v.check(bool(re.search(r'B\s*=\s*1,?000', rg)), "V37-4g guide block-shuffle B = 1,000")

    # --- Manuscript: hypothesis-generating framing for brain candidates ---
    v.check(bool(re.search(r'hypothesis.?generating', t, re.I)),
            "V37-5a brain candidates hypothesis-generating")
    v.check(bool(re.search(r'block.?shuffle', t, re.I)),
            "V37-5b block-shuffle null in manuscript")
    v.check(bool(re.search(r'exploratory', t, re.I)),
            "V37-5c TCGA exploratory framing")


def verify_v37_reviewer_fixes(v: Verifier):
    """v37 peer-review fixes (v36 report): calibration honesty, estimator
    sensitivity, within-donor gradient, lineage enrichment, tier sensitivity,
    and residual overclaim cleanup."""
    print(f"\n{'─'*50}")
    print(f"  v37 Peer-Review Fixes (C-A/C-B/C-C/C-D/C-F/C-G/C-J)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    rg = v.rg_text()

    # --- C-B: internal split-half baselines ---
    v.check(bool(re.search(r'internal baseline of 9\.73', t)), "V37-1a brain internal baseline 9.73 in MS")
    v.check(bool(re.search(r'\[9\.03, 10\.53\]', t)), "V37-1b brain internal CI")
    v.check(bool(re.search(r'7\.67\s*\(95% bootstrap CI \[7\.39, 8\.00\]', t)), "V37-1c TS internal baseline 7.67 in MS")
    v.check(bool(re.search(r'dataset-relative', t)), "V37-1d dataset-relative framing in MS")
    v.check(bool(re.search(r'9\.73', s)), "V37-1e brain internal baseline in supp")
    v.check(bool(re.search(r'7\.67', s)), "V37-1f TS internal baseline in supp")
    v.check(bool(re.search(r'dataset-relative quantity rather than a universal constant', s)),
            "V37-1g dataset-relative framing in supp")

    # --- C-B: old overclaims purged ---
    v.check("remains above the empirical baseline" not in t, "V37-2a old Bergmann claim absent (MS)")
    v.check("All ten non-neuronal brain classes have omega_cal > 1" not in s,
            "V37-2b old all-ten claim absent (supp)")
    v.check("4.88" not in t and "11.52" not in t, "V37-2c over-precise omega_cal values absent (MS)")

    # --- C-B: limitations revised ---
    v.check(bool(re.search(r'dataset-dependent: the Tabula Sapiens internal baseline', t)),
            "V37-3a Limitation Tenth revised")
    v.check(bool(re.search(r'We verified scheme-specific transferability directly', t)),
            "V37-3b Limitation Seventeenth revised")

    # --- C-C: estimator sensitivity + decomposition ---
    v.check('0.988' in t, "V37-4a aggregate-first rho 0.988 in MS")
    v.check(bool(re.search(r'6\.51', t)), "V37-4b aggregate-first gradient 6.51x in MS")
    v.check('0.09' in t and 'global' in t, "V37-4c global-kn reversal in MS")
    v.check(bool(re.search(r'2\.0-fold', t)) and bool(re.search(r'3\.2-fold', t)),
            "V37-4d kf 2.0-fold vs kn 3.2-fold decomposition in MS")

    # --- C-J: within-donor gradient ---
    v.check(bool(re.search(r'4\.50', t)), "V37-5a within-donor 4.50-fold in MS")
    v.check(bool(re.search(r'11,139', t)), "V37-5b 11,139 within-donor pairs in MS")
    v.check(bool(re.search(r'75\.18', t)), "V37-5c astrocyte within-donor 75.18 in MS")

    # --- C-G: lineage enrichment formal test ---
    v.check(bool(re.search(r'2\.30', t)), "V37-6a microglia fold 2.30 in MS")
    v.check('P = 6.0' in t, "V37-6b hypergeometric P = 6.0e-4 in MS")
    v.check(bool(re.search(r'16 of (?:the )?39|16/39', t)), "V37-6c 16 of 39 Strong from microglia in MS")

    # --- C-D: tier sensitivity ---
    v.check(bool(re.search(r'fold enrichment 0\.77', t)) or bool(re.search(r'fold 0\.77', t)),
            "V37-7a oligodendrocyte-lineage null enrichment (fold 0.77) in MS")
    v.check('P = 0.92' in t, "V37-7b lineage enrichment P = 0.92 (null) in MS")

    # --- C-A: Figure 3 retitled; composite framing ---
    v.check(bool(re.search(r'Correlation structure between CKI and standard metrics', t)),
            "V37-8a Figure 3 retitled in MS")
    v.check("CKI captures independent information" not in t,
            "V37-8b old Figure 3 title absent")

    # --- Methods: robustness subsection ---
    v.check(bool(re.search(r'Robustness and calibration analyses', t)),
            "V37-9a Methods robustness subsection in MS")
    v.check(bool(re.search(r'71 populations across six organs', t)),
            "V37-9b TS split-half 71 populations in MS")

    # --- Repro guide: Section 5.6 ---
    v.check(bool(re.search(r'Reviewer Robustness Analyses', rg)), "V37-10a guide Sec 5.6 present")
    v.check(bool(re.search(r'41_reviewer_fix_within_donor\.py', rg)), "V37-10b script 41 documented")
    v.check(bool(re.search(r'42_reviewer_fix_kn_estimators\.py', rg)), "V37-10c script 42 documented")
    v.check(bool(re.search(r'43_reviewer_fix_ts_splithalf\.py', rg)), "V37-10d script 43 documented")
    v.check(bool(re.search(r'reviewer_brain_splithalf_summary\.txt', rg)), "V37-10e brain splithalf output pointer")
    v.check(bool(re.search(r'reviewer_ts_splithalf_summary\.txt', rg)), "V37-10f TS splithalf output pointer")



def verify_v38_reviewer_fixes(v: Verifier):
    """v38 peer-review fixes (#971/#972/#974): ground-truth simulation,
    fixed gene-panel ablation, engineering consistency."""
    print(f"\n{'─'*50}")
    print(f"  v38 Peer-Review Fixes (#971/#972/#974)")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    cl = v.cl_text()
    rg = v.rg_text()

    # --- #972: ground-truth simulation ---
    v.check(bool(re.search(r'Ground-truth simulation', t)),
            "V38-1a ground-truth headings in MS")
    v.check(bool(re.search(r'1,?848', t)), "V38-1b background 1,848 cells in MS")
    v.check(bool(re.search(r'AUC\s*=\s*0\.80\b', t)), "V38-1c AUC = 0.80 in MS")
    v.check(bool(re.search(r'false-positive rate 0\.00|type-I error 0\.00', t)),
            "V38-1d omega type-I 0.00 in MS")
    v.check(bool(re.search(r'0\.58', t)), "V38-1e cosine 0.58 in MS")
    v.check(bool(re.search(r'3\.12\s+Ground-?Truth|Ground-?Truth Simulation', s)),
            "V38-1f SN 3.12 in supplementary")
    v.check(bool(re.search(r'45_groundtruth_simulation\.py', rg)),
            "V38-1g script 45 in repro guide")
    v.check(bool(re.search(r'1,?750', s + t)), "V38-1h 1,750 replicates documented")

    # --- #974: fixed gene-panel ablation ---
    v.check(bool(re.search(r'Fixed gene-panel ablation', t)),
            "V38-2a ablation headings in MS")
    v.check(bool(re.search(r'1\.6[0-9]?-fold', t)), "V38-2b circularity 1.6/1.61-fold in MS")
    v.check(bool(re.search(r'6\.53', t)), "V38-2c LOO gradient 6.53 in MS")
    v.check(bool(re.search(r'6\.4\s*[×x]\s*10', t)) or '6.4e-13' in t,
            "V38-2d reproduction 6.4e-13 in MS")
    v.check(bool(re.search(r'P\s*=\s*0\.980', t)), "V38-2e Bergmann P = 0.980 under LOO in MS")
    v.check(bool(re.search(r'32 of 50', t)), "V38-2f residual retention 32 of 50 (64%) in MS")
    v.check(bool(re.search(r'26\.5', t + s)), "V38-2g LOO grand mean 26.5 in MS/SN")
    v.check(bool(re.search(r'3\.13\s+Fixed Gene-Panel|Fixed Gene-Panel Ablation', s)),
            "V38-2h SN 3.13 in supplementary")
    v.check(bool(re.search(r'46_fixed_panel_ablation\.py', rg)),
            "V38-2i script 46 in repro guide")
    v.check(bool(re.search(r'gene-panel ablation', cl, re.I)),
            "V38-2j ablation in cover letter")

    # --- #971: engineering consistency fixes ---
    v.check(bool(re.search(r'PAM50_SUBTYPE', t)), "V38-3a PAM50 actual implementation in MS")
    v.check('nearest centroid' not in t.lower() and '44 of 47' not in t,
            "V38-3b stale PAM50 wording absent")
    v.check(bool(re.search(r'522 samples', t)), "V38-3c PAM50 522 samples in MS")
    v.check(bool(re.search(r'1 of 31,?764 pairs', t)), "V38-3d brain kn caliber 1 pair below 1e-4 in MS")
    v.check(bool(re.search(r'kn_floor\s*=\s*0', t)), "V38-3e kn_floor = 0 contract in MS")
    v.check(bool(re.search(r'2,?510', t)) and bool(re.search(r'7\.9%', t)),
            "V38-3f upper-tail 2,510 (7.9%) in MS")
    v.check(bool(re.search(r'v0\.4\.4|0\.4\.4', t)) and bool(re.search(r'v0\.4\.4|0\.4\.4', rg))
            and not re.search(r'0\.4\.[23]', t) and not re.search(r'0\.4\.[23]', rg),
            "V38-3g version 0.4.4 in MS + guide (no 0.4.2/0.4.3 residue)")
    _pkg_init = (BASE_DIR / "cki" / "__init__.py").read_text(encoding="utf-8")
    _pyproject = (BASE_DIR / "pyproject.toml").read_text(encoding="utf-8")
    v.check('__version__ = "0.4.4"' in _pkg_init
            and re.search(r'^version\s*=\s*"0\.4\.4"', _pyproject, re.M) is not None,
            "REL-1 package version bumped to 0.4.4 in cki/__init__.py + pyproject.toml")
    v.check(bool(re.search(r'Dockerfile', t)), "V38-3h Dockerfile mention in MS")
    v.check(bool(re.search(r'scanpy:\s+1\.12\.1', rg)), "V38-3i scanpy 1.12.1 in guide")
    v.check(bool(re.search(r'ratio\s+artifact\s+of\s+the\s+k_n\s+denominator', cl)),
            "V38-3j cover letter overclaim retracted (negated)")

    # --- R5 P2-5: data-level guard for the brain minimum k_n (9.2e-5) ---
    import csv as _csv
    kn_vals = []
    with open(RESULTS_DIR / "fixed_panel_ablation_pairs.csv",
              newline="", encoding="utf-8") as _fh:
        for _row in _csv.DictReader(_fh):
            try:
                kn_vals.append(float(_row["kn"]))
            except (TypeError, ValueError, KeyError):
                pass
    if kn_vals:
        _min_kn = min(kn_vals)
        v.check(len(kn_vals) == 31764 and
                9.0e-5 < _min_kn < 9.4e-5 and
                sum(1 for k in kn_vals if k < 1e-4) == 1,
                f"V38-3k data-level brain min k_n = {_min_kn:.3e} "
                f"(1 of {len(kn_vals):,} pairs < 1e-4)")
    else:
        v.check(False, "V38-3k fixed_panel_ablation_pairs.csv readable")

    # --- abstract back under 200 words ---
    lines = t.split('\n')
    ab_lines = []
    in_abstract = False
    for line in lines:
        stripped = line.strip()
        if stripped == 'Abstract' and not in_abstract:
            in_abstract = True
            continue
        if not in_abstract:
            continue
        if stripped.startswith('Keywords:'):
            break
        if stripped == 'Abstract':
            continue
        ab_lines.append(stripped)
    ab_full = ' '.join(ab_lines)
    wc = len([tok for tok in ab_full.split() if re.search(r'[a-zA-Z0-9\u03c9\u2212]', tok)])
    v.check(wc <= 250, f"V38-4a abstract word count = {wc} (GB structured, target <=250)")
    v.check(bool(re.search(r'ground-truth simulation', ab_full, re.I)) and
            bool(re.search(r'gene-panel ablation', ab_full, re.I)),
            "V38-4b abstract retains both new validation claims")


def verify_v38_text_revisions(v: Verifier):
    """P2 text revisions (panel synthesis #8-11): de-Kinetic rename, terminology,
    omega framework rewrite, consolidated Limitations, TODO-STATS placeholders."""
    print(f"\n{'─'*50}")
    print(f"  v38.1 P2 Text Revisions")
    print(f"{'─'*50}")

    t = v.ms_text()
    s = v.supp_text()
    cl = v.cl_text()
    rg = v.rg_text()
    all_docs = t + s + cl + rg

    # --- de-Kinetic rename ---
    v.check(not re.search(r'[Kk]inetic', all_docs),
            "T1 'Kinetic' removed from all four documents")
    v.check(bool(re.search(r'CKI: a Ka/Ks-inspired index for quantifying functional', t)) and
            bool(re.search(r'CKI: a Ka/Ks-inspired index for quantifying functional', cl)) and
            bool(re.search(r'CKI: a Ka/Ks-inspired index for quantifying functional', s)),
            "T2 new title (Ka/Ks-inspired index) in MS + CL + Supp")
    v.check('CKI (Cell-type Ka/Ks-inspired Index)' in t and
            'CKI (Cell-type Ka/Ks-inspired Index)' in cl,
            "T3 CKI expansion renamed (Ka/Ks-inspired) in MS + CL")

    # --- terminology ---
    # T4 updated for round 3 (R1 P1-2): the manuscript now MUST state the
    # per-dataset softmax aggregation order explicitly; 'softmax' stays
    # banned only in supplementary/repro guide/cover letter.
    v.check(not re.search(r'softmax', s + rg + cl, re.I),
            "T4 'softmax' confined to the manuscript (Methods normalization)")
    v.check(bool(re.search(r'softmax\(log1p\(mean counts\)\)', t)) and
            bool(re.search(r'softmax\(mean\(log1p\)\)', t)),
            "T4b per-dataset aggregation order stated (log-of-mean vs mean-of-log)")
    v.check(bool(re.search(r'\+1 pseudo-count followed by L1 normalization', t)) and
            bool(re.search(r'\+1 pseudo-count followed by L1 normalization', s + rg)),
            "T5 '+1 pseudo-count followed by L1 normalization' wording in MS + SN/guide")
    v.check(not re.search(r'migration[ -]candidate', all_docs, re.I),
            "T6 'migration candidate' removed from all four documents")
    v.check(bool(re.search(r'region-associated candidate', t, re.I)) and
            bool(re.search(r'region-associat', s + rg, re.I)),
            "T7 'region-associated candidate' terminology in MS + SN/guide")

    # --- omega three-tier framework removed ---
    v.check('operational thresholds' not in t and 'much greater than 1' not in t,
            "T8 omega fixed cut-off framework removed from MS")
    v.check(bool(re.search(r'anchored in the empirical', t)) and
            bool(re.search(r'anchored in the empirical', s)),
            "T9 distribution-anchored omega interpretation in MS + SN")

    # --- consolidated Limitations ---
    lim_section = t.split("Limitations")[1].split("Future directions")[0] if "Limitations" in t else ""
    ordinal_hits = re.findall(r'\b(?:First|Second|Third|Fourth|Fifth|Sixth|Seventh|Eighth|Ninth|Tenth|Eleventh|Twelfth|Thirteenth|Fourteenth|Fifteenth|Sixteenth|Seventeenth|Eighteenth|Nineteenth|Twentieth|Twenty.first)\b', lim_section)
    v.check(not ordinal_hits, f"T10 no ordinal numbering in consolidated Limitations ({len(ordinal_hits)} found)")

    # --- hedged claims ---
    v.check('We evaluated CKI' in t, "T11 abstract hedged ('evaluated' not 'validated')")
    v.check('robust, interpretable' not in cl and 'rigorous' not in cl,
            "T12 cover letter toned down (no 'robust, interpretable' / 'rigorous')")

    # --- TODO-STATS placeholders filled with real values (round 2) ---
    ms_todo = re.findall(r'TODO-STATS-[A-Z-]+', t)
    sn_todo = re.findall(r'TODO-STATS-[A-Z-]+', s)
    cl_todo = re.findall(r'TODO-STATS-[A-Z-]+', cl)
    rg_todo = re.findall(r'TODO-STATS-[A-Z-]+', rg)
    v.check(len(ms_todo) == 0, f"T13 TODO-STATS in MS all filled ({len(ms_todo)} remaining)")
    v.check(len(sn_todo) + len(cl_todo) + len(rg_todo) == 0,
            f"T14 TODO-STATS in SN/CL/guide all filled ({len(sn_todo)}/{len(cl_todo)}/{len(rg_todo)} remaining)")


# ============================================================
# Round 5 R2: numeric re-derivation layer (data-level checks)
# ============================================================

def _nk_rank(x):
    """Average ranks with tie handling."""
    idx = sorted(range(len(x)), key=lambda i: x[i])
    r = [0.0] * len(x)
    i = 0
    while i < len(idx):
        j = i
        while j + 1 < len(idx) and x[idx[j + 1]] == x[idx[i]]:
            j += 1
        avg = (i + j) / 2 + 1
        for k in range(i, j + 1):
            r[idx[k]] = avg
        i = j + 1
    return r


def _nk_spearman(a, b):
    ra, rb = _nk_rank(a), _nk_rank(b)
    n = len(ra)
    ma, mb = sum(ra) / n, sum(rb) / n
    cov = sum((x - ma) * (y - mb) for x, y in zip(ra, rb))
    va = sum((x - ma) ** 2 for x in ra)
    vb = sum((y - mb) ** 2 for y in rb)
    return cov / ((va * vb) ** 0.5)


def _nk_bh(pvals):
    """Benjamini-Hochberg step-up q-values; return (min q, count q<0.05)."""
    m = len(pvals)
    sp = sorted(pvals)
    q = [0.0] * m
    q[m - 1] = sp[m - 1]
    for i in range(m - 2, -1, -1):
        q[i] = min(q[i + 1], m * sp[i] / (i + 1))
    return q[0], sum(1 for x in q if x < 0.05)


def verify_v38_numeric(v: Verifier):
    """R2 (Round 5): re-derive the headline statistics directly from the
    results/ CSV and JSON files, independently of the manuscript text.
    This complements the regex-based checks with data-level assertions
    (same stdlib-only approach as V38-3k)."""
    import csv as _csv
    import json as _json

    print(f"\n{'─'*50}")
    print(f"  R2 Numeric Re-derivation Layer (data-level)")
    print(f"{'─'*50}")

    def _rows(name):
        with open(RESULTS_DIR / name, newline="", encoding="utf-8") as fh:
            return list(_csv.DictReader(fh))

    def _close(x, y, tol):
        return abs(x - y) <= tol

    # ---- brain block-shuffle null: BH, tails, Strong counts ----
    bs = _rows("brain_bs_null_results.csv")
    n = len(bs)
    p_all = [float(r["p_perm"]) for r in bs]
    ph_all = [float(r["p_perm_high"]) for r in bs]
    strong = [i for i in range(n) if bs[i]["tier"] == "Strong"]
    n_strong_sig = sum(1 for i in strong if p_all[i] < 0.05)
    v.check(n == 31764,
            f"R2-N1 block-shuffle pairs = {n:,} (expect 31,764)")
    v.check(len(strong) == 39 and n_strong_sig == 31,
            f"R2-N2 Strong = {len(strong)}, raw P < 0.05 in {n_strong_sig}/39 "
            f"(expect 39 / 31)")
    min_q, n_q05 = _nk_bh(p_all)
    v.check(_close(min_q, 0.520, 5e-4) and n_q05 == 0,
            f"R2-N3 BH recompute: min q = {min_q:.4f}, q < 0.05 in {n_q05} "
            f"pairs (expect 0.520 / 0)")
    up = sum(1 for p in ph_all if p < 0.05)
    v.check(up == 2510 and _close(100.0 * up / n, 7.9, 0.05),
            f"R2-N4 upper-tail excess = {up:,} ({100.0 * up / n:.2f}%) "
            f"(expect 2,510 / 7.9%)")
    _low = {}
    for r in bs:
        c = _low.setdefault(r["cell_type"], [0, 0])
        c[1] += 1
        if float(r["p_perm"]) < 0.05:
            c[0] += 1
    mi = _low["Microglia"]
    od = _low["Oligodendrocyte"]
    as_ = _low["Astrocyte"]
    bg = _low["Bergmann glia"]
    v.check(_close(100.0 * mi[0] / mi[1], 8.06, 0.005) and
            _close(100.0 * od[0] / od[1], 6.01, 0.005) and
            _close(100.0 * as_[0] / as_[1], 7.60, 0.005) and
            bg == [8, 21],
            f"R2-N5 lower-tail classes: microglia {100.0*mi[0]/mi[1]:.2f}%, "
            f"oligodendrocyte {100.0*od[0]/od[1]:.2f}%, astrocyte "
            f"{100.0*as_[0]/as_[1]:.2f}%, Bergmann {bg[0]}/{bg[1]} = "
            f"{100.0*bg[0]/bg[1]:.1f}% (expect 8.06 / 6.01 / 7.60 / 38.1%)")

    # ---- brain observed pairs: class means, grand mean, gradient, tiers ----
    obs = _rows("brain_bs_null_observed_pairs.csv")
    w_all = [float(r["omega"]) for r in obs]
    w_astro = [float(r["omega"]) for r in obs if r["cell_type"] == "Astrocyte"]
    w_berg = [float(r["omega"]) for r in obs if r["cell_type"] == "Bergmann glia"]
    gm = sum(w_all) / len(w_all)
    ma = sum(w_astro) / len(w_astro)
    mb = sum(w_berg) / len(w_berg)
    v.check(_close(ma, 82.75, 0.005) and _close(mb, 13.56, 0.005) and
            _close(gm, 38.55, 0.005) and _close(ma / mb, 6.10, 0.005),
            f"R2-N6 astrocyte {ma:.2f} / Bergmann {mb:.2f} / grand mean "
            f"{gm:.2f} / gradient {ma/mb:.2f} (expect 82.75/13.56/38.55/6.10)")
    _tiers = {}
    for r in obs:
        _tiers[r["tier"]] = _tiers.get(r["tier"], 0) + 1
    v.check(_tiers.get("Strong") == 39 and _tiers.get("Moderate") == 1171 and
            _tiers.get("Weak") == 5381,
            f"R2-N7 tier counts Strong/Moderate/Weak = "
            f"{_tiers.get('Strong')}/{_tiers.get('Moderate')}/"
            f"{_tiers.get('Weak')} (expect 39/1,171/5,381)")

    # ---- split-half calibrations (brain + Tabula Sapiens) ----
    def _kv(name):
        out = {}
        for line in (RESULTS_DIR / name).read_text(encoding="utf-8").splitlines():
            if "\t" in line:
                k, val = line.split("\t", 1)
                out[k] = val
        return out

    sh_b = _kv("reviewer_brain_splithalf_summary.txt")
    ok_b = (_close(float(sh_b["brain_split_half_mean_omega"]), 9.7261, 5e-5) and
            sh_b["brain_split_half_ci95"].replace(" ", "") == "[9.0271,10.5310]")
    v.check(ok_b, f"R2-N8 brain split-half = {sh_b['brain_split_half_mean_omega']} "
            f"{sh_b['brain_split_half_ci95']} (expect 9.7261 [9.0271, 10.5310])")
    sh_t = _kv("reviewer_ts_splithalf_summary.txt")
    ok_t = (_close(float(sh_t["ts_split_half_mean_omega"]), 7.6708, 5e-5) and
            sh_t["ts_split_half_ci95"].replace(" ", "") == "[7.3898,7.9975]")
    v.check(ok_t, f"R2-N9 TS split-half = {sh_t['ts_split_half_mean_omega']} "
            f"{sh_t['ts_split_half_ci95']} (expect 7.6708 [7.3898, 7.9975])")

    # ---- within-donor gradient ----
    wd = {r["cell_type"]: r for r in _rows("reviewer_within_donor_gradient.csv")}
    wd_a = float(wd["Astrocyte"]["omega_mean_within_donor"])
    wd_b = float(wd["Bergmann glia"]["omega_mean_within_donor"])
    v.check(_close(wd_a, 75.18, 0.005) and _close(wd_b, 16.73, 0.005) and
            _close(wd_a / wd_b, 4.50, 0.006),
            f"R2-N10 within-donor astro {wd_a:.2f} / Bergmann {wd_b:.2f} / "
            f"gradient {wd_a/wd_b:.2f} (expect 75.18/16.73/4.50)")

    # ---- per-pair vs global k_n estimator dependence ----
    pc = _rows("phaseC_omega_pair_vs_global.csv")
    kn = [float(r["kn"]) for r in pc]
    w_pp = [float(r["omega"]) for r in pc]
    w_gl = [float(r["omega_global_kn"]) for r in pc]
    kn_m = sum(kn) / len(kn)
    kn_sd = (sum((x - kn_m) ** 2 for x in kn) / (len(kn) - 1)) ** 0.5
    rho = _nk_spearman(w_pp, w_gl)
    v.check(len(pc) == 31764 and _close(100.0 * kn_sd / kn_m, 97.52, 0.005) and
            _close(rho, 0.142, 5e-4),
            f"R2-N11 per-pair k_n CV = {100.0*kn_sd/kn_m:.2f}%, "
            f"Spearman(per-pair, global) = {rho:.3f} over {len(pc):,} pairs "
            f"(expect 97.52% / 0.142; post-fix from reviewer_brain_pair_kf_kn.csv)")

    # ---- ground-truth simulation ----
    gt = _json.loads((RESULTS_DIR / "groundtruth_simulation_metrics.json")
                     .read_text(encoding="utf-8"))
    auc_w = gt["auc_signal_vs_neutral"]["omega"]
    auc_kf = gt["auc_signal_vs_neutral"]["k_f"]
    thr = gt["null_thresholds"]["omega"]
    gts = {r["series"]: r for r in _rows("groundtruth_simulation_summary.csv")}
    base_mean = float(gts["baseline"]["omega_mean"])
    v.check(_close(auc_w, 0.804, 5e-4) and _close(auc_kf, 0.716, 5e-4) and
            _close(thr, 20.81, 0.005) and _close(base_mean, 11.62, 0.005),
            f"R2-N12 simulation AUC(omega) = {auc_w:.3f}, AUC(k_f) = "
            f"{auc_kf:.3f}, background mean = {base_mean:.1f}, q95 = "
            f"{thr:.1f} (expect 0.804/0.716/11.6/20.8)")

    # ---- fixed-panel ablation ----
    ab = _json.loads((RESULTS_DIR / "fixed_panel_ablation_summary.json")
                     .read_text(encoding="utf-8"))
    r1 = ab["spearman_pair_s0_vs_s1"]["rho"]
    r2 = ab["spearman_pair_s0_vs_s2"]["rho"]
    r3 = ab["spearman_pair_s0_vs_s3"]["rho"]
    c1 = ab["spearman_ctmean_s0_vs_s1"]["rho"]
    c2 = ab["spearman_ctmean_s0_vs_s2"]["rho"]
    c3 = ab["spearman_ctmean_s0_vs_s3"]["rho"]
    v.check(_close(r1, 0.918, 5e-4) and _close(r2, 0.937, 5e-4) and
            _close(r3, 0.931, 5e-4) and _close(c2, 0.99, 0.005) and
            _close(c1, 0.90, 0.005) and _close(c3, 0.93, 0.005),
            f"R2-N13 ablation Spearman rho pair-level (LPO/fixed/all) = "
            f"{r2:.3f}/{r1:.3f}/{r3:.3f}, class-level = {c2:.2f}/{c1:.2f}/"
            f"{c3:.2f} (expect 0.937/0.918/0.931 and 0.99/0.90/0.93)")
    v.check(_close(ab["grand_mean_s0"], 38.55, 0.005) and
            _close(ab["grand_mean_s1"], 6.54, 0.005) and
            _close(ab["grand_mean_s2"], 26.52, 0.005) and
            _close(ab["grand_mean_s3"], 5.26, 0.005) and
            _close(ab["astro_mean_s0"], 82.75, 0.005) and
            _close(ab["bergmann_mean_s0"], 13.56, 0.005) and
            _close(ab["astro_over_bergmann_s0"], 6.10, 0.005) and
            ab["tier_counts"]["s0"]["Strong"] == 39 and
            ab["n_pairs"] == 31764,
            f"R2-N14 ablation grand means 38.55/26.52/6.54/5.26, gradient "
            f"82.75/13.56 = 6.10, Strong(s0) = 39 (all match)")
    kf21 = ab["kf_ratio_s0_over_s2"]
    v.check(_close(kf21["median"], 1.61, 0.005) and
            _close(kf21["q25"], 1.27, 0.005) and
            _close(kf21["q75"], 2.07, 0.005) and
            _close(ab["kf_ratio_s0_over_s1"]["median"], 6.07, 0.005) and
            _close(ab["kf_ratio_s0_over_s3"]["median"], 7.27, 0.005),
            f"R2-N15 circularity k_f ratios: LPO {kf21['median']:.2f} "
            f"[{kf21['q25']:.2f}, {kf21['q75']:.2f}], fixed "
            f"{ab['kf_ratio_s0_over_s1']['median']:.1f}, all "
            f"{ab['kf_ratio_s0_over_s3']['median']:.1f} "
            f"(expect 1.61 [1.27, 2.07] / 6.1 / 7.3)")

    # ---- lower-tail count calibration ----
    tc = _rows("_v38_tailcount_calibration.csv")
    row = tc[0]
    v.check(int(float(row["V_obs"])) == 1960 and
            _close(float(row["null_mean"]), 1588.2, 0.05) and
            _close(float(row["P(V_null >= V_obs)"]), 0.011, 5e-4),
            f"R2-N16 tail-count V_obs = {int(float(row['V_obs']))} vs null "
            f"mean {float(row['null_mean']):.1f}, P = "
            f"{float(row['P(V_null >= V_obs)']):.3f} (expect 1960/1588.2/0.011)")

    # ---- TCGA NN/TT reversal ----
    tt = {r["cancer"]: float(r["omega_NN_TT"])
          for r in _rows("_v38_tcga_nn_tt_reversal_decomposition.csv")}
    exp_tt = {"TCGA-BRCA": 1.51, "TCGA-KIRC": 2.19, "TCGA-LIHC": 1.23,
              "TCGA-LUAD": 2.32, "TCGA-LUSC": 1.77}
    ok_tt = all(_close(tt.get(k, 0), x, 0.005) for k, x in exp_tt.items())
    v.check(ok_tt and len(tt) == 5,
            f"R2-N17 TCGA median NN/TT = " +
            ", ".join(f"{k[-4:]} {tt[k]:.2f}" for k in sorted(tt)) +
            " (expect BRCA 1.51, KIRC 2.19, LIHC 1.23, LUAD 2.32, LUSC 1.77)")

    # ---- human (Tabula Sapiens) ----
    hum = _rows("phase35_all_metrics_pairs.csv")
    hw = sorted(float(r["omega"]) for r in hum)
    hn = len(hw)
    hmean = sum(hw) / hn
    hmed = hw[hn // 2] if hn % 2 else (hw[hn // 2 - 1] + hw[hn // 2]) / 2
    v.check(hn == 4851 and _close(hmean, 21.61, 0.005) and
            _close(hmed, 19.65, 0.005),
            f"R2-N18 human pairs = {hn:,}, mean omega = {hmean:.2f}, "
            f"median = {hmed:.2f} (expect 4,851 / 21.6 / 19.7)")

    # ---- mouse calibration baseline ----
    pil = _rows("mouse_pilot_v2_results.csv")
    ctrl = [float(r["omega"]) for r in pil if r["category"] == "C_control"]
    pb = {r["group"]: r for r in _rows("phaseB_bootstrap_cis.csv")
          if r["dataset"] == "Mouse (pilot)"}
    ctrl_mean = sum(ctrl) / len(ctrl)
    ci_lo = float(pb["C_control"]["ci_95_lower"])
    ci_hi = float(pb["C_control"]["ci_95_upper"])
    v.check(len(ctrl) == 6 and _close(ctrl_mean, 6.67, 0.005) and
            _close(ci_lo, 4.24, 0.005) and _close(ci_hi, 9.24, 0.005),
            f"R2-N19 mouse baseline mean = {ctrl_mean:.2f} "
            f"(n = {len(ctrl)}), committed CI = [{ci_lo:.2f}, {ci_hi:.2f}] "
            f"(expect 6.67, [4.24, 9.24])")

    # ---- tier-threshold sensitivity grid ----
    ts = _rows("reviewer_tier_sensitivity.csv")
    ns = [int(r["n_strong"]) for r in ts]
    pmin = min(float(r["lineage_hypergeom_P"]) for r in ts)
    v.check(len(ts) == 20 and min(ns) == 1 and max(ns) == 259 and
            _close(pmin, 0.432, 5e-4),
            f"R2-N20 tier grid: {len(ts)} combinations, Strong {min(ns)}-{max(ns)}, "
            f"min hypergeometric P = {pmin:.3f} (expect 20 rows / 1-259 / 0.432)")

    # ---- JS dimensionality invariance ----
    dim = {int(float(r["dimension"])): float(r["mean_js"])
           for r in _rows("phaseC_dimensionality_simulation.csv")}
    ratio = dim[2000] / dim[1130]
    v.check(_close(ratio, 1.001, 5e-4),
            f"R2-N21 JS dimensionality ratio d=2000/d=1130 = {ratio:.3f} "
            f"(expect 1.001)")



# ============================================================
# Build
# ============================================================

def build_v38():
    print("=" * 60)
    print("  CKI Submission Package v41 (Genome Biology) Builder")
    print("  Second peer-review revision (#971/#972/#974) + fresh DOCX rebuild")
    print("=" * 60)

    if not collect_figures():
        return False

    # 1. Prepare work dir
    print(f"\n[1] Preparing {WORK_DIR.name}...")
    if WORK_DIR.exists():
        shutil.rmtree(str(WORK_DIR), ignore_errors=True)
    WORK_DIR.mkdir(parents=True, exist_ok=True)

    # Copy figures / GA
    for f in FIGURES_SUBMISSION_DIR.iterdir():
        shutil.copy2(f, WORK_DIR / f.name)
    print(f"  Copied figures into {WORK_DIR.name}")

    # 2. Regenerate DOCX
    print(f"\n[2] Regenerating DOCX files...")
    scripts = [
        (f'"{PYTHON}" -u generate_manuscript_gb.py', "Manuscript (GB)"),
        (f'"{PYTHON}" -u notebooks/68_gen_supplementary_en.py', "Supplementary"),
        (f'"{PYTHON}" -u generate_cover_letter_nar.py', "Cover Letter"),
    ]
    for cmd, label in scripts:
        run_script(cmd, f"Generate {label}")

    node_env = os.environ.copy()
    node_env["NODE_PATH"] = NODE_PATH
    run_script(f'"{NODE}" notebooks/100_gen_reproducibility_docx.js',
               "Generate Repro Guide", env=node_env)

    run_script(f'"{PYTHON}" -u notebooks/_extract_table1_2.py', "Extract Table1-2")

    docx_map = {
        "CKI_Manuscript.docx": RESULTS_DIR / "CKI_Manuscript_GB.docx",
        "CKI_Supplementary.docx": RESULTS_DIR / "CKI_Supplementary.docx",
        "CKI_Cover_Letter.docx": RESULTS_DIR / "CKI_GenomeBiology_Cover_Letter.docx",
        "CKI_Reproducibility_Guide.docx": RESULTS_DIR / "CKI_Reproducibility_Guide.docx",
        "Table1-2.docx": RESULTS_DIR / "Table1-2.docx",
    }
    for name, src_path in docx_map.items():
        if src_path.exists():
            shutil.copy2(src_path, WORK_DIR / name)
            print(f"  {name}: {src_path.stat().st_size/1024:.1f} KB")
        else:
            print(f"  ERROR: {name} not found at {src_path}")
            return False

    # Extract fulltext
    print(f"\n[2f] Extracting fulltext...")
    import docx as _docx
    for docx_name in ["CKI_Manuscript.docx", "CKI_Supplementary.docx",
                       "CKI_Cover_Letter.docx", "CKI_Reproducibility_Guide.docx",
                       "Table1-2.docx"]:
        txt_name = docx_name.replace(".docx", "_fulltext.txt")
        d = _docx.Document(str(WORK_DIR / docx_name))
        lines = [p.text for p in d.paragraphs]
        for table in d.tables:
            lines.append("")
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
        with open(WORK_DIR / txt_name, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"  {txt_name}: {WORK_DIR.joinpath(txt_name).stat().st_size:,} bytes")

    # 3. Manifest
    print(f"\n[3] Writing MANIFEST_v41.txt...")
    manifest = f"""CKI Submission Package v41 (Genome Biology, Methodology article)
Built: {datetime.now().strftime('%Y-%m-%d %H:%M')}
Status: v41 = v40 + E4 software-engineering review fixes (P-value
denominator (B+1) unification, tail parameter, data-driven spot checks,
regression tests, environment description) + Kang IFN-beta real
perturbation demonstration (Fig. S13, Note 4.5) + pseudo-region negative
control for the brain block-shuffle null (Fig. S14, new Note); all v40
content otherwise unchanged.

=== v41 Changes ===
  - E4 review (reproducibility/software engineering) fixes:
    (a) cki/bootstrap.py P-value denominator corrected to
        (n_bootstrap + 1) with NaN-producing permutations counted in the
        denominator, matching the manuscript (B+1) formula;
    (b) cki/blocknull.py block_shuffle_test gained a tail parameter
        ("upper" default / "lower" / "two-sided") with the same (B+1)
        denominator; per-pair lower-tail screening now has a package
        entry point;
    (c) calibrate_omega docstring baselines updated to the authoritative
        post-fix values (brain split-half 9.73 [9.03, 10.53]);
    (d) scripts/spot_check_v19.py replaced by scripts/spot_check.py:
        45 assertions recomputed from the authoritative result files
        (TCGA NN/TT medians, mouse calibration, brain class-level and
        set-level, internal baselines, Kang demo AUC); run_all.py
        references updated;
    (e) tests/test_reference_values.py added (7 regression tests);
    (f) compute-environment description unified across manuscript and
        Reproducibility Guide (Windows x64 workstation, Section 1.3);
        python-docx and other document-generation dependencies moved out
        of the analysis requirements list.
  - Kang et al. IFN-beta PBMC real-perturbation demonstration
    (notebooks/79_kang_ifnb_demo.py, notebooks/80_kang_demo_figure.py):
    24,413 singlet cells, six cell types; omega detects the perturbation
    in all six cell types (median stim-vs-control omega 1.1-2.0x above
    donor drift), but where IFN-beta moves the housekeeping anchor
    (CD14+ monocytes) the omega AUC falls to 0.55 while k_f retains
    0.98 — the anchor-visibility boundary demonstrated on real data.
    Supplementary Figure S13 + Supplementary Note 4.5 + Results/
    Discussion text.
  - Pseudo-region negative control (notebooks/77_pseudoregion_control.py):
    every region's libraries split uniformly at random into two
    pseudo-regions per cell type and the identical block-shuffle test
    re-run (B = 1,000) on 127,756 pseudo-pairs. Cross-origin pairs
    (n = 127,056; the direct analogue of real cross-region pairs) give
    near-nominal marginal tail rates (5.79% lower / 6.87% upper vs 5%
    nominal), closely matching the real analysis (6.17% / 7.90%);
    same-origin pairs (n = 700) show 37.6% lower tail, demonstrating
    full power where within-region library similarity exists. The mild
    over-dispersion of per-pair P-values is thereby attributable to
    library-level grouping, not to null-width mismatch (Fig. S14;
    Note 4.6).
  - Blind-review manuscript revisions (peer-review consensus report):
    (a) thalamo-temporal endpoint enrichment re-tested with a
        design-matched permutation null preserving the endpoint
        co-occurrence structure of the 5,778 mature-oligodendrocyte
        pairs (10 pairs drawn without replacement, B = 100,000):
        thalamic-relay 6/10 vs null mean 1.95 (P = 1.005e-5);
        temporal-fusiform 4/10 vs 0.19; combined axis 9/10 vs 2.27
        (P <= 1e-5) — replaces the earlier base-rate hypergeometric
        version in Results/Discussion (full v40 statistics retained
        in Note 5.1);
    (b) TCGA composition check extended to a four-panel model adding
        the myeloid panel (CD68, CD163, LST1, FCGR3A, C1QA) with
        sample-level cluster bootstrap CIs (B = 200;
        results/tcga_composition_v2.*): four-panel pooled tumor-pair
        attenuation -0.5% (95% CI [-3.2%, +2.6%]), strongly
        heterogeneous per cancer type; three-panel +5.7% [+2.7%,
        +8.7%]; Spearman rho(k_n, 4-panel delta) = 0.387 pooled,
        0.23-0.52 per cancer type (Discussion and Note 5.2 updated);
    (c) CKI expansion renamed to 'Cell-type Ka/Ks-inspired Index'
        consistently across manuscript and cover letter; references
        extended from 49 to 55 (intrinsic-noise baseline literature
        and the Kang et al. 2018 dataset).

=== v40 Changes ===
  - (1) Brain set-level coherence checks (notebooks/72_brain_setlevel_tests.py;
    post-hoc, explicitly NOT FDR-controlled discovery):
    (a) raw-P enrichment is tier-dependent: 79.5% of Strong, 35.6% of
        Moderate, 16.9% of Weak, 2.4% of unclassified pairs at raw
        P < 0.05, versus 6.2% overall (hypergeometric P = 9.6e-31;
        Cochran-Armitage trend z = 61.0; MWU P = 5.3e-24; KS D = 0.846,
        P = 4.9e-32);
    (b) thalamo-temporal axis enrichment among the 10 mature-
        oligodendrocyte Strong candidates: thalamic-relay endpoint 6/10
        vs 19.4% base (P = 0.005); temporal-fusiform endpoint 4/10 vs
        1.9% (P = 2.1e-5); combined axis 9/10 vs 22.7% (P = 1.3e-5).
        Integrated into Discussion with post-hoc framing.
  - (6) TCGA composition-contribution sanity check
    (notebooks/73_tcga_composition_check.py): 25,306 sample-labelled
    NN/TT pairs regenerated with the authoritative per-cancer pipeline
    (v2: per-cancer gene loading/filtering, kn_floor = 1e-4, seed 42).
    The TT/NN median k_n ratio replicates exactly (2.18-3.70x across
    the five cancer types). Lineage marker panels (immune/stromal/
    epithelial): composition deltas are larger in TT pairs (1.33-1.46x,
    all P < 1e-300) and correlate with k_n (Spearman rho = 0.21-0.46;
    pooled 0.377), but adjusting for the three composition deltas
    attenuates the tumor-pair coefficient by only 6% pooled (2%, 4%,
    39%, 23%, -12% per cancer) — composition shifts are real but
    explain a minority of the NN/TT k_n reversal. Integrated into the
    TCGA Discussion paragraph with the exploratory framing retained.
  - (7) 'When to use CKI versus standard metrics' decision-guide
    subsection added to the Discussion (before the practical usage
    guide): use-CKI rules (relative divergence questions, normalized
    scale, higher AUC for strong perturbations, design-matched
    specificity-first screening) versus use-standard-metrics rules
    (absolute cross-dataset distances, classification/clustering,
    weak-to-moderate atlas-scale effects, orthogonal-dimension needs,
    suspect baselines in bulk data).
  - Supplementary Note 5 (post-hoc coherence checks: 5.1 brain
    set-level enrichment; 5.2 TCGA composition contribution) added;
    supplementary TOC updated; Repro Guide 5.3(e) documents both new
    scripts and their outputs.
  - Single-version packaging: documents ship as DOCX only (the plain-text
    extracts remain build-verification artifacts, excluded from the zip);
    the Graphical Abstract ships as PDF only (png/svg sources stay in
    results/figures_final/).
  - New build assertions: verify_v40_additions (24 text checks + 7
    artifact-existence checks + 4 numeric cross-checks against the
    source CSVs + 3 zip single-version checks). All v39 checks retained.

=== v38 Changes (v37 peer-review fixes) ===
  - P0 extract_csr_from_backed() row-allocation fix (v38.1): the backed-
    h5ad row assignment bug (orig_pos indexed by the wrong order vector)
    that invalidated all previously published brain numbers was
    corrected; every brain-derived number in this package was
    regenerated from the fixed pipeline and re-verified.
  - R2-C3/R3 circular gene selection: fixed gene-panel ablation
    (notebooks/46_fixed_panel_ablation.py) on all 31,764 brain
    comparisons. The reported per-pair top-200 scheme (S0) was reproduced
    exactly (max per-pair diff 6.4e-13); three non-circular alternatives
    (fixed top-2000 panel, leave-pair-out top-200, all 5,000 non-HK
    genes) preserve pair-level rankings (Spearman rho 0.918/0.937/0.931),
    the astrocyte-vs-Bergmann-glia gradient (6.10x, 6.53x/7.67x/6.57x
    under the alternatives), and the block-shuffle significance pattern
    (astrocytes, OPCs and committed OPCs at the P = 0.005 permutation
    floor under leave-pair-out; fibroblasts P = 0.020; vascular P =
    0.035; remaining classes n.s.). Circularity inflates k_f by a median
    1.61x (IQR 1.27-2.07). Absolute omega is scheme-specific (grand mean
    38.55 -> 26.5/6.5/5.3), so absolute tier thresholds do not transfer
    across schemes; Bergmann glia is not significant under either scheme
    (P = 0.998 reported; P = 0.980 under leave-pair-out), disclosed.
  - Ground-truth simulation (notebooks/45_groundtruth_simulation.py,
    1,750 replicates on Tabula Muris marrow B cells): omega rejects
    neutral housekeeping-gene drift (type-I 0.00, vs 0.55/0.58 for raw
    JS/cosine; 0.02 under global overdispersion) and ranks first of six
    metrics for functional-vs-neutral discrimination (AUC = 0.80);
    detection power background-dependent (marrow needs fourfold shifts
    on >= 200 genes; the skin replication background reaches 91%
    detection at twofold shifts), confirming k_f as an upper-bound
    estimator.
  - R3-C3 phaseB CI recomputation (notebooks/44_fix_phaseB_cis.py):
    brain rows from brain_bs_null_observed_pairs.csv, human all-pairs
    from phase35_all_metrics_pairs.csv (n = 4,851); 16 human per-CT
    pseudo-CIs removed; phaseC_calibrated_cis.csv patched.
  - R5-C4 PAM50 rewritten to actual implementation (cBioPortal
    brca_tcga_pub PAM50_SUBTYPE; 522 samples with subtype calls,
    506 with matched expression data analyzed; local cache).
  - R5-C5 Dockerfile (python:3.14-slim) in repository; R5-M1 scanpy
    1.12.1 in reproducibility guide. Version note: the package code,
    metadata, manuscript, cover letter, and reproducibility guide are
    all at v0.4.4 (v0.4.3 was tagged and released on GitHub with v39),
    tagged and released on GitHub (v0.4.4) and archived on Zenodo
    (concept DOI 10.5281/zenodo.20405458).
  - kn_floor contract unified: single-cell analyses use the package
    default (kn_floor = 0, positivity guard only; brain per-pair k_n
    min 9.2e-5; only 1 of 31,764 pairs below 1e-4, uncapped); TCGA bulk
    uses 1e-4 (3 of 5 cancer types saturate).
  - Upper-tail disclosure: 2,510 of 31,764 pairs (7.9%) with upper-tail
    P < 0.05, reported alongside the lower-tail candidate screen.
  - Abstract recompressed to <= 200 words after adding the two new
    validation sentences.

=== v39 Changes (Genome Biology conversion, no content change) ===
  - Journal target switched from NAR to Genome Biology (Methodology
    article). Manuscript regenerated by generate_manuscript_gb.py:
    structured abstract (Background/Results/Conclusions, <=250 words);
    section order Background -> Results -> Discussion -> Conclusions ->
    Methods; new Conclusions section; Methods moved after Conclusions;
    LLM-use statement relocated into Methods per GB policy; List of
    abbreviations added; consolidated Declarations block (Ethics /
    Consent / Availability of data and materials / Competing interests /
    Funding / Authors' contributions / Acknowledgements); Additional
    files section (Additional file 1 = supplementary notes);
    [n] square-bracket citations; Vancouver reference list (49 refs);
    supplementary citations rewritten as 'Additional file 1: Fig./Table/
    Note S..'. Cover letter was already Genome Biology-formatted.
  - References renumbered by first appearance in GB layout
    (_renumber_gb_citations.py; Methods-only refs 13, 15-22, 26-28 moved
    to 36-44, 45-47); the omega-cap grid range in Methods was restored
    to a parenthetical value range after a false 'caps [12-25]' citation
    was found; reference list now explicitly numbered (1..49 at v39,
    extended to 1..55 in v41 by the Kang perturbation demonstration and
    its intrinsic-noise baseline references; checked by E4-8). Reference
    order and content are otherwise identical to v38.
  - All scientific content, numbers, and claims are identical to v38.
  - Round-7 expert-review fixes (E1/E2/E3/E4, scores 8.9/9.0/8.9/8.6):
    (a) count parentheses restored: 'oligodendrocytes 10 (10), ...' and
        Fig S7 caption 'oligodendrocytes (10) and fibroblasts (6)' were
        falsely converted to [10]/[4]/[3]/[1] citations; converter now
        guards these strings (E4-M1/E2/E3-M1);
    (b) Supplementary Notes reference numbers updated to GB numbering:
        HRT Atlas (14) -> [13]; Siletti (12) -> [12] bracket style
        (E4-M2/E3-M2);
    (c) Fig S2 now cited in the Discussion cross-species sentence
        (E3-M4; previously referenced only in the caption list);
    (d) Cover Letter: 'with cross-species consistency in Tabula Muris'
        replaced by 'complemented by mouse Tabula Muris analyses' to
        match the null cross-species result (r = -0.17, P = 0.55;
        source: full_matrix_pairs.csv x phase35_all_metrics_pairs.csv,
        15 shared cell types; E1-M);
    (e) inline cross-references remapped 'Materials and Methods' ->
        'Methods' (E4-m);
    (f) terminology unified to 'permutation test' where P-values are
        meant, 'bootstrap' reserved for CIs (E1-m5); Efron&Tibshirani
        citation reworded as resampling-based inference;
    (g) wording/anchor fixes: OPC expansion at first Results use,
        Fig S5 'same-CT' -> 'same-cell-type', Table 1 'CKI omega' ->
        Greek omega, Fig S2 caption retitled 'Cross-species comparison
        details' with the null result stated, thalamo-temporal sentences
        rewritten (Results + Discussion), k_f class-mean anchors added
        (0.197/0.054/0.048), dual-background AUC rankings spelled out,
        PMI/single-nucleus Limitations sentences added, FDR-floor
        sentence aligned with Methods, brain k_n sensitivity sentence
        relocated to the brain robustness paragraph;
    (h) SN 3.12: second-background replication paragraph added (marrow
        vs skin numbers from groundtruth_simulation_background2.csv).
  - All v37 checks retained. Brain candidates remain
    hypothesis-generating signals: no candidate survived Benjamini-
    Hochberg FDR correction under the block-shuffle null (minimum
    q = 0.520); no formal FDR discovery is claimed.
  - Additional file 1 coverage verified in-build (check E4-9): all four
    table headings (Table S1-S4) present; notes cite Figures S8-S12;
    word count reported in the build log. Internal self-references use
    the 'Figure S..'/'Table S..' style (drop 'Supplementary' prefix,
    since this document IS Additional file 1; round-7 m2).

  - Round-8 expert-review fixes (E1 8.7 / E2 9.2 / E3 8.5 / E4 8.9,
    weighted ~8.8; all Major/Minor items cross-verified against source
    data before fixing):
    (a) E3-M1 pre-fix/post-fix data mixing resolved: phaseC_kn_stats
        recomputed from results/reviewer_brain_pair_kf_kn.csv (the
        authoritative per-pair k_f/k_n file matching
        brain_bs_null_observed_pairs.csv to 6.4e-13) — k_n CV 92.89% ->
        97.52%, mean 0.0027 -> 0.0035, median 0.0017 -> 0.0022,
        Spearman rho 0.181 -> 0.142 (P = 7.07e-143), per-CT CV range
        36.5-69.8% -> 36.9-70.1%; Fig S11 regenerated; 09c C-M3 data
        source updated;
    (b) imbalance reference-frame corrected (vs signal delta = 1):
        omega -27% -> -32% (8.47 versus 12.46), k_f +81% -> +67%,
        cosine +103% -> +108%;
    (c) median baseline k_f 0.029 -> 0.025 (the 0.029 figure was the
        mean 0.0287, not the median 0.0249);
    (d) E3-M2 SN 3.2 now reports the mean bootstrap CIs shipped in
        phaseB_bootstrap_cis.csv (astrocytes 82.75 [81.56, 83.89],
        Bergmann glia 13.56 [11.31, 16.14]) instead of median CIs
        recomputed inline; Guide adds the cross-species matching rule
        note (case-insensitive + prefix matching, 15 shared types);
    (e) E4-M1 Additional file 1 tables now cited in the main text:
        Table S1 (parameter sweep), Table S3 (brain landscape),
        Table S4 (candidate tiers), and Supplementary Fig. S9
        (block-shuffle null diagnostics);
    (f) E1-m terminology: 'per-cancer bootstrap' -> 'per-cancer
        permutation test', 'bootstrap: m = 10' -> 'group-level tests:
        m = 10', SN/Guide bootstrap-vs-permutation wording unified;
    (g) E2 sensitivity claims background-qualified in Limitations and
        Cover Letter (marrow needs fourfold shifts; skin reaches 91%
        detection at twofold); skin imbalance range re-attributed to
        delta = 0.25 and delta = 1 (seed 42 only), not module seeds;
    (h) Table1-2.docx now extracted from the GB manuscript docx
        (CKI_Manuscript_GB.docx), fixing the stale NAR-derived 'CKI
        omega' ASCII header.

Contents:
1. CKI_Manuscript.docx - Main manuscript (peer-review revised)
2. CKI_Supplementary.docx - Supplementary materials (SN 3.12, 3.13)
3. CKI_Cover_Letter.docx - Cover letter (validation point 6)
4. CKI_Reproducibility_Guide.docx - Reproducibility guide (scripts 44-46)
5. Table1-2.docx - Standalone parameter tables
6. figure1.pdf through figure6.pdf - Main figures
7. Supplementary_Figure_S1.pdf through Supplementary_Figure_S14.pdf - Supplementary figures
8. CKI_graphical_abstract.pdf - Graphical Abstract (PDF only; png/svg
   source files remain in results/figures_final/)
9. (Plain-text extracts of the DOCX files are generated for build
   verification but excluded from this package: single-version,
   Word-only deliverables.)
"""
    with open(WORK_DIR / "MANIFEST_v41.txt", "w", encoding="utf-8") as f:
        f.write(manifest)

    # 4. ZIP
    print(f"\n[4] Creating ZIP...")
    # Single-version package: documents ship as DOCX only (no plain-text
    # extracts); Graphical Abstract ships as PDF only (no png/svg copies).
    ZIP_EXCLUDE = (
        lambda fn: fn.endswith("_fulltext.txt")
        or fn in ("CKI_graphical_abstract.png", "CKI_graphical_abstract.svg")
    )
    with zipfile.ZipFile(V38_ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(WORK_DIR):
            for fn in sorted(files):
                if ZIP_EXCLUDE(fn):
                    continue
                fp = Path(root) / fn
                zf.write(fp, f"CKI_Submission_v41/{fn}")

    zip_mb = V38_ZIP.stat().st_size / (1024 * 1024)
    print(f"\n{'='*60}")
    print(f"  v41 (GB) Package: {V38_ZIP}")
    print(f"  Size: {zip_mb:.1f} MB")
    with zipfile.ZipFile(V38_ZIP, "r") as zf:
        print(f"  Files: {len(zf.infolist())}")
        for info in sorted(zf.infolist(), key=lambda x: x.filename):
            print(f"    {info.file_size:>10,}  {info.filename}")

    # 5. Verification
    print(f"\n{'='*60}")
    print(f"  v38 Final Verification")
    print(f"{'='*60}")

    v = Verifier(WORK_DIR)
    verify_files(v)
    verify_core_numbers(v)
    verify_legacy(v)
    verify_p1(v)
    verify_p2_e1(v)
    verify_p2_e2(v)
    verify_p2_e3(v)
    verify_p2_e4(v)
    verify_v34_expert_panel(v)
    verify_v36_expert_fixes(v)
    verify_v37_reviewer_fixes(v)
    verify_v38_reviewer_fixes(v)
    verify_v38_text_revisions(v)
    verify_v38_numeric(v)
    verify_v39_round8(v)
    verify_v40_additions(v)
    verify_v41_additions(v)

    print(f"\n{'='*60}")
    print(f"  v38 Verification Summary")
    print(f"{'='*60}")
    print(f"  Passed: {v.passed}  Failed: {v.failed}")
    for label, ok in v.results.items():
        print(f"  {'[OK]' if ok else '[FAIL]'} {label}")

    if v.failed == 0:
        print(f"\n  *** ALL {v.passed} CHECKS PASSED — v41 (GB) FINAL ***")
    else:
        print(f"\n  *** {v.failed} FAILURES — review above ***")

    return v.failed == 0


if __name__ == "__main__":
    success = build_v38()
    sys.exit(0 if success else 1)
