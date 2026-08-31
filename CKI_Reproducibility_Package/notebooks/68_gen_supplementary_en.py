"""
Generate English Supplementary Materials DOCX with continuous line numbers.
Replaces the Chinese supplementary with English translation in Chinese-researcher style.
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from _load_manuscript_data import get_manuscript_data
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from lxml import etree

# Load all manuscript data dynamically
DATA = get_manuscript_data()
_br = DATA['brain']
_h = DATA['human']
_mc = DATA['mouse_calibration']

import json
import numpy as np
import pandas as pd
from scipy import stats

# Brain block-shuffle observed pairs (current pipeline: 08d/08e)
_obs = pd.read_csv(Path(__file__).resolve().parent.parent / "results" / "brain_bs_null_observed_pairs.csv")

# Region-associated candidates for S4 top-5 (Strong tier from the block-shuffle re-analysis)
_mig = _obs
_strong = _mig[_mig['tier'] == 'Strong'].sort_values('residual').head(5)
_candidates = _mig[_mig['tier'].isin(['Strong', 'Moderate', 'Weak'])]

# Per-pair k_n variability (phaseC)
_knv = json.loads((Path(__file__).resolve().parent.parent / "results" / "phaseC_kn_variability.json").read_text())
_kn_overall = _knv['brain_overall']
_kn_per_ct = _knv['per_cell_type']
_kn_rho = _knv['omega_correlation']['spearman_rho']
_kn_rho_p = _knv['omega_correlation']['p_value']
_kn_cv_lo = min(_kn_per_ct.items(), key=lambda kv: kv[1]['kn_cv'])
_kn_cv_hi = max(_kn_per_ct.items(), key=lambda kv: kv[1]['kn_cv'])

# Scheme-matched split-half internal baselines (reviewer C-B; notebooks 42/43)
_bs_txt = (Path(__file__).resolve().parent.parent / "results" / "reviewer_brain_splithalf_summary.txt").read_text()
_bs = dict(kv.split(None, 1) for kv in (ln.strip() for ln in _bs_txt.splitlines()) if '\t' in kv or ' ' in kv)
_brain_sh_mean = float(_bs['brain_split_half_mean_omega'])
_brain_sh_ci = [float(x) for x in _bs['brain_split_half_ci95'].strip('[]').split(',')]

_ts_txt = (Path(__file__).resolve().parent.parent / "results" / "reviewer_ts_splithalf_summary.txt").read_text()
_ts = dict(kv.split(None, 1) for kv in (ln.strip() for ln in _ts_txt.splitlines()) if '\t' in kv or ' ' in kv)
_ts_sh_mean = float(_ts['ts_split_half_mean_omega'])
_ts_sh_ci = [float(x) for x in _ts['ts_split_half_ci95'].strip('[]').split(',')]

# Normality tests on the current-pipeline omega distributions
_norm_brain_p = float(stats.normaltest(_obs['omega'])[1])
_norm_human_p = float(stats.shapiro(pd.read_csv(
    Path(__file__).resolve().parent.parent / "results" / "phase35_all_metrics_pairs.csv")['omega'])[1])
_norm_mouse_p = float(stats.shapiro(pd.read_csv(
    Path(__file__).resolve().parent.parent / "results" / "mouse_pilot_v2_results.csv")['omega'])[1])
_norm_mouse_skew = float(stats.skew(pd.read_csv(
    Path(__file__).resolve().parent.parent / "results" / "mouse_pilot_v2_results.csv")['omega']))

# Bootstrap 95% CIs for the median (pair-level resampling, B = 10,000) — examples for SN3.2
_rng = np.random.default_rng(42)


def _median_ci(values, B=10000):
    v = np.asarray(values, dtype=float)
    meds = np.array([np.median(_rng.choice(v, size=len(v), replace=True)) for _ in range(B)])
    return float(np.median(v)), float(np.percentile(meds, 2.5)), float(np.percentile(meds, 97.5))


_ci_astro = _median_ci(_obs[_obs['cell_type'] == 'Astrocyte']['omega'])
_ci_berg = _median_ci(_obs[_obs['cell_type'] == 'Bergmann glia']['omega'])

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Page margins (NAR: 2.5cm)
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

# Heading styles
for lvl, size in [(1, 16), (2, 14), (3, 12)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Arial'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


# ===== TITLE PAGE =====
add_heading('Supplementary Materials', 1)
add_para('CKI: A Cell-type Identity Index for Quantifying Baseline-Normalized Divergence')
add_para('Li Zhang')
add_para('')

add_heading('Table of Contents', 2)
toc = [
    'Supplementary Note 1: CKI Mathematical Derivation',
    'Supplementary Note 2: CKI Algorithm Pseudocode',
    'Supplementary Note 3: Statistical Testing Details',
    'Supplementary Note 4: Dataset Quality Control and Filtering Criteria',
    'Supplementary Table 1: Parameter Sweep Results (Phase 3.2)',
    'Supplementary Table 2: Cross-Organ Conservation Data (Phase 3.5)',
    'Supplementary Table 3: Human Brain Non-neuronal Cell Regional CKI Data',
    'Supplementary Table 4: Inter-regional Region-Associated Candidate Data',
    'Supplementary Data 1: Complete Analysis Script Index',
]
for item in toc:
    add_para(item)

doc.add_page_break()

# ===== SN1: Mathematical Derivation =====
add_heading('Supplementary Note 1: CKI Mathematical Derivation', 2)

add_para('1.1 Jensen-Shannon Divergence', bold=True)
add_para(
    'The Jensen-Shannon (JS) divergence is a symmetrized and smoothed version of the '
    'Kullback-Leibler divergence. For two probability vectors p and q: '
    'JS(p, q) = 1/2 D(p||m) + 1/2 D(q||m), where m = 1/2(p+q), '
    'and D(p||q) = \u03a3 p_i log2(p_i/q_i). When using the base-2 logarithm, '
    'the JS divergence is bounded in [0, 1]. This bound is important for interpreting '
    'omega: when both k_n and k_f approach 1, omega = k_f/k_n may still vary. '
    'A small floor value (1e-4) on k_n is applied only in the TCGA bulk RNA-seq analysis '
    '(where pseudobulk averaging across millions of cells compresses HK gene variance and '
    'drives aggregate k_n toward zero) to prevent inflated omega from near-zero denominators; '
    'all single-cell analyses (mouse, Tabula Sapiens, brain atlas) apply no floor '
    '(kn_floor = 0, positivity guard only; minimum observed per-pair k_n: '
    '1.1e-4, mouse; 6.3e-4, human; 9.2e-5, brain; under the reported brain pipeline, '
    'only 1 of 31,764 pairs had k_n below 1e-4, and these values entered omega uncapped). '
    'Before computing JS divergence, both pseudobulk vectors are normalized to probability '
    'distributions by adding a +1 pseudo-count followed by L1 normalization '
    '(p_i = (x_i + 1)/\u03a3(x_j + 1), applied to log1p-transformed counts).'
)

add_para('1.2 Baseline Divergence Rate k_n', bold=True)
add_para(
    'Housekeeping (HK) genes are defined as genes that maintain stable expression '
    'across cell types and conditions. Let H = {g1, ..., gM} be the set of HK gene '
    'indices. Given pseudobulk vectors \u03bc_A and \u03bc_B (length G, total number of genes), '
    'the baseline divergence rate is: k_n = JS(norm(\u03bc_A[H]), norm(\u03bc_B[H])), '
    'where norm() denotes +1 pseudo-count addition followed by L1 normalization. '
    'Rationale: HK genes should not exhibit systematic differences between biologically '
    'identical cell populations. The JS divergence observed on HK genes therefore reflects '
    'baseline noise: technical variation, stochastic transcriptional bursting, and '
    'individual-level physiological differences. k_n thus provides an internal baseline, '
    'heuristically analogous to Ks (synonymous substitution rate) in molecular evolution. '
    'HK gene set selection: HK genes were loaded from the HRT Atlas v1.0 reference '
    '(1,130 human-mouse conserved HK genes) (13). For mouse datasets, the mouse ortholog '
    'column is used; for human datasets (Tabula Sapiens, TCGA, brain atlas), the human '
    'gene column is used. The CKI package also supports data-driven auto-detection via '
    'detect_housekeeping_genes() (combined criterion: detection rate > 0.9 and CV < 30th '
    'percentile, use_reference = False), but all reported analyses use the pre-specified '
    'HRT Atlas reference. Sensitivity analysis indicates that CKI results are robust to HK set '
    'selection: using the top 10% lowest-variance genes as an alternative constrained set '
    'yields omega correlations of r > 0.95.'
)

add_para('1.3 Functional Divergence Rate k_f', bold=True)
add_para(
    'Identity genes I are defined as genes that capture cell-type-specific functional '
    'programs. In the default configuration (w1 = 1.0, w2 = 0.0), I consists of the '
    'top-N highly variable genes (HVG), excluding HK genes. The functional divergence '
    'rate is: k_f = JS(norm(\u03bc_A[I]), norm(\u03bc_B[I])). Extended configurations '
    'can incorporate additional gene sets: (1) regulon activity genes \u2014 genes enriched '
    'for cell-type-specific transcription factor motifs; (2) pathway enrichment genes '
    '\u2014 genes from MSigDB pathways differentially active between the two groups; '
    '(3) macro-gene embeddings \u2014 gene-level embeddings from protein language models '
    '(e.g., ESM-2). These extensions use a weighted formulation: '
    'k_f = w1*JS(HVG) + w2*JS(pathway) + w3*JS(macro). Parameter sweep (Phase 3.2) '
    f'showed that the pure identity gene configuration (w1=1.0, w2=w3=0.0) achieved '
    f'optimal cell type discrimination (AUC = {DATA["sweep"]["identity_auc"]:.3f}, n = {DATA["sweep"]["n_pairs"]:,} '
    f'mouse cell-type pairs); this was therefore adopted as '
    'the default scheme.'
)

add_para('1.4 Omega Ratio and Its Interpretation', bold=True)
add_para(
    'omega = k_f/k_n. Interpretation is anchored in the empirical omega distribution '
    'within each dataset rather than in fixed ratio cut-offs: values close to the '
    'equivalent-population calibration baseline are consistent with baseline '
    'expectation, with no evidence of functional reprogramming; values far above the '
    'baseline indicate functional divergence exceeding baseline drift, i.e. evidence of '
    'functional transcriptional reprogramming beyond baseline; values far below the '
    'baseline indicate functional constraint, the two groups being more similar in '
    'functional genes than expected from baseline drift (rare in practice). '
    'The Ka/Ks analogy is structurally similar but mathematically non-equivalent. '
    'Key differences: (1) Ka/Ks operates on sequence alignments with explicit codon '
    'models, while CKI operates on continuous expression vectors; (2) the neutral '
    'reference in Ka/Ks has a mechanistic basis in the genetic code (synonymous changes '
    'are assumed neutral), whereas HK genes in CKI are empirically defined; (3) Ka/Ks '
    'uses explicit evolutionary models (e.g., PAML), while CKI uses empirical bootstrap '
    'inference.'
)

add_para('1.5 Bootstrap Permutation Test', bold=True)
add_para(
    'Statistical inference is performed by generating a null distribution of omega '
    'under the null hypothesis that the group labels are exchangeable between the two '
    'cell populations. Procedure: (1) Annotate all cells in the pooled dataset with their '
    'original group labels (A or B); (2) Randomly permute group labels B times '
    '(B=1,000): cell labels for mouse and human, sample (tumor/normal) labels for TCGA; '
    'for the brain atlas the authoritative null is the library-to-region block-shuffle '
    'permutation rather than label permutation; pseudobulk vectors and omega_null are '
    'recomputed each time, with the top-N identity genes re-selected on the permuted '
    'pseudobulks in the mouse, human, and brain pipelines, while the TCGA per-cancer '
    'bootstrap holds a fixed HVG panel across permutations (anti-conservative relative '
    'to re-selection); '
    '(3) Empirical P-value (one-sided): '
    'P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1), with the '
    '+1 term avoiding P = 0; (4) Effect size: SES = (omega_obs - '
    'mean(omega_null))/sd(omega_null). Benjamini-Hochberg FDR correction is '
    'applied within each dataset to control the false discovery rate. '
    'Because the test is one-sided, rejection at a given alpha is decided '
    'directly by the empirical P-value (P < alpha); no separate two-sided '
    'critical values are derived from the null distribution, and the null '
    'quantiles are not confidence intervals for omega itself. '
    'Permutation testing was performed for all four datasets: label permutation for the '
    'mouse pilot (15 cell-type pairs), human Tabula Sapiens, and TCGA (B=1,000 each), and '
    'the block-shuffle null for the brain atlas (B=1,000). For the larger-scale analyses, results are supplemented with non-parametric '
    'statistical tests (Spearman correlation, Mann-Whitney U, Kruskal-Wallis, '
    'Jonckheere-Terpstra) and descriptive statistics.'
)

add_para('1.6 Pseudobulk Construction', bold=True)
add_para(
    'Raw count matrices X (cells x genes) are preprocessed as follows: '
    '(1) Library size normalization: X_norm = 10,000 * (X/colSums(X)); '
    '(2) log1p transformation: X_log = log1p(X_norm), stabilizing variance and reducing '
    'the influence of high-expression outliers; (3) Pseudobulk: mu = column-wise mean '
    'of X_log for all cells with the same cell type annotation, with a minimum of 10 '
    'cells per group. For TCGA bulk RNA-seq data, TPM normalization is used instead: '
    'TPM values from UCSC Xena, followed by log2(TPM + 1) transformation. No pseudobulk step '
    'is needed as each sample is already a bulk expression profile.'
)

doc.add_page_break()

# ===== SN2: Algorithm Pseudocode =====
add_heading('Supplementary Note 2: CKI Algorithm Pseudocode', 2)
add_para('Algorithm 1: CKI Core Computation', bold=True)
add_para(
    'Input: Two cell populations A and B (expression matrices), HK gene set H, '
    'identity gene set I (default top-N HVG excluding H). '
    'Output: omega, P-value, SES, null distribution.'
)
pseudo = [
    ' 1. X_A, X_B <- library-normalize and log1p-transform A and B',
    ' 2. mu_A <- mean(X_A, axis=0); mu_B <- mean(X_B, axis=0)  // pseudobulk',
    ' 3. mu_A_H <- mu_A[H]; mu_B_H <- mu_B[H]',
    ' 4. k_n <- JS_divergence(norm(mu_A_H), norm(mu_B_H))  // norm(x) = (x+1)/sum(x+1) on log1p counts',
    ' 5. mu_A_I <- mu_A[I]; mu_B_I <- mu_B[I]',
    ' 6. k_f <- JS_divergence(norm(mu_A_I), norm(mu_B_I))',
    ' 7. if (bulk RNA-seq analysis) and k_n < 1e-4: k_n <- 1e-4  // TCGA-only floor; single-cell analyses apply no floor',
    ' 8. omega <- k_f / k_n',
    ' 9. // Permutation test',
    '10. labels <- concatenate([A]*n_A, [B]*n_B)  // for bulk RNA-seq, one label per sample',
    '11. for b = 1 to B (B = 1,000 for all datasets):',
    '12.     labels_perm <- random_permutation(labels)',
    '13.     mu_perm1 <- mean(pooled[labels_perm[:n_A]], axis=0); mu_perm2 <- mean(pooled[labels_perm[n_A:]], axis=0)',
    '14.     I_b <- top-N identity genes selected on (mu_perm1, mu_perm2)  // re-selected at every permutation',
    '15.     omega_null[b] <- CKI_core(mu_perm1, mu_perm2, H, I_b)',
    '16. // Inference (one-sided permutation test)',
    '17. P <- (count(omega_null >= omega_obs) + 1) / (B + 1)',
    '18. d <- (omega - mean(omega_null)) / sd(omega_null)',
]
for line in pseudo:
    p = add_para(line)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)

add_para(
    'Note: the released bootstrap_test() API resolves H and I once and holds them fixed '
    'across permutations; the reported mouse, human, and brain analyses re-select I at '
    'every permutation as shown in line 14, and the brain analysis uses a library-level '
    'block-shuffle null instead of label permutation (see 1.5 Bootstrap Permutation Test).'
)
add_para('')
add_para('Algorithm 2: Pairwise Identity Gene Selection (Tabula Sapiens Extension)', bold=True)
add_para(
    'Unlike the Tabula Muris full pairwise matrix (global HVG set for Fig. 2), Tabula Sapiens and all pilot analyses (mouse calibration, human, TCGA, brain) employ pairwise identity '
    'gene selection to avoid dilution of HVG across 102 cell types.'
)
pseudo2 = [
    'Input: Pseudobulk vectors mu_A, mu_B; HK set H; top-N parameter N (default 200)',
    '1. Delta <- |mu_A - mu_B|  // per-gene absolute expression difference',
    '2. I <- indices of top-N genes ranked by descending Delta, excluding H',
    '3. k_f <- JS(norm(mu_A[I]), norm(mu_B[I]))',
    '',
    'Note: k_n uses the global HK set (same for all pairs); k_f uses pairwise top-N genes.',
]
for line in pseudo2:
    p = add_para(line)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ===== SN3: Statistical Testing =====
add_heading('Supplementary Note 3: Statistical Testing Details', 2)
add_para('3.1 Tests Performed and Correction Strategy', bold=True)
add_para(
    'The following statistical tests were used in this study: Mann-Whitney U test '
    '(two-sided) for independent comparisons between two groups; Kruskal-Wallis test '
    'for multi-group comparisons (e.g., BRCA PAM50 subtypes); Jonckheere-Terpstra trend '
    'test for ordered categorical variables (e.g., LIHC Edmondson grade); Spearman rank '
    'correlation for correlations between metrics; Bootstrap permutation test (B=1,000) '
    'for CKI omega significance inference; ROC-AUC for cell type classification '
    'performance assessment.'
)

add_para('3.2 Bootstrap Details', bold=True)
add_para(
    'Bootstrap iterations: B=1,000 for all datasets (mouse pilot study with '
    '15 cell-type pairs, human Tabula Sapiens via script 08b, TCGA via script 08a, '
    'and brain atlas via scripts 08d-08e, which implement the block-shuffle null). '
    'Bootstrap permutation testing was '
    'performed for all four datasets. Benjamini-Hochberg FDR correction is '
    'applied within each dataset to control the false discovery rate. '
    'For the calibration '
    'experiment, empirical P-values '
    'are computed as: P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1). '
    'Standardized effect size = (\u03c9_obs - mean(\u03c9_null)) / sd(\u03c9_null). '
    'In all CKI results, standardized effect sizes are typically > 1.0 for '
    'biologically meaningful comparisons, indicating large effects relative to '
    'the null distribution.'
)
add_para(
    'Bootstrap confidence intervals (95% CI) for all key \u03c9 estimates were '
    'computed by pair-level resampling with B=10,000 iterations. For each cell '
    'type, observed pair-level \u03c9 values were resampled with replacement and the '
    'median was computed; the 2.5th and 97.5th percentiles of the resulting '
    'distribution define the 95% CI. A region-clustered block bootstrap (B = 2,000; the 108 '
    'regions resampled with replacement, all statistics recomputed per resample) yields wider, '
    'cluster-aware intervals for landscape-level quantities: gradient 6.10 [5.55, 9.63], '
    'grand mean \u03c9 38.55 [36.35, 40.73], Strong-candidate count 39 [12, 74]. '
    'Confidence interval widths scale inversely '
    'with the number of contributing pairs: well-sampled cell types (e.g., '
    f'astrocytes, 5,778 pairs) yield narrow intervals ([{_ci_astro[1]:.2f}, {_ci_astro[2]:.2f}], '
    f'median {_ci_astro[0]:.2f}), whereas cell types with fewer comparisons '
    f'(e.g., Bergmann glia, 21 pairs) produce wider intervals '
    f'([{_ci_berg[1]:.2f}, {_ci_berg[2]:.2f}], median {_ci_berg[0]:.2f}).'
)

add_para('3.3 Multiple Testing Correction', bold=True)
_bs = _br['bs_null']
add_para(
    'Bootstrap permutation testing was performed for all four datasets with B=1,000: '
    'mouse pilot (15 cell-type pairs), human Tabula Sapiens, TCGA, and '
    'brain atlas. Benjamini-Hochberg FDR correction is applied within each '
    'dataset to control the false discovery rate, with the number of tests '
    'defined at the level at which each analysis is performed. For the human '
    f'atlas, {_h["n_ct_analyzed"]} of 102 cell-type entries passed the pairwise-analysis filters '
    f'(at least 20 cells per entry and a donor with at least 10 cells), yielding '
    f'C({_h["n_ct_analyzed"]}, 2) = {_h["n_pairs_total"]:,} pairs; for the brain atlas, 31,764 '
    'same-cell-type cross-region pairs were tested. For the larger-scale analyses, bootstrap results '
    'are supplemented with non-parametric '
    'statistical tests and descriptive statistics (median, IQR, effect sizes). For TCGA stratified '
    'analyses (BRCA PAM50, LIHC Edmondson) involving 4-5 groups, omnibus tests '
    '(Kruskal-Wallis, Jonckheere-Terpstra) are used. Effect sizes are reported alongside '
    'all significance statements to distinguish statistical significance from biological '
    'magnitude.'
)
add_para(
    f'For the brain region-association screen, per-pair empirical P-values were computed with a '
    f'block-shuffle permutation null (B = {_bs["B"]:,}), in which 10x Chromium libraries '
    '(sample_id) were treated as blocks and the library-to-region assignment was randomly permuted '
    '(preserving the observed per-region library-count structure), after which region pseudobulks, all '
    f'31,764 pair \u03c9 values, and multiplicative residuals were recomputed. The one-sided lower-tail '
    f'P-value is P = (count(\u03c9_null \u2264 \u03c9_obs) + 1)/(B + 1); the complementary upper-tail P-value '
    f'P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1) was also computed for every pair, and both tails are '
    f'reported here. Of the {_br["n_strong"]} Strong-tier candidates '
    f'(residual < 0.3), {_br.get("n_significant", 31)} showed raw P < 0.05, but after Benjamini-Hochberg '
    f'correction across all m = 31,764 pairs, none reached q < 0.05 (minimum q = {_br["min_q_fdr"]:.3f}). '
    f'Globally, {_bs["n_p_lt_05"]:,} of 31,764 pairs (6.2%) showed raw lower-tail P < 0.05, above the '
    f'permutation null mean of 1,588 and beyond its 99th percentile (1,956): calibrated against the '
    f'design-matched null distribution of the count itself, this excess is significant (P = 0.011; 0.022 '
    f'after Bonferroni correction across the two tails tested), and it is contributed chiefly by classes '
    f'whose mean \u03c9 sits below the null expectation\u2014microglia (8.06% of class pairs at lower-tail '
    f'P < 0.05, +173 pairs above the 5% null rate), oligodendrocytes (6.01%, +58), and Bergmann glia '
    f'(8 of 21 pairs, 38.1%, the family nominally significant under stratified BH)\u2014but also by '
    f'astrocytes (7.60%, +150 pairs), whose wide, right-skewed \u03c9 distributions generate pairs far '
    f'below the null even though the class mean lies far above it. The complementary upper tail '
    f'was also tested and shows a clear excess in the opposite direction: '
    f'{_bs.get("n_p_high_lt_05", 2510):,} pairs ({_bs.get("pct_p_high_lt_05", 7.9):.1f}%) had upper-tail '
    f'P < 0.05 (observed \u03c9 above the block-shuffle null), concentrated in the most regionally structured '
    f'classes (astrocytes 16.2%, committed OPCs 9.7%, OPCs 9.6% of class pairs; Bergmann glia 0%). The '
    f'two-sided excesses are the pair-level signature of the same regional structure detected at the '
    f'cell-class level\u2014classes with upward-shifted \u03c9 distributions contribute the upper-tail excess, '
    f'while the lower-tail excess, though concentrated in the below-null classes, also draws a nearly equal '
    f'contribution from the wide astrocyte distribution\u2014and the lower-tail excess in aggregate is not '
    f'evidence for anomalously low-\u03c9 pairs beyond what the class-level shifts predict. '
    'The null FDR outcome for the candidate screen is structural rather than purely empirical: with m = 31,764 tests, '
    'the BH threshold for the smallest ordered P-value (0.05/31,764 \u2248 1.6 \u00d7 10\u207b\u2076) lies roughly '
    '600-fold below the smallest resolvable permutation P (9.99 \u00d7 10\u207b\u2074 at B = 1,000), so '
    'q < 0.05 would require either B \u2248 6 \u00d7 10\u2075 permutations or at least ~635 of the 31,764 '
    'P-values at the permutation floor\u2014neither condition is approached in these data, as the '
    'selection-rule analysis quantifies (the null alone would pass the complete '
    'Strong rule 148.3 times in expectation, 3.8-fold above the 39 observed candidates; for the rule '
    'restricted to raw P < 0.05 the null expectation of 131.2 exceeds the observed 31 4.2-fold). '
    'We therefore report the Strong candidates as a prioritized hypothesis-generating list rather than '
    'FDR-controlled discoveries. An earlier implementation of this test, which shuffled cell-type '
    'labels within each region pair (B = 10,000), produced anti-conservative P-values (36.3% of pairs at '
    'the P-value floor) because per-pair shuffling ignores the block structure of 10x libraries; that '
    'implementation was superseded by the block-shuffle null reported here. Per-signal tests are not '
    'independent (the same cell type or region pair appears in multiple comparisons); interpretation is '
    'therefore restricted to the predefined Strong tier. (Supplementary Figure S8: \u03c9 distribution '
    'characterization; Supplementary Figure S9: block-shuffle null distribution for the residual model.)'
)

add_para('3.4 Reporting Conventions', bold=True)
add_para(
    'Summary statistics are reported as mean +/- standard deviation (range) or median '
    '[interquartile range]. Boxplots display: median (center line), IQR (box), '
    '1.5x IQR (whiskers), with data points beyond the whiskers shown as outliers. '
    'All P-values from non-bootstrap tests are two-sided; bootstrap permutation '
    'P-values are one-sided (see SN 1.5). Correlation '
    'coefficients (Spearman rho) are reported with P-values. Effect sizes (standardized effect size, SES = (\u03c9_obs \u2212 \u03bc_null) / \u03c3_null) '
    'are reported as descriptive measures of magnitude; because the \u03c9 distribution '
    'is right-skewed and non-normal for the large datasets (brain: D\u2019Agostino-Pearson '
    f'P < 2.2 \u00d7 10\u207b\u00b9\u2076; human: Shapiro-Wilk P = {_norm_human_p:.1e}; the mouse pilot at n = 15 '
    f'pairs does not reject normality, Shapiro-Wilk P = {_norm_mouse_p:.3f}), SES should be '
    'interpreted as a '
    'non-parametric descriptive statistic rather than a parametric test result.'
)

add_para('3.5 Calibrated Omega Normalization', bold=True)
add_para(
    'The theoretical baseline omega = 1 (k_f = k_n) is never observed in practice because '
    'identity-gene selection (per-pair top-200 DE genes or HVGs) systematically inflates k_f '
    'relative to k_n. '
    f'Empirical calibration on split-half equivalent populations (mouse, n = {_mc["control_n"]}) yielded a '
    f'mean omega = {_mc["control_mean"]:.2f} (95% bootstrap CI [4.24, 9.24], B = 10,000 resamples of the '
    f'control omega values). We introduce calibrated omega: omega_cal = omega_obs / {_mc["control_mean"]:.2f}, which '
    'rescales all values so that equivalent populations yield omega_cal ~ 1.0. Given the width '
    'of the baseline CI, calibrated values are reported with one significant figure and should '
    'be read as order-of-magnitude estimates rather than precise quantities. Under this '
    f'mouse-derived calibration: mouse controls yield omega_cal = 1.0, the brain global mean becomes omega_cal '
    f'\u2248 {_br["global_mean"] / _mc["control_mean"]:.0f} (raw {_br["global_mean"]:.2f}; range 4.2\u20139.1 across the baseline CI), and the most divergent brain cell type '
    f'({_br["gradient_highest_ct"].lower()}s) yields omega_cal '
    f'\u2248 {_br["gradient_highest_omega"] / _mc["control_mean"]:.0f} (raw {_br["gradient_highest_omega"]:.2f}). '
    f'A scheme-matched split-half calibration performed inside the brain atlas itself (29 populations '
    f'with at least 200 nuclei; B = 50 random splits per population) gave an internal baseline of '
    f'{_brain_sh_mean:.2f} (95% bootstrap CI [{_brain_sh_ci[0]:.2f}, {_brain_sh_ci[1]:.2f}]), approximately 1.5-fold '
    'higher than the mouse-derived factor, indicating that the mouse-derived calibration '
    'overstates omega_cal in the brain dataset. Under the brain-internal baseline, the brain '
    f'global mean corresponds to omega_cal \u2248 {_br["global_mean"] / _brain_sh_mean:.1f}, the most divergent class to '
    f'omega_cal \u2248 {_br["gradient_highest_omega"] / _brain_sh_mean:.1f}, and the most constrained class '
    f'({_br["gradient_lowest_ct"].lower()}) to omega_cal \u2248 {_br["gradient_lowest_omega"] / _brain_sh_mean:.2f} (raw '
    f'{_br["gradient_lowest_omega"]:.2f}): the most regionally constrained class sits at, not above, the '
    'split-half expectation for equivalent brain populations, so the statement that all ten '
    'non-neuronal classes diverge beyond the split-half baseline holds only under the '
    'mouse-derived calibration and is not supported under brain-internal calibration. '
    'In Tabula Sapiens, the analogous scheme-matched internal baseline was '
    f'{_ts_sh_mean:.2f} (95% bootstrap CI [{_ts_sh_ci[0]:.2f}, {_ts_sh_ci[1]:.2f}]; 71 populations from the largest donor per group), which lies inside the '
    'mouse-derived CI [4.24, 9.24], so the mouse-derived calibration factor is transferable to '
    'the Tabula Sapiens dataset but not to the brain atlas, and omega_cal should be treated as a '
    'dataset-relative quantity rather than a universal constant. '
    'The calibrate_omega() function is available in the CKI package (cki.calibrate_omega). '
    'Both raw and calibrated omega values are reported in all key results. (Supplementary Figure S12.)'
)

add_para('3.6 JS Divergence Dimensionality Invariance', bold=True)
add_para(
    'Because k_n is computed on ~1,130 housekeeping (HK) genes and k_f on 200-2,000 highly '
    'variable genes (HVGs), we verified that JS divergence is not systematically biased by '
    'gene set dimensionality. A simulation of 2,000 random Dirichlet distribution pairs '
    'across dimensions ranging from 50 to 5,000 showed that mean JS divergence is effectively '
    'constant across all dimensions tested (range: 0.155-0.159; ratio = 1.001 between d = 1,130 '
    'and d = 2,000). This confirms that the systematic inflation of k_f relative to k_n '
    '(omega = 6.67 for equivalent populations) arises from HVG selection bias (selecting genes '
    'with high cross-cell variance) rather than from dimensional mismatch between the HK and '
    'HVG gene sets. We note that this simulation addresses dimensionality per se (random '
    'probability vectors of different lengths) but does not simulate the variance-based '
    'gene selection mechanism that generates the \u03c9 inflation. A more complete validation '
    'would test whether the inflation magnitude scales with the stringency of variance '
    'filtering rather than gene count. The calibrated omega (omega_cal = omega / 6.67) '
    'absorbs this bias into the empirical baseline, and the permutation null distribution '
    '- constructed using the same gene sets as the observed data - ensures internal '
    'consistency. (Supplementary Figure S10.)'
)

add_para('3.7 Pair-Specific k_n Variability', bold=True)
add_para(
    'In all reported analyses (mouse pilot, human, TCGA, and brain), k_n is computed '
    'per pair on the shared HK gene set: for each comparison, the HK-gene JS divergence '
    'between the two pseudobulk vectors is evaluated on that pair alone (a single HK '
    'reference is shared across all pairs, keeping k_n on a consistent scale, but the '
    'divergence itself is pair-specific). '
    f'Analysis of per-pair k_n across {_kn_overall["n_pairs"]:,} brain comparisons revealed substantial cross-pair '
    f'variability: overall k_n CV = {_kn_overall["kn_cv"]:.2%} (mean = {_kn_overall["kn_mean"]:.4f}, median = {_kn_overall["kn_median"]:.4f}), with per-cell-type '
    f'CVs ranging from {_kn_cv_lo[1]["kn_cv"]:.1%} ({_kn_cv_lo[0].lower()}s) to {_kn_cv_hi[1]["kn_cv"]:.1%} ({_kn_cv_hi[0].lower()}s). '
    'As a sensitivity analysis, we contrasted this per-pair k_n estimator with a global-k_n '
    'variant (k_n computed once from the full gene-by-cell-type pseudobulk matrix): the Spearman '
    f'correlation between the resulting omega rankings was only \u03c1 = {_kn_rho:.3f} (P = {_kn_rho_p:.2e}), so the '
    'global-k_n simplification would preserve only ~3% of the variance in omega orderings '
    '(\u03c1\u00b2 \u2248 0.03). This justifies the per-pair k_n approach used throughout and '
    'highlights that fine-grained omega orderings should be interpreted with the estimator '
    'choice in mind. (Supplementary Figure S11.)'
)

add_para('3.8 TCGA Exploratory Analysis Caveats', bold=True)
add_para(
    'The TCGA pan-cancer analysis (Results Section: Cancer analysis) is exploratory in nature '
    'due to several inherent limitations of bulk RNA-seq data. First, bulk RNA-seq confounds '
    'cell-composition shifts (tumor purity, stromal infiltration, immune cell infiltration) '
    'with genuine transcriptional divergence; the observed NN/TT > 1.0 pattern may partly '
    'reflect shared cell-composition changes across tumors rather than true transcriptional '
    'convergence. Second, peritumoral inflammation and desmoplastic reactions are shared '
    'across tumors and could contribute to apparent convergence. Third, systematic RNA quality '
    'differences between tumor and normal specimens may introduce technical bias. Fourth, '
    'the paired tumor-normal analysis (n = 2-5 per cancer type) lacks statistical power for '
    'formal hypothesis testing (minimum two-sided P approx 0.33 for n = 2); we therefore '
    'report paired comparisons as descriptive statistics only. Fifth, PAM50 and Edmondson '
    'stratification includes small subgroups (Normal-like n = 7; Edmondson G4 n = 11) whose '
    'rankings are unreliable. Sixth, the PAM50 gradient (aggressive subtypes have lower omega) '
    'may be driven by proliferative fraction differences rather than shared transcriptional '
    'states. These caveats do not invalidate the descriptive findings but preclude causal '
    'inference from bulk-level data alone.'
)

add_para('3.9 Cross-Organ Sample Size Considerations', bold=True)
add_para(
    'The cross-organ conservation ranking (Results Section: CKI ranks cell types by cross-organ '
    'conservation) includes cell types with varying numbers of cross-organ pairs. Several cell '
    'types have very few pairs (n = 1-3; e.g., Memory B cells n = 1, Smooth muscle n = 1), '
    'making their mean omega estimates unreliable. We recommend interpreting rankings of cell '
    'types with n < 5 as suggestive only. Bootstrap 95% confidence intervals for cell types '
    'with n >= 5 are provided in Supplementary Table S2. The Spearman correlations between '
    'CKI omega and standard metrics are reported with bootstrap 95% CIs (B = 10,000 resamples).'
)

add_para('3.10 One-Sided Permutation Test Justification', bold=True)
add_para(
    'All bootstrap P-values use a one-sided test: P = (count(omega_null >= omega_obs) + 1)/(B + 1). '
    'The one-sided formulation is appropriate because our hypothesis is directional: we test '
    'whether observed omega exceeds the null expectation (equivalent populations), not whether '
    'it differs in either direction. A two-sided test would be appropriate if we were testing '
    'for any departure from the null (either elevated or suppressed omega). However, the '
    'biological questions addressed here (functional divergence exceeding baseline, Strong '
    'region-associated candidates showing anomalously low omega) are inherently directional.'
)

add_para('3.11 Parameter Justification', bold=True)
add_para(
    'Key CKI parameters and their rationale: (1) +1 pseudo-count followed by L1 '
    'normalization: converts expression vectors to probability distributions for JS '
    'divergence computation; adding one pseudo-count before L1 normalization preserves '
    'relative magnitude information while ensuring non-negativity and sum-to-one. '
    '(2) Pseudocount epsilon = 1e-9: added to avoid log(0) in '
    'JS divergence; the value is small enough to not affect results materially but large enough '
    'to prevent numerical instability. (3) Top-200 DE genes for k_f (mouse): selected per pair '
    'to adaptively capture the most informative identity genes; the parameter sweep (Phase 3.2) '
    'tested N_HVG in {500, 1000, 2000, 3000, 5000} and found 2000 optimal for AUC. (4) HVG '
    'count of 2,000 (human/brain): standard Scanpy default; Seurat flavor used for consistency '
    'with the Scanpy ecosystem. (5) Log-base 2 for JS divergence: standard choice giving JS '
    'range [0, 1]; the base does not affect omega = k_f/k_n since it cancels in the ratio. '
    '(6) B = 1,000 for bootstrap: justified by adaptive permutation analysis showing minimum '
    'P = 9.99 \u00d7 10\u207b\u2074 (= 1/(B+1) = 1/1001) is well below BH thresholds for cell-type-level tests (Phase B).'
)

add_para('3.12 Ground-Truth Simulation', bold=True)
add_para(
    'Semi-synthetic benchmark. To provide a known ground truth, perturbations of known '
    'magnitude were injected into a real single-cell background (Tabula Muris FACS marrow '
    'B cells, 1,848 cells). Each replicate resampled two independent groups of 200 cells '
    'from the same cell type, so the true functional divergence before injection is zero. '
    'The kept gene set mirrored the brain pipeline (1,064 matched HK genes plus the 5,000 '
    'non-HK genes with highest global means). Functional signal: multiplicative shift of '
    '2^delta (delta = 0.125-2) on a fixed module of 200 non-HK genes in group B. Neutral '
    'perturbations, injected separately: (i) 2^eta shift on HK genes in group A '
    '(eta = 0.25-1, neutral drift), (ii) extra Poisson noise across all genes in group A '
    '(epsilon = 0.3-1, technical batch noise). Six metrics were computed with the identical '
    'code path as the brain analysis: omega, k_f, k_n, raw JS over the full kept gene set, '
    'cosine distance, and k_f/k_total. Signal scenarios used three independent module draws '
    '(seeds 42, 137, 2024). Detection thresholds were the 95th percentile of 200 baseline '
    'replicates (pure cell resampling) per metric, so type-I error and power refer to a '
    'common nominal 5% level. 1,750 replicates total; script: '
    'notebooks/45_groundtruth_simulation.py (runtime about 1 minute).'
)
add_para(
    'Type-I error under neutral perturbation. Under pure HK drift (eta = 0.25-1, pooled), '
    'exceedance rates were: omega 0.000, k_f 0.007, k_n 0.813, raw JS 0.553, cosine 0.580, '
    'k_f/k_total 0.000. Exact Clopper-Pearson intervals: omega 0/150 under pure HK drift '
    '(one-sided 95% upper bound 0.0198; 2/250 pooled neutral replicates, 95% CI '
    '[0.001, 0.029]); raw JS 0.553 (95% CI [0.470, 0.635]); cosine 0.580 (95% CI '
    '[0.497, 0.660]). '
    'Under global overdispersion noise (epsilon = 0.3-1, pooled): '
    'omega 0.020 (95% CI [0.002, 0.070]), k_f 0.020, k_n 0.060, raw JS 0.040, cosine 0.040. Standard unnormalized '
    'metrics cannot distinguish neutral drift from functional divergence; omega and the '
    'ratio variants reject neutral drift almost perfectly.'
)
add_para(
    'Power for injected functional signal (detection rate, pooled over three module seeds; '
    'thresholds calibrated at the 95th percentile of the baseline distribution). '
    'delta = 0.125: omega 0.013, k_f 0.013, raw JS 0.040, cosine 0.060. delta = 0.25: '
    'omega 0.000, k_f 0.013, raw JS 0.053, cosine 0.060. delta = 0.5: omega 0.000, k_f '
    '0.000, raw JS 0.140, cosine 0.100. delta = 1: omega 0.000 (all three module seeds), '
    'k_f 0.013, raw JS 0.993, cosine 0.667. delta = 2: omega 0.133 (per-seed 0.12-0.16), '
    'k_f 0.153, raw JS 1.000, cosine 1.000. The omega detection floor is structural: the '
    'per-pair top-200 selection saturates with noise under the null (median null k_f = '
    '0.029), so weak module shifts do not lift k_f above the selection floor, and omega '
    'inherits this as an upper-bound estimator. The baseline omega distribution on this '
    'background (mean 11.6, 95th percentile 20.8) reproduces the scale of the brain-internal '
    'split-half calibration baseline (9.73), supporting external validity.'
)
add_para(
    'Confounded scenarios (delta = 0.5 plus neutral HK drift). Detection rates: '
    'eta = 0.25: omega 0.000, raw JS 0.260, cosine 0.140; eta = 0.5: omega 0.000, raw JS '
    '0.960, cosine 0.940; eta = 1: omega 0.000, raw JS 1.000, cosine 1.000. Neutral '
    'drift suppresses omega further through the inflated denominator while inflating the '
    'unnormalized metrics (which count drift as divergence).'
)
add_para(
    'Discrimination and robustness. ROC AUC for separating functional (delta >= 0.25) from '
    'neutral (HK drift or global overdispersion) replicates: omega 0.804 > k_f 0.716 > '
    'raw JS 0.640 > cosine 0.584 > k_f/k_total 0.437 > k_n 0.213. Under a fourfold '
    'cell-count imbalance (n_B = 50, delta = 1), omega changed by -27% (9.15 versus 12.46) '
    'while k_f inflated by +81% and cosine by +103%; 30% dropout and a twofold depth '
    'difference produced no systematic shift in omega. Module-size sensitivity at '
    'delta = 1: m = 50 (omega mean 11.99), m = 200 (12.46), m = 500 (7.31; large modules '
    'shift total library composition, which leaks into k_n and suppresses omega). '
    'Interpretation: omega is a specificity-first screen; its construction rejects neutral '
    'drift, at the cost of bounded power for weak-to-moderate functional signals.'
)

add_para('3.13 Fixed Gene-Panel Ablation', bold=True)
add_para(
    'Design. To test whether the brain conclusions depend on the per-pair circular '
    'selection of k_f genes (top-200 genes ranked by the absolute difference of the '
    'same two pseudobulks on which k_f is computed), the entire observed brain '
    'landscape was recomputed under four gene-selection schemes with identical keep '
    'gene set, pseudobulks, and k_n: S0, the reported per-pair top-200 scheme '
    '(circular selection; reference); S1, a fixed panel of the 2,000 non-HK genes '
    'with the highest global mean expression; S2, a leave-pair-out panel in which the '
    'top-200 genes for a pair are selected by the mean absolute difference over all '
    'other region pairs of the same cell type (adaptive but not circular for the '
    'tested pair); S3, all 5,000 non-HK genes of the keep set (no selection). The '
    'reference implementation reproduced the reported landscape exactly (maximum '
    'per-pair omega difference 6.4e-13 over 31,764 pairs). A scheme-matched '
    'block-shuffle null (B = 200) was rerun under S2. Script: '
    'notebooks/46_fixed_panel_ablation.py (runtime about 28 minutes).'
)
add_para(
    'Rank robustness. Pair-level Spearman correlation with S0: rho = 0.918 (S1), '
    '0.937 (S2), 0.931 (S3). Ordering of the ten class means: rho = 0.90 (S1), 0.99 '
    '(S2), 0.93 (S3); under S2 the only change was a swap of the adjacent '
    'Bergmann-glia and vascular means. Astrocyte-to-Bergmann-glia ratio: 6.10 (S0) versus 7.67 '
    '(S1), 6.53 (S2), 6.57 (S3) - the gradient is preserved under the '
    'non-circular adaptive panel, because non-circular panels deflate omega more '
    'strongly in transcriptionally constrained classes. Under the S2-matched '
    'block-shuffle null, astrocytes, OPCs and committed OPCs reached the permutation floor '
    '(P = 0.005); fibroblasts remained significant (P = 0.020, versus 0.030 '
    'reported); vascular cells retained significance (P = 0.035, versus 0.115 '
    'reported); ependymal cells were not significant (P = 0.164); and the remaining '
    'classes stayed clearly non-significant (microglia P = 0.796, oligodendrocytes '
    'P = 0.876, choroid plexus P = 0.562, Bergmann glia P = 0.980).'
)
add_para(
    'Circularity inflation and scale. On the same pairs, the circular panel '
    'inflated k_f by a median of 1.61-fold relative to the leave-pair-out panel '
    '(IQR 1.27-2.07) and by 6.1- to 7.3-fold relative to the fixed and unselected '
    'panels (which also differ in panel composition). The absolute omega scale is '
    'therefore scheme-specific (grand mean 38.55 for S0 versus 6.5, 26.5, and 5.3 '
    'for S1-S3). Rank-based conclusions are robust: the multiplicative-residual '
    'ranking correlated at rho = 0.86-0.88 across schemes, and S2 retained 32 of '
    '50 (64%) of the S0 residual < 0.3 candidate pairs. Absolute tier cutoffs '
    '(omega < 15/25/35) are calibrated to the S0 scale and do not transfer across '
    'schemes: candidate lists defined by absolute omega thresholds are not '
    'comparable across gene-selection schemes. Outputs: '
    'results/fixed_panel_ablation_pairs.csv, results/fixed_panel_ablation_ct.csv, '
    'results/fixed_panel_ablation_null_<CT>.npy, results/fixed_panel_ablation_summary.json.'
)

add_para('3.14 Region Glossary (Siletti et al. Dissection Nomenclature)', bold=True)
add_para(
    'All brain-region abbreviations used in the manuscript follow the dissection '
    'nomenclature of the Siletti et al. atlas (12), which is derived from the adult '
    'human brain structural ontology of Ding et al. We extract the ROI-to-dissection '
    'mapping verbatim from the dataset metadata (data/brain/Nonneurons.h5ad, obs '
    'roi and dissection fields; script notebooks/_v38_region_glossary.py, output '
    'results/v38_region_glossary.csv). Candidate-screen endpoints discussed in the '
    'text: A13, caudal intermediate orbital gyrus (orbitofrontal cortex); A14, gyrus '
    'rectus (medial orbitofrontal cortex); A38, temporopolar area; A40, supramarginal '
    'gyrus; A43, parietal operculum (gustatory cortex); A46, middle frontal gyrus; '
    'A5-A7, posterosuperior parietal cortex; TF, temporal area TF of the '
    'occipitotemporal (fusiform) gyrus; Pro, area prostriata; Cla, claustrum; '
    'LG, lateral geniculate nucleus; MG, medial geniculate nuclei; LP, lateral '
    'posterior nucleus; Pul, pulvinar; VA, ventral anterior nucleus; MD, '
    'mediodorsal nucleus; MD-Re, mediodorsal plus reuniens nuclei; CM-Pf, '
    'centromedian and parafascicular nuclei; VPL, ventral posterolateral nucleus '
    '(all thalamic except where noted); STH, subthalamic nucleus (grouped under '
    'thalamus in the dissection ontology but not a thalamic relay nucleus); GPe, '
    'external segment of the globus pallidus; NAC, nucleus accumbens; SI, '
    'substantia innominata; BL, basolateral amygdaloid nucleus; BM, basomedial '
    'amygdaloid nucleus; CMN, corticomedial amygdaloid nuclear group; CA1C-CA3C, '
    'caudal hippocampus (cornu ammonis fields CA1-CA3); CBL, lateral hemisphere of '
    'the cerebellum; CBV, cerebellar vermis; IC, inferior colliculus; DTg, dorsal '
    'tegmental nucleus; PAG-DR, periaqueductal gray and dorsal raphe nucleus; PN, '
    'pontine nucleus; PnAN, afferent nuclei of cranial nerves in pons; PnRF, '
    'pontine reticular formation; MoRF-MoEN, medullary reticular formation and '
    'efferent nuclei of cranial nerves in the medulla oblongata; HTHso-HTHtub, '
    'supraoptic and tuberal hypothalamus. A13 and A14 are cortical '
    '(Brodmann-designated orbitofrontal) areas, not thalamic nuclei; the numbered '
    'A-series labels in the 108-region set are cortical areas, whereas thalamic '
    'nuclei carry their conventional abbreviations (LG, MG, Pul, LP, VPL, VA, MD, '
    'CM-Pf, and related labels). The full 108-region glossary '
    '(abbreviation to full dissection path) ships as '
    'results/v38_region_glossary.csv in the reproducibility package.'
)

doc.add_page_break()
add_heading('Supplementary Note 4: Dataset Quality Control and Filtering Criteria', 2)

add_para('4.1 Tabula Muris FACS (Mouse)', bold=True)
add_para(
    'Downloaded from GEO (GSE109774). FACS-sorted cells (not droplet-based) were used '
    'to ensure high per-cell gene detection. QC filtering: cells with < 500 detected '
    'genes were removed; cells with > 10% mitochondrial gene expression were removed; '
    'genes detected in < 3 cells were removed. Result: 15,057 cells x 22,308 genes '
    '(post-QC). Cell type annotation: 32 cell type entries with >= 10 cells per group '
    'were retained for pseudobulk construction, spanning 6 organs (Liver, Kidney, '
    'Spleen, Lung, Heart, Bone Marrow).'
)

add_para('4.2 Tabula Sapiens (Human)', bold=True)
add_para(
    'Downloaded from CZ CELLxGENE Discover. QC filtering: cells with < 500 detected '
    'genes were removed; cells with > 20% mitochondrial gene expression were removed. '
    'Result: 108,136 cells retained (6 h5ad files total), with 51,852 genes (filtered from the '
    'original 58,870). Cell type entries: 102 entries across 6 organs (Liver, Kidney, '
    'Heart, Bone Marrow, Spleen, Lung). Cell types included in pairwise omega analysis '
    f'were required to have >= 10 cells in at least one donor and at least 20 cells per '
    f'entry ("unknown" annotations excluded); 99 of the 102 entries passed these filters, '
    f'yielding C(99, 2) = 4,851 analyzed pairs. Pairwise identity gene '
    'selection (top-200 genes by |Delta expression| ranking) ensures that each comparison '
    'uses the most informative genes for that specific pair. Human HK genes: HRT Atlas '
    'v1.0 reference (1,130 genes; human column, 1,129 matched to data).'
)

add_para('4.3 TCGA Bulk RNA-seq', bold=True)
add_para(
    'Data were obtained from the NCI Genomic Data Commons. Five cancer types were selected: '
    'LUAD (495 tumor + 76 normal), LUSC (567 + 58), LIHC (365 + 57), KIRC (755 + 82), '
    'BRCA (1,032 + 109), totaling n = 3,596 samples. Normalization: TPM values from UCSC Xena, '
    'followed by log2(TPM + 1) transformation. For paired analysis, tumor-normal pairs were '
    'matched by patient barcode (TCGA-XX-XXXX format). Clinical metadata for stratified '
    'analyses were obtained from GDC (via the TCGAbiolinks R package) and the cBioPortal API.'
)

add_para('4.4 Highly Variable Gene (HVG) Selection', bold=True)
add_para(
    'Tabula Muris: Global HVG selection was performed using '
    'scanpy.pp.highly_variable_genes, with parameters flavor="seurat" and '
    'n_top_genes=2,000. The global HVG set was used for all pairwise comparisons '
    'in Phases 3.1-3.2. Tabula Sapiens: Pairwise HVG selection. For each cell type '
    'pair (CT_i, CT_j), the top-200 genes ranked by |mu_i - mu_j| (absolute log1p '
    'expression difference) were selected as identity genes, excluding HK genes. This '
    'avoids the dilution effect of HVG across comparisons involving 102 cell types. '
    'HVG count sensitivity: the parameter sweep (Phase 3.2) tested N_HVG in '
    '{50, 100, 200, 500, 1,000, 2,000}. The global scheme (mouse) achieved peak AUC '
    'at N=2,000, while the pairwise scheme (human) used N=200 to maintain discriminative '
    'power with computational efficiency.'
)

doc.add_page_break()

# ===== Supplementary Tables =====
add_heading('Supplementary Table 1: Parameter Sweep Results', 2)
add_para(
    f'Phase 3.2 parameter sweep on Tabula Muris mouse data (n = 703 cell type pairs, '
    f'6 organs). The pure identity gene configuration (w1 = 1.0, w2 = 0.0) achieved '
    f'optimal cell type discrimination (AUC = {DATA["sweep"]["identity_auc"]:.3f}). Data file: '
    'results/phase32_sweep_results.csv. Visualization: results/phase32_sweep_barplot.png.'
)

add_para('')
add_heading('Supplementary Table 2: Cross-Organ Conservation Data', 2)
add_para(
    'Complete dataset of 59 same-cell-type cross-organ pairs in Tabula Sapiens, '
    'including omega, Jensen-Shannon divergence, Spearman distance, Cosine distance, '
    'and Marker Jaccard distance values. Data file: '
    'results/phase35_cross_organ_conservation.csv.'
)

add_para('')
add_heading('Supplementary Table 3: Human Brain Non-neuronal Cell Regional CKI Data', 2)
add_para(
    f'Complete results of CKI brain region analysis for non-neuronal cells from the '
    f'Siletti et al. (2023) human brain atlas, comprising {_br["total_pairs"]:,} pairwise cross-region '
    f'comparisons across 10 cell types (n = {_br["n_nuclei"]:,} nuclei, {_br["n_regions"]} regions, '
    f'{_br["n_genes"]:,} genes). Summary statistics for each cell type (omega mean, '
    f'median, SD, range, k_n and k_f components) are provided in Supplementary Table 3. '
    f'Raw data file: results/brain_bs_null_observed_pairs.csv ({_br["total_pairs"]:,} rows, '
    f'block-shuffle re-analysis pipeline). '
    f'Summary file: results/brain_bs_null_ct_test.csv (10-row summary). '
    f'Analysis scripts: notebooks/08d_brain_blockshuffle_null.py and notebooks/08e_brain_blockshuffle_results.py. '
    f'Figure generation: notebooks/_fig6_clean.py (Figure 6).'
)

add_para('')
add_heading('Supplementary Table 4: Inter-regional Region-Associated Candidate Data', 2)

# Build dynamic S4 text
n_candidates = len(_candidates)
pct_candidates = n_candidates / len(_mig) * 100
n_strong = int(_br['n_strong'])
n_moderate = int(_br['n_moderate'])
n_weak = int(_br['n_weak'])
pct_strong = float(_br['pct_strong'])
pct_moderate = float(_br['pct_moderate'])
pct_weak = float(_br['pct_weak'])

# Build top-5 list (Strong tier, lowest multiplicative residuals; strip 'Human ' prefix)
top5_lines = []
for i, (_, r) in enumerate(_strong.iterrows()):
    ra = str(r['region_a']).replace('Human ', '')
    rb = str(r['region_b']).replace('Human ', '')
    top5_lines.append(
        f'{r["cell_type"]} {ra} vs. {rb} '
        f'(omega = {r["omega"]:.2f}, residual = {r["residual"]:.3f})'
    )

s4_text = (
    f'Results of multiplicative model-based detection of candidate inter-regional cell '
    f'region-association signals (block-shuffle re-analysis pipeline). Of the {_br["total_pairs"]:,} total pairwise cross-region comparisons, '
    f'{n_candidates:,} pairs ({pct_candidates:.1f}%) were classified as threshold-passing candidates '
    f'(residual < 0.75): '
    f'{n_strong} strong signals (residual < {DATA["brain"]["residual_thresholds"]["strong"]}, '
    f'{pct_strong:.2f}%), '
    f'{n_moderate:,} moderate signals (residual < {DATA["brain"]["residual_thresholds"]["moderate"]}, '
    f'{pct_moderate:.2f}%), '
    f'and {n_weak:,} weak signals (residual < {DATA["brain"]["residual_thresholds"]["weak"]}, '
    f'{pct_weak:.2f}%). '
    f'Under the block-shuffle permutation null (B = {_bs["B"]:,}), {_br.get("n_significant", 31)} of the {n_strong} '
    f'Strong candidates showed raw one-sided P < 0.05, but none survived Benjamini-Hochberg '
    f'correction across all 31,764 pairs (minimum q = {_br["min_q_fdr"]:.3f}); the candidate list is '
    f'therefore hypothesis-generating only. '
    f'Top-5 strongest signals by cell type (ranked by residual): '
    f'1) {top5_lines[0]}, '
    f'2) {top5_lines[1]}, '
    f'3) {top5_lines[2]}, '
    f'4) {top5_lines[3]}, '
    f'5) {top5_lines[4]}. '
    f'Complete candidate dataset: results/brain_bs_null_observed_pairs.csv '
    f'({n_candidates:,} threshold-passing rows of 31,764 total). '
    f'Analysis scripts: notebooks/08d_brain_blockshuffle_null.py and notebooks/08e_brain_blockshuffle_results.py.'
)
add_para(s4_text)

add_para('')
add_heading('Supplementary Data 1: Analysis Script Index', 2)
add_para(
    'Complete analysis scripts used in this study are organized in the notebooks/ '
    'directory of the GitHub repository (github.com/zhanglknt/CKI-cell-type-identification). '
    'The package can be installed via: pip install git+https://github.com/zhanglknt/CKI-cell-type-identification.git. '
    'Key scripts include: notebooks/04_phase32_sweep.py (parameter sweep and calibration), '
    'notebooks/06_phase34_v2.py (TCGA pan-cancer analysis), '
    'notebooks/05_phase33_v3_fixed.py (Tabula Sapiens cross-organ analysis), '
    'notebooks/07c_brain_siletti_v3.py (brain regional CKI analysis and region-associated candidate detection), '
    'and notebooks/30_genome_biology_figures.py with notebooks/_fig1_clean.py to _fig6_clean.py (figure generation).'
)

# ===== Add line numbers (continuous, every line) =====
for sec in doc.sections:
    sect_pr = sec._sectPr
    ln_num = etree.SubElement(sect_pr, qn('w:lnNumType'))
    ln_num.set(qn('w:countBy'), '1')
    ln_num.set(qn('w:start'), '1')
    ln_num.set(qn('w:restart'), 'continuous')

out_path = 'results/CKI_Supplementary.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'Paragraphs: {len(doc.paragraphs)}')
