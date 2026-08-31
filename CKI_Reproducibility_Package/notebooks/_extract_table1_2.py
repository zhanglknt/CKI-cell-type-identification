#!/usr/bin/env python3
"""Extract Table 1 and Table 2 from CKI_Manuscript.docx into a separate Table1-2.docx."""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "results" / "CKI_Manuscript.docx"
OUT = ROOT / "results" / "Table1-2.docx"


def copy_table(src_table, dst_doc):
    """Copy a table preserving structure, text, and basic formatting."""
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    new_table = dst_doc.add_table(rows=rows, cols=cols)
    # Try to apply the same style; fall back to default
    try:
        new_table.style = src_table.style
    except Exception:
        pass

    for i, src_row in enumerate(src_table.rows):
        dst_row = new_table.rows[i]
        for j, src_cell in enumerate(src_row.cells):
            dst_cell = dst_row.cells[j]
            # Copy text preserving runs for bold/italic if needed
            dst_para = dst_cell.paragraphs[0]
            dst_para.clear()
            for src_para in src_cell.paragraphs:
                for run in src_para.runs:
                    dst_run = dst_para.add_run(run.text)
                    dst_run.font.name = run.font.name or "Arial"
                    dst_run.font.size = Pt(run.font.size.pt if run.font.size else 10)
                    dst_run.bold = run.bold
                    dst_run.italic = run.italic
                    dst_run.font.color.rgb = run.font.color.rgb
    return new_table


def main():
    if not SRC.exists():
        print(f"Source manuscript not found: {SRC}")
        sys.exit(1)

    src_doc = Document(str(SRC))
    body = src_doc.element.body

    # Iterate body children in order, collecting tables and their preceding captions
    tables_in_order = []
    current_caption = None
    for child in body:
        if child.tag.endswith('p'):
            # paragraph element
            para = next((p for p in src_doc.paragraphs if p._element is child), None)
            if para is not None:
                text = para.text.strip()
                if text.startswith("Table 1.") or text.startswith("Table 2."):
                    current_caption = text
        elif child.tag.endswith('tbl'):
            tbl = next((t for t in src_doc.tables if t._element is child), None)
            if tbl is not None and current_caption is not None:
                tables_in_order.append((current_caption, tbl))
                current_caption = None

    if not tables_in_order:
        print("No tables with captions found in manuscript.")
        sys.exit(1)

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Supplementary Tables")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    doc.add_paragraph()

    for caption, src_table in tables_in_order:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap_run = cap_para.add_run(caption)
        cap_run.font.name = "Arial"
        cap_run.font.size = Pt(10)
        cap_run.font.bold = True

        copy_table(src_table, doc)
        doc.add_paragraph()

    doc.save(str(OUT))
    print(f"Saved: {OUT} ({len(tables_in_order)} tables extracted)")


if __name__ == "__main__":
    main()
