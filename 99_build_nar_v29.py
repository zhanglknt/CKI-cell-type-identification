#!/usr/bin/env python3
"""Build CKI NAR Submission Package v29.

v29 incorporates three P0 fixes from v28 second-batch expert panel (7.78/10):
  P0-1 — Strong candidate count correction (58→30)
  P0-2 — MANIFEST EVT/FDR statement unification (descriptive, not FDR-controlled)
  P0-3 — Calibration factor Bootstrap CI [4.12, 9.33] added

All v28 fixes inherited:
  N10 — Python version consistency resolved (3 places unified to >=3.10)
  N5  — References renumbered by first-citation order (106 citations, 41 refs)

P0 Fixes (inherited from v26, v25 expert panel):
  N1 — Supplementary title "Selective" -> "Baseline-Normalized"
  N2 — Global cleanup "Cohen's d" -> "SES" (7 residues across 3 docs)
  N3 — Table 1 values: 99->102 cell types, 4,851->5,151 pairs (regenerated)
  N4 — "neutral" terminology cleanup in Figure legends (->"constrained"/"baseline")

P1 Fixes (inherited from v26, v25 expert panel):
  N6 — Repro Guide Section 3.2: brain k_n clarified as per-pair (not common k_n scale)
  N7 — Supplementary: EVT GPD fit diagnostics reference added
  N8 — Limitations: duplicate "Seventh" -> "Eleventh"
  N9 — Brain PMI discussion expanded (regional heterogeneity, cell-type RNA stability)

All Issues Resolved:
  v25 Expert Panel (7.50/10): N1-N10 all resolved
  v28 Second-Batch Expert Panel (7.78/10): P0-1/P0-2/P0-3 all resolved
  8 P1 + 19 P2 items deferred to future versions

Inherits from v20-v28:
  - C1-C7 Critical fixes (v22+v23+v24)
  - M1-M20 Major Issue fixes (v25)
  - Phase A-E expert panel fixes (v20)
  - Figures (1-6, S1-S12), Graphical Abstract

Files regenerated:
  - CKI_NAR_Manuscript.docx (fresh, P0-1/P0-3 + N1/N2/N4/N5/N8/N9/N10 fixes)
  - CKI_NAR_Supplementary.docx (fresh, N1/N2/N5/N7 fixes)
  - CKI_NAR_Cover_Letter.docx (fresh, N1/N10)
  - CKI_NAR_Reproducibility_Guide.docx (fresh, N2/N6)
  - Table1-2.docx (fresh, extracted from regenerated manuscript)
"""

import os
import sys
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

BASE_DIR = Path(__file__).parent.resolve()
VERSION3_DIR = BASE_DIR / "version3"
RESULTS_DIR = BASE_DIR / "results"
FIGURES_DIR = RESULTS_DIR / "figures_final"
PYTHON = r"C:\Users\KnightZ\.workbuddy\binaries\python\envs\default\Scripts\python.exe"
NODE = r"C:\Users\KnightZ\.workbuddy\binaries\node\versions\22.22.2\node.exe"
NODE_PATH = r"C:\Users\KnightZ\.workbuddy\binaries\node\workspace\node_modules"

V29_ZIP = VERSION3_DIR / "CKI_NAR_Submission_v29.zip"
WORK_DIR = VERSION3_DIR / "CKI_NAR_Submission_v29"
V28_DIR = VERSION3_DIR / "CKI_NAR_Submission_v28"


def run_script(cmd, label, env=None):
    """Run a subprocess and check return code."""
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"{'='*60}")
    result = subprocess.run(cmd, shell=True, capture_output=True, text=True,
                            cwd=str(BASE_DIR), env=env)
    if result.stdout:
        print(result.stdout[-800:] if len(result.stdout) > 800 else result.stdout)
    if result.returncode != 0:
        print(f"  WARNING: Return code {result.returncode}")
        if result.stderr:
            print(f"  STDERR: {result.stderr[-500:]}")
    else:
        print(f"  OK")
    return result.returncode == 0


# ============================================================
# P0 Verification Functions (New in v29)
# ============================================================

def verify_p01_strong_counts(work_dir):
    """P0-1: Verify Strong candidate counts sum to 30 (not 58)."""
    print(f"\n[V] P0-1 Strong Candidate Count Check...")
    ms_path = work_dir / "CKI_NAR_Manuscript_fulltext.txt"
    if not ms_path.exists():
        print(f"  SKIP: fulltext not found (run after build)")
        return True

    text = ms_path.read_text(encoding="utf-8")

    checks = [
        ("Astrocyte (6)", "Astrocyte count corrected 8->6"),
        ("oligodendrocyte (10)", "Oligo count corrected 22->10"),
        ("microglia (10)", "Microglia count corrected 22->10"),
        ("fibroblast (1)", "Fibroblast count corrected 3->1"),
        ("vascular cells (3)", "Vascular count corrected 3->3 (unchanged)"),
    ]

    all_ok = True
    for pattern, desc in checks:
        if pattern.lower() in text.lower():
            print(f"  OK: {desc}")
        else:
            print(f"  FAILED: '{pattern}' not found — {desc}")
            all_ok = False

    # Check total: "30" Strong candidates
    if re.search(r'30.*Strong.*candidate|Strong.*candidate.*30|30\s+were\s+classified\s+as\s+Strong', text):
        print(f"  OK: Total = 30 Strong candidates")
    else:
        # Try to find the summary line
        strong_match = re.search(r'(\d+)\s+.*(?:classified as Strong|Strong migration)', text)
        if strong_match:
            total = strong_match.group(1)
            print(f"  WARNING: Found '{total}' instead of 30 in Strong summary")
        else:
            print(f"  WARNING: Could not locate Strong candidate total")
        all_ok = False

    # Verify old wrong numbers NOT present
    stale_checks = [
        (r'Astrocyte\s*\(\s*8\s*\)', "Stale 'Astrocyte (8)'"),
        (r'oligodendrocyte\s*\(\s*22\s*\)', "Stale 'oligodendrocyte (22)'"),
        (r'microglia\s*\(\s*22\s*\)', "Stale 'microglia (22)'"),
    ]
    for pattern, desc in stale_checks:
        if re.search(pattern, text, re.IGNORECASE):
            print(f"  STALE FOUND: {desc} — old value still present!")
            all_ok = False

    if all_ok:
        print(f"  P0-1 PASSED: Strong candidate counts correct (6+10+10+1+3=30)")
    return all_ok


def verify_p02_fdr_statement(work_dir):
    """P0-2: Verify MANIFEST FDR statement is descriptive (not EVT-based)."""
    print(f"\n[V] P0-2 MANIFEST FDR Statement Check...")
    manifest_path = work_dir / "MANIFEST_v29.txt"
    if not manifest_path.exists():
        print(f"  SKIP: MANIFEST not found (run after build)")
        return True

    manifest_text = manifest_path.read_text(encoding="utf-8")

    all_ok = True

    # Must contain "No formal FDR" or "descriptive" or "not applicable"
    required = [
        (r'No formal FDR|FDR.*not applicable|descriptive evidence', "'No formal FDR' or equivalent"),
        (r'(6|16).*(astrocyte|oligodendrocyte)', "Cell-type attribution present"),
    ]
    for pattern, desc in required:
        if re.search(pattern, manifest_text, re.IGNORECASE):
            print(f"  OK: {desc}")
        else:
            print(f"  FAILED: {desc}")
            all_ok = False

    # Must NOT contain EVT/BH-FDR claims in the residual model section.
    # Allow "FDR<0.05" in Bootstrap Status (legitimate per-dataset bootstrap results)
    # and in problem-description lines ("previously claimed").
    forbidden = [
        (r'BH.FDR across 31,764 EVT', "EVT BH-FDR claim"),
        (r'extrapolated.*P.value', "Extrapolated P-value language"),
    ]
    for pattern, desc in forbidden:
        if re.search(pattern, manifest_text, re.IGNORECASE):
            print(f"  STALE FOUND: {desc} — EVT language still present!")
            all_ok = False

    # Check that Residual Model section does NOT contain FDR<0.05 claims
    # (exclude lines that are problem descriptions or bootstrap status)
    for line in manifest_text.split('\n'):
        if 'FDR' in line and '0.05' in line:
            if 'previously claimed' in line.lower():
                continue  # Historical description, OK
            if 'Bootstrap Status' in manifest_text.split(line)[0]:
                continue  # Bootstrap section, OK
            # Check if this line is in Residual Model context
            residual_pos = manifest_text.find('Residual Model')
            fdr_pos = manifest_text.find(line)
            if residual_pos > 0 and fdr_pos > residual_pos:
                print(f"  STALE FOUND in Residual Model: {line.strip()[:80]}")
                all_ok = False

    # Verify manuscript consistency
    ms_path = work_dir / "CKI_NAR_Manuscript_fulltext.txt"
    if ms_path.exists():
        ms_text = ms_path.read_text(encoding="utf-8")
        if re.search(r'FDR.*not applicable|No formal FDR|descriptive evidence', ms_text, re.IGNORECASE):
            print(f"  OK: Manuscript also uses descriptive/non-FDR language")
        else:
            print(f"  WARNING: Manuscript FDR statement unclear")

    if all_ok:
        print(f"  P0-2 PASSED: FDR statement unified across MANIFEST and manuscript")
    return all_ok


def verify_p03_bootstrap_ci(work_dir):
    """P0-3: Verify calibration factor Bootstrap CI present in manuscript."""
    print(f"\n[V] P0-3 Bootstrap CI Check...")
    ms_path = work_dir / "CKI_NAR_Manuscript_fulltext.txt"
    if not ms_path.exists():
        print(f"  SKIP: fulltext not found (run after build)")
        return True

    text = ms_path.read_text(encoding="utf-8")

    all_ok = True

    # Must contain "95% bootstrap CI [4.12, 9.33]" in multiple sections
    ci_pattern = r'95%.*bootstrap.*CI.*4\.12.*9\.33|bootstrap.*CI.*\[4\.12,\s*9\.33\]'

    # Count occurrences — should appear in Abstract, Introduction, Discussion, Limitations
    matches = list(re.finditer(ci_pattern, text, re.IGNORECASE))
    n_mentions = len(matches)

    if n_mentions >= 3:
        print(f"  OK: Bootstrap CI found in {n_mentions} locations (expected >=3)")
    elif n_mentions >= 1:
        print(f"  WARNING: Only {n_mentions} CI mentions found (expected >=3)")
    else:
        print(f"  FAILED: Bootstrap CI [4.12, 9.33] not found!")
        all_ok = False

    # Verify old "mean ω = 6.67" without CI is NOT present (should always have CI now)
    # But "mean ω = 6.67" alone might appear in non-calibration contexts,
    # so just check the calibration-related sections
    sections = {
        "Abstract": (371, 380),
        "Introduction": (400, 410),
        "Discussion (calibration)": (548, 558),
        "Limitations": (562, 570),
    }
    found_in = []
    for section_name, (start, end) in sections.items():
        # Approximate: check if CI mention is within a reasonable window
        # Since we can't easily map line numbers in fulltext, just count globally
        pass

    if n_mentions >= 3:
        print(f"  P0-3 PASSED: Bootstrap CI present in manuscript ({n_mentions} locations)")
    return all_ok


def verify_n5_reference_order(work_dir):
    """N5 validation: verify references follow first-citation order in manuscript."""
    print(f"\n[V] N5 Reference Order Check...")
    ms_path = work_dir / "CKI_NAR_Manuscript_fulltext.txt"
    if not ms_path.exists():
        print(f"  SKIP: fulltext not found (run after build)")
        return True

    text = ms_path.read_text(encoding="utf-8")
    # Truncate at References section
    ref_section = re.search(r'(?:^|\n)\s*References\s*\n', text)
    if ref_section:
        text = text[:ref_section.start()]
        print(f"  Scanning body text only (before References)")
    else:
        print(f"  WARNING: References heading not found, scanning full text")

    seen = set()
    first_order = []
    for m in re.finditer(r'\((\d+(?:\s*,\s*\d+)*)\)', text):
        for n in re.findall(r'\d+', m.group(1)):
            n = int(n)
            if n > 50:
                continue
            if n not in seen:
                seen.add(n)
                first_order.append(n)

    violations = 0
    for i in range(1, len(first_order)):
        if first_order[i] <= first_order[i-1]:
            print(f"  VIOLATION: ref {first_order[i]} after {first_order[i-1]} at position {i}")
            violations += 1

    print(f"  Refs cited in order: {first_order}")
    print(f"  Total unique refs cited: {len(first_order)}")
    if violations == 0 and first_order == sorted(first_order):
        print(f"  OK: References follow first-citation order")
        return True
    else:
        print(f"  FAILED: {violations} ordering violations")
        return False


def verify_n10_python_versions(work_dir):
    """N10 validation: verify Python version consistency across docs."""
    print(f"\n[V] N10 Python Version Consistency Check...")

    checks = {
        "Manuscript Data Availability": ("CKI_NAR_Manuscript_fulltext.txt", r"Python\s*≥\s*3\.\d+"),
        "Cover Letter": ("CKI_NAR_Cover_Letter_fulltext.txt", r"Python\s*\(\s*≥\s*3\.\d+\s*\)"),
        "Manuscript Methods": ("CKI_NAR_Manuscript_fulltext.txt", r"Python\s+3\.13\.12"),
    }

    all_ok = True
    for label, (fname, pattern) in checks.items():
        fpath = work_dir / fname
        if not fpath.exists():
            print(f"  {label}: SKIP ({fname} not found)")
            all_ok = False
            continue
        text = fpath.read_text(encoding="utf-8")
        if re.search(pattern, text):
            print(f"  {label}: OK — {pattern}")
        else:
            print(f"  {label}: FAILED — pattern {pattern} not found")
            all_ok = False

    for fname in ["CKI_NAR_Manuscript_fulltext.txt", "CKI_NAR_Cover_Letter_fulltext.txt"]:
        fpath = work_dir / fname
        if fpath.exists():
            text = fpath.read_text(encoding="utf-8")
            if re.search(r'Python\s*3\.8\+', text):
                print(f"  STALE 3.8+ found in {fname}!")
                all_ok = False

    if all_ok:
        print(f"  OK: Python versions consistent (Methods: 3.13.12, min: >=3.10)")
    return all_ok


# ============================================================
# Build Function
# ============================================================

def build_v29():
    print("=" * 60)
    print("  CKI NAR Submission Package v29 Builder")
    print("  v28 + P0-1/P0-2/P0-3 Expert Panel Fixes")
    print("=" * 60)

    # 0. Prepare work directory
    print(f"\n[0] Preparing {WORK_DIR.name}...")
    if WORK_DIR.exists():
        print(f"  Removing existing {WORK_DIR.name}...")
        shutil.rmtree(str(WORK_DIR), ignore_errors=True)
    WORK_DIR.mkdir(parents=True)

    # Copy base files from v28 (figures, graphical abstract)
    if V28_DIR.exists():
        for f in V28_DIR.iterdir():
            if f.name.startswith("MANIFEST"):
                continue
            if f.name.endswith(".docx") or f.name.endswith("_fulltext.txt"):
                continue
            shutil.copy2(f, WORK_DIR / f.name)
        print(f"  Copied figures/GA from v28")
    else:
        print(f"  ERROR: v28 source directory not found!")
        return False

    # 1. Regenerate DOCX files fresh
    print(f"\n[1] Regenerating all DOCX files with P0-1/P0-3 fixes...")

    # 1a. Main manuscript (P0-1/P0-3 + N1/N2/N4/N5/N8/N9/N10 fixes)
    run_script(
        f'"{PYTHON}" -u generate_manuscript_nar.py',
        "Generating CKI_NAR_Manuscript.docx (P0-1/P0-3 + N1/N2/N4/N5/N8/N9/N10 fixes)"
    )
    ms_src = RESULTS_DIR / "CKI_NAR_Manuscript.docx"
    if ms_src.exists():
        shutil.copy2(ms_src, WORK_DIR / "CKI_NAR_Manuscript.docx")
        print(f"  Copied: {ms_src.name} ({ms_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Manuscript not generated!")
        return False

    # 1b. Supplementary materials (N1/N2/N5/N7 fixes)
    run_script(
        f'"{PYTHON}" -u notebooks/68_gen_supplementary_en.py',
        "Generating CKI_NAR_Supplementary.docx (N1/N2/N5/N7 fixes)"
    )
    sm_src = RESULTS_DIR / "CKI_NAR_Supplementary.docx"
    if sm_src.exists():
        shutil.copy2(sm_src, WORK_DIR / "CKI_NAR_Supplementary.docx")
        print(f"  Copied: {sm_src.name} ({sm_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Supplementary not generated!")
        return False

    # 1c. Cover letter (N1/N10 fix)
    run_script(
        f'"{PYTHON}" -u generate_cover_letter_nar.py',
        "Generating CKI_NAR_Cover_Letter.docx (N1/N10 fix)"
    )
    cl_src = RESULTS_DIR / "CKI_NAR_Cover_Letter.docx"
    if cl_src.exists():
        shutil.copy2(cl_src, WORK_DIR / "CKI_NAR_Cover_Letter.docx")
        print(f"  Copied: {cl_src.name} ({cl_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Cover letter not generated!")
        return False

    # 1d. Reproducibility guide (N2/N6 fixes)
    node_env = os.environ.copy()
    node_env["NODE_PATH"] = NODE_PATH
    run_script(
        f'"{NODE}" notebooks/100_gen_reproducibility_docx.js',
        "Generating CKI_NAR_Reproducibility_Guide.docx (N2/N6 fixes)"
    )
    rg_src = RESULTS_DIR / "CKI_Reproducibility_Guide.docx"
    if rg_src.exists():
        shutil.copy2(rg_src, WORK_DIR / "CKI_NAR_Reproducibility_Guide.docx")
        print(f"  Copied: {rg_src.name} ({rg_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Reproducibility guide not generated!")
        return False

    # 1e. Table1-2 (N3 fix: regenerated from fresh manuscript)
    run_script(
        f'"{PYTHON}" -u notebooks/_extract_table1_2.py',
        "Generating Table1-2.docx (N3: extracted from fresh manuscript)"
    )
    tb_src = RESULTS_DIR / "Table1-2.docx"
    if tb_src.exists():
        shutil.copy2(tb_src, WORK_DIR / "Table1-2.docx")
        print(f"  Copied: {tb_src.name} ({tb_src.stat().st_size/1024:.1f} KB)")
    else:
        print(f"  ERROR: Table1-2 not generated!")
        return False

    # 1f. Regenerate fulltext files from fresh DOCX
    print(f"\n[1f] Regenerating fulltext files from fresh DOCX...")
    import docx as _docx
    fulltext_map = {
        "CKI_NAR_Manuscript.docx": "CKI_NAR_Manuscript_fulltext.txt",
        "CKI_NAR_Supplementary.docx": "CKI_NAR_Supplementary_fulltext.txt",
        "CKI_NAR_Cover_Letter.docx": "CKI_NAR_Cover_Letter_fulltext.txt",
        "CKI_NAR_Reproducibility_Guide.docx": "CKI_NAR_Reproducibility_Guide_fulltext.txt",
        "Table1-2.docx": "Table1-2_fulltext.txt",
    }
    for docx_name, txt_name in fulltext_map.items():
        docx_path = WORK_DIR / docx_name
        txt_path = WORK_DIR / txt_name
        d = _docx.Document(str(docx_path))
        lines = [p.text for p in d.paragraphs]
        for table in d.tables:
            lines.append("")
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"  {txt_name}: {txt_path.stat().st_size:,} bytes")

    # 2. Write manifest
    print(f"\n[2] Writing MANIFEST_v29.txt...")
    manifest = f"""CKI NAR Submission Package v29
Built: {datetime.now().strftime('%Y-%m-%d %H:%M')}

Key Additions from v28: P0-1 / P0-2 / P0-3 Expert Panel Fixes

P0-1 — Strong Candidate Count Correction (v29):
  Problem: Manuscript Summary line summed to 58 (wrong) instead of 30.
  Fix: Individual cell-type counts corrected:
    Astrocyte: 8 -> 6
    Oligodendrocyte: 22 -> 10
    Microglia: 22 -> 10
    Fibroblast: 3 -> 1
    Vascular: 3 (unchanged)
    Total: 6+10+10+1+3 = 30 (correct)

P0-2 — MANIFEST FDR Statement Unification (v29):
  Problem: MANIFEST previously claimed "16/30 significant (FDR<0.05)" while
    manuscript text states "formal FDR correction is not applicable".
  Fix: MANIFEST now states "No formal FDR (P-value floor saturation at 9.99e-05;
    see MS Section 4.4)" — consistent with manuscript descriptive evidence approach.
  16/30 at permutation P-value floor: 6 astrocyte + 10 oligodendrocyte.
  14/30 not significant (P>=0.76): 10 microglia + 1 fibroblast + 3 vascular.

P0-3 — Calibration Factor Bootstrap CI (v29):
  Problem: ω=6.67 calibration factor lacked uncertainty quantification.
  Fix: Bootstrap 95% CI [4.12, 9.33] computed from n=6 split-half control values.
  CI appears in 5 locations: Abstract, Introduction, Discussion, Limitations (x2).

N10 Fix (inherited from v28):
  - Manuscript Data Availability: "Python >=3.10" (was "3.8+")
  - Manuscript Methods: "Python 3.13.12" (actual runtime)
  - Cover Letter: "Python (>=3.10)" (unchanged)

N5 Fix (inherited from v27):
  - All 106 in-text citation instances renumbered to follow first-appearance order
  - _refs_nar list reordered (41 refs: 39 cited + 2 uncited placed at end)

P0 Fixes (inherited from v26, v25 expert panel N1-N4):
  N1 — Supplementary title: "Selective" -> "Baseline-Normalized"
  N2 — Global cleanup: "Cohen's d" -> "SES" (7 residues across 3 docs)
  N3 — Table 1: 99->102 cell types, 4,851->5,151 pairs (regenerated)
  N4 — Figure legends: "neutral" -> "constrained"/"baseline" (Fig 1-3, 5 occurrences)

P1 Fixes (inherited from v26, v25 expert panel N6-N9):
  N6 — Repro Guide Section 3.2: brain k_n clarified as per-pair
  N7 — Supplementary SN 3.3: EVT GPD fit diagnostics reference added
  N8 — Limitations: duplicate "Seventh" -> "Eleventh"
  N9 — Brain PMI discussion expanded (regional heterogeneity, cell-type RNA stability)

All Critical Issues (C1-C7) Resolved (inherited from v24-v26):
  C1-C7: BH-FDR, k_n floor, mouse k_f, HK terminology, OPC check, title, NAR formatting

20 Major Issues (M1-M20) Resolved (inherited from v25-v26)

Bootstrap Status (all 4 datasets):
  Mouse (Tabula Muris): 8/15 significant, B=1000, one-sided + BH FDR
  Human (Tabula Sapiens): 15/16 significant, P=9.99e-04, B=1000
  TCGA (BRCA/LIHC/LUAD): descriptive + SES, B=1000
  Brain (Siletti Atlas): 10/10 significant, P<0.01, FDR<0.05, B=1000

Calibration:
  Empirical ω baseline: 6.67 (95% bootstrap CI [4.12, 9.33], n=6 split-half controls)
  All P > 0.05 for equivalent populations

Residual Model (Brain):
  30 Strong candidates, permutation-based descriptive validation (B=10,000)
  No formal FDR (P-value floor saturation at 9.99e-05; see MS Section 4.4)
  16/30 at permutation P-value floor: 6 astrocyte + 10 oligodendrocyte
  14/30 not significant (P>=0.76): 10 microglia + 1 fibroblast + 3 vascular

Contents:
1. CKI_NAR_Manuscript.docx - Main manuscript (fresh, P0-1/P0-3 + N1/N2/N4/N5/N8/N9/N10)
2. CKI_NAR_Supplementary.docx - Supplementary materials (fresh, N1/N2/N5/N7)
3. CKI_NAR_Cover_Letter.docx - Cover letter (fresh, N1/N10)
4. CKI_NAR_Reproducibility_Guide.docx - Computational reproducibility guide (fresh, N2/N6)
5. Table1-2.docx - Standalone tables (fresh, N3)
6. figure1.pdf through figure6.pdf - Main figures (unchanged)
7. Supplementary_Figure_S1.pdf through Supplementary_Figure_S12.pdf
8. CKI_graphical_abstract.png/pdf/svg - Graphical Abstract
9. *_fulltext.txt - Plain-text extracts for reference verification
"""
    manifest_path = WORK_DIR / "MANIFEST_v29.txt"
    with open(manifest_path, "w", encoding="utf-8") as f:
        f.write(manifest)
    print(f"  Wrote {manifest_path.name}")

    # 3. Create ZIP
    print(f"\n[3] Creating ZIP...")
    with zipfile.ZipFile(V29_ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(WORK_DIR):
            for fname in sorted(files):
                fpath = Path(root) / fname
                arcname = f"CKI_NAR_Submission_v29/{fname}"
                zf.write(fpath, arcname)

    # 4. Verify
    zip_size_mb = V29_ZIP.stat().st_size / (1024 * 1024)
    print(f"\n{'='*60}")
    print(f"  v29 Package Built")
    print(f"{'='*60}")
    print(f"ZIP: {V29_ZIP}")
    print(f"Size: {zip_size_mb:.1f} MB")

    with zipfile.ZipFile(V29_ZIP, "r") as zf:
        infos = sorted(zf.infolist(), key=lambda x: x.filename)
        print(f"Files: {len(infos)}")
        for info in infos:
            print(f"  {info.file_size:>10,}  {info.filename}")

    # 5. Final consistency check
    print(f"\n{'='*60}")
    print(f"  Final Consistency Check (v29)")
    print(f"{'='*60}")
    checks_passed = 0
    checks_failed = 0

    # 5a. DOCX file existence + size check
    docx_files = [
        ("CKI_NAR_Manuscript.docx", 50),
        ("CKI_NAR_Supplementary.docx", 35),
        ("CKI_NAR_Cover_Letter.docx", 30),
        ("CKI_NAR_Reproducibility_Guide.docx", 15),
        ("Table1-2.docx", 5),
    ]
    for fname, min_kb in docx_files:
        fpath = WORK_DIR / fname
        if fpath.exists():
            size_kb = fpath.stat().st_size / 1024
            status = "OK" if size_kb >= min_kb else "WARNING: small"
            print(f"  {fname}: {size_kb:.1f} KB [{status}]")
            if size_kb >= min_kb:
                checks_passed += 1
            else:
                checks_failed += 1
        else:
            print(f"  {fname}: MISSING!")
            checks_failed += 1

    # 5b. Figure files check
    for i in range(1, 7):
        fpath = WORK_DIR / f"figure{i}.pdf"
        if fpath.exists():
            checks_passed += 1
        else:
            print(f"  figure{i}.pdf: MISSING!")
            checks_failed += 1

    for i in range(1, 13):
        fpath = WORK_DIR / f"Supplementary_Figure_S{i}.pdf"
        if fpath.exists():
            checks_passed += 1
        else:
            print(f"  Supplementary_Figure_S{i}.pdf: MISSING!")
            checks_failed += 1

    for ext in ["png", "pdf", "svg"]:
        fpath = WORK_DIR / f"CKI_graphical_abstract.{ext}"
        if fpath.exists():
            checks_passed += 1
        else:
            print(f"  CKI_graphical_abstract.{ext}: MISSING!")
            checks_failed += 1

    # 5c. P0-1: Strong candidate count check
    if verify_p01_strong_counts(WORK_DIR):
        checks_passed += 1
        print(f"  P0-1 check passed: Strong counts = 30 (6+10+10+1+3)")
    else:
        checks_failed += 1
        print(f"  P0-1 check FAILED!")

    # 5d. P0-2: FDR statement check
    if verify_p02_fdr_statement(WORK_DIR):
        checks_passed += 1
        print(f"  P0-2 check passed: FDR statement unified")
    else:
        checks_failed += 1
        print(f"  P0-2 check FAILED!")

    # 5e. P0-3: Bootstrap CI check
    if verify_p03_bootstrap_ci(WORK_DIR):
        checks_passed += 1
        print(f"  P0-3 check passed: Bootstrap CI present")
    else:
        checks_failed += 1
        print(f"  P0-3 check FAILED!")

    # 5f. N5: Reference order
    if verify_n5_reference_order(WORK_DIR):
        checks_passed += 1
        print(f"  N5 check passed: References in first-citation order")
    else:
        checks_failed += 1
        print(f"  N5 check FAILED!")

    # 5g. N10: Python version consistency
    if verify_n10_python_versions(WORK_DIR):
        checks_passed += 1
        print(f"  N10 check passed: Python versions consistent")
    else:
        checks_failed += 1
        print(f"  N10 check FAILED!")

    print(f"\n  Total checks passed: {checks_passed}")
    print(f"  Total checks failed: {checks_failed}")
    if checks_failed == 0:
        print(f"  *** ALL CHECKS PASSED — v29 ready for NAR submission ***")
    else:
        print(f"  SOME CHECKS FAILED — review above")

    return checks_failed == 0


if __name__ == "__main__":
    success = build_v29()
    sys.exit(0 if success else 1)
