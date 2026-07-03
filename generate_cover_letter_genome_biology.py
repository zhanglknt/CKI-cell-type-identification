"""
Generate Genome Biology Cover Letter
 CKI Project — Genome Biology submission
"""

import sys
from pathlib import Path
from datetime import date
from docx.shared import Pt, Inches

PROJECT_ROOT = Path(__file__).parent
OUTPUT_DIR = PROJECT_ROOT / "results"
OUTPUT_DIR.mkdir(exist_ok=True)

# ============================================================
# Genome Biology Cover Letter Content
# ============================================================

ADDRESS = [
    "Li Zhang",
    "Chinese Academy of Medical Sciences & Peking Union Medical College",
    "Institute of Blood Transfusion",
    "Chengdu, China",
    "",
    "Chinese Institute for Brain Research, Beijing",
    "Beijing, China",
    "",
    "Email: knightz@pumc.edu.cn",
]

RECIPIENT = [
    "The Editor-in-Chief",
    "Genome Biology",
    "BioMed Central / Springer Nature",
]

# ============================================================
# Build Cover Letter
# ============================================================

def p(text, doc, style_name=None):
    para = doc.add_paragraph(text, style=style_name)
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_after = Pt(0)
    return para

def run():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Sender address (right-aligned) ---
    for line in ADDRESS:
        if line == "":
            p("", doc)
        else:
            para = p(line, doc)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p("", doc)
    p(date.today().strftime("%B %d, %Y"), doc)

    # --- Recipient ---
    p("", doc)
    for line in RECIPIENT:
        p(line, doc)

    p("", doc)
    p("Re: Submission of Original Research Article", doc)
    p("", doc)

    # --- Title ---
    title_para = p("", doc)
    run_title = title_para.add_run(
        "CKI: A Cell-state Kinetic Index for Quantifying Selective "
        "Transcriptomic Remodeling"
    )
    run_title.bold = True
    run_title.font.size = Pt(13)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p("", doc)

    # --- Body ---
    p(
        "Dear Editor,",
        doc,
    )

    p(
        "On behalf of my co-author, Dr. Xianming Wu (Chinese Institute for Brain "
        "Research, Beijing), I am pleased to submit our original research article, "
        "\u201cCKI: A Cell-state Kinetic Index for Quantifying Selective "
        "Transcriptomic Remodeling\u201d, for consideration as a Research Article "
        "in Genome Biology.",
        doc,
    )

    p(
        "Cell identity is fundamentally defined by its transcriptome. However, a "
        "standardized, assumption-free metric to quantify transcriptomic divergence "
        "between cell states has been lacking. Drawing inspiration from evolutionary "
        "biology\u2019s Ka/Ks ratio, we propose CKI (Cell-state Kinetic Index), "
        "which decomposes Jensen\u2013Shannon divergence into two orthogonal components: "
        "k_n (neutral offset rate, from housekeeping gene expression) and k_f "
        "(functional conversion rate, from identity gene expression). The ratio "
        "\u03c9 = k_f/k_n provides a robust, interpretable measure of selective "
        "transcriptomic remodeling.",
        doc,
    )

    p(
        "We validated CKI \u03c9 across three dimensions: (1) orthogonal information\u2014"
        "CKI \u03c9 captures an independent information dimension, showing negative "
        "correlation with all four standard distance metrics (Spearman r = \u22120.36 "
        "to \u22120.46, all P < 0.001), proving it measures something fundamentally "
        "different from Cosine similarity, JS divergence, and other existing approaches; "
        "(2) cross-species consistency\u2014mouse orthologs show strong correlation with "
        "human CKI \u03c9, confirming evolutionary consistency; and (3) two biological "
        "applications\u2014pan-cancer analysis revealing that tumors are more transcriptionally "
        "homogeneous than normal tissues (median NN/TT ratio 1.40\u20132.83), and "
        "brain regional analysis identifying 30 cell-type-specific migration candidates "
        "among 31,764 cross-region comparisons through a multiplicative residual model.",
        doc,
    )

    p(
        "Genome Biology is our first choice because this work sits at the intersection "
        "of genomics, transcriptomics, and computational biology\u2014the exact scope of "
        "the journal. The proposed \u03c9 metric provides the community with a new genomic "
        "toolkit for comparing cell states, with direct applications in development, "
        "cancer, and neuroscience. We believe the readership of Genome Biology will "
        "find both the method and its biological insights of significant interest.",
        doc,
    )

    p(
        "Dr. Xianming Wu (first author) and I (corresponding author) have approved "
        "the manuscript and declare no competing interests. "
        "This work has not been published elsewhere and is not under consideration "
        "by any other journal. All data and code are publicly available (see Data "
        "Availability Statement in the manuscript).",
        doc,
    )

    p(
        "I suggest the following potential reviewers:\n"
        "1. Prof. Sarah A. Teichmann (Wellcome Sanger Institute, st9@sanger.ac.uk)\n"
        "2. Prof. Peter V. Kharchenko (Harvard Medical School, peter_kharchenko@hms.harvard.edu)\n"
        "3. Prof. Cole Trapnell (University of Washington, coletrap@uw.edu)",
        doc,
    )

    p(
        "Thank you for considering our work for publication in Genome Biology. "
        "We look forward to hearing from you.",
        doc,
    )

    p("", doc)
    p("Sincerely,", doc)
    p("", doc)
    p("Li Zhang (Corresponding Author)", doc)
    p("Xianming Wu (First Author)", doc)
    p("", doc)
    p("Li Zhang", doc)
    p("Institute of Blood Transfusion, CAMS & PUMC, Chengdu, China", doc)
    p("Chinese Institute for Brain Research, Beijing, China", doc)
    p("Email: knightz@pumc.edu.cn | ORCID: 0000-0002-0698-0754", doc)

    # --- Save ---
    out = str(OUTPUT_DIR / "CKI_GenomeBiology_Cover_Letter_v2.docx")
    doc.save(out)
    print(f"Saved: {out}")
    return out


if __name__ == "__main__":
    run()
