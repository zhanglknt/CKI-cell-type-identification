"""
Generate CKI_GenomeBiology_Manuscript_v1.docx — Genome Biology Research article.

Genome Biology formatting compliance:
- Single continuous document (no page breaks)
- Structured abstract (Background / Results / Conclusions subheadings)
- Vancouver numbered references [1]
- Double spacing
- Sections: Abstract → Keywords → Background → Results → Discussion →
  Conclusions → Methods → Abbreviations → Declarations → References →
  Figure Legends → Tables
- All text black; Arial/Helvetica font
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from pathlib import Path

doc = Document()

# == Genome Biology formatting ==
# Double spacing
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.line_spacing = 2.0  # double spacing
style.paragraph_format.space_after = Pt(0)

# Add line numbers and page numbers would need a template; 
# Genome Biology accepts docx and applies styles during production.
# We set double spacing and will note in cover letter.

# Page margins: 2.5 cm (Genome Biology default)
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# == Helpers (same as NAR version but ref format changed) ==
def set_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_superscript(run):
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn('w:vertAlign')):
        rPr.remove(old)
    va = rPr.makeelement(qn('w:vertAlign'), {qn('w:val'): 'superscript'})
    rPr.append(va)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
        set_black(run)
        run.font.size = Pt([16, 14, 12][level-1])
        rPr = run._element.get_or_add_rPr()
        for old in rPr.findall(qn('w:b')):
            rPr.remove(old)
        b = rPr.makeelement(qn('w:b'), {})
        rPr.append(b)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def p(text, bold=False, italic=False, size=10):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    set_black(run)
    run.bold = bold
    run.italic = italic
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.first_line_indent = Cm(0)
    return para

# == Table helpers (reuse from NAR version) ==
def add_table_1(doc):
    """Table 1: AUC for cell-type classification."""
    para = doc.add_paragraph()
    run = para.add_run('Table 1: Classification AUC of five metrics on Tabula Sapiens (99 cell types, 4,851 pairs).')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'ROC-AUC'
    data = [
        ('Cosine distance', '0.887'),
        ('Raw JS divergence', '0.849'),
        ('Marker Jaccard distance', '0.801'),
        ('Spearman distance', '0.690'),
        ('CKI omega', '0.680'),
    ]
    for i, (metric, auc) in enumerate(data):
        table.rows[i+1].cells[0].text = metric
        table.rows[i+1].cells[1].text = auc
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_table_2(doc):
    """Table 2: Cross-organ conservation ranking by cell type (Tabula Sapiens, n=59 same-cell-type cross-organ pairs)."""
    para = doc.add_paragraph()
    run = para.add_run('Table 2: Cross-organ conservation ranking by cell type (Tabula Sapiens, n=59 same-cell-type cross-organ pairs).')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    table = doc.add_table(rows=18, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Cell type'
    hdr[1].text = 'Mean \u03c9'
    hdr[2].text = 'SD'
    hdr[3].text = 'n pairs'
    data = [
        ('Hepatocyte', '8.57', '\u2014', '1'),
        ('B cell', '9.36', '\u2014', '1'),
        ('CD8+ T cell', '9.93', '4.2', '6'),
        ('Plasma cell', '10.20', '1.9', '6'),
        ('Hematopoietic stem cell', '11.12', '\u2014', '1'),
        ('Smooth muscle cell', '12.22', '\u2014', '1'),
        ('Neutrophil', '14.22', '7.8', '6'),
        ('Monocyte', '15.19', '\u2014', '1'),
        ('Macrophage', '15.49', '8.2', '15'),
        ('CD4+ T cell', '16.71', '\u2014', '1'),
        ('NK cell', '17.28', '3.0', '10'),
        ('Classical monocyte', '17.50', '\u2014', '1'),
        ('Naive B cell', '21.06', '\u2014', '1'),
        ('Intermediate monocyte', '21.78', '\u2014', '1'),
        ('Memory B cell', '22.36', '\u2014', '1'),
        ('Endothelial cell', '26.25', '7.0', '3'),
        ('Erythrocyte', '29.36', '18.8', '3'),
    ]
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# == Vancouver reference list ==
# Genome Biology uses numbered Vancouver style: [1], [2], etc.
# We use BMC/Vancouver format: Author AB, Author CD. Title. Journal. Year;Volume:Pages.

_refs_gb = [
    '[1]  Korsunsky I, Millard N, Fan J, Slowikowski K, Zhang F, Baglaenko Y, et al. Fast, sensitive and accurate integration of single-cell data with Harmony. Nat Methods. 2019;16:1289–1296.',
    '[2]  Lopez R, Regier J, Cole MB, Jordan MI, Yosef N. Deep generative modeling for single-cell transcriptomics. Nat Methods. 2018;15:1053–1058.',
    '[3]  Rosen Y, Brbic M, Roohani Y, Swanson K, Li Z, Leskovec J. Toward universal cell embeddings: integrating single-cell RNA-seq datasets across species with SATURN. Nat Methods. 2024;21:1492–1500.',
    '[4]  Hounkpel B, Chen J, Gosline SJC, Domeniconi C, Jiang D. HRT Atlas v1.0 database: redefining human and mouse housekeeping genes and candidate reference transcripts by mining massive RNA-seq datasets. Nucleic Acids Res. 2021;49:D947–D955.',
    '[5]  Tabula Muris Consortium. Single-cell transcriptomics of 20 mouse organs creates a Tabula Muris. Nature. 2018;562:367–372.',
    '[6]  Tabula Sapiens Consortium. The Tabula Sapiens: a multiple-organ, single-cell transcriptomic atlas of humans. Science. 2022;376:eabl4896.',
    '[7]  Cancer Genome Atlas Research Network. Comprehensive molecular profiling of lung adenocarcinoma. Nature. 2014;511:543–550.',
    '[8]  Cancer Genome Atlas Network. Comprehensive molecular portraits of human breast tumours. Nature. 2012;490:61–70.',
    '[9]  Siletti K, Hodge R, Mossi Albiach A, Lee KW, Ding SL, Hu L, et al. Transcriptomic diversity of cell types across the adult human brain. Science. 2023;382:eadl7046.',
    '[10] Edmondson HA, Steiner PE. Primary carcinoma of the liver: a study of 100 cases among 48,900 necropsies. Cancer. 1954;7:462–503.',
    '[11] Perou CM, Sørlie T, Eisen MB, van de Rijn M, Jeffrey SS, Rees CA, et al. Molecular portraits of human breast tumours. Nature. 2000;406:747–752.',
    '[12] Parker JS, Mullins M, Cheang MCU, Leung S, Voduc D, Vickery T, et al. Supervised risk predictor of breast cancer based on intrinsic subtypes. J Clin Oncol. 2009;27:1160–1167.',
    '[13] Tsai HH, Niu J, Munji R, Davalos D, Chang J, Zhang H, et al. Oligodendrocyte precursors migrate along vasculature in the developing nervous system. Science. 2016;351:379–384.',
    '[14] Cerami E, Gao J, Dogrusoz U, Gross BE, Sumer SO, Aksoy BA, et al. The cBio cancer genomics portal: an open platform for exploring multidimensional cancer genomics data. Cancer Discov. 2012;2:401–404.',
    '[15] Tarashansky AJ, Musser JM, Khariton M, Li P, Arendt D, Quake SR, et al. Mapping single-cell atlases throughout Metazoa unravels cell type evolution. Elife. 2021;10:e66747.',
    '[16] Regev A, Teichmann SA, Lander ES, Amit I, Benoist C, Birney E, et al. The Human Cell Atlas. Elife. 2017;6:e27041.',
    '[17] Akay LA, Effenberger AH, Tsai LH. Astrocyte endfoot formation controls the termination of oligodendrocyte precursor cell perivascular migration. Neuron. 2022;111:190–201.e8.',
    '[18] Endo F, Kasai A, Cui W, Tanaka KF, Hashimoto H. Astrocyte allocation during brain development is controlled by Tcf4-mediated fate restriction. EMBO J. 2024;43:4423–4447.',
    '[19] Yang L, Zhao Z, Li Y, Wang J, Chen X, Liu Z. Single-cell multi-omics analysis of lineage development and spatial organization in the human fetal cerebellum. Cell Discov. 2024;10:25.',
    '[20] Menassa DA, Muntslag TAO, Martin-Estebane M, Barry-Carroll L, Chapman MA, Adorjan I, et al. The spatiotemporal dynamics of microglia across the human lifespan. Dev Cell. 2022;57:1910–1927.e10.',
    '[21] Yang Z. PAML 4: phylogenetic analysis by maximum likelihood. Mol Biol Evol. 2007;24:1586–1591.',
    '[22] Jiang J, Li J, Huang Y, Wang Y, Chen L, Zhang X. CACIMAR: cross-species analysis of cell identities, markers, regulations, and interactions. Brief Bioinform. 2024;25:bbae283.',
    '[23] Tan YL, Yuan Y, Tian L. Microglial regional heterogeneity and its role in the brain. Mol Psychiatry. 2020;25:351–367.',
    '[24] Storey JD, Tibshirani R. Statistical significance for genomewide studies. Proc Natl Acad Sci USA. 2003;100:9440–9445.',
    '[25] Wolf FA, Angerer P, Theis FJ. SCANPY: large-scale single-cell gene expression data analysis. Genome Biol. 2018;19:15.',
    '[26] Hao Y, Hao S, Andersen-Nissen E, Mauck WM, Zheng S, Butler A, et al. Integrated analysis of multimodal single-cell data. Cell. 2021;184:3573–3587.',
    '[27] Hao Y, Stuart T, Kowalski MH, Choudhary S, Hoffman P, Hartman A, et al. Dictionary learning for integrative, multimodal and scalable single-cell analysis. Nat Biotechnol. 2024;42:293–304.',
    '[28] Weinstein JN, Collisson EA, Mills GB, Shaw KRM, Ozenberger BA, Ellrott K, et al. The Cancer Genome Atlas Pan-Cancer analysis project. Nat Genet. 2013;45:1113–1120.',
    '[29] Colaprico A, Silva TC, Olsen C, Garofano L, Cava C, Garolini D, et al. TCGAbiolinks: an R/Bioconductor package for integrative analysis of TCGA data. Nucleic Acids Res. 2016;44:e71.',
    '[30] Luecken MD, Theis FJ. Current best practices in single-cell RNA-seq analysis: a tutorial. Mol Syst Biol. 2019;15:e8746.',
    '[31] Nei M, Gojobori T. Simple methods for estimating the numbers of synonymous and nonsynonymous nucleotide substitutions. Mol Biol Evol. 1986;3:418–426.',
    '[32] Tran HTN, Ang KS, Chevrier M, Zhang X, Lee NYS, Goh M, Chen J. A benchmark of batch-effect correction methods for single-cell RNA sequencing data. Genome Biol. 2020;21:12.',
    '[33] CZI Cell Science Program. CZ CELLxGENE Discover: a single-cell data platform for scalable exploration, analysis and modeling of aggregated data. Nucleic Acids Res. 2025;53:D886–D900.',
    '[34] Liberzon A, Birger C, Thorvaldsdóttir H, Ghandi M, Mesirov JP, Tamayo P. The Molecular Signatures Database Hallmark Gene Set Collection. Cell Syst. 2015;1:417–425.',
    '[35] Foerster S, Floriddia EM, Neumann B, Agirre E, Castelo-Branco G, Franklin RJM. Developmental origin of oligodendrocytes determines their function in the adult brain. Nat Neurosci. 2024;27:1155–1166.',
    '[36] Wälchli T, Ghobrial M, Schwab ME, Takada S, Zhong H, Le J, et al. Single-cell atlas of the human brain vasculature across development, adulthood and disease. Nature. 2024;632:603–613.',
    '[37] Shemer A, Jung S. The molecular determinants of microglial developmental colonization. Nat Rev Neurosci. 2024;25:414–427.',
    '[38] Jones HE, Coelho-Santos V, Bonney SK, Abrams SR, Shih AY, Siegenthaler JA. Meningeal origins and dynamics of perivascular fibroblast development. Development. 2023;150:dev201805.',
    '[39] Schaffenrath J, Huang SF, Wyss T, Delorenzi M, Keller A. Characteristics of blood-brain barrier heterogeneity between brain regions. Nat Neurosci. 2024;27:1851–1865.',
    '[40] Barry-Carroll L, Greulich P, Marshall AR, Riecken K, Fehse B, Askew KE, et al. Microglia colonize the developing brain by clonal expansion of highly proliferative progenitors. Cell Rep. 2023;42:113453.',
    '[41] Reeber SL, Arancillo M, Sillitoe RV. Bergmann glia are patterned into topographic molecular zones in the cerebellum. Cerebellum. 2015;14:392–403.',
]

def ref_p_gb(text):
    """Add a reference paragraph in Vancouver format."""
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 1.0  # references single-spaced
    para.paragraph_format.space_after = Pt(2)
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        set_black(run)
    return para

# ============================================================
# TITLE PAGE
# ============================================================
t = doc.add_heading('CKI: A Cell-state Kinetic Index for Quantifying Selective Transcriptomic Remodeling', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t.runs:
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(16)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Xianming Wu (first author)
run = sub.add_run('Xianming Wu')
run.font.name = 'Arial'
run.font.size = Pt(12)
set_black(run)
run.italic = True
for ch in ['1']:
    r = sub.add_run(ch)
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    set_black(r)
    set_superscript(r)
    r.italic = True

sub.add_run(', ')

# Li Zhang (corresponding author)
run = sub.add_run('Li Zhang')
run.font.name = 'Arial'
run.font.size = Pt(12)
set_black(run)
run.italic = True
for ch in ['1', '2', '*']:
    r = sub.add_run(ch)
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    set_black(r)
    set_superscript(r)
    if ch in ('1', '2'):
        r.italic = True

# Affiliations (in order of author appearance)
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = auth.add_run('1')
run1.font.name = 'Arial'
run1.font.size = Pt(10)
set_black(run1)
set_superscript(run1)
run2 = auth.add_run('Chinese Institute for Brain Research, Beijing, China')
run2.font.name = 'Arial'
run2.font.size = Pt(10)
set_black(run2)

auth2 = doc.add_paragraph()
auth2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = auth2.add_run('2')
run3.font.name = 'Arial'
run3.font.size = Pt(10)
set_black(run3)
set_superscript(run3)
run4 = auth2.add_run('Institute of Blood Transfusion, Chinese Academy of Medical Sciences & Peking Union Medical College, Chengdu, China')
run4.font.name = 'Arial'
run4.font.size = Pt(10)
set_black(run4)


# Correspondence
cor = doc.add_paragraph()
cor.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cor.add_run('* Correspondence: knightz@pumc.edu.cn')
run.font.name = 'Arial'
set_black(run)
run.font.size = Pt(10)
run.italic = True

# ORCID
orcid = doc.add_paragraph()
orcid.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = orcid.add_run('ORCID: 0000-0002-0698-0754')
run.font.name = 'Arial'
set_black(run)
run.font.size = Pt(9)

# ============================================================
# ABSTRACT (structured per Genome Biology format)
# ============================================================
heading('Abstract', level=1)

heading('Background', level=2)
p(
    'Comparing cell populations is a fundamental task in single-cell genomics, '
    'yet standard distance metrics treat all gene expression changes equally, '
    'conflating neutral variation with functional adaptation. CKI (Cell-state '
    'Kinetic Index) addresses this by decomposing transcriptomic divergence into '
    'neutral (k_n, from housekeeping genes) and functional (k_f, from identity '
    'genes) components, with \u03c9 = k_f/k_n quantifying selective remodeling '
    'analogous to Ka/Ks in molecular evolution.'
)
heading('Results', level=2)
p(
    'We validated CKI across four datasets: Tabula Muris mouse atlas (15,057 cells), '
    'Tabula Sapiens human atlas (108,136 cells), TCGA pan-cancer (10,535 samples), '
    'and a human brain single-nucleus atlas (888,263 nuclei). Calibration on Tabula '
    'Muris confirmed neutral behavior for biologically equivalent populations '
    '(mean \u03c9 = 1.54, all P > 0.05). CKI \u03c9 was negatively correlated with all '
    'four standard distance metrics (Spearman r = \u22120.36 to \u22120.46, all P < 0.001), '
    'proving it captures an independent information dimension. Cancer analysis '
    'revealed that tumors are more transcriptionally homogeneous than normal tissues '
    '(median NN/TT ratio = 1.40\u20132.83 across five cancer types). Brain regional '
    'analysis revealed a 6.06-fold \u03c9 gradient across 10 non-neuronal cell classes '
    'and identified 30 cell-type-specific signals among 31,764 cross-region comparisons, '
    'systematically validated against developmental neuroscience literature as embryonic origin, '
    'colonization route, and specification signatures rather than active migration.'
)
heading('Conclusions', level=2)
p(
    'CKI provides a principled framework for quantifying cell-type-specific selective '
    'pressure from transcriptomic data, complementing existing distance metrics with '
    'an orthogonal measure of functional remodeling. The method is available as an '
    'open-source Python package (v0.3.1, MIT License) at '
    'https://github.com/zhanglknt/CKI-cell-type-identification.'
)

# ============================================================
# KEYWORDS
# ============================================================
heading('Keywords', level=1)
p('Transcriptomic divergence, cell-state kinetics, selective remodeling, brain cell atlas, housekeeping genes, Ka/Ks analogy')

# ============================================================
# BACKGROUND
# ============================================================
heading('Background', level=1)

p('Single-cell transcriptomics has transformed how we study cells. Comparing two cell populations is one of the most common tasks: we want to know how different they are. Researchers typically choose a standard metric: Euclidean distance, cosine similarity, Pearson or Spearman correlation, or Jensen-Shannon divergence. These metrics are convenient, but they have a fundamental limitation: they treat all gene expression differences equally.')

p('This matters because not all expression changes have the same biological meaning. A twofold change in GAPDH expression might reflect technical noise; a twofold change in a transcription factor might reflect a functional shift in cell state. Standard metrics cannot tell these apart. This problem is especially acute in large single-cell atlases, where donor-level and batch-level variation often dominates over cell-type identity. Methods such as Harmony [1], scVI [2], and SATURN [3] have been developed specifically to remove such nuisance variation. But after correction, a key question remains: how much of the remaining difference between two populations represents functional adaptation, and how much is simply neutral drift?')

p('We realized that this question mirrors a problem solved in molecular evolution forty years ago. When comparing two DNA sequences, the Ka/Ks ratio (also called dN/dS) distinguishes nonsynonymous changes (Ka, which alter the protein and may be shaped by selection) from synonymous changes (Ks, which are silent and presumed neutral). The elegance of Ka/Ks lies in its use of synonymous sites as an internal baseline: the same mutational process produces both types of change, so their ratio reveals whether selection is acting.')

p('We applied this logic to transcriptomic comparisons. CKI defines two rates: a neutral offset rate k_n, estimated from housekeeping (HK) gene expression; and a functional conversion rate k_f, estimated from cell-type identity genes. The ratio \u03c9 = k_f/k_n quantifies selective transcriptomic remodeling: \u03c9 near 1 means the observed differences are consistent with neutral expectation; \u03c9 much greater than 1 means functional gain exceeds background variation; \u03c9 much less than 1 means strong functional constraint.')

p('We validated CKI across three scales, each testing a different aspect of the method. First, we calibrated CKI on Tabula Muris mouse data [5], confirming that random splits of the same cell population yield \u03c9 close to 1 (neutral behavior). Second, we extended CKI to Tabula Sapiens human data [6] and found that CKI \u03c9 is negatively correlated with all four standard distance metrics, proving it captures an orthogonal information dimension. Third, we applied CKI to TCGA cancer data [7,8], revealing that tumors are paradoxically more transcriptionally homogeneous than normal tissues. Fourth, we used CKI to analyze a human brain single-nucleus atlas [9], measuring how the same cell types differ across brain regions and demonstrating that CKI can detect persistent developmental signatures—including developmental origin heterogeneity, colonization route boundaries, and a single postnatal migration event—from adult transcriptomic data.')

# ============================================================
# RESULTS
# ============================================================
heading('Results', level=1)

# --- Result 1 ---
heading('Decomposing transcriptomic variation', level=2)

p('CKI takes two cell populations as input, each represented as a pseudobulk expression vector (the mean expression across all cells in that population). The computation has three steps, all of which use the same metric (Jensen-Shannon divergence) on the same underlying expression matrix, ensuring the ratio is internally calibrated (Fig. 1).')

p('Step 1: Compute the neutral offset rate k_n. We restrict the pseudobulk vectors to housekeeping (HK) gene indices and apply softmax normalization, which converts expression values to probabilities. k_n is the JS divergence between these two HK gene probability distributions. Because HK genes should not differ systematically between biologically equivalent populations [4], k_n captures baseline noise: technical variation, stochastic transcriptional bursting, and individual-level physiological differences.')

p('Step 2: Compute the functional conversion rate k_f. We restrict the pseudobulk vectors to identity gene indices—genes that define cell-type-specific functions. In the default configuration, identity genes are the top-2,000 highly variable genes (HVGs; Seurat v3 flavor), excluding HK genes to maintain k_n/k_f independence. k_f is the JS divergence between these two identity gene probability distributions.')

p('Step 3: \u03c9 = k_f/k_n. For statistical inference, we perform bootstrap permutation testing (B = 500). Cell labels are randomly shuffled and \u03c9 recalculated to generate a null distribution. The empirical P-value is the fraction of permuted \u03c9 values that exceed the observed \u03c9 (with a +1 pseudocount to avoid P = 0), and effect size is reported as Cohen\'s d. All reported P-values are raw bootstrap P-values without multiple testing correction.')

p('We ran a parameter sweep on Tabula Muris mouse data (703 cell-type pairs across 6 organs) to test whether adding pathway enrichment scores to k_f would improve performance. We found that the identity-only configuration (w_identity = 1.0, w_pathway = 0.0) achieved the best cell-type discrimination (AUC = 0.847, Extended Data Fig. 1). CKI does not require external pathway databases to produce biologically meaningful results—partitioning the expression data into neutral and identity gene sets is sufficient.')

# --- Result 2 ---
heading('Calibration confirms neutral behavior at baseline', level=2)

p('We calibrated CKI on the Tabula Muris FACS dataset [5] (SmartSeq2, 15,057 cells, 22,308 genes, 6 organs). Housekeeping genes were auto-detected from data using a combined criterion (detection rate > 0.9 and CV below the 30th percentile), supplemented with 1,130 human-mouse conserved reference HK genes from the HRT Atlas [4] for human and mouse datasets. Identity genes were the top-2,000 highly variable genes (HVGs; Seurat v3), excluding HK genes (Fig. 2).')

p('The calibration confirmed correct neutral behavior. We performed six control comparisons in which we randomly split the same cell population into two halves. The mean \u03c9 was 1.54 (median 1.42, range 1.09–2.10), and none of the six comparisons reached statistical significance (all P > 0.05, two-sided bootstrap test). This confirms that CKI recognizes biologically equivalent cell populations as having no selective remodeling.')

p('Beyond controls, \u03c9 values increased monotonically with biological distance. Same cell type across different organs (S category: mean \u03c9 = 4.03, n = 4 pairs) had lower \u03c9 than different cell types within the same organ (D category: mean \u03c9 = 13.18, n = 3 pairs). The component-level analysis confirmed that k_f was the driver: k_f increased roughly 1,000-fold from controls to inter-cell-type comparisons, while k_n increased only about 100-fold. This establishes that CKI measures selective remodeling, not just total difference.')

# --- Result 3 ---
heading('CKI captures information that standard metrics miss', level=2)

p('We extended CKI to the Tabula Sapiens human atlas [6] (108,136 cells; 6 h5ad files total, 102 cell-type entries, 6 organs: liver, kidney, heart, bone marrow, spleen, lung). For human data, we used a hybrid scheme: k_n was computed once globally (using the full gene-by-cell-type pseudobulk matrix with the shared HK gene set), while k_f was computed per pair using the top-200 differentially expressed genes for that specific pair. This hybrid approach keeps k_n on a consistent scale (all cell types share the same HK gene set), while k_f adaptively selects the most informative identity genes for each pair. Critically, since \u03c9 = k_f/k_n is a ratio of JS divergences computed from the same underlying pseudobulk expression space, the normalization remains internally valid despite the different gene selection strategies. HK genes were auto-detected (combined criterion, with optional HRT Atlas enhancement) (Fig. 3).')

p('Human \u03c9 values ranged from 1.35 to 87.69 (mean 21.61, median 19.65, n = 4,851 pairs), substantively higher than mouse (mean 7.62). This difference likely reflects both the larger number of cell types (102 vs. ~30) and greater donor heterogeneity in human data (multiple donors vs. inbred mouse strains). Despite this, the biological hierarchy was preserved: same cell type across organs (mean \u03c9 = 8.65, n = 59 pairs) was lower than different cell types within the same organ (mean \u03c9 = 16.00, n = 1,140 pairs).')

p('The critical finding was that CKI captures a largely independent information dimension. We computed five metrics on all 4,851 human cell-type pairs: CKI \u03c9, raw JS divergence (all genes), Spearman distance, cosine distance, and marker Jaccard distance. CKI \u03c9 was negatively correlated with all four standard metrics (Spearman r = -0.36 to -0.46, all P < 0.001). In contrast, the four standard metrics formed a tight positive cluster (pairwise r = 0.57–0.94). This negative correlation is the strongest evidence that CKI measures something fundamentally different from all existing distance metrics.')

add_table_1(doc)
p('As expected by design, CKI was not optimized for cell-type classification (AUC = 0.680 vs. cosine AUC = 0.887; Table 1). But CKI was the only metric where same-organ pairs had higher values than different-organ pairs (mean \u03c9 16.00 vs. 13.66, Mann-Whitney U test, P < 0.001). All four standard metrics showed the opposite pattern (same-organ < different-organ). This reversal reflects CKI\'s sensitivity to functional specialization within shared microenvironments, a signal that standard metrics systematically obscure.')

# --- Result 4 ---
heading('Cancer analysis reveals unexpected transcriptional convergence', level=2)

p('We applied CKI to TCGA bulk RNA-seq data across five cancer types (LUAD, LUSC, LIHC, KIRC, BRCA) [7,8], totalling 10,535 samples. We asked a simple question: when cancer develops, how much selective transcriptional remodeling occurs? CKI provides a principled answer by comparing tumor-tumor (TT), normal-normal (NN), and tumor-normal (TN) \u03c9 values (Fig. 4).')

p('The most striking finding was that tumors are more transcriptionally homogeneous than normal tissues. In all five cancer types, the median NN/TT \u03c9 ratio exceeded 1.0, meaning that normal individuals differ more from each other than tumors differ from each other. Breast cancer (BRCA) showed the smallest contrast (median NN/TT = 1.40), while liver cancer (LIHC) showed the largest (median NN/TT = 2.83), with intermediate values for lung adenocarcinoma (LUAD 1.60), lung squamous (LUSC 1.43), and kidney clear cell (KIRC 1.98). This convergence toward shared transcriptional states may represent common vulnerabilities across genetically diverse tumors.')

p('Paired tumor-normal comparisons yielded higher \u03c9 than unpaired comparisons in four of five cancer types (paired/unpaired ratio = 0.99–3.25, Mann-Whitney P = 0.024 for LIHC, not significant for others). However, the small number of patients with paired tumor and normal samples (n = 2–5 per cancer type) limits statistical power and precludes definitive conclusions about within-patient versus between-patient variation.')

p('We then asked whether \u03c9 tracks with clinical severity within cancer types. In liver cancer, \u03c9 decreased with increasing Edmondson grade [10]: G1 (101.8 \u00b1 46.8, n = 39) > G2 (100.2 \u00b1 63.9, n = 133) > G3 (96.8 \u00b1 58.2, n = 105) > G4 (90.0 \u00b1 57.8, n = 11; Jonckheere-Terpstra trend test, P < 0.001). In breast cancer, PAM50 subtype analysis [11,12] revealed a gradient of transcriptional heterogeneity: Luminal A tumors had the highest intratumoral \u03c9 (344.5 \u00b1 323.4, n = 224), followed by Luminal B (313.6 \u00b1 282.7, n = 123), HER2-enriched (263.0 \u00b1 255.6, n = 55), and Basal-like tumors (223.4 \u00b1 183.7, n = 97), with Normal-like tumors having the lowest \u03c9 (108.0 \u00b1 65.5, n = 7; Kruskal-Wallis, P = 0.0002). Lung adenocarcinoma mutation stratification showed significant differences (Kruskal-Wallis, P = 0.017), with EGFR-mutant (285.3 \u00b1 180.1, n = 61) and KRAS-mutant tumors (284.6 \u00b1 227.9, n = 120) exhibiting higher \u03c9 than wild-type tumors (237.6 \u00b1 195.4, n = 311).')

# --- Result 5 ---
heading('CKI ranks cell types by cross-organ conservation', level=2)

add_table_2(doc)
p('Among the 4,851 Tabula Sapiens cell-type pairs, 59 are same-cell-type cross-organ comparisons. These pairs allowed us to ask: which cell types maintain their transcriptional identity regardless of where they reside, and which are strongly shaped by their organ environment (Fig. 5; Table 2)?')

p('The cross-organ ω ranking reveals a broad spectrum of conservation across 17 cell types (Table 2). Hepatocytes and B cells were among the most conserved (mean ω = 8.57 and 9.36, respectively, n = 1 each), followed by CD8+ T cells (mean 9.93 ± 4.2, n = 6) and plasma cells (mean 10.20 ± 1.9, n = 6). Macrophages, the most abundant cell type in cross-organ comparisons (n = 15), showed intermediate conservation (mean 15.49 ± 8.2). At the divergent end of the spectrum, endothelial cells (mean 26.25 ± 7.0, n = 3) and erythrocytes (mean 29.36 ± 18.8, n = 3) were the most organ-specific cell types. Endothelial cells are known to express organ-specific gene programs tailored to local vascular needs [32]. We note that several cell types (particularly those with n = 1 or 3) have small sample sizes, and their rankings should be interpreted with appropriate caution.')

p('The cross-organ conservation ranking from CKI showed little agreement with rankings from standard metrics (Spearman r = -0.40 to +0.02, n = 59 pairs). This is because CKI explicitly normalizes: two cell populations might share similar highly expressed genes (yielding high Jaccard similarity), but if their neutral baseline k_n is low, even modest functional differences can produce a high \u03c9. This normalization reveals patterns that raw expression similarity misses.')

# --- Result 6 ---
heading('Brain regional analysis reveals cell-type differentiation gradients', level=2)

p('We applied CKI to the Siletti et al. human brain single-nucleus RNA-seq atlas [9], which profiles ~3.3 million nuclei across ~100 brain regions. This dataset allowed us to ask: for a given cell type, how much functional divergence exists between the same cells residing in different brain regions? We focused on the 888,263 non-neuronal nuclei spanning 10 major cell classes (astrocytes, oligodendrocytes, oligodendrocyte precursors, microglia, vascular cells, fibroblasts, ependymal cells, choroid plexus, committed oligodendrocyte precursors, and Bergmann glia), and computed CKI \u03c9 for all same-cell-type cross-region comparisons (31,764 pairs total) (Fig. 6).')

p('The analysis revealed a striking differentiation gradient spanning 6.06-fold. Bergmann glia showed the lowest mean \u03c9 (2.37 \u00b1 1.14, n = 21 pairs across 7 regions), followed by committed oligodendrocyte precursor cells (3.17 \u00b1 1.47, n = 1,326 pairs across 52 regions), and fibroblasts (3.99 \u00b1 1.90, n = 3,403 pairs across 83 regions). Vascular cells (3.40 \u00b1 1.24, n = 3,321 pairs across 82 regions) and ependymal cells (4.13 \u00b1 1.73, n = 780 pairs across 40 regions) showed similarly low divergence. Microglia exhibited moderate divergence (mean \u03c9 = 8.02 \u00b1 4.93, n = 5,671 pairs across 107 regions). Oligodendrocytes and their precursors showed intermediate divergence (mean \u03c9 = 8.66 \u00b1 4.44 and 7.65 \u00b1 4.03, respectively). Astrocytes were the most regionally divergent cell type (mean \u03c9 = 14.36 \u00b1 8.68, n = 5,778 pairs across 108 regions), a 6.06-fold increase over Bergmann glia.')

p('This gradient aligns with known cell biology. Vascular cells and fibroblasts encounter relatively uniform extracellular environments across the brain—the blood-brain barrier and meningeal structures impose similar constraints regardless of anatomical location. Their low \u03c9 values suggest a conserved core transcriptional program with limited regional adaptation. Microglia, the brain\'s resident immune cells, showed intermediate divergence: while microglial phenotypes vary regionally, their core surveillance and phagocytic machinery is shared. Oligodendrocytes must myelinate diverse axonal populations and adjust internode lengths regionally, explaining moderate \u03c9 values. Astrocytes showed the highest \u03c9, consistent with extensive literature showing that astrocytes express region-specific sets of ion channels, neurotransmitter transporters, and secreted factors tailored to local neuronal circuit demands.')

p('The cross-region \u03c9 gradient provides a computational framework for inferring cell migration history. Cell types that recently migrated to new regions or that continuously exchange between regions should show low inter-regional \u03c9 because insufficient time has elapsed for transcriptional drift and local adaptation to accumulate. Conversely, cell types that have stably resided in specific regions for long periods should accumulate regional transcriptional signatures, yielding higher \u03c9. Under this framework, the low \u03c9 values for vascular cells and fibroblasts are consistent with continuous turnover and exchange through the circulatory and meningeal systems, while the high \u03c9 for astrocytes reflect long-term regional residence and functional specialization. This approach complements lineage tracing and developmental studies by providing an orthogonal transcriptomic readout of migration history.')

heading('CKI detects putative inter-regional migration events', level=2)

p('Low CKI \u03c9 between a cell type across two brain regions indicates transcriptomic similarity beyond neutral expectation. We systematically cross-validated Strong candidate signals against the developmental neuroscience literature and identified three distinct biological mechanisms underlying low inter-regional \u03c9: (i) developmental origin heterogeneity—cells from different embryonic progenitor pools (e.g., dorsal vs. ventral) retain distinct transcriptomic identities in adulthood; (ii) embryonic colonization route boundaries—immune cells that entered the brain through different developmental entry points show residual transcriptomic discontinuities; and (iii) postnatal cell migration—cells that physically relocate between regions through active motility. Mechanism assignment uses the multiplicative residual model to flag anomalous cell-type/region-pair combinations, with the biological interpretation corroborated by the known developmental biology of each cell type.')

p('To formalize migration inference, we designed a multiplicative model: for each (cell_type, region_pair) combination, expected_\u03c9 = \u03bc_ct \u00d7 \u03bc_pair / \u03bc_grand, where \u03bc_ct is the cell type\'s global mean \u03c9, \u03bc_pair is the region pair\'s mean \u03c9, and \u03bc_grand is the global mean (8.01). The multiplicative residual = observed / expected: a residual substantially below 1 indicates that the cell type is far less differentiated between those two regions than expected from both its own global plasticity and the region pair\'s overall divergence—a signature of shared transcriptional state potentially reflecting recent migration. We defined three confidence tiers: Strong (residual < 0.3, \u03c9 < 15, lowest \u03c9 in the region pair, and pair median \u03c9 > 20), Moderate (residual < 0.5, \u03c9 < 25), and Weak (residual < 0.75, \u03c9 < 35).')

p('Among 31,764 cross-region comparisons, 30 (0.09%) were classified as Strong migration candidates: Astrocyte (6), fibroblast (1), microglia (10), oligodendrocyte (10), and vascular cells (3). Another 1,247 pairs (3.93%) were Moderate candidates, and 6,567 (20.67%) were Weak candidates. By cell type, the Strong candidate counts reflect a combination of regional adjacency, shared developmental origins, and ongoing cellular interchange.')

heading('OPCs: key negative control validates method specificity', level=3)
p('Oligodendrocyte precursor cells (OPCs) are the most actively migrating cells in the adult CNS, continuously surveilling their environment along vascular scaffolds [13,17]. Yet CKI detected 0 Strong signals among 5,671 OPC cross-region comparisons—a finding that provides a critical orthogonal validation of the multiplicative residual model. The model is not simply detecting high \u03c9 values or absolute transcriptional differences; it identifies cell-type/region-pair combinations where the observed selective remodeling is strikingly below what the cell type\'s global plasticity and the region pair\'s background divergence would jointly predict. OPCs have a high global mean \u03c9 (7.65) because their transcriptional program includes both progenitor and differentiation states; their 52 Moderate signals (residual < 0.5) likely reflect the balance between shared developmental origins and ongoing regional maturation [35]. The complete absence of Strong signals despite OPCs being the brain\'s most motile cell type demonstrates that the residual model differentiates between broad baseline motility and specific transcriptional signatures of developmental history.')

heading('Oligodendrocytes: developmental origin rather than migration', level=3)
p('Mature oligodendrocytes contributed 10 Strong signals (residual 0.237–0.292), yet the prevailing view is that adult oligodendrocytes do not migrate between brain regions. We systematically cross-validated all 10 Strong pairs against the developmental neurobiology literature. Strikingly, all 10 signals involved cortex/thalamus (A13/A14/A19/A32/A40/Idg vs. Pul/LP) or brainstem-internal (MoRF-MoEN vs. PnRF) pairings—precisely the anatomical boundaries between dorsal- and ventral-derived oligodendrocyte populations. Foerster et al. [35] demonstrated through dorsal oligodendrocyte lineage ablation that >90% of adult cortical oligodendrocytes are dorsally derived (from cortical radial glia), while thalamic and brainstem oligodendrocytes are ventrally derived (MGE/LGE precursors). Ventral-derived cells fail to adopt cortical transcriptional programs even when transplanted into the cortex, indicating persistent cell-autonomous transcriptional identity. Our CKI analysis detects this developmental origin signature: dorsal vs. ventral oligodendrocyte populations are far less selectively diverged than expected from the oligodendrocyte global \u03c9 (8.66), because their transcriptional differences reflect shared generic myelination programs rather than region-specific functional specialization. LP (lateral posterior nucleus) and Pul (pulvinar), thalamic relay nuclei, contrast with cortical Brodmann areas A13-A40, forming the most consistent developmental boundary detected by our analysis. This reinterpretation—that CKI Strong signals for oligodendrocytes detect persistent developmental origin signatures rather than migration—is fully consistent with Foerster et al.\'s experimental data and provides, to the best of our knowledge, the first transcriptome-wide metric to distinguish dorsal and ventral oligodendrocyte populations without requiring lineage tracing.')

heading('Astrocytes: regional specialization with developmental origins', level=3)
p('Astrocytes showed the highest global \u03c9 (14.36) yet contributed 6 Strong signals, all concentrated in thalamic subnuclei (VLN-VPL, CM-VPL, Pul-VPL, LP-VPL-MN), hippocampal subfields (CA2-3 vs. DG-CA4), and cerebellar lobules (CBL vs. CBV). The thalamic signals are particularly informative: the ventroposterior lateral nucleus (VPL) appears in 4 of 6 Strong pairs, suggesting conserved astrocyte programs across thalamic relay nuclei that share a common developmental origin. Regionalized astrogenesis, driven by subnucleus-specific transcriptional programs, has been shown to produce persistent thalamic astrocyte heterogeneity that is detectable in adult tissue. Our finding that thalamic astrocyte pairs have \u03c9 values 5–6-fold below expectation indicates that these developmental signatures are selectively constrained\u2014astrocytes in functionally related thalamic nuclei retain transcriptional similarity beyond what would be predicted from astrocyte global plasticity alone. The cerebellar CBL vs. CBV signal (residual = 0.274) reflects the molecular topographic zones of Bergmann glia and cerebellar astrocytes described by Reeber et al. [41]. Endo et al. [18] demonstrated that Tcf4 controls astrocyte allocation during cortical development; our results extend this principle to subcortical structures, showing that compartmentalized astrogenesis leaves persistent transcriptional signatures detectable by CKI across the entire brain.')

heading('Bergmann glia: cerebellar molecular topography', level=3)
p('Bergmann glia had the lowest global \u03c9 (2.37) and only one Strong signal (CBL vs. CBV, residual = 0.274), consistent with their developmentally fixed, transcriptionally constrained state in the adult cerebellum. Bergmann glia are patterned into topographic molecular zones that align with cerebellar functional compartments [41], and their low global \u03c9 reflects their specialized role in maintaining Purkinje cell layer architecture with minimal regional transcriptional variation. The CBL (cerebellar lobule) vs. CBV (cerebellar vermis) signal likely reflects the established molecular topography difference between lateral cerebellar hemispheres and the midline vermis, rather than any migratory event.')

heading('Microglia: developmental colonization wave boundaries', level=3)
p('Microglia contributed 10 Strong signals despite the prevailing view that adult microglia are maintained by local self-renewal with minimal inter-regional migration [23]. We interpret these signals as developmental colonization residuals—transcriptomic boundaries between microglial populations that entered the brain through distinct embryonic entry routes. Shemer and Jung [37] comprehensively reviewed the molecular determinants of microglial developmental colonization, identifying three major invasion routes: the pial surface, the ventricular zone, and the vasculature. Critically, colonization proceeds in a rostral-to-caudal wave, with forebrain regions colonized earlier than midbrain and hindbrain structures. Menassa et al. [20] confirmed this spatiotemporal gradient in human fetal brain, showing that microglial density peaks in the cortical subplate at mid-gestation before declining as cells redistribute. Barry-Carroll et al. [40] further demonstrated that microglial colonization occurs through clonal expansion of highly proliferative precursors, with spatial constraints limiting clone mixing between regions. Our CKI Strong signals systematically trace this rostral-to-caudal colonization gradient: 7 of 10 Strong microglia signals involve cortex (A5-A7, A14, A23, A29-A30) paired with midbrain/pontine structures (IC, DTg, PnAN, PnEN, MoRF-MoEN). The inferior colliculus (IC), a midbrain auditory relay, appears in 4 of 10 Strong pairs and we propose that IC represents a contact zone where forebrain-derived and hindbrain-derived microglial colonization waves meet. DTg (dorsal tegmental nucleus) vs. SN (substantia nigra), the single strongest signal (residual = 0.132), lies entirely within the mesencephalon—a region that receives microglial progenitors from both the pial surface and ventricular routes, potentially creating mixed-origin populations with shared transcriptional programs. This hypothesis—that CKI detects colonization wave boundaries from adult transcriptomic data—is directly testable through lineage-tracing or spatial transcriptomic analysis of developing human brain tissue.')

heading('Vascular cells: blood-brain barrier regional identity', level=3)
p('Vascular cells contributed 3 Strong signals, all involving ITG (inferior temporal gyrus, cortical) paired with brainstem structures (PB, PnRF, SN). Adult brain endothelial cells have a proliferation rate of ~0.4% and form a physically continuous network, making active migration an unlikely explanation. Schaffenrath et al. [39] systematically characterized blood-brain barrier (BBB) heterogeneity across brain regions, identifying hundreds of region-specific molecular differences in endothelial cells, pericytes, and vascular smooth muscle cells. These differences are established during development through region-specific transcription factors (Foxf2, Foxq1, Barhl2) that sculpt the BBB in response to local neuronal and glial signals. Our Strong signals likely reflect this developmentally fixed BBB regional identity: cortical (ITG) and brainstem vasculature have systematically different transporter repertoires, tight junction compositions, and metabolic enzyme profiles. The low inter-regional \u03c9 values (1.75–1.80) indicate that the core vascular transcriptional program is highly conserved, with selective remodeling detectable only as modest deviations from the global vascular \u03c9 baseline.')

heading('Fibroblast: the sole postnatal migration signal', level=3)
p('Perivascular fibroblasts contributed 1 Strong signal (A40 vs. SN, residual = 0.299)—the only Strong candidate among all 30 that can be parsimoniously attributed to postnatal cell migration. Jones et al. [38] demonstrated that brain perivascular fibroblasts originate from the meninges and migrate along penetrating vessels into the brain parenchyma during postnatal development. This migration is active, coordinated with perivascular macrophage colonization, and establishes a population of fibroblasts in the Virchow-Robin spaces that persists in adulthood. The A40 (prefrontal cortex) vs. SN (substantia nigra) signal likely reflects fibroblasts that entered these anatomically distant regions along different branches of the cerebral vasculature from a shared meningeal origin. The brain-wide fibroblast population is small (8,156 nuclei across 83 regions), and the detection of a single Strong signal from this sparse population demonstrates the multiplicative residual model\'s sensitivity to genuine migration events with limited statistical power.')

# ============================================================
# DISCUSSION
# ============================================================
heading('Discussion', level=1)

p('CKI introduces a conceptual shift in transcriptomic comparison: from measuring absolute distance to quantifying selective remodeling. The key insight is that housekeeping genes provide an internal neutral reference, analogous to synonymous substitutions in Ka/Ks analysis. This decomposition lets us ask not just "how different are these two populations?" but "how much of that difference is functional versus neutral?"')

p('CKI is a perturbation index, not a classifier—and this is by design. Classifying cell types from transcriptomic data is largely a solved problem. CKI answers a complementary question: regardless of cell-type labels, how much selective remodeling separates two populations? The negative correlation with all standard metrics proves that CKI captures information that existing approaches miss.')

p('The Ka/Ks analogy is productive but not perfect. Ka/Ks operates on DNA sequence alignments with explicit codon models [21]; CKI operates on continuous expression vectors. The neutral reference in Ka/Ks has a mechanistic basis in the genetic code, whereas HK genes are defined empirically by two data-driven criteria—high detection rate and low expression variance—and may vary by tissue context. These differences suggest directions for future theoretical development.')

p('CKI complements rather than replaces existing methods. SAMap [15] and SATURN [3] excel at cross-species alignment; CACIMAR [22] provides conservation scoring that could be reinterpreted through the CKI lens. More broadly, CKI provides a principled null model for any transcriptomic comparison: before concluding that two populations are meaningfully different, ask whether the difference exceeds neutral expectation.')

p('The TCGA finding that tumors are more homogeneous than normal tissues (median NN/TT > 1.0) has implications for cancer biology. If genetically diverse tumors converge on shared transcriptional states, that convergence may point to common vulnerabilities that transcend individual mutations. The PAM50 analysis reinforces this: aggressive subtypes show the strongest convergence, consistent with proliferation programs overriding tissue-specific expression.')

p('The cross-organ and cross-brain-region analyses establish CKI as a general tool for measuring functional differentiation at multiple spatial scales. The brain analysis revealed a 6.06-fold \u03c9 gradient across 10 cell classes (from 2.37 in Bergmann glia to 14.36 in astrocytes), demonstrating that CKI can detect regional functional specialization even among cells of the same type. Critically, the multiplicative residual model detected 30 Strong candidate signals and systematic cross-validation against the developmental neuroscience literature revealed that these signals predominantly reflect three distinct biological processes rather than active cell migration: (i) developmental origin heterogeneity—oligodendrocyte dorsal/ventral origin differences [35] explain all 10 oligodendrocyte Strong signals as cortex vs. thalamus/brainstem boundaries; (ii) embryonic colonization route discontinuities—microglial rostral-to-caudal colonization waves [20,37,40] create transcriptomic boundaries at the forebrain-midbrain interface, with the inferior colliculus identified as a candidate contact zone; and (iii) compartmentalized developmental astrogenesis and vascular specification [18,39]. Importantly, OPCs—the most actively migrating cells in the adult CNS [13,17]—yielded 0 Strong signals among 5,671 comparisons, providing a powerful orthogonal validation that the residual model specifically detects fixed developmental signatures rather than ongoing cell motility. The sole exception is the perivascular fibroblast A40-SN signal, which is consistent with known postnatal meningeal-to-parenchymal fibroblast migration [38]. Together, these results demonstrate that CKI detects persistent transcriptional signatures of developmental history—origin, colonization route, and specification—embedded in adult transcriptomic data, rather than inferring active migration per se. This reframing opens new applications for adult single-cell atlases as archives of developmental information.')

p('Limitations should be noted. First, CKI currently operates at the pseudobulk level; single-cell extensions would need to address sparsity and dropout. Second, the choice of HK gene set influences results; CKI uses data-driven auto-detection (combined detection-rate and CV filtering) as the universal default, with optional HRT Atlas enhancement [4] for human and mouse. Sensitivity analysis showed that CKI results are robust to alternative HK definitions (using the lowest 10% variable genes as a neutral set yielded \u03c9 correlations r > 0.95). Third, TCGA analysis was limited to bulk RNA-seq resolution; single-cell or spatial transcriptomic data would enable finer perturbation quantification. Fourth, the brain analysis uses post-mortem tissue; developmental time courses would provide stronger evidence for migration inference.')

p('Future directions include developmental biology (quantifying functional differentiation between developmental stages), drug response profiling (measuring selectivity of drug-induced transcriptional changes), aging research (tracking age-related neutral vs. functional transcriptional drift), and evolutionary cell biology (quantifying conservation and divergence of cell-type programs across the tree of life). The CKI Python package (v0.3.1) and all analysis notebooks are available at https://github.com/zhanglknt/CKI-cell-type-identification under the MIT License.')

# ============================================================
# CONCLUSIONS
# ============================================================
heading('Conclusions', level=1)

p('CKI provides a principled framework for decomposing transcriptomic divergence into neutral and functional components, enabling quantitative assessment of selective remodeling analogous to Ka/Ks in molecular evolution. Key conclusions include:')
p('(i) CKI \u03c9 is negatively correlated with all standard distance metrics (r = -0.36 to -0.46), proving it captures an orthogonal information dimension not measured by existing approaches.')
p('(ii) Validation on Tabula Muris confirms that biologically equivalent cell populations yield \u03c9 close to 1 (mean 1.54, all P > 0.05), establishing a calibrated neutral baseline.')
p('(iii) TCGA analysis reveals that tumors are more transcriptionally homogeneous than normal tissues (median NN/TT = 1.40–2.83 across five cancer types), suggesting convergent transcriptional states across genetically diverse tumors.')
p('(iv) Brain regional analysis of 888,263 nuclei reveals a 6.06-fold \u03c9 gradient across 10 non-neuronal cell classes, with astrocytes showing the highest regional divergence (mean \u03c9 = 14.36) and Bergmann glia the lowest (mean \u03c9 = 2.37).')
p('(v) The multiplicative residual model detects 30 candidate signals among 31,764 cross-region comparisons; cross-validation against developmental neuroscience literature reveals that these primarily reflect developmental origin heterogeneity (oligodendrocytes, 10/30), embryonic colonization route boundaries (microglia, 10/30), and compartmentalized developmental specification (astrocytes and vascular cells, 9/30), with only a single postnatal migration signal (perivascular fibroblasts).')
p('CKI is available as an open-source Python package (v0.3.1, MIT License) at https://github.com/zhanglknt/CKI-cell-type-identification, with all analysis notebooks and reproduction scripts provided.')

# ============================================================
# METHODS (must appear AFTER Conclusions for Genome Biology)
# ============================================================
heading('Methods', level=1)

heading('CKI computation', level=2)
p('We normalize raw count matrices to 10,000 counts per cell and apply log1p transformation. Pseudobulk vectors are computed by averaging expression across cells sharing the same cell-type annotation, requiring at least 10 cells per group. Housekeeping (HK) genes are auto-detected from data using a combined criterion: detection rate > 0.9 (expressed in >90% of cells) and coefficient of variation below the 30th percentile among well-expressed genes (mean expression > 0.5). For human and mouse datasets, the HRT Atlas v1.0 consensus set (1,130 human-mouse shared HK genes) [4] is optionally used as supplementary enhancement (union with detected set). For any other species, detection is purely data-driven without external references.')

p('For populations A and B with pseudobulk vectors \u03b5_A and \u03b5_B, k_n = JS(softmax(\u03b5_A[H]), softmax(\u03b5_B[H])), where H is the set of HK gene indices. k_f = JS(softmax(\u03b5_A[I]), softmax(\u03b5_B[I])), where I is the set of top-2,000 highly variable genes (HVGs; Seurat v3 flavor) excluding HK genes. \u03c9 = k_f/k_n. JS divergence uses base-2 logarithm (range [0,1]). Softmax normalization converts expression vectors to probability distributions.')

heading('Bootstrap permutation test', level=2)
p('We randomly permute cell labels between the two populations (B = 500), recompute pseudobulk vectors, and calculate \u03c9_null for each permutation. We apply a two-sided bootstrap test: Empirical P = (count(|\u03c9_null - 1| >= |\u03c9_obs - 1|) + 1) / (B + 1). Cohen\'s d = (\u03c9_obs - mean(\u03c9_null)) / sd(\u03c9_null). All reported P-values are raw bootstrap P-values without multiple testing correction.')

heading('Datasets', level=2)
p('Tabula Muris FACS SmartSeq2 [5]: 15,057 cells, 22,308 genes, 6 organs (liver, kidney, spleen, lung, heart, bone marrow). Post-QC: 32 cell-type entries (each with at least 10 cells). Highly variable genes selected using scanpy [25] with flavor="seurat" [26,27] and n_top_genes=2,000.')

p('Tabula Sapiens v1.0 [6]: accessed via CZ CELLxGENE Discover. Post-QC: 108,136 cells (6 h5ad files total), 51,852 genes, 99 cell-type entries across 6 organs. Human HK genes: auto-detected (combined detection-rate/CV criterion), with optional enhancement from 1,130 HRT Atlas v1.0 genes mapped by gene symbol.')

p('TCGA bulk RNA-seq [28]: five cancer types from NCI Genomic Data Commons, accessed via TCGAbiolinks [29] and cBioPortal [14] APIs. LUAD: 495 tumor + 76 normal; LUSC: 567 + 58; LIHC: 365 + 57; KIRC: 755 + 82; BRCA: 1,032 + 109. FPKM values, log2(x+1) transformed. PAM50 classification [11,12]: nearest centroid (Pearson correlation), 44 of 47 PAM50 genes matched. LIHC Edmondson grade [10]: from cBioPortal, 289 tumors. LUAD mutations: from cBioPortal, 497 samples (61 EGFR, 121 KRAS, 312 WT).')

p('Human brain atlas [9]: Siletti et al. (2023) single-nucleus RNA-seq from CZ CELLxGENE Discover. We used the Nonneurons.h5ad dataset (888,263 nuclei, 59,480 genes, 108 brain regions). Cell types were classified by supercluster_term annotation, generating 10 major non-neuronal classes: astrocytes (155,025 nuclei), oligodendrocytes (490,246), oligodendrocyte precursors (110,454 total including committed), microglia (91,838), vascular cells (8,932), fibroblasts (8,156), ependymal cells (5,882), choroid plexus (7,689), and Bergmann glia. We required >= 20 nuclei per (region, cell_type) group and >= 50 nuclei per region. Normalization: Scanpy normalize_total (target_sum = 10,000) followed by log1p transformation. Pseudobulk vectors were computed as the mean log-normalized expression per group. CKI \u03c9 was computed for all same-cell-type cross-region comparisons (31,764 pairs total), using the hybrid scheme described above. HK genes were auto-detected per dataset (combined detection-rate/CV criterion). Top-200 identity genes were selected per comparison, excluding HK genes.')

heading('Method comparison', level=2)
p('We computed five metrics on all 4,851 Tabula Sapiens cell-type pairs: CKI \u03c9 (hybrid scheme), raw JS divergence (all genes), Spearman distance (1 - \u03c1), cosine distance (1 - cos \u03b8), and marker Jaccard distance (1 - Jaccard index of top-200 expressed genes). Inter-metric Spearman correlations and cell-type classification ROC-AUC were computed using scikit-learn.')

heading('Multiplicative residual model for brain regional analysis', level=2)
p('For the brain regional analysis, we designed a multiplicative model to detect cell-type/region-pair combinations with anomalously low \u03c9. For each (cell_type, region_pair) combination, expected_\u03c9 = \u03bc_ct \u00d7 \u03bc_pair / \u03bc_grand, where \u03bc_ct is the cell type\'s global mean \u03c9, \u03bc_pair is the region pair\'s mean \u03c9, and \u03bc_grand is the global mean (8.01). The multiplicative residual = observed / expected. A residual substantially below 1 indicates the cell type is far less differentiated between those two regions than expected from both its own global plasticity and the region pair\'s overall divergence. We defined three confidence tiers: Strong (residual < 0.3, \u03c9 < 15, lowest \u03c9 in the region pair, and pair median \u03c9 > 20), Moderate (residual < 0.5, \u03c9 < 25), and Weak (residual < 0.75, \u03c9 < 35). Strong candidate signals were systematically cross-validated against the developmental neuroscience literature to assign each signal to one of four biological mechanisms: developmental origin heterogeneity (DO), embryonic colonization route boundaries (CR), compartmentalized developmental specification (DS), or postnatal cell migration (PM).')

heading('Clinical severity analysis', level=2)
p('For TCGA clinical severity analysis, we computed intratumoral \u03c9 for samples within each clinical stratum using the hybrid scheme (global k_n from shared HK genes, per-pair k_f from top-200 DE genes). BRCA PAM50 subtypes were classified using nearest centroid (Pearson correlation) against the published PAM50 centroids [11,12]; 44 of 47 PAM50 genes matched to TCGA gene symbols. LIHC Edmondson histological grades [10] were obtained from cBioPortal (n = 288 tumors with both grade and expression data). LUAD mutation status (EGFR, KRAS, WT) was retrieved from cBioPortal (n = 492 samples). Between-stratum differences were tested with Kruskal-Wallis for PAM50 subtypes and LUAD mutations; trend across Edmondson grades was tested with Jonckheere-Terpstra. Paired vs. unpaired tumor-normal comparisons used Mann-Whitney U test.')

heading('Computational environment', level=2)
p('All analyses were performed in Python 3.12 with scanpy >= 1.9.0 [25], scipy >= 1.10.0, numpy >= 1.23.0, pandas >= 1.5.0, matplotlib >= 3.6.0, seaborn >= 0.12.0, and scikit-learn >= 1.2.0. All random seeds were fixed at 42.')

heading('Statistical reporting', level=2)
p('We report summary statistics as mean \u00b1 s.d. (range) or median [IQR] as noted. Box plots display median, IQR, and 1.5\u00d7 IQR whiskers. All P-values are two-sided unless otherwise specified. Bootstrap inference uses B = 500 permutations; empirical P-values use the +1 pseudocount formula. All reported P-values are raw bootstrap P-values without multiple testing correction. Effect sizes are reported as Cohen\'s d; d > 0.8 indicates a large effect. Correlation coefficients (Spearman \u03c1) are reported with associated P-values. Omnibus tests (Kruskal-Wallis, Jonckheere-Terpstra) use P < 0.05 without additional correction.')

# ============================================================
# ABBREVIATIONS
# ============================================================
heading('Abbreviations', level=1)
p('CKI: Cell-state Kinetic Index')
p('HK: Housekeeping (genes)')
p('JS: Jensen-Shannon (divergence)')
p('AUC: Area Under the Curve')
p('ROC: Receiver Operating Characteristic')
p('HVG: Highly Variable Gene')
p('SNN: Shared Nearest Neighbor')
p('TCGA: The Cancer Genome Atlas')
p('NN/TT: Normal-Normal / Tumor-Tumor (omega ratio)')

# ============================================================
# DECLARATIONS (mandatory for Genome Biology)
# ============================================================
heading('Declarations', level=1)

heading('Ethics approval and consent to participate', level=2)
p('Not applicable. This study used only publicly available, previously published datasets: Tabula Muris [5], Tabula Sapiens [6], TCGA [7,8], and the human brain atlas [9]. No new human subjects research was conducted.')

heading('Consent for publication', level=2)
p('Not applicable.')

heading('Availability of data and materials', level=2)
p('Tabula Muris data: GEO accession GSE109774. Tabula Sapiens data: CZ CELLxGENE Discover (https://cellxgene.cziscience.com/). TCGA data: NCI Genomic Data Commons (https://portal.gdc.cancer.gov/). HRT Atlas (optional human/mouse HK reference): https://www.housekeeping.unicamp.br. Human brain atlas: CZ CELLxGENE Discover (collection ID as referenced in [9]). PAM50 centroids: from Parker et al. [11]. MSigDB Hallmark gene sets: from Liberzon et al. [34]. The CKI source code (v0.3.1, MIT License) is publicly available at https://github.com/zhanglknt/CKI-cell-type-identification (tag v0.3.1). A permanent archival copy has been deposited at Zenodo (DOI: 10.5281/zenodo.15670808). All analysis notebooks and processed data matrices are included in the Supplementary Materials.')

heading('Competing interests', level=2)
p('The authors declare no competing interests.')

heading('Funding', level=2)
p('This work was supported by the National Natural Science Foundation of China (NSFC) under grant number 32370682, and the Prevention and Control of Emerging and Major Infectious Diseases — National Science and Technology Major Project (grant number 2026ZD01910500).')

heading('Authors\' contributions', level=2)
p('XW and LZ conceived the study and designed the computational framework. LZ developed the CKI algorithm, performed all analyses, and prepared all figures. XW contributed to data curation, validation, and manuscript writing. Both authors read and approved the final manuscript.')

heading('Acknowledgements', level=2)
p('We thank the Tabula Muris Consortium, Tabula Sapiens Consortium, TCGA Research Network, and the Siletti et al. brain atlas team for making their data publicly available.')

# ============================================================
# FIGURE LEGENDS
# ============================================================
heading('Figure legends', level=1)

p('Figure 1. The CKI framework. (a) Conceptual analogy between Ka/Ks in molecular evolution and CKI in transcriptomics. Ka/Ks uses synonymous substitution rate (Ks) as a neutral baseline; \u03c9 = Ka/Ks > 1 indicates positive selection. CKI uses housekeeping gene divergence for k_n and identity gene divergence for k_f; \u03c9 = k_f/k_n > 1 indicates selective transcriptomic remodeling. (b) Computational pipeline: raw count matrix \u2192 pseudobulk \u2192 JS divergence on HK genes (k_n) and identity genes (k_f) \u2192 \u03c9 = k_f/k_n. (c) Bootstrap \u03c9 distribution (B = 500 permutations) with median indicated by dashed line. (d) Scatter plot of k_n vs. k_f, showing that functional variation dominates neutral baseline. (e) \u03c9 distribution with \u03c9 = 1 (neutral) marked by dashed line.')

p('Figure 2. CKI calibration on Tabula Muris mouse data. (a) k_n calibration across six Tabula Muris cell types from control comparisons (C category: random split of same population). k_n values are stable across cell types, confirming neutral baseline behavior. (b) Component decomposition of k_n and k_f across four comparison categories: C (same cell type), S (same sub-organ), D (different cell type within same organ), and X (cross-organ). k_f increases monotonically with biological distance, while k_n remains relatively constrained. (c) Spearman correlation between CKI \u03c9 and four standard metrics (Cosine, Raw JS, Marker Jaccard, Spearman) on Tabula Muris data. All show negative correlation, confirming \u03c9 captures orthogonal information. (d) Pathway enrichment in the k_f component, showing fold change (k_f/k_n) for top enriched pathways. Stars indicate significance: *** P < 0.001, ** P < 0.01, * P < 0.05.')

p('Figure 3. CKI captures independent information. (a) Spearman correlation heatmap of five metrics on n = 4,851 Tabula Sapiens pairs. CKI \u03c9 is negatively correlated with all four standard metrics. Standard metrics form a positive cluster. (b) Scatter plot of CKI \u03c9 vs. k_n (neutral rate) on Tabula Sapiens data (n = 4,851 pairs), colored by same-organ vs. cross-organ pairs, with Spearman r. (c) ROC curves for cell-type classification across five metrics on Tabula Sapiens data. (d) \u03c9 by comparison category: C (same cell type), S (same sub-organ), D (different cell type, same organ), X (cross-organ). Box plots show log10(\u03c9) distribution. (e) AUC (cell-type classification) vs. interpretability (decomposability) comparison across five metrics. CKI \u03c9 has the lowest AUC but is the only fully decomposable metric.')

p('Figure 4. TCGA pan-cancer perturbation analysis. (a) Median NN/TT \u03c9 ratio across five cancer types. NN > TT in all cancers (all ratios > 1.0), indicating that normal individuals differ more from each other than tumors differ from each other. (b) \u03c9 distributions for normal-normal (NN) vs. tumor-tumor (TT) comparisons per cancer type. (c) Median TN/NN \u03c9 ratio (tumor-normal vs. normal-normal) across five cancer types. (d) Bootstrap Cohen\'s d effect sizes for NN vs. TT comparisons per cancer type. (e) Tissue-level pairwise \u03c9 matrix heatmap across five cancer types.')

p('Figure 5. Cross-organ cell-type conservation. (a) CKI \u03c9 ranking of 38 shared cell types between human and mouse. Immune cells (macrophage, T cells) rank lowest; structural cells (erythrocyte, endothelial) rank highest. (b) \u03c9 distribution across all cell-type pairs, showing conserved vs. variable cell types. (c) Cross-organ \u03c9 gradient (mean \u00b1 SD) showing systematic variation across organ pairs. (d) Table of top conservative cell-type pairs with the lowest cross-organ \u03c9 values.')

p('Figure 6. Brain regional cell-type differentiation and migration inference. (a) Brain region map schematic showing anatomical regions analyzed in the Siletti et al. atlas. (b) \u03c9 gradient across 10 non-neuronal cell classes. Bergmann glia (mean \u03c9 = 2.37) shows the lowest regional divergence; astrocytes (mean \u03c9 = 14.36) show the highest (6.06-fold gradient). Box plots: center line, median; box, IQR; whiskers, 1.5\u00d7 IQR. (c) Astrocyte \u03c9 across brain regions, with hierarchical clustering. Cortex regions cluster together, as do thalamic and brainstem regions. (d) Migration candidate detection: multiplicative residual model identifies 30 Strong candidates (residual < 0.3, \u03c9 < 15, lowest \u03c9 in pair, pair median \u03c9 > 20) among 31,764 cross-region comparisons. (e) Observed vs. expected \u03c9 for migration candidates, showing the gap between observed regional similarity and the cell-type-specific expectation.')

# ============================================================
# EXTENDED DATA FIGURE LEGENDS
# ============================================================
heading('Extended Data Figure legends', level=1)

p('Extended Data Figure 1. Parameter sweep and pathway analysis. (a) k_n stability as a function of housekeeping gene set size, showing convergence at ~200-300 HK genes. (b) k_f component contribution per pathway: decomposition of k_n vs. k_f for representative pathways. (c) Weight sweep for multi-component k_f. Identity-only (w_identity = 1.0, w_pathway = 0.0) achieves optimal cell-type discrimination (AUC = 0.847, n = 703 mouse cell-type pairs, 6 organs).')

p('Extended Data Figure 2. Cross-species validation details. (a) Cross-species \u03c9 conservation: scatter plot of human vs. mouse \u03c9 for shared cell types, with Spearman r and P-value. (b) HK gene set detection stability: overlap between human and mouse HK gene sets detected by the combined criterion. (c) \u03c9 distribution comparison between mouse (n = 15 shared cell types) and human (n = 4,851 pairs).')

p('Extended Data Figure 3. TCGA per-cancer matrices. Pairwise \u03c9 matrices for six cancer types (BRCA, KIRC, LIHC, LUAD, COAD, HNSC) showing tissue-level transcriptomic divergence structure within each cancer cohort.')

p('Extended Data Figure 4. Method comparison performance. ROC-AUC bar plot for cell-type classification across five metrics.')

p('Extended Data Figure 5. Cross-organ conservation raw data. Complete table of 59 same-CT cross-organ pairs (Extended Data Table 2).')

p('Extended Data Figure 6. Brain regional analysis details. (a) Cell type nuclei counts per brain region, showing the distribution of non-neuronal nuclei across sampled regions. (b) k_n/k_f decomposition per cell class, showing that \u03c9 variance is predominantly driven by k_f. (c) \u03c9 vs. number of regions (n_regions) per cell type, showing no significant correlation, ruling out sampling bias. (d) Region-region \u03c9 matrix for astrocytes across brain regions, with hierarchical clustering. (e) Top migration candidates by tier (Strong, Moderate, Weak), ranked by multiplicative residual.')

p('Extended Data Figure 7. Developmental signature detection. (a) Multiplicative residual distribution for all 31,764 cross-region pairs, with Strong (residual < 0.3), Moderate (residual < 0.5), and Weak (residual < 0.75) tiers shaded. (b) Strong candidate counts by cell type. OPCs (0 Strong despite highest motility) provide key negative control, validating that the model detects developmental origin signatures rather than general motility. (c) Top 10 Strong candidates ranked by multiplicative residual, annotated with cell type, region pair, and biological mechanism: DO (developmental origin), CR (colonization route), DS (developmental specification), or PM (postnatal migration). (d) Migration candidates by cell type and confidence tier, showing tier distribution across all 10 non-neuronal cell classes.')

# ============================================================
# REFERENCES (Vancouver numbered format)
# ============================================================
heading('References', level=1)

for ref in _refs_gb:
    ref_p_gb(ref)

# == Save ==
PROJECT_ROOT = Path(__file__).resolve().parent
out = str(PROJECT_ROOT / "results" / "CKI_GenomeBiology_Manuscript_v1.docx")
doc.save(out)
print(f'Saved: {out}')
