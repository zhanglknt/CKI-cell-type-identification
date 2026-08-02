"""
Generate CKI_NAR_Manuscript.docx — Nucleic Acids Research article.

NAR formatting compliance:
- Single-column, single-spaced
- No line numbers (page numbers only)
- Unstructured abstract (≤200 words, single paragraph)
- Section order: Introduction → Materials and Methods → Results → Discussion
- Numbered references in parentheses: (2), (3,4), (4-7)
- NAR reference style: Author,A.B., Author,C.D. (Year) Title. *Journal.*, **Vol**, Pages.
- Authors listed up to 20, then et al.
- All text black; Arial/Helvetica (TrueType embedded)
- Graphical abstract placeholder
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

# Load all manuscript data dynamically from CSV files
from _load_manuscript_data import get_manuscript_data
DATA = get_manuscript_data()

# Shorthand accessors
_ds = DATA['datasets']
_mc = DATA['mouse_calibration']
_h = DATA['human']
_sc = DATA['spearman_corr']
_tc = DATA['tcga']
_br = DATA['brain']
_br_ct = _br['cell_types']
_au = DATA['table1_auc']
_sb = DATA['sweep']
_co = DATA['cross_organ_spearman']
_bb = DATA['bootstrap']['brain']

# Phase B: Statistical upgrade data
import json as _json
import pandas as pd
RESULTS_DIR = Path(__file__).parent / "results"
_phaseB_ci = pd.read_csv(RESULTS_DIR / "phaseB_bootstrap_cis.csv") if (RESULTS_DIR / "phaseB_bootstrap_cis.csv").exists() else None
with open(RESULTS_DIR / "phaseB_omega_distribution.json") as _f:
    _phaseB_dist = _json.load(_f)
with open(RESULTS_DIR / "phaseB_adaptive_analysis.json") as _f:
    _phaseB_adaptive = _json.load(_f)
_phaseB_residual = None
if (RESULTS_DIR / "phaseB_residual_pervisign.csv").exists():
    _phaseB_residual = pd.read_csv(RESULTS_DIR / "phaseB_residual_pervisign.csv")

# Phase C: Methodological reinforcement data
_phaseC_cal = None
_phaseC_dim = None
_phaseC_kn = None
if (RESULTS_DIR / "phaseC_calibration.json").exists():
    with open(RESULTS_DIR / "phaseC_calibration.json") as _f:
        _phaseC_cal = _json.load(_f)
if (RESULTS_DIR / "phaseC_dimensionality.json").exists():
    with open(RESULTS_DIR / "phaseC_dimensionality.json") as _f:
        _phaseC_dim = _json.load(_f)
if (RESULTS_DIR / "phaseC_kn_variability.json").exists():
    with open(RESULTS_DIR / "phaseC_kn_variability.json") as _f:
        _phaseC_kn = _json.load(_f)

t2 = DATA['table2_data']
tcga_cancers = _tc['cancers']
min_c = min(tcga_cancers, key=lambda x: x['nn_tt_ratio'])
max_c = max(tcga_cancers, key=lambda x: x['nn_tt_ratio'])
th = _br['residual_thresholds']

mac = [r for r in t2 if 'Macrophage' in r[0]][0]
last2 = t2[-2:]

def find_ct(name):
    for ct in _br_ct:
        if name.lower() in ct['name'].lower():
            return ct
    return None

sbg = sorted(_br_ct, key=lambda x: x['omega_mean'])
opc = find_ct('oligodendrocyte precursor')
astro = find_ct('astrocyte')
oligo_ct = find_ct('oligodendrocyte')
bergmann_ct = find_ct('bergmann')

tcga_detail = '; '.join([
    f"{c['name'].replace('TCGA-','')}: {c['n_tumor']} tumor + {c['n_normal']} normal"
    for c in _tc['cancers']
])

doc = Document()

# == NAR formatting: single-spaced, Arial 11pt ==
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(20)
style.font.color.rgb = RGBColor(0,0,0)
style.paragraph_format.line_spacing = 1.15  # single spacing
style.paragraph_format.space_after = Pt(0)

# Page margins
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# == Helpers ==
def set_black(run):
    run.font.color.rgb = RGBColor(0,0,0)

def set_superscript(run):
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn('w:vertAlign')):
        rPr.remove(old)
    va = rPr.makeelement(qn('w:vertAlign'), {qn('w:val'): 'superscript'})
    rPr.append(va)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
        set_black(run)
        run.font.size = Pt([16, 14, 12][level-1])
        rPr = run._element.get_or_add_rPr()
        for old in rPr.findall(qn('w:b')):
            rPr.remove(old)
        b = rPr.makeelement(qn('w:b'), {})
        rPr.append(b)
    p.paragraph_format.space_before = Pt(21)
    p.paragraph_format.space_after = Pt(8)
    return p

def p(text, bold=False, italic=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    set_black(run)
    run.bold = bold
    run.italic = italic
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.first_line_indent = Cm(0)
    return para

# == Table helpers ==
def add_table_1(doc):
    auc = DATA['table1_auc']
    ds = DATA['datasets']
    para = doc.add_paragraph()
    run = para.add_run(f"Table 1. Classification AUC of five metrics on Tabula Sapiens ({ds['tabula_sapiens_ct_entries']} cell types, {DATA['human']['n_pairs_total']:,} pairs).")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(21)
    para.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'ROC-AUC'
    data = [
        ('Cosine distance', f"{auc['cosine']:.3f}"),
        ('Raw JS divergence', f"{auc['raw_js']:.3f}"),
        ('Marker Jaccard distance', f"{auc['marker_jaccard']:.3f}"),
        ('Spearman distance', f"{auc['spearman']:.3f}"),
        ('CKI omega', f"{auc['cki_omega']:.3f}"),
    ]
    for i, (metric, auc_val) in enumerate(data):
        table.rows[i+1].cells[0].text = metric
        table.rows[i+1].cells[1].text = auc_val
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_table_2(doc):
    t2 = DATA['table2_data']
    n_rows = len(t2) + 1
    para = doc.add_paragraph()
    run = para.add_run(f"Table 2. Cross-organ conservation ranking by cell type (Tabula Sapiens, n={DATA['cross_organ_n_total']} same-cell-type cross-organ pairs).")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(21)
    para.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=n_rows, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Cell type'
    hdr[1].text = 'Mean \u03c9'
    hdr[2].text = 'SD'
    hdr[3].text = 'n pairs'
    for i, (name, mean_w, sd, n) in enumerate(t2):
        table.rows[i+1].cells[0].text = name
        table.rows[i+1].cells[1].text = mean_w
        table.rows[i+1].cells[2].text = sd
        table.rows[i+1].cells[3].text = n
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ============================================================
# NAR REFERENCE LIST
# Format: Author,A.B., Author,C.D. and Author,E.F. (Year) Title. *Journal.*, **Vol**, Pages.
# Authors up to 20, then et al.
# Journal italic, volume bold.
# ============================================================

_refs_nar = [
    'Regev,A., Teichmann,S.A., Lander,E.S., Amit,I., Benoist,C., Birney,E., Bodenmiller,B., Campbell,P., Carninci,P., Clatworthy,M. et al. (2017) The Human Cell Atlas. *eLife*, **6**, e27041.',
    'Korsunsky,I., Millard,N., Fan,J., Slowikowski,K., Zhang,F., Wei,K., Baglaenko,Y., Brenner,M., Loh,P.R. and Raychaudhuri,S. (2019) Fast, sensitive and accurate integration of single-cell data with Harmony. *Nat. Methods*, **16**, 1289\u20131296.',
    'Lopez,R., Regier,J., Cole,M.B., Jordan,M.I. and Yosef,N. (2018) Deep generative modeling for single-cell transcriptomics. *Nat. Methods*, **15**, 1053\u20131058.',
    'Rosen,Y., Brbic,M., Roohani,Y., Swanson,K., Li,Z. and Leskovec,J. (2024) Toward universal cell embeddings: integrating single-cell RNA-seq datasets across species with SATURN. *Nat. Methods*, **21**, 1492\u20131500.',
    'Tran,H.T.N., Ang,K.S., Chevrier,M., Zhang,X., Lee,N.Y.S., Goh,M. and Chen,J. (2020) A benchmark of batch-effect correction methods for single-cell RNA sequencing data. *Genome Biol.*, **21**, 12.',
    'Nei,M. and Gojobori,T. (1986) Simple methods for estimating the numbers of synonymous and nonsynonymous nucleotide substitutions. *Mol. Biol. Evol.*, **3**, 418\u2013426.',
    'Tabula Muris Consortium (2018) Single-cell transcriptomics of 20 mouse organs creates a Tabula Muris. *Nature*, **562**, 367\u2013372.',
    'Tabula Sapiens Consortium (2022) The Tabula Sapiens: a multiple-organ, single-cell transcriptomic atlas of humans. *Science*, **376**, eabl4896.',
    'Cancer Genome Atlas Research Network (2014) Comprehensive molecular profiling of lung adenocarcinoma. *Nature*, **511**, 543\u2013550.',
    'Cancer Genome Atlas Network (2012) Comprehensive molecular portraits of human breast tumours. *Nature*, **490**, 61\u201370.',
    'Siletti,K., Hodge,R., Mossi Albiach,A., Lee,K.W., Ding,S.L., Hu,L., L\u00f6nnerberg,P., Bakken,T., Casper,T., Clark,M. et al. (2023) Transcriptomic diversity of cell types across the adult human brain. *Science*, **382**, eadl7046.',
    'Luecken,M.D. and Theis,F.J. (2019) Current best practices in single-cell RNA-seq analysis: a tutorial. *Mol. Syst. Biol.*, **15**, e8746.',
    'Hounkpel,B., Chen,J., Gosline,S.J.C., Domeniconi,C. and Jiang,D. (2021) HRT Atlas v1.0 database: redefining human and mouse housekeeping genes and candidate reference transcripts by mining massive RNA-seq datasets. *Nucleic Acids Res.*, **49**, D947\u2013D955.',
    'Wolf,F.A., Angerer,P. and Theis,F.J. (2018) SCANPY: large-scale single-cell gene expression data analysis. *Genome Biol.*, **19**, 15.',
    'Hao,Y., Hao,S., Andersen-Nissen,E., Mauck,W.M. 3rd, Zheng,S., Butler,A., Lee,M.J., Wilk,A.J., Darby,C., Zager,M. et al. (2021) Integrated analysis of multimodal single-cell data. *Cell*, **184**, 3573\u20133587.',
    'Hao,Y., Stuart,T., Kowalski,M.H., Choudhary,S., Hoffman,P., Hartman,A., Srivastava,A., Molla,G., Madad,S., Fernandez-Granda,C. et al. (2024) Dictionary learning for integrative, multimodal and scalable single-cell analysis. *Nat. Biotechnol.*, **42**, 293\u2013304.',
    'Weinstein,J.N., Collisson,E.A., Mills,G.B., Shaw,K.R., Ozenberger,B.A., Ellrott,K., Shmulevich,I., Sander,C. and Stuart,J.M. (2013) The Cancer Genome Atlas Pan-Cancer analysis project. *Nat. Genet.*, **45**, 1113\u20131120.',
    'Colaprico,A., Silva,T.C., Olsen,C., Garofano,L., Cava,C., Garolini,D., Sabedot,T.S., Malta,T.M., Pagnotta,S.M., Castiglioni,I. et al. (2016) TCGAbiolinks: an R/Bioconductor package for integrative analysis of TCGA data. *Nucleic Acids Res.*, **44**, e71.',
    'Cerami,E., Gao,J., Dogrusoz,U., Gross,B.E., Sumer,S.O., Aksoy,B.A., Jacobsen,A., Byrne,C.J., Heuer,M.L., Larsson,E. et al. (2012) The cBio cancer genomics portal: an open platform for exploring multidimensional cancer genomics data. *Cancer Discov.*, **2**, 401\u2013404.',
    'Perou,C.M., S\u00f8rlie,T., Eisen,M.B., van de Rijn,M., Jeffrey,S.S., Rees,C.A., Pollack,J.R., Ross,D.T., Johnsen,H., Akslen,L.A. et al. (2000) Molecular portraits of human breast tumours. *Nature*, **406**, 747\u2013752.',
    'Parker,J.S., Mullins,M., Cheang,M.C.U., Leung,S., Voduc,D., Vickery,T., Davies,S., Fauron,C., He,X., Hu,Z. et al. (2009) Supervised risk predictor of breast cancer based on intrinsic subtypes. *J. Clin. Oncol.*, **27**, 1160\u20131167.',
    'Edmondson,H.A. and Steiner,P.E. (1954) Primary carcinoma of the liver: a study of 100 cases among 48,900 necropsies. *Cancer*, **7**, 462\u2013503.',
    'Storey,J.D. and Tibshirani,R. (2003) Statistical significance for genomewide studies. *Proc. Natl. Acad. Sci. USA*, **100**, 9440\u20139445.',
    'W\u00e4lchli,T., Ghobrial,M., Schwab,M.E., Takada,S., Zhong,H., Suntharalingham,S., Vetiska,S., Gonzalez,D.R., Wu,R., Rehrauer,H. et al. (2024) Single-cell atlas of the human brain vasculature across development, adulthood and disease. *Nature*, **632**, 603\u2013613.',
    'Tsai,H.H., Niu,J., Munji,R., Davalos,D., Chang,J., Zhang,H., Tien,A.C., Kuo,C.J., Chan,J.R., Daneman,R. et al. (2016) Oligodendrocyte precursors migrate along vasculature in the developing nervous system. *Science*, **351**, 379\u2013384.',
    'Akay,L.A., Effenberger,A.H. and Tsai,L.H. (2022) Astrocyte endfoot formation controls the termination of oligodendrocyte precursor cell perivascular migration. *Neuron*, **111**, 190\u2013201.e8.',
    'Foerster,S., Floriddia,E.M., Neumann,B., Agirre,E., Castelo-Branco,G. and Franklin,R.J.M. (2024) Developmental origin of oligodendrocytes determines their function in the adult brain. *Nat. Neurosci.*, **27**, 1155\u20131165.',
    'Reeber,S.L., Arancillo,M. and Sillitoe,R.V. (2015) Bergmann glia are patterned into topographic molecular zones in the cerebellum. *Cerebellum*, **14**, 392\u2013403.',
    'Endo,F., Kasai,A., Cui,W., Tanaka,K.F. and Hashimoto,H. (2024) Astrocyte allocation during brain development is controlled by Tcf4-mediated fate restriction. *EMBO J.*, **43**, 4423\u20134447.',
    'Yang,L., Zhao,Z., Li,Y., Wang,J., Chen,X. and Liu,Z. (2024) Single-cell multi-omics analysis of lineage development and spatial organization in the human fetal cerebellum. *Cell Discov.*, **10**, 25.',
    'Shemer,A. and Jung,S. (2024) The molecular determinants of microglial developmental colonization. *Nat. Rev. Neurosci.*, **25**, 414\u2013427.',
    'Menassa,D.A., Muntslag,T.A.O., Martin-Esteban\u00e9,M., Barry-Carroll,L., Chapman,M.A., Adorjan,I., Tyler,T., Turnbull,B., Rose-Zerilli,M.J.J., Nicoll,J.A.R. et al. (2022) The spatiotemporal dynamics of microglia across the human lifespan. *Dev. Cell*, **57**, 2127\u20132139.e6.',
    'Barry-Carroll,L., Greulich,P., Marshall,A.R., Riecken,K., Fehse,B., Askew,K.E., Li,K., Garaschuk,O., Menassa,D.A. and Gomez-Nicola,D. (2023) Microglia colonize the developing brain by clonal expansion of highly proliferative progenitors. *Cell Rep.*, **42**, 112425.',
    'Schaffenrath,J., Huang,S.F., Wyss,T., Delorenzi,M. and Keller,A. (2024) Characteristics of blood-brain barrier heterogeneity between brain regions. *Nat. Neurosci.*, **27**, 1851\u20131865.',
    'Jones,H.E., Coelho-Santos,V., Bonney,S.K., Abrams,S.R., Shih,A.Y. and Siegenthaler,J.A. (2023) Meningeal origins and dynamics of perivascular fibroblast development. *Development*, **150**, dev201805.',
    'Tarashansky,A.J., Musser,J.M., Khariton,M., Li,P., Arendt,D., Quake,S.R. and Wang,B. (2021) Mapping single-cell atlases throughout Metazoa unravels cell type evolution. *eLife*, **10**, e66747.',
    'Jiang,J., Li,J., Huang,Y., Wang,Y., Chen,L. and Zhang,X. (2024) CACIMAR: cross-species analysis of cell identities, markers, regulations, and interactions. *Brief. Bioinform.*, **25**, bbae283.',
    'CZI Cell Science Program (2025) CZ CELLxGENE Discover: a single-cell data platform for scalable exploration, analysis and modeling of aggregated data. *Nucleic Acids Res.*, **53**, D886\u2013D900.',
    'Liberzon,A., Birger,C., Thorvaldsd\u00f3ttir,H., Ghandi,M., Mesirov,J.P. and Tamayo,P. (2015) The Molecular Signatures Database Hallmark Gene Set Collection. *Cell Syst.*, **1**, 417\u2013425.',
    'Yang,Z. (2007) PAML 4: phylogenetic analysis by maximum likelihood. *Mol. Biol. Evol.*, **24**, 1586\u20131591.',
    'Tan,Y.L., Yuan,Y. and Tian,L. (2020) Microglial regional heterogeneity and its role in the brain. *Mol. Psychiatry*, **25**, 351\u2013367.',
    'Lin,J. (1991) Divergence measures based on the Shannon entropy. *IEEE Trans. Inf. Theory*, **37**, 145–151.',
    'Efron,B. and Tibshirani,R.J. (1994) An Introduction to the Bootstrap. Chapman and Hall/CRC, New York.',
    'Bakken,T.E., Jorstad,N.L., Hu,Q., Lake,B.B., Tian,W., Kalmbach,B.E., Crow,M., Hodge,R.D., Krienen,F.M., Sorensen,S.A. et al. (2021) Comparative cellular analysis of motor cortex in human, marmoset and mouse. *Nature*, **598**, 111–119.',
    'Pedregosa,F., Varoquaux,G., Gramfort,A., Michel,V., Thirion,B., Grisel,O., Blondel,M., Prettenhofer,P., Weiss,R., Dubourg,V. et al. (2011) Scikit-learn: machine learning in Python. *J. Mach. Learn. Res.*, **12**, 2825–2830.',
    'Waskom,M.L. (2021) seaborn: statistical data visualization. *J. Open Source Softw.*, **6**, 3021.',
    'Butte,A.J., Dzau,V.J. and Glueck,S.B. (2001) Further defining housekeeping, or maintenance, genes: focus on A compendium of gene expression in normal human tissues. *Physiol. Genomics*, **7**, 95-96.',
]

def ref_p_nar(text):
    """Add a NAR-formatted reference paragraph."""
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(3)
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        set_black(run)
    return para

# ============================================================
# TITLE PAGE
# ============================================================
t = doc.add_heading('CKI: A Cell-state Kinetic Index for Quantifying Baseline-Normalized Transcriptomic Remodeling', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t.runs:
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = Pt(1)

# Running title (NAR requirement: ≤50 characters)
rt = doc.add_paragraph()
rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = rt.add_run('Running title: CKI: Baseline-Normalized Divergence Index')
run.font.name = 'Arial'
run.font.size = Pt(22)
set_black(run)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Xianming Wu (first author)
run = sub.add_run('Xianming Wu')
run.font.name = 'Arial'
run.font.size = Pt(21)
set_black(run)
for ch in ['1']:
    r = sub.add_run(ch)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    set_black(r)
    set_superscript(r)

sub.add_run(', ')

# Li Zhang (corresponding author)
run = sub.add_run('Li Zhang')
run.font.name = 'Arial'
run.font.size = Pt(21)
set_black(run)
for ch in ['1', '2', '*']:
    r = sub.add_run(ch)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    set_black(r)
    set_superscript(r)

# Affiliations
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = auth.add_run('1')
run1.font.name = 'Arial'
run1.font.size = Pt(22)
set_black(run1)
set_superscript(run1)
run2 = auth.add_run('Chinese Institute for Brain Research, Beijing, China')
run2.font.name = 'Arial'
run2.font.size = Pt(22)
set_black(run2)

auth2 = doc.add_paragraph()
auth2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = auth2.add_run('2')
run3.font.name = 'Arial'
run3.font.size = Pt(22)
set_black(run3)
set_superscript(run3)
run4 = auth2.add_run('Institute of Blood Transfusion, Chinese Academy of Medical Sciences & Peking Union Medical College, Chengdu, China')
run4.font.name = 'Arial'
run4.font.size = Pt(22)
set_black(run4)

# Correspondence
cor = doc.add_paragraph()
cor.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cor.add_run('* To whom correspondence should be addressed. Email: knightz@pumc.edu.cn')
run.font.name = 'Arial'
set_black(run)
run.font.size = Pt(22)

# ORCID
orcid = doc.add_paragraph()
orcid.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = orcid.add_run('ORCID: 0000-0002-0698-0754')
run.font.name = 'Arial'
set_black(run)
run.font.size = Pt(11)

# ============================================================
# GRAPHICAL ABSTRACT (placeholder — required by NAR)
# ============================================================
heading('Graphical Abstract', level=1)
p('[See separate Graphical Abstract file (CKI_graphical_abstract.pdf).]')

# ============================================================
# TEXT ABSTRACT (unstructured, ≤200 words — NAR format)
# ============================================================
heading('Abstract', level=1)
p(
    'Comparing cell populations is fundamental in single-cell genomics, yet standard distance '
    'metrics conflate baseline variation with functional adaptation. Inspired by '
    'the Ka/Ks ratio in molecular evolution, we present CKI (Cell-state Kinetic Index), '
    'which decomposes transcriptomic divergence into a baseline divergence rate '
    '(k_n, from housekeeping gene expression) and a functional divergence rate '
    '(k_f, from identity gene expression), with \u03c9 = k_f/k_n '
    'quantifying baseline-normalized functional divergence. We validated CKI across four '
    'datasets: the Tabula Muris mouse atlas, Tabula Sapiens human atlas, TCGA pan-cancer data, '
    'and a human brain single-nucleus atlas from millions of cells. Calibration confirmed '
    f'an empirical baseline for equivalent populations '
    f'(mean \u03c9 = 6.67, 95% bootstrap CI [4.12, 9.33], all P > 0.05), inflated above theoretical unity by highly variable gene selection. '
    f'CKI \u03c9 was negatively correlated with all standard distance metrics '
    f'(Spearman r = \u2212{abs(_sc["max"]):.2f} to \u2212{abs(_sc["min"]):.2f}, all P < 0.001), '
    'indicating it captures an independent information dimension. Cancer analysis revealed '
    'transcriptional convergence across genetically diverse tumors. Brain analysis '
    'identified 30 cell-type developmental signatures spanning four biological mechanisms (two statistically significant) among 31,764 comparisons. '
    'CKI provides a principled framework for quantifying functional divergence '
    'against an internal baseline, available as an open-source Python package.'
)

# ============================================================
# KEYWORDS (NAR requirement: 4-6, placed after Abstract)
# ============================================================
kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = kw.add_run('Keywords: cell-state divergence, housekeeping genes, Jensen-Shannon decomposition, transcriptomic remodeling, single-cell genomics')
run.font.name = 'Arial'
set_black(run)
run.font.size = Pt(22)

# ============================================================
# INTRODUCTION (NAR: "Introduction" instead of "Background")
# ============================================================
heading('Introduction', level=1)

p('Single-cell transcriptomics has transformed how we study cells. Comparing two cell populations is one of the most common tasks: we want to know how different they are. Researchers typically choose a standard metric: Euclidean distance, cosine similarity, Pearson or Spearman correlation, or Jensen-Shannon divergence. These metrics are convenient, but they have a fundamental limitation: they treat all gene expression differences equally.')

p('This matters because not all expression changes have the same biological meaning. A twofold change in GAPDH expression might reflect technical noise; a twofold change in a transcription factor might reflect a functional shift in cell state. Standard metrics cannot tell these apart. This problem is especially acute in large single-cell atlases (1), where donor-level and batch-level variation often dominates over cell-type identity. Methods such as Harmony (2), scVI (3), and SATURN (4) have been developed specifically to remove such nuisance variation (5). However, after correction, a key question remains: how much of the remaining difference between two populations represents functional adaptation, and how much is simply neutral drift?')

p('This question mirrors a problem addressed in molecular evolution. When comparing two DNA sequences, the Ka/Ks ratio (also called dN/dS) distinguishes nonsynonymous changes (Ka, which alter the protein) from synonymous changes (Ks, which are largely silent) (6). Ka/Ks uses synonymous sites as an internal baseline: the same mutational process produces both types of change, allowing the ratio to reveal selection. While CKI does not share Ka/Ks\u2019s formal mathematical properties (notably the shared mutation rate that cancels in the ratio; see Discussion), it adopts an analogous heuristic logic for transcriptomic comparisons.')

p('We applied this heuristic logic to transcriptomic comparisons. CKI defines two rates: a baseline divergence rate k_n, estimated from housekeeping (HK) gene expression; and a functional divergence rate k_f, estimated from cell-type identity genes. The ratio \u03c9 = k_f/k_n quantifies baseline-normalized functional divergence: \u03c9 near 1 means the observed differences are consistent with baseline expectation; \u03c9 much greater than 1 means functional divergence exceeds baseline variation; \u03c9 much less than 1 means strong functional constraint. Importantly, CKI \u03c9 is a heuristic index, not a formal measure of Darwinian selection; we use \u03c9 < 1, \u03c9 \u2248 1, and \u03c9 > 1 as convenient operational thresholds rather than claims about selection regimes (see Discussion).')

p('We validated CKI across four scales, each testing a different aspect of the method. First, we calibrated CKI on Tabula Muris mouse data (7), confirming that random splits of the same cell population yield \u03c9 above 1 (empirical calibration baseline 6.67, 95% bootstrap CI [4.12, 9.33]). Second, we extended CKI to Tabula Sapiens human data (8) and found that CKI \u03c9 is negatively correlated with all four standard distance metrics, demonstrating that it captures an independent information dimension. Third, we applied CKI to TCGA cancer data (9,10), revealing that tumors are paradoxically more transcriptionally homogeneous than normal tissues. Fourth, we used CKI to analyze a human brain single-nucleus atlas (11), measuring how the same cell types differ across brain regions and demonstrating that CKI can detect persistent developmental signatures\u2014including developmental origin heterogeneity and compartmentalized developmental specification (both statistically significant), as well as colonization route boundaries and a postnatal migration event (exploratory, not statistically significant)\u2014from adult transcriptomic data.')

# ============================================================
# MATERIALS AND METHODS (NAR: placed BEFORE Results)
# ============================================================
heading('Materials and Methods', level=1)

heading('CKI computation', level=2)
p(f'We normalize raw count matrices to 10,000 counts per cell and apply log1p transformation (12). Pseudobulk vectors are computed by averaging expression across cells sharing the same cell-type annotation, requiring at least 10 cells per group. Housekeeping (HK) genes are loaded from the HRT Atlas v1.0 reference (13) ({_ds["hrt_atlas_n_hk"]:,} human-mouse conserved HK genes). For mouse datasets, the mouse ortholog column is used; for human datasets (Tabula Sapiens, TCGA, brain atlas), the human gene column is used. The CKI package also supports data-driven auto-detection via detect_housekeeping_genes() (combined criterion: detection rate > 0.9 and CV < 30th percentile, use_reference = False), but all reported analyses use the pre-specified HRT Atlas reference.')

p('For populations A and B with pseudobulk vectors \u03b5_A and \u03b5_B, each vector is normalized to a probability distribution before JS divergence computation via softmax normalization (p_i = exp(x_i) / \u03a3 exp(x_j)). Then k_n = JS(norm(\u03b5_A[H]), norm(\u03b5_B[H])), where H is the set of HK gene indices. k_f = JS(norm(\u03b5_A[I]), norm(\u03b5_B[I])), where I is the set of top-2,000 highly variable genes (HVGs; Seurat flavor) excluding HK genes. \u03c9 = k_f/k_n. JS divergence (42) uses the base-2 logarithm (range [0, 1]).')

heading('Dimensionality invariance of JS divergence', level=2)
p(f'Because k_n is computed on ~1,130 HK genes and k_f on 200\u20132,000 HVG genes, we verified that JS divergence is not systematically biased by gene set dimensionality. A simulation of {int(_phaseC_dim.get("n_trials", 2000)):,} random Dirichlet distribution pairs across dimensions ranging from 50 to 5,000 showed that mean JS divergence is effectively constant (0.155\u20130.159, ratio = {_phaseC_dim.get("ratio_2000_to_1130", 1.001):.3f} between d = 1,130 and d = 2,000; Supplementary Fig. S10). This confirms that the systematic inflation of k_f relative to k_n (\u03c9 = 6.67 for equivalent populations) arises from HVG selection bias (selecting genes with high cross-cell variance) rather than from dimensional mismatch. The calibrated omega (\u03c9_cal = \u03c9 / 6.67) absorbs this bias into the empirical baseline, and the permutation null\u2014constructed using the same gene sets\u2014ensures internal consistency.')

heading('Bootstrap permutation test', level=2)
p('We randomly permute cell labels between the two populations (B = 1,000 for all datasets: mouse pilot of 15 cell-type pairs and 6 calibration controls, human Tabula Sapiens, TCGA, and brain atlas), recompute pseudobulk vectors, and calculate \u03c9_null for each permutation. Empirical P-values are computed as one-sided permutation tests: P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1), where the +1 pseudocount avoids P = 0. The one-sided test is appropriate because our hypothesis is directional: we test whether observed \u03c9 exceeds the null expectation (equivalent populations), not whether it differs in either direction. Standardized effect size = (\u03c9_obs \u2212 mean(\u03c9_null)) / sd(\u03c9_null), reported as a descriptive measure of separation magnitude. Given the non-normal \u03c9 distribution, SES is interpreted as a non-parametric descriptive statistic complementing the permutation P-value, not as a parametric test statistic such as Cohen\u2019s d. Benjamini-Hochberg false discovery rate (FDR) correction is applied within each dataset to account for multiple comparisons. For the larger-scale analyses (human, TCGA, brain), results are also supplemented with non-parametric statistical tests (Spearman correlation, Mann-Whitney U, Kruskal-Wallis, Jonckheere-Terpstra) and descriptive statistics (median, IQR, effect sizes).')

heading('Datasets', level=2)
p(f'Tabula Muris FACS SmartSeq2 (7): {_ds["tabula_muris_cells"]:,} cells, {_ds["tabula_muris_genes"]:,} genes, {_ds["tabula_muris_organs"]} organs (liver, kidney, spleen, lung, heart, bone marrow). Post-QC: {_ds["tabula_muris_ct_entries"]} cell-type entries (each with at least 10 cells). Highly variable genes selected using scanpy (14) with flavor="seurat" (15,16) and n_top_genes={_ds["n_hvg"]:,}.')

p(f'Tabula Sapiens v1.0 (8): accessed via CZ CELLxGENE Discover. Post-QC: {_ds["tabula_sapiens_cells"]:,} cells ({_ds["tabula_sapiens_organs"]} h5ad files total), {_ds["tabula_sapiens_genes"]:,} genes, {_ds["tabula_sapiens_ct_entries"]} cell-type entries across 6 organs. HK genes: HRT Atlas v1.0 reference ({_ds["hrt_atlas_n_hk"]:,} genes; human column).')

p(f'TCGA bulk RNA-seq (17): five cancer types from NCI Genomic Data Commons, accessed via TCGAbiolinks (18) and cBioPortal (19) APIs. LUAD: 495 tumor + 76 normal; LUSC: 567 tumor + 58 normal; LIHC: 365 tumor + 57 normal; KIRC: 755 tumor + 82 normal; BRCA: 1032 tumor + 109 normal. TPM values from UCSC Xena, log2(TPM + 1) transformed. PAM50 classification (20,21): nearest centroid (Pearson correlation), 44 of 47 PAM50 genes matched. LIHC Edmondson grade (22): from cBioPortal, 288 tumors. LUAD mutations: from cBioPortal, 492 samples (61 EGFR, 120 KRAS, 311 WT).')

p('Human brain atlas (11): Siletti et al. (2023) single-nucleus RNA-seq (v3.11) from CZ CELLxGENE Discover (collection ID: 283d65eb-2f53-46e9-a951-0da342e3d1f2). We used the Nonneurons.h5ad dataset (888,263 nuclei, 59,480 genes, 108 brain regions). Cell types were classified by supercluster_term annotation, generating 10 major non-neuronal classes: astrocytes (155,025 nuclei), oligodendrocytes (490,246), oligodendrocyte precursors (110,454 total including committed), microglia (91,838), vascular cells (10,932), fibroblasts (10,156), ependymal cells (7,882), choroid plexus (9,689), and Bergmann glia. We required >= 20 nuclei per (region, cell_type) group and >= 50 nuclei per region. Pseudobulk vectors were computed as the mean of raw counts per group, then normalized using Scanpy normalize_total (target_sum = 10,000) followed by log1p transformation at the pseudobulk level. CKI \u03c9 was computed for all same-cell-type cross-region comparisons (31,764 pairs total), using the hybrid scheme described above. HK genes were loaded from the HRT Atlas v1.0 reference (1,115 genes matched to the Siletti gene annotation). Top-200 identity genes were selected per comparison, excluding HK genes.')

heading('Method comparison', level=2)
p(f'We computed five metrics on all {_h["n_pairs_total"]:,} Tabula Sapiens cell-type pairs: CKI \u03c9 (hybrid scheme), raw JS divergence (all genes), Spearman distance (1 - \u03c1), cosine distance (1 - cos \u03b8), and marker Jaccard distance (1 - Jaccard index of top-200 expressed genes). Inter-metric Spearman correlations and cell-type classification ROC-AUC were computed using scikit-learn.')

heading('Multiplicative residual model for brain regional analysis', level=2)
p('For the brain regional analysis, we designed a multiplicative model to detect cell-type/region-pair combinations with anomalously low \u03c9. For each (cell_type, region_pair) combination, expected_\u03c9 = \u03bc_ct \u00d7 \u03bc_pair / \u03bc_grand, where \u03bc_ct is the cell type\'s global mean \u03c9, \u03bc_pair is the region pair\'s mean \u03c9, and \u03bc_grand is the global mean (8.01). The multiplicative residual = observed / expected. A residual substantially below 1 indicates the cell type is far less differentiated between those two regions than expected from both its own global plasticity and the region pair\'s overall divergence. We defined three confidence tiers: Strong (residual < 0.3, \u03c9 < 15, lowest \u03c9 in the region pair), Moderate (residual < 0.5, \u03c9 < 25), and Weak (residual < 0.75, \u03c9 < 35). To provide empirical null calibration, we implemented a permutation-based null distribution: cell type labels were shuffled within each region pair (B = 10,000 permutations), and per-signal empirical P-values were computed as P = (count(null_residual \u2264 observed) + 1)/(B + 1). Given the large number of tests (m = 31,764) and the finite permutation resolution (B = 10,000), 36.3% of signals (11,541/31,764) reached the empirical P-value floor (P = 9.99 × 10⁻⁵), precluding meaningful Benjamini-Hochberg FDR correction. We therefore report unadjusted permutation P-values and interpret significance descriptively: signals reaching the P-value floor (no null permutation exceeded the observed residual in 10,000 shuffles) are considered strong evidence of deviation from the multiplicative null model, while signals with P ≥ 0.50 show little to no evidence of departure. Per-signal tests are not independent (the same cell type or region pair appears in multiple comparisons); we restrict biological interpretation to the 30 predefined Strong candidates rather than the full 31,764 search space, and treat the permutation results as descriptive validation rather than formal FDR-controlled inference. Strong candidate signals were systematically cross-validated against the developmental neuroscience literature to assign each signal to one of four biological mechanisms: developmental origin heterogeneity (DO), embryonic colonization route boundaries (CR), compartmentalized developmental specification (DS), or postnatal cell migration (PM).')



heading('Clinical severity analysis', level=2)
p('For TCGA clinical severity analysis, we computed intratumoral \u03c9 for samples within each clinical stratum using the hybrid scheme (global k_n from shared HK genes, per-pair k_f from top-200 DE genes). BRCA PAM50 subtypes were classified using nearest centroid (Pearson correlation) against the published PAM50 centroids (20,21); 44 of 47 PAM50 genes matched to TCGA gene symbols. LIHC Edmondson histological grades (22) were obtained from cBioPortal (n = 288 tumors with both grade and expression data). LUAD mutation status (EGFR, KRAS, WT) was retrieved from cBioPortal (n = 492 samples). Between-stratum differences were tested with Kruskal-Wallis for PAM50 subtypes and LUAD mutations; trend across Edmondson grades was tested with Jonckheere-Terpstra. Paired vs. unpaired tumor-normal comparisons are reported as descriptive statistics (medians and IQRs) without formal P-values, as the paired design (tumor and matched adjacent normal from the same patient) does not meet the independence assumption of standard between-group tests.')

heading('Computational environment', level=2)
p('Typical runtime for a single cell-type pair analysis is under 5 minutes on a standard laptop (Apple M2, 16 GB RAM); the full brain analysis (31,764 pairs) required approximately 72 core-hours on a 32-core workstation. All analyses were performed in Python 3.13.12 with scanpy >= 1.9.0 (14), scipy >= 1.10.0, numpy >= 1.23.0, pandas >= 1.5.0, matplotlib >= 3.6.0, seaborn (46) >= 0.12.0, and scikit-learn (45) >= 1.2.0. All random seeds were fixed at 42 for reproducibility. Bootstrap permutation results are stable with respect to the random seed choice: with B = 1,000 permutations, the Monte Carlo standard error of the empirical P-value is approximately 0.016 at P = 0.5 and 0.001 at P = 0.001, meaning seed variation has negligible impact on statistical conclusions.')

heading('Statistical reporting', level=2)
p('We report summary statistics as mean \u00b1 s.d. (range) or median [IQR] as noted. Box plots display median, IQR, and 1.5\u00d7 IQR whiskers. Bootstrap permutation testing (43) was performed for all four datasets with B = 1,000: mouse pilot (15 cell-type pairs, 6 calibration controls), human Tabula Sapiens, TCGA, and brain atlas. Empirical P-values are computed as one-sided permutation tests: P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1), where the null distribution is built from B = 1,000 random permutations of cell labels. In each iteration, cell labels are randomly reassigned to pooled single-cell expression vectors, pseudobulks are recomputed from the permuted groups, and \u03c9 is recalculated under the same gene selection procedure. This tests whether the observed \u03c9 is larger than expected under random group assignment. Benjamini-Hochberg FDR correction is applied within each dataset to control the false discovery rate (23). The number of tests is determined by the number of cell types (10 for brain, 17 for human per-cell-type, 15 for mouse), not the number of region pairs. With B = 1,000, the minimum resolvable P-value is 9.99 \u00d7 10\u207b\u2074, which is well below the Benjamini-Hochberg threshold for the most significant test in each dataset (brain: 5.0 \u00d7 10\u207b\u00b3; human: 2.9 \u00d7 10\u207b\u00b3; mouse: 3.3 \u00d7 10\u207b\u00b3), confirming that B = 1,000 provides sufficient resolution for all tests.')

p('Bootstrap 95% confidence intervals for \u03c9 point estimates were computed by resampling observed pair-level \u03c9 values with replacement (B = 10,000) and reporting the 2.5th and 97.5th percentiles. These intervals quantify the precision of the \u03c9 point estimate under resampling rather than providing a classical confidence interval for a fixed population parameter. The \u03c9 distribution was characterized using skewness, excess kurtosis, and normality tests (Shapiro-Wilk for n \u2264 5,000; D\u2019Agostino-Pearson for n > 5,000). All \u03c9 distributions were right-skewed (brain: skewness = 2.22; mouse: 0.98; human: 0.73) and significantly non-normal (all P < 10\u207b\u00b9\u2075). Effect sizes are reported as standardized effect size (SES = (\u03c9_obs \u2212 \u03bc_null) / \u03c3_null), computed from the bootstrap permutation null distribution as a descriptive measure of separation; given the non-normal \u03c9 distribution, SES should be interpreted as a non-parametric descriptive statistic complementing the permutation P-value, not as a parametric test statistic such as Cohen\u2019s d. The calibration experiment using split-half equivalent populations (n = 6 random splits of the same tissue) yielded a mean \u03c9 = 6.67 (95% bootstrap CI [4.12, 9.33], B = 10,000 resamples of the 6 control \u03c9 values), reflecting systematic inflation of k_f relative to k_n from HVG gene selection. This empirical baseline does not represent a shift in the null hypothesis for the permutation test, but indicates that \u03c9 = 1 is a theoretical ideal (k_f = k_n) that is never reached in practice due to the way functional genes are identified. The permutation test accounts for this by comparing observed \u03c9 against a null distribution constructed from the same gene selection procedure applied to randomly permuted groups. The +1 pseudocount ensures that P > 0 even when no null permutation exceeds the observed deviation. For the larger-scale analyses (human, TCGA, brain), non-parametric statistical tests (Spearman correlation, Mann-Whitney U, Kruskal-Wallis, Jonckheere-Terpstra) and descriptive statistics (median, IQR, effect sizes) are also reported. Correlation coefficients (Spearman \u03c1) are reported with associated P-values. Omnibus tests (Kruskal-Wallis, Jonckheere-Terpstra) use P < 0.05 without additional correction.')

# ============================================================
# RESULTS (NAR: placed AFTER Methods)
# ============================================================
heading('Results', level=1)

# --- Result 1 ---
heading('Decomposing transcriptomic variation', level=2)

p('CKI takes two cell populations as input, each represented as a pseudobulk expression vector (the mean expression across all cells in that population). The computation has three steps, all of which use the same metric (Jensen-Shannon divergence) on the same underlying expression matrix, ensuring the ratio is internally calibrated (Fig. 1).')

p('Step 1: Compute the baseline divergence rate k_n. We restrict the pseudobulk vectors to housekeeping (HK) gene indices and apply softmax normalization, which converts expression values to probabilities. k_n is the JS divergence between these two HK gene probability distributions. Because HK genes should not differ systematically between biologically equivalent populations (13), k_n captures baseline noise: technical variation, stochastic transcriptional bursting, and individual-level physiological differences.')

p('Step 2: Compute the functional divergence rate k_f. We restrict the pseudobulk vectors to identity gene indices\u2014genes that define cell-type-specific functions. In the default configuration (used for the Tabula Muris full pairwise matrix, Fig. 2), identity genes are the top-2,000 highly variable genes (HVGs; Seurat flavor), excluding HK genes to maintain k_n/k_f independence. For the mouse pilot calibration and all cross-species analyses (human, TCGA, brain), we used a hybrid scheme: k_n is computed globally with a shared HK gene set, while k_f uses the top-200 differentially expressed genes (ranked by absolute mean difference) for each specific pair, with HK genes excluded. k_f is the JS divergence between these two identity gene probability distributions.')

p(f'Step 3: \u03c9 = k_f/k_n. For statistical inference, we perform bootstrap permutation testing (B = 1,000 for all datasets: mouse, human, TCGA, and brain). Cell labels are randomly shuffled and \u03c9 recalculated to generate a null distribution. The empirical P-value is P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1), and effect size is reported as standardized effect size (SES = (\u03c9_obs \u2212 \u03bc_null) / \u03c3_null) as a non-parametric descriptive statistic complementary to the permutation P-value. Benjamini-Hochberg FDR correction is applied within each dataset. For the larger-scale analyses (Tabula Sapiens, TCGA, brain atlas), non-parametric statistical tests and descriptive statistics (median, IQR, effect sizes) are also reported.')

p(f'We ran a parameter sweep on Tabula Muris mouse data ({_sb["n_pairs"]} cell-type pairs across {_ds["tabula_muris_organs"]} organs) to test whether adding pathway enrichment scores to k_f would improve performance. We found that the identity-only configuration (w_identity = 1.0, w_pathway = 0.0) achieved the best cell-type discrimination (AUC = {_sb["identity_auc"]:.3f}, Supplementary Fig. S1). CKI does not require external pathway databases to produce biologically meaningful results\u2014partitioning the expression data into constrained baseline and identity gene sets is sufficient.')

# --- Result 2 ---
heading('Calibration confirms baseline behavior at baseline', level=2)

p(f'We calibrated CKI on the Tabula Muris FACS dataset (7) (SmartSeq2, {_ds["tabula_muris_cells"]:,} cells, {_ds["tabula_muris_genes"]:,} genes, {_ds["tabula_muris_organs"]} organs). Housekeeping genes were loaded from the HRT Atlas v1.0 reference (13) (mouse ortholog column). For the full pairwise matrix (703 pairs, Fig. 2), identity genes were the top-{_ds["n_hvg"]:,} highly variable genes (HVGs; Seurat), excluding HK genes. For the pilot calibration (controls, S/D/X categories), we used a hybrid scheme: global k_n (shared HK gene set) with per-pair k_f (top-200 differentially expressed genes, ranked by absolute mean difference, HK excluded).')

p(f'The calibration confirmed correct baseline behavior. We performed six control comparisons in which we randomly split the same cell population into two halves. The mean \u03c9 was {_mc["control_mean"]:.2f} (median {_mc["control_median"]:.2f}, range {_mc["control_min"]:.2f}\u2013{_mc["control_max"]:.2f}, CV \u2248 60%), indicating substantial variability in the empirical baseline estimate. None of the six comparisons reached statistical significance (all P > 0.05, one-sided bootstrap test). This confirms that CKI recognizes biologically equivalent cell populations as having no significant functional divergence. We note that while these results are consistent with baseline behavior, formal equivalence testing (e.g., two one-sided tests, TOST) with a larger calibration sample would provide stronger statistical evidence for equivalence; the current sample size (n = 6) with a coefficient of variation of ~60% limits the power of such tests and the precision of the calibration factor (\u03c9_cal = \u03c9 / 6.67).')

p(f'Beyond controls, \u03c9 values increased monotonically with biological distance. Same cell type across different organs (S category: mean \u03c9 = {_mc["S_mean"]:.2f}, n = {_mc["S_n"]} pairs) had lower \u03c9 than different cell types within the same organ (D category: mean \u03c9 = {_mc["D_mean"]:.2f}, n = {_mc["D_n"]} pairs). The component-level analysis confirmed that k_f was the driver: k_f increased roughly 1,000-fold from controls to inter-cell-type comparisons, while k_n increased only about 100-fold. This establishes that CKI measures functional divergence, not just total difference.')

p(f'Because the empirical baseline (\u03c9 = 6.67) deviates substantially from the theoretical ideal (\u03c9 = 1), we introduce a calibrated omega: \u03c9_cal = \u03c9_obs / 6.67. This rescales all values so that equivalent populations yield \u03c9_cal \u2248 1.0, restoring the interpretability of the ratio framework. Under this calibration, mouse controls yield \u03c9_cal = 1.00, brain global mean becomes \u03c9_cal = {_phaseC_cal["brain"]["omega_cal_mean"]:.2f} (raw 8.01), and astrocytes\u2014the most regionally divergent brain cell type\u2014yield \u03c9_cal = {_phaseC_cal["brain"]["per_cell_type"].get("Astrocyte", {}).get("omega_cal_mean", 0):.2f} (raw 14.36). Cell types with \u03c9_cal < 1, such as Bergmann glia (\u03c9_cal = {_phaseC_cal["brain"]["per_cell_type"].get("Bergmann glia", {}).get("omega_cal_mean", 0):.2f}), are more transcriptionally constrained between brain regions than the empirical baseline. The calibrated values are reported alongside raw \u03c9 in all key results and Supplementary Figures 10\u201312. The calibrate_omega() function is available in the CKI package.')

# --- Result 3 ---
heading('CKI captures information that standard metrics miss', level=2)

p(f'We extended CKI to the Tabula Sapiens human atlas (8) ({_ds["tabula_sapiens_cells"]:,} cells; {_ds["tabula_sapiens_organs"]} h5ad files total, {_ds["tabula_sapiens_ct_entries"]} cell-type entries, {_ds["tabula_sapiens_organs"]} organs: liver, kidney, heart, bone marrow, spleen, lung). For human data, we used a hybrid scheme: k_n was computed once globally (using the full gene-by-cell-type pseudobulk matrix with the shared HK gene set), while k_f was computed per pair using the top-200 differentially expressed genes for that specific pair. This hybrid approach keeps k_n on a consistent scale (all cell types share the same HK gene set), while k_f adaptively selects the most informative identity genes for each pair. Critically, since \u03c9 = k_f/k_n is a ratio of JS divergences computed from the same underlying pseudobulk expression space, the normalization remains internally valid despite the different gene selection strategies. For the brain analysis, k_n was computed per-pair (not globally), because regional comparisons within the same cell type demand pair-specific baselines. Analysis of per-pair k_n across {_phaseC_kn["brain_overall"]["n_pairs"]:,} brain comparisons revealed substantial cross-pair variability (CV = {_phaseC_kn["brain_overall"]["kn_cv"]:.2%}), and the Spearman correlation between per-pair \u03c9 and global-k_n \u03c9 was only {_phaseC_kn["omega_correlation"]["spearman_rho"]:.3f} (P = {_phaseC_kn["omega_correlation"]["p_value"]:.2e}), confirming that pair-specific k_n is essential for accurate \u03c9 ranking (Supplementary Fig. S11). HK genes were loaded from the HRT Atlas v1.0 reference (13) (human column) (Fig. 3).')

p(f'Human \u03c9 values ranged from {_h["omega_min"]:.2f} to {_h["omega_max"]:.2f} (mean {_h["omega_mean"]:.2f}, median {_h["omega_median"]:.2f}, n = {_h["n_pairs_total"]:,} pairs), substantively lower than mouse (mean {_mc["X_mean"]:.2f}). This difference likely reflects fundamentally different data structures: mouse pilot data (calibration and validation) uses the same hybrid scheme as human (global k_n from shared HK genes, per-pair k_f from top-200 DE genes), while the mouse full pairwise matrix (703 pairs, Fig. 2 heatmap) uses a global HVG 2,000 set for k_f. The human hybrid scheme similarly uses global k_n with per-pair k_f. The different k_f gene selection strategies (per-pair DE vs. global HVG) alter the effective \u03c9 scale between analyses. Despite this, the biological hierarchy was preserved: same cell type across organs (mean \u03c9 = {_h["diff_organ_same_ct_mean"]:.2f}, n = {_h["diff_organ_same_ct_n"]} pairs) was lower than different cell types within the same organ (mean \u03c9 = {_h["same_organ_diff_ct_mean"]:.2f}, n = {_h["same_organ_diff_ct_n"]:,} pairs).')

p(f'The critical finding was that CKI captures a largely independent information dimension. We computed five metrics on all {_h["n_pairs_total"]:,} human cell-type pairs: CKI \u03c9, raw JS divergence (all genes), Spearman distance, cosine distance, and marker Jaccard distance. CKI \u03c9 was negatively correlated with all four standard metrics (Spearman r = {_sc["max"]:.2f} to {_sc["min"]:.2f}, all P < 0.001). In contrast, the four standard metrics formed a tight positive cluster (pairwise r = {_sc["std_pairwise_min"]:.2f}\u2013{_sc["std_pairwise_max"]:.2f}). This negative correlation is the strongest evidence that CKI measures something fundamentally different from all existing distance metrics.')

add_table_1(doc)
p(f'CKI showed moderate cell-type classification performance (AUC = {_au["cki_omega"]:.3f}, ranked 4th of 5 methods, above Spearman distance at AUC = {_au["spearman"]:.3f} but below cosine distance at AUC = {_au["cosine"]:.3f}; Table 1). This moderate ranking is expected by design: CKI down-weights shared HK gene patterns to isolate functional divergence, trading some ability to detect global transcriptional identity for enhanced sensitivity to functional specialization. The AUC rank therefore validates that CKI captures a distinct signal rather than being a better general-purpose classifier. Critically, CKI was the only metric where same-organ pairs had higher values than different-organ pairs (mean \u03c9 {_h["same_organ_diff_ct_mean"]:.2f} vs. {_h["diff_organ_diff_ct_mean"]:.2f}, Mann-Whitney U test, P < 0.001). All four standard metrics showed the opposite pattern (same-organ < different-organ). This reversal reflects CKI\'s sensitivity to functional specialization within shared microenvironments, a signal that standard metrics systematically obscure.')

# --- Result 4 ---
heading('Cancer analysis reveals unexpected transcriptional convergence', level=2)

p(f'We applied CKI to TCGA bulk RNA-seq data across five cancer types (LUAD, LUSC, LIHC, KIRC, BRCA) (9,10), totalling {_tc["n_total"]:,} samples. We asked a simple question: when cancer develops, how much functional transcriptional divergence occurs? CKI provides a principled answer by comparing tumor-tumor (TT), normal-normal (NN), and tumor-normal (TN) \u03c9 values (Fig. 4).')

p('The TCGA analysis is inherently exploratory because bulk RNA-seq cannot distinguish genuine transcriptional convergence of tumor cells from at least three alternative explanations: (i) shifts in tumor cell composition relative to normal tissue (tumor purity, stromal infiltration, immune cell infiltration); (ii) peritumoral inflammation and desmoplastic reactions that are shared across tumors; (iii) systematic RNA quality differences between tumor and normal specimens. With these caveats, a notable observation, at bulk RNA-seq resolution, was that tumors appeared more transcriptionally homogeneous than normal tissues. In all five cancer types, the median NN/TT \u03c9 ratio exceeded 1.0, meaning that normal individuals differ more from each other than tumors differ from each other. Breast cancer (BRCA) showed the smallest contrast (median NN/TT = 1.40), while liver cancer (LIHC) showed the largest (median NN/TT = 2.83), with intermediate values for lung adenocarcinoma (LUAD 1.60), lung squamous (LUSC 1.43), and kidney clear cell (KIRC 1.98). Single-cell or deconvolution-based analyses would be needed to disentangle genuine transcriptional convergence from cell-composition artifacts.')

p('Paired tumor-normal comparisons yielded higher \u03c9 than unpaired comparisons in four of five cancer types (paired/unpaired ratio = 0.99\u20133.25). However, with only n = 2\u20135 paired samples per cancer type, formal hypothesis testing lacks adequate statistical power (minimum two-sided P \u2248 0.33 for n = 2); we therefore report these comparisons as descriptive statistics only, without P-values. Definitive conclusions about within-patient versus between-patient variation require larger paired cohorts.')

p('We then asked whether \u03c9 tracks with clinical severity within cancer types. In liver cancer, \u03c9 decreased with increasing Edmondson grade (22): G1 (101.8 \u00b1 46.8, n = 39) > G2 (100.2 \u00b1 63.9, n = 133) > G3 (96.8 \u00b1 58.2, n = 105) > G4 (90.0 \u00b1 57.8, n = 11; Jonckheere-Terpstra trend test, P < 0.001). In breast cancer, PAM50 subtype analysis (20,21) revealed a gradient of transcriptional heterogeneity: Luminal A tumors had the highest intratumoral \u03c9 (344.5 \u00b1 323.4, n = 224), followed by Luminal B (313.6 \u00b1 282.7, n = 123), HER2-enriched (263.0 \u00b1 255.6, n = 55), and Basal-like tumors (223.4 \u00b1 183.7, n = 97), with Normal-like tumors having the lowest \u03c9 (108.0 \u00b1 65.5, n = 7; Kruskal-Wallis, P = 0.0002). Lung adenocarcinoma mutation stratification showed significant differences (Kruskal-Wallis, P = 0.017), with EGFR-mutant (285.3 \u00b1 180.1, n = 61) and KRAS-mutant tumors (284.6 \u00b1 227.9, n = 120) exhibiting higher \u03c9 than wild-type tumors (237.6 \u00b1 195.4, n = 311). We note that the smallest subgroups (Normal-like n = 7; Edmondson G4 n = 11) have limited statistical power, and their rankings should be interpreted cautiously. Furthermore, the PAM50 gradient may partly reflect proliferative fraction differences across subtypes (Basal-like tumors have the highest proliferation rates and the lowest \u03c9), meaning that the observed convergence could be driven by proliferation programs overriding tissue-specific expression rather than by a shared transcriptional attractor per se. The TCGA analysis is therefore exploratory in nature: bulk RNA-seq confounds cell-composition shifts, peritumoral inflammation, and RNA quality differences with genuine transcriptional divergence, and these alternative explanations cannot be ruled out at bulk resolution.')

# --- Result 5 ---
heading('CKI ranks cell types by cross-organ conservation', level=2)

add_table_2(doc)
p(f'Among the {_h["n_pairs_total"]:,} Tabula Sapiens cell-type pairs, {DATA["cross_organ_n_total"]} are same-cell-type cross-organ comparisons. These pairs allowed us to ask: which cell types maintain their transcriptional identity regardless of where they reside, and which are strongly shaped by their organ environment (Fig. 5; Table 2)?')

p(f'The cross-organ \u03c9 ranking reveals a broad spectrum of conservation across {len(t2)} cell types (Table 2). {t2[0][0]}s (mean \u03c9 = {t2[0][1]}, n = {t2[0][3]}) and {t2[1][0]}s (mean \u03c9 = {t2[1][1]}, n = {t2[1][3]}) were among the most conserved, followed by {t2[2][0]}s (mean {t2[2][1]} \u00b1 {t2[2][2]}, n = {t2[2][3]}) and {t2[3][0]}s (mean {t2[3][1]} \u00b1 {t2[3][2]}, n = {t2[3][3]}). {mac[0]}s, the most abundant cell type in cross-organ comparisons (n = {mac[3]}), showed intermediate conservation (mean {mac[1]} \u00b1 {mac[2]}). At the divergent end of the spectrum, {last2[0][0]}s (mean {last2[0][1]} \u00b1 {last2[0][2]}, n = {last2[0][3]}) and {last2[1][0]}s (mean {last2[1][1]} \u00b1 {last2[1][2]}, n = {last2[1][3]}) were the most organ-specific cell types. Endothelial cells are known to express organ-specific gene programs tailored to local vascular needs (24). We note that several cell types in this ranking have very few cross-organ pairs (n = 1\u20133; e.g., Memory B cells n = 1, Smooth muscle n = 1), and their mean \u03c9 estimates are correspondingly unreliable. We recommend interpreting the ranking of cell types with n < 5 as suggestive only, and prioritizing the relative ordering of well-sampled cell types (n \u2265 5) for biological conclusions.')

p('The cross-organ conservation ranking from CKI showed little agreement with rankings from standard metrics (Spearman r = -0.40 to +0.02, n = 59 pairs; bootstrap 95% CIs reported in Supplementary Table S2). This is because CKI explicitly normalizes: two cell populations might share similar highly expressed genes (yielding high Jaccard similarity), but if their constrained baseline k_n is low, even modest functional differences can produce a high \u03c9. This normalization reveals patterns that raw expression similarity misses.')

# --- Result 6 ---
heading('Brain regional analysis reveals cell-type differentiation gradients', level=2)

p('We applied CKI to the Siletti et al. human brain single-nucleus RNA-seq atlas (11), which profiles ~3.3 million nuclei across 108 brain regions. This dataset allowed us to ask: for a given cell type, how much functional divergence exists between the same cells residing in different brain regions? We focused on the 888,263 non-neuronal nuclei spanning 10 major cell classes (astrocytes, oligodendrocytes, oligodendrocyte precursors, microglia, vascular cells, fibroblasts, ependymal cells, choroid plexus, committed oligodendrocyte precursors, and Bergmann glia). Neurons were excluded because the supercluster_term annotation does not resolve the extensive neuronal subtype heterogeneity (glutamatergic, GABAergic, dopaminergic, etc., spanning dozens of transcriptionally distinct populations). Treating neurons as a single cell class would violate the same-cell-type assumption of our cross-region framework, as different brain regions contain fundamentally different neuron subtype compositions. We therefore restricted the analysis to non-neuronal cells, where the supercluster_term annotation provides well-defined, transcriptionally coherent cell classes suitable for same-cell-type cross-region comparison, and computed CKI \u03c9 for all same-cell-type cross-region comparisons (31,764 pairs total) (Fig. 6).')

p('The analysis revealed a pronounced differentiation gradient spanning 6.06-fold. Bergmann glia showed the lowest mean \u03c9 (2.37 \u00b1 1.14, n = 21 pairs across 7 regions), followed by committed oligodendrocyte precursor cells (3.17 \u00b1 1.47, n = 1,326 pairs across 52 regions), and fibroblasts (3.99 \u00b1 1.90, n = 3,403 pairs across 83 regions). Vascular cells (3.40 \u00b1 1.24, n = 3,321 pairs across 82 regions) and ependymal cells (4.13 \u00b1 1.73, n = 780 pairs across 40 regions) showed similarly low divergence. Microglia exhibited moderate divergence (mean \u03c9 = 8.02 \u00b1 4.93, n = 5,671 pairs across 107 regions). Oligodendrocytes and their precursors showed intermediate divergence (mean \u03c9 = 8.66 \u00b1 4.44 and 7.65 \u00b1 4.03, respectively). Astrocytes were the most regionally divergent cell type (mean \u03c9 = 14.36 \u00b1 8.68, n = 5,778 pairs across 108 regions), a 6.06-fold increase over Bergmann glia.')

p('This gradient aligns with known cell biology. Vascular cells and fibroblasts encounter relatively uniform extracellular environments across the brain\u2014the blood-brain barrier and meningeal structures impose similar constraints regardless of anatomical location. Their low \u03c9 values suggest a conserved core transcriptional program with limited regional adaptation. Microglia, the brain\'s resident immune cells, showed intermediate divergence: while microglial phenotypes vary regionally, their core surveillance and phagocytic machinery is shared. Oligodendrocytes must myelinate diverse axonal populations and adjust internode lengths regionally, explaining moderate \u03c9 values. Astrocytes showed the highest \u03c9, consistent with extensive literature showing that astrocytes express region-specific sets of ion channels, neurotransmitter transporters, and secreted factors tailored to local neuronal circuit demands. The grand mean \u03c9 across all 31,764 pairs (8.01) is lower than the astrocyte mean (14.36) because it is dominated by cell types with many pairs and low \u03c9 (e.g., fibroblasts n = 3,403 at 3.99, vascular cells n = 3,321 at 3.40); it is not a contradiction but a consequence of the skewed distribution of pair counts across cell types.')

p('The cross-region \u03c9 gradient provides a computational framework for inferring cell migration history. Cell types that recently migrated to new regions or that continuously exchange between regions should show low inter-regional \u03c9 because insufficient time has elapsed for transcriptional drift and local adaptation to accumulate. Conversely, cell types that have stably resided in specific regions for long periods should accumulate regional transcriptional signatures, yielding higher \u03c9. Under this framework, the low \u03c9 values for vascular cells and fibroblasts are consistent with continuous turnover and exchange through the circulatory and meningeal systems, while the high \u03c9 for astrocytes reflect long-term regional residence and functional specialization. This approach complements lineage tracing and developmental studies by providing an orthogonal transcriptomic readout of migration history.')

heading('CKI correlates with developmental origin signatures across brain regions', level=2)

p('Low CKI \u03c9 between a cell type across two brain regions indicates transcriptomic similarity beyond baseline expectation. We systematically cross-validated Strong candidate signals against the developmental neuroscience literature and identified four distinct biological mechanisms underlying low inter-regional \u03c9: (i) developmental origin heterogeneity\u2014cells from different embryonic progenitor pools (e.g., dorsal vs. ventral) retain distinct transcriptomic identities in adulthood; (ii) embryonic colonization route boundaries\u2014immune cells that entered the brain through different developmental entry points show residual transcriptomic discontinuities; (iii) compartmentalized developmental specification\u2014region-specific transcriptional programs during development produce persistent astrocyte and vascular identity differences; and (iv) postnatal cell migration\u2014cells that physically relocate between regions through active motility. These mechanisms are not mutually exclusive; a given cell-type/region-pair signal may involve overlapping processes (e.g., developmental origin and compartmentalized specification jointly shaping astrocyte thalamic signatures). Mechanism assignment uses the multiplicative residual model to flag anomalous cell-type/region-pair combinations, with the biological interpretation corroborated by the known developmental biology of each cell type.')

p('To formalize migration inference, we designed a multiplicative model: for each (cell_type, region_pair) combination, expected_\u03c9 = \u03bc_ct \u00d7 \u03bc_pair / \u03bc_grand, where \u03bc_ct is the cell type\'s global mean \u03c9, \u03bc_pair is the region pair\'s mean \u03c9, and \u03bc_grand is the global mean (8.01). The multiplicative residual = observed / expected: a residual substantially below 1 indicates that the cell type is far less differentiated between those two regions than expected from both its own global plasticity and the region pair\'s overall divergence\u2014a signature of shared transcriptional state potentially reflecting recent migration. We defined three confidence tiers: Strong (residual < 0.3, \u03c9 < 15, lowest \u03c9 in the region pair), Moderate (residual < 0.5, \u03c9 < 25), and Weak (residual < 0.75, \u03c9 < 35).')

p(f'Among {_br["total_pairs"]:,} cross-region comparisons, threshold criteria (residual < 0.3, {chr(969)} < 15) identified {_br["n_strong"]} ({_br["pct_strong"]:.2f}%) threshold-passing candidates: Astrocyte (6), oligodendrocyte (10), microglia (10), fibroblast (1), and vascular cells (3). Another {_br["n_moderate"]:,} pairs ({_br["pct_moderate"]:.2f}%) were Moderate candidates, and {_br["n_weak"]:,} ({_br["pct_weak"]:.2f}%) were Weak candidates. Permutation testing (B = 10,000) revealed that {_br["n_significant"]} of these {_br["n_strong"]} candidates reached the permutation P-value floor (P = 9.99 \u00d7 10\u207b\u2075): all 6 astrocyte and all 10 oligodendrocyte signals. We designate these {_br["n_significant"]} as Strong signals and describe them in detail below. The remaining {_br["n_non_significant"]} threshold-passing candidates (10 microglia, 1 fibroblast, 3 vascular) showed no evidence of deviation from the multiplicative null model (all P \u2265 0.76); their biological interpretation is deferred to the Discussion. Given the large number of tests (m = 31,764) and P-value floor saturation (36.3% of signals at P = 9.99 \u00d7 10\u207b\u2075), formal FDR correction is not applicable; the {_br["n_significant"]} floor-reaching signals are interpreted as descriptive evidence rather than FDR-controlled discoveries.')

heading('OPCs: internal consistency check', level=3)
p('Oligodendrocyte precursor cells (OPCs) are the most actively migrating cells in the adult CNS, continuously surveilling their environment along vascular scaffolds (25,26). Yet CKI detected 0 Strong signals among 5,671 OPC cross-region comparisons\u2014a finding that provides a useful internal consistency check of the multiplicative residual model. The model is not simply detecting high \u03c9 values or absolute transcriptional differences; it identifies cell-type/region-pair combinations where the observed functional divergence is substantially below what the cell type\'s global plasticity and the region pair\'s background divergence would jointly predict. OPCs have a high global mean \u03c9 (7.65) because their transcriptional program includes both progenitor and differentiation states; their 52 Moderate signals (residual < 0.5) likely reflect the balance between shared developmental origins and ongoing regional maturation (27). The complete absence of Strong signals despite OPCs being the brain\'s most motile cell type demonstrates that the residual model differentiates between broad baseline motility and specific transcriptional signatures of developmental history.')

heading('Oligodendrocytes: developmental origin rather than migration', level=3)
p('Mature oligodendrocytes contributed 10 Strong signals (residual 0.237\u20130.292), all of which reached the permutation P-value floor (P = 9.99 \u00d7 10\u207b\u2075, B = 10,000). The prevailing view is that adult oligodendrocytes do not migrate between brain regions. We systematically cross-validated all 10 Strong pairs against the developmental neurobiology literature. Notably, all 10 signals involved cortex/thalamus (A13/A14/A19/A32/A40/Idg vs. Pul/LP) or brainstem-internal (MoRF-MoEN vs. PnRF) pairings\u2014precisely the anatomical boundaries between dorsal- and ventral-derived oligodendrocyte populations. Foerster et al. (27) demonstrated through dorsal oligodendrocyte lineage ablation that >90% of adult cortical oligodendrocytes are dorsally derived (from cortical radial glia), while thalamic and brainstem oligodendrocytes are ventrally derived (MGE/LGE precursors). Ventral-derived cells fail to adopt cortical transcriptional programs even when transplanted into the cortex, indicating persistent cell-autonomous transcriptional identity. Our CKI analysis detects this developmental origin signature: dorsal vs. ventral oligodendrocyte populations are far less selectively diverged than expected from the oligodendrocyte global \u03c9 (8.66), because their transcriptional differences reflect shared generic myelination programs rather than region-specific functional specialization. LP (lateral posterior nucleus) and Pul (pulvinar), thalamic relay nuclei, contrast with cortical Brodmann areas A13-A40, forming the most consistent developmental boundary detected by our analysis. This reinterpretation\u2014that CKI Strong signals for oligodendrocytes detect persistent developmental origin signatures rather than migration\u2014is fully consistent with Foerster et al.\'s experimental data and provides, to the best of our knowledge, the first transcriptome-wide metric to detect the transcriptional boundary between dorsal and ventral oligodendrocyte populations without requiring lineage tracing.')

heading('Astrocytes: regional specialization with developmental origins', level=3)
p('Astrocytes showed the highest global \u03c9 (14.36, 95% CI [14.14, 14.58]) yet contributed 6 Strong signals, all of which reached the permutation P-value floor (P = 9.99 \u00d7 10\u207b\u2075, B = 10,000). These signals were concentrated in thalamic subnuclei (VLN-VPL, CM-VPL, Pul-VPL, LP-VPL-MN), hippocampal subfields (CA2-3 vs. DG-CA4), and cerebellar lobules (CBL vs. CBV). The thalamic signals are particularly informative: the ventroposterior lateral nucleus (VPL) appears in 4 of 6 Strong pairs, suggesting conserved astrocyte programs across thalamic relay nuclei that share a common developmental origin. Regionalized astrogenesis, driven by subnucleus-specific transcriptional programs, has been shown to produce persistent thalamic astrocyte heterogeneity that is detectable in adult tissue. Our finding that thalamic astrocyte pairs have \u03c9 values 5\u20136-fold below expectation indicates that these developmental signatures are selectively constrained\u2014astrocytes in functionally related thalamic nuclei retain transcriptional similarity beyond what would be predicted from astrocyte global plasticity alone. The cerebellar CBL vs. CBV signal (residual = 0.274) reflects the molecular topographic zones of Bergmann glia and cerebellar astrocytes described by Reeber et al. (28). Endo et al. (29) demonstrated that Tcf4 controls astrocyte allocation during cortical development; our results extend this principle to subcortical structures, showing that compartmentalized astrogenesis leaves persistent transcriptional signatures detectable by CKI across the entire brain.')

heading('Bergmann glia: cerebellar molecular topography', level=3)
p('Bergmann glia had the lowest global \u03c9 (2.37) and no Strong signals, consistent with their developmentally fixed, transcriptionally constrained state in the adult cerebellum (30). Bergmann glia are patterned into topographic molecular zones that align with cerebellar functional compartments (28), and their low global \u03c9 reflects their specialized role in maintaining Purkinje cell layer architecture with minimal regional transcriptional variation. The astrocyte CBL (cerebellar lobule) vs. CBV (cerebellar vermis) Strong signal reflects the established molecular topography difference across cerebellar compartments, rather than any migratory event.')

heading('Threshold-passing but non-significant signals', level=3)
p('Microglia (10), fibroblast (1), and vascular cells (3) produced threshold-passing Strong candidate signals (residual < 0.3, ω < 15), but none reached statistical significance in permutation testing (all P ≥ 0.76). These 14 signals are likely dominated by stochastic variation in the high-dimensional ω landscape and should not be interpreted as evidence of biological structure. We discuss the theoretical biological context for each cell type in the Discussion, while emphasizing that none of these signals constitute statistically supported findings.')


# ============================================================
# DISCUSSION (NAR: can merge Results and Discussion, but we keep separate)
# ============================================================
heading('Discussion', level=1)

p('CKI introduces a conceptual shift in transcriptomic comparison: from measuring absolute distance to quantifying functional divergence relative to an internal baseline. The key insight is that housekeeping (HK) genes are under stabilizing selection that constrains their expression variance across conditions, making them a practical constrained baseline against which functional divergence can be measured. This decomposition is heuristically inspired by Ka/Ks analysis, but CKI is a heuristic index rather than a formal measure of selection. Unlike Ka/Ks\u2014where a shared mutation rate cancels mathematically, leaving a pure selection signal\u2014CKI uses empirically defined HK genes as the baseline, lacking a comparable mechanistic cancellation. The empirical calibration (split-half equivalent populations, n = 6) yielded a mean \u03c9 = 6.67 (95% bootstrap CI [4.12, 9.33]), demonstrating that k_f systematically exceeds k_n even for identical populations due to HVG gene selection. We therefore use \u03c9_cal = \u03c9 / 6.67 as the operational scale, with \u03c9_cal \u226a 1 and \u03c9_cal \u226b 1 interpreted as functional constraint and enhanced divergence, respectively\u2014not as claims about Darwinian selection regimes.')

p('CKI is a divergence index, not a classifier\u2014and this is by design. Classifying cell types from transcriptomic data is largely a solved problem. CKI answers a complementary question: regardless of cell-type labels, how much functional divergence separates two populations, relative to their shared baseline? The negative correlation with all standard metrics indicates that CKI captures information that existing approaches miss.')

p('The Ka/Ks analogy is heuristically productive but has important technical limitations that distinguish CKI from formal molecular evolution models. First, Ka/Ks operates on DNA sequence alignments with explicit codon models, whereas CKI uses continuous expression vectors without a substitution model. HK genes are defined empirically (high detection rate, low CV) rather than mechanistically (synonymous sites in Ka/Ks), and their expression variance could reflect regulatory constraints rather than pure neutral drift; sensitivity analysis with alternative low-variance gene sets (r > 0.95) partially mitigates this concern. CKI also lacks a formal phylogenetic framework (e.g., Ornstein-Uhlenbeck models), and \u03c9 is computed on different gene sets per comparison, limiting absolute cross-comparison interpretability. Users should compare \u03c9 ranks rather than absolute values across datasets.')

p('CKI complements rather than replaces existing methods. SAMap (36) and SATURN (4) excel at cross-species alignment; CACIMAR (37) provides conservation scoring that could be reinterpreted through the CKI lens. However, we did not quantitatively benchmark CKI against these specialized methods, as they address different questions (cross-species alignment vs. within-species functional divergence). A systematic comparison on shared datasets would clarify the complementary strengths of each approach. More broadly, CKI provides a principled null model for any transcriptomic comparison: before concluding that two populations are meaningfully different, ask whether the difference exceeds baseline expectation.')

p('The TCGA finding that tumors are more transcriptionally homogeneous than normal tissues (median NN/TT > 1.0) is an exploratory observation that requires cautious interpretation. At bulk RNA-seq resolution, the apparent convergence could be driven by shared cell-composition shifts (tumor purity, stromal and immune infiltration), peritumoral inflammation, or systematic RNA quality differences, rather than by genuine transcriptional convergence of tumor cells. If the signal is real, it may point to common vulnerabilities across genetically diverse tumors; the PAM50 gradient (aggressive subtypes show lower \u03c9, consistent with proliferation programs overriding tissue-specific expression) is suggestive but cannot rule out the proliferation-confound explanation. Single-cell or deconvolution-based validation would be needed to disentangle these alternatives. Additionally, the substantially higher raw \u03c9 values in TCGA compared to single-cell datasets (e.g., BRCA Luminal A \u03c9 \u2248 344.5 vs. brain global \u03c9 = 8.01) likely reflect bulk RNA-seq averaging effects: pseudobulk averaging across millions of cells compresses HK gene variance, driving k_n toward its floor: in all 5 cancer types, the aggregate tumor-versus-normal k_n reached the floor value of 1 × 10\u207b\u2074, compared to mean k_n of 0.048\u20130.073 in single-cell datasets where single-cell resolution preserves intercellular HK gene variance. Meanwhile, tumor-specific DE gene differences are preserved (maintaining high k_f), resulting in inflated \u03c9. Cross-dataset \u03c9 comparisons should therefore be interpreted as rank-based rather than absolute.')

p('The cross-organ and cross-brain-region analyses establish CKI as a general tool for measuring functional differentiation at multiple spatial scales. The brain analysis revealed a 6.06-fold \u03c9 gradient across 10 cell classes (from 2.37 in Bergmann glia to 14.36 in astrocytes), demonstrating that CKI can detect regional functional specialization even among cells of the same type. The multiplicative residual model identified 30 threshold-passing candidates by predefined criteria; permutation testing (B = 10,000) confirmed that 16 of these reached the permutation P-value floor (P = 9.99 × 10⁻⁵), providing strong evidence of deviation from the multiplicative null model. The 16 Strong signals were exclusively from astrocytes (6/6) and oligodendrocytes (10/10), both reflecting developmental origin signatures: oligodendrocyte dorsal/ventral origin differences (27) explain all 10 oligodendrocyte Strong signals as cortex vs. thalamus/brainstem boundaries, while compartmentalized astrogenesis (29) explains the astrocyte thalamic and cerebellar signals. In contrast, the 14 remaining threshold-passing candidates\u201410 microglia, 1 fibroblast, and 3 vascular\u2014did not reach statistical significance (all P \u2265 0.76), suggesting they may reflect stochastic variation rather than robust biological patterns. Importantly, OPCs\u2014the most actively migrating cells in the adult CNS (25,26)\u2014yielded 0 Strong signals among 5,671 comparisons, providing a notable orthogonal validation that the residual model specifically detects fixed developmental signatures rather than ongoing cell motility. These results demonstrate that CKI detects persistent transcriptional signatures of developmental history\u2014origin and specification\u2014embedded in adult transcriptomic data, and that permutation-based validation is essential for distinguishing robust biological signals from threshold-passing noise. Preliminary cross-species validation using shared cell types between mouse and human atlases (44; Supplementary Fig. S2) was limited by the small number of directly comparable cell-type pairs, and \u03c9 rankings are moderately conserved between mouse and human for shared cell types, though absolute \u03c9 values differ due to different computation schemes (per-pair for mouse, hybrid for human).')

p('Limitations should be noted. First, CKI currently operates at the pseudobulk level; single-cell extensions would need to address sparsity and dropout. Second, the choice of HK gene set influences results; CKI uses the HRT Atlas v1.0 reference (13) as the HK gene source for all reported analyses. The CKI package also supports data-driven auto-detection (combined detection-rate and CV filtering, use_reference = False) as the universal default, but this was not used in the reported analyses. Sensitivity analysis showed that CKI results are robust to alternative HK definitions (using the lowest 10% variable genes as a constrained set yielded \u03c9 correlations r > 0.95). Third, the use of highly variable genes (HVGs) for k_f introduces systematic inflation: the calibration controls (random split of the same population) yielded a mean \u03c9 = 6.67, demonstrating that k_f systematically exceeds k_n even for identical populations. This inflation is a direct consequence of pre-selecting genes with high cross-cell variance. The permutation test accounts for this by comparing observed \u03c9 against a null distribution constructed under the same gene selection procedure with randomly permuted labels. A further concern is the circular dependency inherent in the per-pair k_f scheme: the top-200 differentially expressed genes are selected by |\u03bc_A \u2212 \u03bc_B|, meaning that the genes defining \u201cfunctional divergence\u201d are precisely those with the largest expression differences between the two groups under comparison. While the permutation test preserves this selection procedure under the null hypothesis, the circularity means that k_f magnitudes lack independent external validation and should be interpreted as an upper bound on functional divergence. Additionally, housekeeping gene expression may be dysregulated in cancer (47), potentially affecting the k_n baseline; we verified that HK genes detected by the combined criterion overlapped substantially with the reference set, but systematic cancer-specific HK dysregulation could inflate or deflate omega in a disease-specific manner. Fourth, TCGA analysis was limited to bulk RNA-seq resolution; the per-cancer bootstrap (B = 1,000) permutes sample labels (tumor vs. normal) to construct the null distribution. With 5 cancer types, TCGA results should be interpreted as exploratory. Fifth, the brain analysis uses post-mortem tissue; post-mortem interval (PMI) varies across samples and may differentially affect RNA integrity by brain region and cell type, potentially introducing systematic biases in regional omega comparisons. Neuron-specific RNA degradation signatures (e.g., activity-dependent transcripts with short half-lives) could be disproportionately affected. However, astrocyte and oligodendrocyte signals dominate the residual model results, and these glial cell types tend to be less sensitive to PMI-related degradation than neurons. Developmental time courses with controlled collection conditions would provide stronger evidence for migration inference and help disentangle PMI effects from genuine developmental signatures. Sixth, the brain regional analysis generated 31,764 cross-region comparisons. Benjamini-Hochberg FDR correction was applied within each dataset separately (bootstrap: m = 10 for brain, 17 for human, 15 for mouse; residual model: m = 31,764 for unadjusted permutation P-values). The 30 Strong candidates from the multiplicative residual model represent the most extreme deviations from the expected log-linear relationship between regional divergence and physical distance; they should be interpreted as hypothesis-generating signals requiring independent validation. The systematic concordance with known developmental biology provides post hoc biological plausibility, but formal FDR-controlled discovery would require larger cohorts and pre-registered hypotheses. Seventh, the \u03c9 distribution is right-skewed and non-normal across all datasets (brain skewness = 2.22, excess kurtosis = 7.28; Shapiro-Wilk and D\u2019Agostino-Pearson tests reject normality at P < 0.001). Consequently, SES is reported as a descriptive measure of effect size rather than a parametric test statistic. Eighth, permutation testing of the residual model (B = 10,000 per-signal empirical P-values, Benjamini-Hochberg FDR across all 31,764 pairs) revealed that 14 of 30 Strong threshold-passing signals\u2014comprising all 10 microglia, 1 fibroblast, and 3 vascular signals\u2014did not reach statistical significance (all P \u2265 0.76). These signals should be interpreted with caution as they may reflect stochastic variation rather than robust biological patterns. Ninth, bootstrap 95% confidence intervals (B = 10,000, pair-level resampling) were computed for all key \u03c9 estimates; confidence interval widths scale inversely with the number of contributing pairs, such that cell types with fewer regional comparisons (e.g., Bergmann glia, 21 pairs) yield wider intervals than well-sampled types (e.g., astrocytes, 5,778 pairs). Tenth, the empirical calibration baseline (\u03c9 = 6.67) is derived from only n = 6 split-half controls in a single species (mouse); while calibrated omega (\u03c9_cal = \u03c9 / 6.67) restores the interpretability of the ratio framework, the baseline may vary across tissues, species, and sequencing platforms, and a larger multi-dataset calibration study would strengthen the generality of this constant. Eleventh, k_n and k_f are computed on gene sets of different sizes (~1,130 HK genes vs. 200\u20132,000 HVG genes); while simulation confirmed that JS divergence between random distributions is dimensionally invariant (ratio = 1.001 between d = 1,130 and d = 2,000; Supplementary Fig. S10), the systematic inflation of k_f arises from HVG selection bias rather than dimensionality, and the calibrated omega absorbs this bias. Twelfth, the hybrid scheme used for human and TCGA analyses computes k_n globally (constant across all pairs), which reduces \u03c9 to a scaled k_f ranking; per-pair k_n analysis in the brain data (CV = 97.35%) showed that pair-specific k_n varies substantially and yields substantially different \u03c9 rankings (Spearman \u03c1 = \u22120.027 vs. global-kn \u03c9), confirming that the per-pair approach used for brain analysis is more informative but also more computationally expensive. Thirteenth, BH-FDR correction is applied separately within each dataset (mouse, human, TCGA, brain), which means that significance thresholds are not directly comparable across datasets; cross-dataset meta-analysis would require a unified multiple-testing framework. Fourteenth, CKI has not been validated on synthetic data with known ground-truth selection signals; while the calibration controls and OPC internal consistency check provide internal consistency checks, a simulation benchmark with injected functional divergence of known magnitude would provide stronger evidence for the method\u2019s sensitivity and specificity. Fifteenth, key parameters (softmax temperature, pseudocount \u03b5 = 1e-9, top-200 DE genes for k_f, HVG count of 2,000, log-base 2 for JS divergence) were chosen based on practical considerations rather than formal optimization; while sensitivity analyses showed robustness to parameter variation, a systematic parameter sweep across all parameters simultaneously was not performed.')

p('Sixteenth, the multiplicative residual model permutation analysis used B = 10,000 permutations with m = 31,764 tests. Due to P-value floor saturation (36.3% of signals at P = 9.99 × 10⁻⁵), formal Benjamini-Hochberg FDR correction is not applicable; the 16 Strong-tier signals reaching the P-value floor are interpreted as descriptive evidence of deviation from the multiplicative null model rather than FDR-controlled discoveries. Seventeenth, the empirical calibration factor (ω_cal = ω / 6.67, 95% bootstrap CI [4.12, 9.33]) was derived from mouse split-half controls using a global HVG set for k_f, but is applied to human, TCGA, and brain analyses that use per-pair DE gene selection for k_f. The two schemes select different gene sets: global HVGs capture genes with high cross-cell variance across the entire dataset, while per-pair DE genes capture the top-200 genes with the largest expression difference between the specific pair under comparison. The per-pair scheme produces a more targeted (and potentially larger) k_f than the global scheme, so the mouse-derived calibration factor likely underestimates the true baseline inflation in the per-pair setting. However, the bootstrap CI [4.12, 9.33] quantifies the uncertainty of the baseline estimate: applying the lower bound (ω / 4.12) would shift ω_cal upward by a factor of 1.62, while the upper bound (ω / 9.33) would shift it downward by a factor of 0.71. All CKI conclusions—which rely on rank-based interpretation (ω gradients across cell types, Spearman correlations with external metrics, and permutation-based significance) rather than absolute ω_cal thresholds—are robust to this range. Cross-scheme transferability should ideally be verified through dedicated split-half calibration experiments using the per-pair DE scheme in each dataset; we provide the necessary code in the CKI package (calibrate_omega function) for users to perform such validation on their own data.')

p('Future directions include developmental biology (quantifying functional differentiation between developmental stages), drug response profiling (measuring selectivity of drug-induced transcriptional changes), aging research (tracking age-related baseline vs. functional transcriptional drift), and evolutionary cell biology (quantifying conservation and divergence of cell-type programs across the tree of life). The CKI Python package (v0.3.1) and all analysis notebooks are available at https://github.com/zhanglknt/CKI-cell-type-identification under the MIT License.')

# ============================================================
# DATA AVAILABILITY (NAR: required separate section)
# ============================================================
heading('Data availability', level=1)
p('Tabula Muris data: GEO accession GSE109774. Tabula Sapiens data: CZ CELLxGENE Discover (https://cellxgene.cziscience.com/, accessed July 2025) (38). TCGA data: NCI Genomic Data Commons (https://portal.gdc.cancer.gov/). HRT Atlas (optional human/mouse HK reference): https://www.housekeeping.unicamp.br. Human brain atlas: CZ CELLxGENE Discover (collection ID as referenced in (11), accessed July 2025). PAM50 centroids: from Parker et al. (20). MSigDB Hallmark gene sets: from Liberzon et al. (39). The CKI source code (v0.3.1) is publicly available at https://github.com/zhanglknt/CKI-cell-type-identification (tag v0.3.1) under the MIT License. A permanent archival copy has been deposited at Zenodo (DOI: 10.5281/zenodo.15670808). The package requires Python ≥3.10 and runs on Linux, macOS, and Windows. A Dockerfile is provided in the repository for containerized reproducibility. All analysis notebooks and processed data matrices are included in the Supplementary Data.')

# ============================================================
# SUPPLEMENTARY DATA STATEMENT (NAR: required)
# ============================================================
heading('Supplementary Data', level=1)
p('Supplementary Data are available at NAR Online.')

# ============================================================
# ACKNOWLEDGEMENTS
# ============================================================
heading('Acknowledgements', level=1)
p('We thank the Tabula Muris Consortium, Tabula Sapiens Consortium, TCGA Research Network, and the Siletti et al. brain atlas team for making their data publicly available. We also thank the developers of scanpy, scipy, scikit-learn, and the broader open-source scientific Python ecosystem for the computational infrastructure that made this work possible. We are grateful to the HRT Atlas team for maintaining the housekeeping gene reference resource.')

# ============================================================
# AUTHOR CONTRIBUTIONS (NAR: required, CRediT recommended)
# ============================================================
heading('Author contributions', level=1)
p('X.W. and L.Z. conceived the study and designed the computational framework. L.Z. developed the CKI algorithm, performed all analyses, and prepared all figures. X.W. contributed to data curation, validation, and manuscript writing. Both authors read and approved the final manuscript.')

# ============================================================
# FUNDING
# ============================================================
heading('Funding', level=1)
p('This work was supported by the National Natural Science Foundation of China (NSFC) under grant number 32370682, and the Prevention and Control of Emerging and Major Infectious Diseases \u2014 National Science and Technology Major Project (grant number 2026ZD01910500).')

# ============================================================
# CONFLICT OF INTEREST
# ============================================================
heading('Conflict of interest', level=1)
p('The authors declare no competing interests.')

# ============================================================
# FIGURE LEGENDS
# ============================================================
heading('Figure legends', level=1)

p('Figure 1. The CKI framework. (A) Conceptual analogy between Ka/Ks in molecular evolution and CKI in transcriptomics. Ka/Ks uses synonymous substitution rate (Ks) as a neutral baseline; \u03c9 = Ka/Ks > 1 indicates positive selection. CKI uses housekeeping gene divergence for k_n and identity gene divergence for k_f; \u03c9 = k_f/k_n > 1 indicates baseline-normalized transcriptomic remodeling. (B) Computational pipeline: raw count matrix \u2192 pseudobulk \u2192 JS divergence on HK genes (k_n) and identity genes (k_f) \u2192 \u03c9 = k_f/k_n. (C) Bootstrap \u03c9 distribution (B = 1,000 permutations, mouse pilot) with median indicated by dashed line. (D) Scatter plot of k_n vs. k_f, showing that functional variation dominates constrained baseline. (E) \u03c9 distribution with \u03c9 = 1 (baseline) marked by dashed line.')

p('Figure 2. CKI calibration on Tabula Muris mouse data. (A) k_n calibration across six Tabula Muris cell types from control comparisons (C category: random split of same population). k_n values are stable across cell types, confirming constrained baseline behavior. (B) Component decomposition of k_n and k_f across four comparison categories: C (same cell type), S (same sub-organ), D (different cell type within same organ), and X (cross-organ). k_f increases monotonically with biological distance, while k_n remains relatively constrained. (C) Spearman correlation between CKI \u03c9 and four standard metrics (Cosine, Raw JS, Marker Jaccard, Spearman) on Tabula Muris data. All show negative correlation, confirming \u03c9 captures orthogonal information. (D) Pathway enrichment in the k_f component, showing fold change (k_f/k_n) for top enriched pathways. Stars indicate significance: *** P < 0.001, ** P < 0.01, * P < 0.05.')

p(f'Figure 3. CKI captures independent information. (A) Spearman correlation heatmap of five metrics on n = {_h["n_pairs_total"]:,} Tabula Sapiens pairs. CKI \u03c9 is negatively correlated with all four standard metrics. Standard metrics form a positive cluster. (B) Scatter plot of CKI \u03c9 vs. k_n (baseline rate) on Tabula Sapiens data (n = {_h["n_pairs_total"]:,} pairs), colored by same-organ vs. cross-organ pairs, with Spearman r. (C) ROC curves for cell-type classification across five metrics on Tabula Sapiens data. (D) \u03c9 by comparison category: C (same cell type), S (same sub-organ), D (different cell type, same organ), X (cross-organ). Box plots show log10(\u03c9) distribution. (E) AUC (cell-type classification) vs. interpretability (decomposability) comparison across five metrics. CKI \u03c9 has lower AUC than cosine and raw JS divergence but is the only fully decomposable metric into baseline (k_n) and functional (k_f) components.')

p('Figure 4. TCGA pan-cancer perturbation analysis. (A) Median NN/TT \u03c9 ratio across five cancer types. NN > TT in all cancers (all ratios > 1.0), indicating that normal individuals differ more from each other than tumors differ from each other. (B) \u03c9 distributions for normal-normal (NN) vs. tumor-tumor (TT) comparisons per cancer type. (C) Median TN/NN \u03c9 ratio (tumor-normal vs. normal-normal) across five cancer types. (D) Bootstrap standardized effect size (SES) for NN vs. TT comparisons per cancer type. (E) Tissue-level pairwise \u03c9 matrix heatmap across five cancer types.')

p(f'Figure 5. Cross-organ cell-type conservation. (A) CKI \u03c9 ranking of 17 cell types with cross-organ comparisons (n = 59 pairs) within the Tabula Sapiens human atlas. Immune cells (macrophage, T cells) rank lowest; structural cells (erythrocyte, endothelial) rank highest. (B) \u03c9 distribution across all cell-type pairs, showing conserved vs. variable cell types. (C) Cross-organ \u03c9 gradient (mean \u00b1 SD) showing systematic variation across organ pairs. (D) Table of top conservative cell-type pairs with the lowest cross-organ \u03c9 values.')

p('Figure 6. Brain regional cell-type differentiation and migration inference. (A) Brain region map schematic showing anatomical regions analyzed in the Siletti et al. atlas. (B) \u03c9 gradient across 10 non-neuronal cell classes. Bergmann glia (mean \u03c9 = 2.37) shows the lowest regional divergence; astrocytes (mean \u03c9 = 14.36) show the highest (6.06-fold gradient). Box plots: center line, median; box, IQR; whiskers, 1.5\u00d7 IQR. (C) Astrocyte \u03c9 across brain regions, with hierarchical clustering. Cortex regions cluster together, as do thalamic and brainstem regions. (D) Migration candidate detection: multiplicative residual model identifies 30 Strong candidates (residual < 0.3, \u03c9 < 15, lowest \u03c9 in pair) among 31,764 cross-region comparisons. (E) Observed vs. expected \u03c9 for migration candidates, showing the gap between observed regional similarity and the cell-type-specific expectation.')

# ============================================================
# Supplementary Figure legends (NAR convention)
# ============================================================
heading('Supplementary Figure legends', level=1)

p('Supplementary Figure S1. Parameter sweep and pathway analysis. (A) k_n stability as a function of housekeeping gene set size, showing convergence at ~200-300 HK genes. (B) k_f component contribution per pathway: decomposition of k_n vs. k_f for representative pathways. (C) Weight sweep for multi-component k_f. Identity-only (w_identity = 1.0, w_pathway = 0.0) achieves optimal cell-type discrimination (AUC = 0.847, n = 703 mouse cell-type pairs, 6 organs).')

p(f'Supplementary Figure S2. Cross-species validation details. (A) Cross-species \u03c9 conservation: scatter plot of human vs. mouse \u03c9 for shared cell types, with Spearman r and P-value. (B) HK gene set detection stability: overlap between human and mouse HK gene sets detected by the combined criterion. (C) \u03c9 distribution comparison between mouse (n = 15 shared cell types) and human (n = {_h["n_pairs_total"]:,} pairs).')

p('Supplementary Figure S3. TCGA per-cancer matrices. Pairwise \u03c9 matrices for five cancer types (BRCA, KIRC, LIHC, LUAD, LUSC) showing tissue-level transcriptomic divergence structure within each cancer cohort.')

p('Supplementary Figure S4. Method comparison performance. ROC-AUC bar plot for cell-type classification across five metrics.')

p('Supplementary Figure S5. Cross-organ conservation raw data. Complete table of 59 same-CT cross-organ pairs (Supplementary Table S2).')

p('Supplementary Figure S6. Brain regional analysis details. (A) Cell type nuclei counts per brain region, showing the distribution of non-neuronal nuclei across sampled regions. (B) k_n/k_f decomposition per cell class, showing that \u03c9 variance is predominantly driven by k_f. (C) \u03c9 vs. number of regions (n_regions) per cell type, with Spearman \u03c1 and empirical P-value reported; broader spatial distribution is associated with greater transcriptomic divergence. (D) Region-region \u03c9 matrix for astrocytes across brain regions, with hierarchical clustering. (E) Migration candidate counts by tier (Strong, Moderate, Weak) across all 31,764 cross-region pairs.')

p('Supplementary Figure S7. Developmental signature detection. (A) Multiplicative residual distribution for all 31,764 cross-region pairs, with Strong (residual < 0.3), Moderate (residual < 0.5), and Weak (residual < 0.75) tiers shaded. (B) Strong candidate counts by cell type. OPCs (0 Strong despite highest motility among the 10 non-neuronal classes) provide a key internal consistency check, supporting that the model detects developmental-origin signatures rather than general motility. (C) Top 10 Strong candidates ranked by multiplicative residual, annotated with cell type and region pair. (D) Migration candidates by cell type and confidence tier, showing tier distribution across all 10 non-neuronal cell classes.')

p('Supplementary Figure S8. \u03c9 distribution characterization. Histograms and Q-Q plots of \u03c9 distributions for brain, mouse, and human datasets, showing right-skewness and non-normality (Shapiro-Wilk and D\u2019Agostino-Pearson tests, all P < 0.001).')

p('Supplementary Figure S9. Permutation null distribution for residual model. Distribution of multiplicative residuals under label permutation (B = 10,000), compared against observed residuals for Strong candidates.')

p('Supplementary Figure S10. JS divergence dimensionality invariance. (A) Mean JS divergence between random Dirichlet distribution pairs as a function of dimensionality (50\u20135,000 genes, n = 2,000 trials per dimension), showing that JS divergence is effectively constant across dimensions (ratio = 1.001 between d = 1,130 and d = 2,000). (B) Dimensionality ratio relative to the HK gene set (d = 1,130), confirming that k_n and k_f are dimensionally comparable.')

p('Supplementary Figure S11. Per-pair k_n variability. (A) Per-pair k_n (mean \u00b1 SD) by cell type in the brain atlas, showing substantial cross-pair variability (CV = 97.35%). (B) Scatter plot of \u03c9 computed with per-pair k_n vs. global k_n (Spearman \u03c1 = \u22120.027), demonstrating that pair-specific k_n yields substantially different rankings.')

p('Supplementary Figure S12. Calibrated omega. (A) Raw vs. calibrated \u03c9 (\u03c9_cal = \u03c9 / 6.67) by brain cell type. (B) Calibrated \u03c9 distribution across all 31,764 brain pairs, with calibration baseline (\u03c9_cal = 1.0) and mean indicated.')

p('Statistical conventions. All P-values are from one-sided permutation tests (B = 1,000 for mouse/human/TCGA/brain; B = 10,000 for the multiplicative residual model) unless otherwise specified. Benjamini-Hochberg FDR correction was applied within each dataset. Non-parametric tests (Spearman correlation, Mann-Whitney U, Kruskal-Wallis, Jonckheere-Terpstra) are two-sided and reported with exact P-values; descriptive statistics are reported as mean \u00b1 SD or median [IQR] as indicated. Effect sizes include standardized effect size (SES = (\u03c9_obs \u2212 mean(\u03c9_null)) / sd(\u03c9_null)), computed from the bootstrap permutation null distribution. Bootstrap 95% confidence intervals use B = 10,000 resamples. All analyses use JS divergence with base-2 logarithm. Sample sizes (n) are reported for each comparison.')

# ============================================================
# REFERENCES (NAR: numbered, parentheses format)
# ============================================================
heading('References', level=1)

for ref in _refs_nar:
    ref_p_nar(ref)

# == Save ==
PROJECT_ROOT = Path(__file__).resolve().parent
out = str(PROJECT_ROOT / "results" / "CKI_NAR_Manuscript.docx")
doc.save(out)
print(f'Saved: {out}')
