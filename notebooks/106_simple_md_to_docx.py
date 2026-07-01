#!/usr/bin/env python3
"""
Generate CKI_Reproducibility_Guide.docx from markdown.
Simplified version - just gets content right.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGNMENT
import os

MD = r"C:\Users\KnightZ\Desktop\细胞受选择\results\CKI_Reproducibility_Guide.md"
OUT = r"C:\Users\KnightZ\Desktop\细胞受选择\results\CKI_Reproducibility_Guide.docx"

FONT = "Arial"
SIZE_BODY = 10
SIZE_H1 = 16
SIZE_H2 = 14
SIZE_H3 = 12

def build_simple():
    doc = Document()
    with open(MD, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        if not line.strip():
            i += 1
            continue
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            _style_run(p, SIZE_H1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            _style_run(p, SIZE_H2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            _style_run(p, SIZE_H3)
        elif line.startswith('|'):
            # Parse table
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                # Skip separator row
                if not all(set(c) <= set('-:| ') for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        cell = table.cell(r, c)
                        cell.text = cell_text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = FONT
                                run.font.size = Pt(SIZE_BODY)
                doc.add_paragraph()
            continue
        elif line.startswith('    ') or line.startswith('\t'):
            # Code block (indented)
            p = doc.add_paragraph()
            p.style.font.name = "Courier New"
            p.add_run(line.strip()).font.size = Pt(8)
        else:
            p = doc.add_paragraph()
            p.add_run(line).font.name = FONT
            p.add_run(line).font.size = Pt(SIZE_BODY)
        
        i += 1
    
    doc.save(OUT)
    print(f"Saved: {OUT}")

def _style_run(p, size):
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

if __name__ == '__main__':
    build_simple()
