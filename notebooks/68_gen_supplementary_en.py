"""
Generate English Supplementary Materials DOCX with continuous line numbers.
Replaces the Chinese supplementary with English translation in Chinese-researcher style.
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from _load_manuscript_data import get_manuscript_data
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from lxml import etree

# Load all manuscript data dynamically
DATA = get_manuscript_data()
_br = DATA['brain']

# Load migration candidates for S4 top-5
import pandas as pd
_mig = pd.read_csv(Path(__file__).resolve().parent.parent / "results" / "brain_siletti_migration_candidates_v3.csv")
_strong = _mig[_mig['tier'] == 'Strong'].sort_values('residual').head(5)
_candidates = _mig[_mig['residual'] < 0.75]

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Page margins (NAR: 2.5cm)
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

# Heading styles
for lvl, size in [(1, 16), (2, 14), (3, 12)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Arial'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


# ===== TITLE PAGE =====
add_heading('Supplementary Materials', 1)
add_para('CKI: A Cell-state Kinetic Index for Quantifying Baseline-Normalized Transcriptomic Remodeling')
add_para('Li Zhang')
add_para('')

add_heading('Table of Contents', 2)
toc = [
    'Supplementary Note 1: CKI Mathematical Derivation',
    'Supplementary Note 2: CKI Algorithm Pseudocode',
    'Supplementary Note 3: Statistical Testing Details',
    'Supplementary Note 4: Dataset Quality Control and Filtering Criteria',
    'Supplementary Table 1: Parameter Sweep Results (Phase 3.2)',
    'Supplementary Table 2: Cross-Organ Conservation Data (Phase 3.5)',
    'Supplementary Table 3: Human Brain Non-neuronal Cell Regional CKI Data',
    'Supplementary Table 4: Inter-regional Cell Migration Candidate Data',
    'Supplementary Data 1: Complete Analysis Script Index',
]
for item in toc:
    add_para(item)

doc.add_page_break()

# ===== SN1: Mathematical Derivation =====
add_heading('Supplementary Note 1: CKI Mathematical Derivation', 2)

add_para('1.1 Jensen-Shannon Divergence', bold=True)
add_para(
    'The Jensen-Shannon (JS) divergence is a symmetrized and smoothed version of the '
    'Kullback-Leibler divergence. For two probability vectors p and q: '
    'JS(p, q) = 1/2 D(p||m) + 1/2 D(q||m), where m = 1/2(p+q), '
    'and D(p||q) = \u03a3 p_i log2(p_i/q_i). When using the base-2 logarithm, '
    'the JS divergence is bounded in [0, 1]. This bound is important for interpreting '
    'omega: when both k_n and k_f approach 1, omega = k_f/k_n may still vary, '
    'and a small floor value (1e-4) is applied to k_n to prevent inflated omega from near-zero denominators. '
    'Before computing JS divergence, both pseudobulk vectors are normalized to probability distributions via softmax normalization '
    '(p_i = exp(x_i)/\u03a3exp(x_j)).'
)

add_para('1.2 Baseline Divergence Rate k_n', bold=True)
add_para(
    'Housekeeping (HK) genes are defined as genes that maintain stable expression '
    'across cell types and conditions. Let H = {g1, ..., gM} be the set of HK gene '
    'indices. Given pseudobulk vectors \u03bc_A and \u03bc_B (length G, total number of genes), '
    'the baseline divergence rate is: k_n = JS(norm(\u03bc_A[H]), norm(\u03bc_B[H])), '
    'where norm() denotes softmax normalization. '
    'Rationale: HK genes should not exhibit systematic differences between biologically '
    'identical cell populations. The JS divergence observed on HK genes therefore reflects '
    'baseline noise: technical variation, stochastic transcriptional bursting, and '
    'individual-level physiological differences. k_n thus provides an internal baseline, '
    'heuristically analogous to Ks (synonymous substitution rate) in molecular evolution. '
    'HK gene set selection: HK genes were loaded from the HRT Atlas v1.0 reference '
    '(1,130 human-mouse conserved HK genes) (4). For mouse datasets, the mouse ortholog '
    'column is used; for human datasets (Tabula Sapiens, TCGA, brain atlas), the human '
    'gene column is used. The CKI package also supports data-driven auto-detection via '
    'detect_housekeeping_genes() (combined criterion: detection rate > 0.9 and CV < 30th '
    'percentile, use_reference = False), but all reported analyses use the pre-specified '
    'HRT Atlas reference. Sensitivity analysis indicates that CKI results are robust to HK set '
    'selection: using the top 10% lowest-variance genes as an alternative constrained set '
    'yields omega correlations of r > 0.95.'
)

add_para('1.3 Functional Divergence Rate k_f', bold=True)
add_para(
    'Identity genes I are defined as genes that capture cell-type-specific functional '
    'programs. In the default configuration (w1 = 1.0, w2 = 0.0), I consists of the '
    'top-N highly variable genes (HVG), excluding HK genes. The functional divergence '
    'rate is: k_f = JS(norm(\u03bc_A[I]), norm(\u03bc_B[I])). Extended configurations '
    'can incorporate additional gene sets: (1) regulon activity genes \u2014 genes enriched '
    'for cell-type-specific transcription factor motifs; (2) pathway enrichment genes '
    '\u2014 genes from MSigDB pathways differentially active between the two groups; '
    '(3) macro-gene embeddings \u2014 gene-level embeddings from protein language models '
    '(e.g., ESM-2). These extensions use a weighted formulation: '
    'k_f = w1*JS(HVG) + w2*JS(pathway) + w3*JS(macro). Parameter sweep (Phase 3.2) '
    'showed that the pure identity gene configuration (w1=1.0, w2=w3=0.0) achieved '
    'optimal cell type discrimination (AUC = 0.847); this was therefore adopted as '
    'the default scheme.'
)

add_para('1.4 Omega Ratio and Its Interpretation', bold=True)
add_para(
    'omega = k_f/k_n. Interpretation follows a Ka/Ks analogy: '
    'omega ~ 1: the observed transcriptomic difference is consistent with baseline '
    'expectation, with no evidence of functional reprogramming; '
    'omega >> 1: functional divergence exceeds baseline drift, indicating evidence of '
    'functional transcriptional reprogramming beyond baseline; '
    'omega << 1: functional constraint, the two groups are more similar in functional '
    'genes than expected from baseline drift (rare in practice). '
    'The Ka/Ks analogy is structurally similar but mathematically non-equivalent. '
    'Key differences: (1) Ka/Ks operates on sequence alignments with explicit codon '
    'models, while CKI operates on continuous expression vectors; (2) the neutral '
    'reference in Ka/Ks has a mechanistic basis in the genetic code (synonymous changes '
    'are assumed neutral), whereas HK genes in CKI are empirically defined; (3) Ka/Ks '
    'uses explicit evolutionary models (e.g., PAML), while CKI uses empirical bootstrap '
    'inference.'
)

add_para('1.5 Bootstrap Permutation Test', bold=True)
add_para(
    'Statistical inference is performed by generating a null distribution of omega '
    'under the null hypothesis that the two cell populations are drawn from the same '
    'distribution. Procedure: (1) Annotate all cells in the pooled dataset with their '
    'original group labels (A or B); (2) Randomly permute labels B times (default '
    'B=1,000 for all datasets (mouse, human, TCGA, and '
    'brain atlas), recomputing pseudobulk vectors and omega_null each time; '
    '(3) Empirical P-value (one-sided): '
    'P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1), with the '
    '+1 term avoiding P = 0; (4) Effect size: SES = (omega_obs - '
    'mean(omega_null))/sd(omega_null). Benjamini-Hochberg FDR correction is '
    'applied within each dataset to control the false discovery rate. '
    'Test critical values at alpha=0.05 are derived from the permutation null '
    'distribution (2.5th and 97.5th percentiles). Note: these are permutation-based '
    'test critical values for rejecting H0, NOT confidence intervals for omega itself. '
    'Bootstrap permutation testing was performed for all four datasets: mouse pilot '
    '(15 cell-type pairs, B=1,000), human Tabula Sapiens (B=1,000), TCGA (B=1,000), and '
    'brain atlas (B=1,000). For the larger-scale analyses, results are supplemented with non-parametric '
    'statistical tests (Spearman correlation, Mann-Whitney U, Kruskal-Wallis, '
    'Jonckheere-Terpstra) and descriptive statistics.'
)

add_para('1.6 Pseudobulk Construction', bold=True)
add_para(
    'Raw count matrices X (cells x genes) are preprocessed as follows: '
    '(1) Library size normalization: X_norm = 10,000 * (X/colSums(X)); '
    '(2) log1p transformation: X_log = log1p(X_norm), stabilizing variance and reducing '
    'the influence of high-expression outliers; (3) Pseudobulk: mu = column-wise mean '
    'of X_log for all cells with the same cell type annotation, with a minimum of 10 '
    'cells per group. For TCGA bulk RNA-seq data, TPM normalization is used instead: '
    'TPM values from UCSC Xena, followed by log2(TPM + 1) transformation. No pseudobulk step '
    'is needed as each sample is already a bulk expression profile.'
)

doc.add_page_break()

# ===== SN2: Algorithm Pseudocode =====
add_heading('Supplementary Note 2: CKI Algorithm Pseudocode', 2)
add_para('Algorithm 1: CKI Core Computation', bold=True)
add_para(
    'Input: Two cell populations A and B (expression matrices), HK gene set H, '
    'identity gene set I (default top-N HVG excluding H). '
    'Output: omega, P-value, SES, null distribution.'
)
pseudo = [
    ' 1. X_A, X_B <- library-normalize and log1p-transform A and B',
    ' 2. mu_A <- mean(X_A, axis=0); mu_B <- mean(X_B, axis=0)  // pseudobulk',
    ' 3. mu_A_H <- mu_A[H]; mu_B_H <- mu_B[H]',
    ' 4. k_n <- JS_divergence(softmax(mu_A_H), softmax(mu_B_H))',
    ' 5. mu_A_I <- mu_A[I]; mu_B_I <- mu_B[I]',
    ' 6. k_f <- JS_divergence(softmax(mu_A_I), softmax(mu_B_I))',
    ' 7. if k_n < 1e-4: k_n <- 1e-4  // floor to prevent inflated omega',
    ' 8. omega <- k_f / k_n',
    ' 9. // Permutation test',
    '10. labels <- concatenate([A]*n_A, [B]*n_B)',
    '11. for b = 1 to B (B = 1,000 for all datasets):',
    '12.     labels_perm <- random_permutation(labels)',
    '13.     mu_perm1 <- mean(pooled[labels_perm[:n_A]], axis=0)',
    '14.     mu_perm2 <- mean(pooled[labels_perm[n_A:]], axis=0)',
    '15.     omega_null[b] <- CKI_core(mu_perm1, mu_perm2, H, I)',
    '16. // Inference (one-sided permutation test)',
    '17. P <- (count(omega_null >= omega_obs) + 1) / (B + 1)',
    '18. d <- (omega - mean(omega_null)) / sd(omega_null)',
]
for line in pseudo:
    p = add_para(line)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)

add_para('')
add_para('Algorithm 2: Pairwise Identity Gene Selection (Tabula Sapiens Extension)', bold=True)
add_para(
    'Unlike the Tabula Muris full pairwise matrix (global HVG set for Fig. 2), Tabula Sapiens and all pilot analyses (mouse calibration, human, TCGA, brain) employ pairwise identity '
    'gene selection to avoid dilution of HVG across 102 cell types.'
)
pseudo2 = [
    'Input: Pseudobulk vectors mu_A, mu_B; HK set H; top-N parameter N (default 200)',
    '1. Delta <- |mu_A - mu_B|  // per-gene absolute expression difference',
    '2. I <- indices of top-N genes ranked by descending Delta, excluding H',
    '3. k_f <- JS(softmax(mu_A[I]), softmax(mu_B[I]))',
    '',
    'Note: k_n uses the global HK set (same for all pairs); k_f uses pairwise top-N genes.',
]
for line in pseudo2:
    p = add_para(line)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ===== SN3: Statistical Testing =====
add_heading('Supplementary Note 3: Statistical Testing Details', 2)
add_para('3.1 Tests Performed and Correction Strategy', bold=True)
add_para(
    'The following statistical tests were used in this study: Mann-Whitney U test '
    '(two-sided) for independent comparisons between two groups; Kruskal-Wallis test '
    'for multi-group comparisons (e.g., BRCA PAM50 subtypes); Jonckheere-Terpstra trend '
    'test for ordered categorical variables (e.g., LIHC Edmondson grade); Spearman rank '
    'correlation for correlations between metrics; Bootstrap permutation test (B=1,000) '
    'for CKI omega significance inference; ROC-AUC for cell type classification '
    'performance assessment.'
)

add_para('3.2 Bootstrap Details', bold=True)
add_para(
    'Bootstrap iterations: B=1,000 for all datasets (mouse pilot study with '
    '15 cell-type pairs, human Tabula Sapiens via script 08b, TCGA via script 08a, '
    'and brain atlas via script 08c). Bootstrap permutation testing was '
    'performed for all four datasets. Benjamini-Hochberg FDR correction is '
    'applied within each dataset to control the false discovery rate. '
    'For the calibration '
    'experiment, empirical P-values '
    'are computed as: P = (count(\u03c9_null \u2265 \u03c9_obs) + 1)/(B + 1). '
    'Standardized effect size = (\u03c9_obs - mean(\u03c9_null)) / sd(\u03c9_null). '
    'In all CKI results, standardized effect sizes are typically > 1.0 for '
    'biologically meaningful comparisons, indicating large effects relative to '
    'the null distribution.'
)
add_para(
    'Bootstrap confidence intervals (95% CI) for all key \u03c9 estimates were '
    'computed by pair-level resampling with B=10,000 iterations. For each cell '
    'type, observed pair-level \u03c9 values were resampled with replacement and the '
    'median was computed; the 2.5th and 97.5th percentiles of the resulting '
    'distribution define the 95% CI. Confidence interval widths scale inversely '
    'with the number of contributing pairs: well-sampled cell types (e.g., '
    'astrocytes, 5,778 pairs) yield narrow intervals ([14.14, 14.58]), whereas '
    'cell types with fewer comparisons (e.g., Bergmann glia, 21 pairs) produce '
    'wider intervals ([1.95, 2.90]).'
)

add_para('3.3 Multiple Testing Correction', bold=True)
add_para(
    'Bootstrap permutation testing was performed for all four datasets with B=1,000: '
    'mouse pilot (15 cell-type pairs), human Tabula Sapiens, TCGA, and '
    'brain atlas. Benjamini-Hochberg FDR correction is applied within each '
    'dataset to control the false discovery rate. For the larger-scale analyses (Tabula Sapiens: '
    '5,151 pairs; brain atlas: 31,764 pairs; TCGA pan-cancer), bootstrap results '
    'are supplemented with non-parametric '
    'statistical tests and descriptive statistics (median, IQR, effect sizes). For TCGA stratified '
    'analyses (BRCA PAM50, LIHC Edmondson) involving 4-5 groups, omnibus tests '
    '(Kruskal-Wallis, Jonckheere-Terpstra) are used. Effect sizes are reported alongside '
    'all significance statements to distinguish statistical significance from biological '
    'magnitude. For the brain atlas analysis, 31,764 cross-region comparisons yielded 30 Strong candidates (residual < 0.3). Per-signal empirical P-values were computed via permutation testing (B = 10,000); however, 36.3% of signals (11,541/31,764) reached the empirical P-value floor (P = 9.99 \u00d7 10\u207b\u2075), precluding meaningful Benjamini-Hochberg FDR correction. We therefore report unadjusted permutation P-values and interpret significance descriptively: 16 of 30 Strong-tier candidates reached the P-value floor (no null permutation exceeded the observed residual in 10,000 shuffles), while 14 signals showed no evidence of departure (all P \u2265 0.76).'
)
add_para(
    'Permutation-based validation of the multiplicative residual model was '
    'performed using B = 10,000 permutations. For each of the 31,764 brain '
    'region pairs, cell type labels were randomly shuffled within region pairs '
    'to construct a null distribution of residuals. Per-signal empirical '
    'P-values were computed as P = (count(null_residual \u2264 observed_residual) + 1) '
    '/ (B + 1). Of the 31,764 signals, 11,541 (36.3%) reached the empirical '
    'P-value floor (P = 9.99 \u00d7 10\u207b\u2075) of the B = 10,000 permutation test, '
    'precluding meaningful Benjamini-Hochberg FDR correction. We therefore '
    'interpret the permutation results descriptively: among the 30 Strong-tier '
    'candidates, 16 signals (6 astrocytes, 10 oligodendrocytes) reached the '
    'P-value floor, indicating strong evidence of deviation from the '
    'multiplicative null model, while 14 signals (10 microglia, 1 fibroblast, '
    '3 vascular) showed no evidence of departure (P \u2265 0.76). Per-signal '
    'tests are not independent (the same cell type or region pair appears in '
    'multiple comparisons); we restrict biological interpretation to the 30 '
    'predefined Strong candidates. (Supplementary Figure 8: \u03c9 distribution '
    'characterization; Supplementary Figure 9: residual null distribution.)'
)

add_para('3.4 Reporting Conventions', bold=True)
add_para(
    'Summary statistics are reported as mean +/- standard deviation (range) or median '
    '[interquartile range]. Boxplots display: median (center line), IQR (box), '
    '1.5x IQR (whiskers), with data points beyond the whiskers shown as outliers. '
    'All P-values from non-bootstrap tests are two-sided; bootstrap permutation '
    'P-values are one-sided (see SN 1.5). Correlation '
    'coefficients (Spearman rho) are reported with P-values. Effect sizes (standardized effect size, SES = (\u03c9_obs \u2212 \u03bc_null) / \u03c3_null) '
    'are reported as descriptive measures of magnitude; because the \u03c9 distribution '
    'is right-skewed and non-normal (Shapiro-Wilk and D\u2019Agostino-Pearson tests reject '
    'normality at P < 0.001 for all datasets), SES should be interpreted as a '
    'non-parametric descriptive statistic rather than a parametric test result.'
)

add_para('3.5 Calibrated Omega Normalization', bold=True)
add_para(
    'The theoretical baseline omega = 1 (k_f = k_n) is never observed in practice because '
    'highly variable gene (HVG) selection systematically inflates k_f relative to k_n. '
    'Empirical calibration on split-half equivalent populations (mouse, n = 6) yielded a '
    'mean omega = 6.67. We introduce calibrated omega: omega_cal = omega_obs / 6.67, which '
    'rescales all values so that equivalent populations yield omega_cal ~ 1.0. Under this '
    'calibration: mouse controls yield omega_cal = 1.00, brain global mean becomes omega_cal '
    '= 1.20 (raw 8.01), and the most divergent brain cell type (astrocytes) yields omega_cal '
    '= 2.15 (raw 14.36). Cell types with omega_cal < 1 (e.g., Bergmann glia, omega_cal = 0.36) '
    'are more transcriptionally constrained between brain regions than the empirical baseline. '
    'The calibrate_omega() function is available in the CKI package (cki.calibrate_omega). '
    'Both raw and calibrated omega values are reported in all key results. (Supplementary Figure S12.)'
)

add_para('3.6 JS Divergence Dimensionality Invariance', bold=True)
add_para(
    'Because k_n is computed on ~1,130 housekeeping (HK) genes and k_f on 200-2,000 highly '
    'variable genes (HVGs), we verified that JS divergence is not systematically biased by '
    'gene set dimensionality. A simulation of 2,000 random Dirichlet distribution pairs '
    'across dimensions ranging from 50 to 5,000 showed that mean JS divergence is effectively '
    'constant across all dimensions tested (range: 0.155-0.159; ratio = 1.001 between d = 1,130 '
    'and d = 2,000). This confirms that the systematic inflation of k_f relative to k_n '
    '(omega = 6.67 for equivalent populations) arises from HVG selection bias (selecting genes '
    'with high cross-cell variance) rather than from dimensional mismatch between the HK and '
    'HVG gene sets. We note that this simulation addresses dimensionality per se (random '
    'probability vectors of different lengths) but does not simulate the variance-based '
    'gene selection mechanism that generates the \u03c9 inflation. A more complete validation '
    'would test whether the inflation magnitude scales with the stringency of variance '
    'filtering rather than gene count. The calibrated omega (omega_cal = omega / 6.67) '
    'absorbs this bias into the empirical baseline, and the permutation null distribution '
    '- constructed using the same gene sets as the observed data - ensures internal '
    'consistency. (Supplementary Figure S10.)'
)

add_para('3.7 Pair-Specific k_n Variability', bold=True)
add_para(
    'In the hybrid scheme used for human and TCGA analyses, k_n is computed once globally '
    '(constant across all pairs), which reduces omega to a scaled k_f ranking. For the brain '
    'analysis, k_n was computed per-pair to capture region-specific baseline variation. '
    'Analysis of per-pair k_n across 31,764 brain comparisons revealed substantial cross-pair '
    'variability: overall k_n CV = 97.35% (mean = 0.0141, median = 0.0086), with per-cell-type '
    'CVs ranging from 37.6% (committed oligodendrocyte precursors) to 81.4% (oligodendrocytes). '
    'The Spearman correlation between omega computed with per-pair k_n and omega computed with '
    'global k_n was only rho = -0.027 (P = 9.96 x 10^-7), confirming that pair-specific k_n '
    'yields substantially different omega rankings than the global-kn simplification. This '
    'justifies the per-pair k_n approach used for the brain analysis and highlights a limitation '
    'of the hybrid scheme used for human/TCGA data. (Supplementary Figure S11.)'
)

add_para('3.8 TCGA Exploratory Analysis Caveats', bold=True)
add_para(
    'The TCGA pan-cancer analysis (Results Section: Cancer analysis) is exploratory in nature '
    'due to several inherent limitations of bulk RNA-seq data. First, bulk RNA-seq confounds '
    'cell-composition shifts (tumor purity, stromal infiltration, immune cell infiltration) '
    'with genuine transcriptional divergence; the observed NN/TT > 1.0 pattern may partly '
    'reflect shared cell-composition changes across tumors rather than true transcriptional '
    'convergence. Second, peritumoral inflammation and desmoplastic reactions are shared '
    'across tumors and could contribute to apparent convergence. Third, systematic RNA quality '
    'differences between tumor and normal specimens may introduce technical bias. Fourth, '
    'the paired tumor-normal analysis (n = 2-5 per cancer type) lacks statistical power for '
    'formal hypothesis testing (minimum two-sided P approx 0.33 for n = 2); we therefore '
    'report paired comparisons as descriptive statistics only. Fifth, PAM50 and Edmondson '
    'stratification includes small subgroups (Normal-like n = 7; Edmondson G4 n = 11) whose '
    'rankings are unreliable. Sixth, the PAM50 gradient (aggressive subtypes have lower omega) '
    'may be driven by proliferative fraction differences rather than shared transcriptional '
    'states. These caveats do not invalidate the descriptive findings but preclude causal '
    'inference from bulk-level data alone.'
)

add_para('3.9 Cross-Organ Sample Size Considerations', bold=True)
add_para(
    'The cross-organ conservation ranking (Results Section: CKI ranks cell types by cross-organ '
    'conservation) includes cell types with varying numbers of cross-organ pairs. Several cell '
    'types have very few pairs (n = 1-3; e.g., Memory B cells n = 1, Smooth muscle n = 1), '
    'making their mean omega estimates unreliable. We recommend interpreting rankings of cell '
    'types with n < 5 as suggestive only. Bootstrap 95% confidence intervals for cell types '
    'with n >= 5 are provided in Supplementary Table S2. The Spearman correlations between '
    'CKI omega and standard metrics are reported with bootstrap 95% CIs (B = 10,000 resamples).'
)

add_para('3.10 One-Sided Permutation Test Justification', bold=True)
add_para(
    'All bootstrap P-values use a one-sided test: P = (count(omega_null >= omega_obs) + 1)/(B + 1). '
    'The one-sided formulation is appropriate because our hypothesis is directional: we test '
    'whether observed omega exceeds the null expectation (equivalent populations), not whether '
    'it differs in either direction. A two-sided test would be appropriate if we were testing '
    'for any departure from the null (either elevated or suppressed omega). However, the '
    'biological questions addressed here (functional divergence exceeding baseline, Strong '
    'migration candidates showing anomalously low omega) are inherently directional.'
)

add_para('3.11 Parameter Justification', bold=True)
add_para(
    'Key CKI parameters and their rationale: (1) Softmax normalization: converts expression '
    'vectors to probability distributions for JS divergence computation; the softmax function '
    'is a standard choice that preserves relative magnitude information while ensuring '
    'non-negativity and sum-to-one. (2) Pseudocount epsilon = 1e-9: added to avoid log(0) in '
    'JS divergence; the value is small enough to not affect results materially but large enough '
    'to prevent numerical instability. (3) Top-200 DE genes for k_f (mouse): selected per pair '
    'to adaptively capture the most informative identity genes; the parameter sweep (Phase 3.2) '
    'tested N_HVG in {500, 1000, 2000, 3000, 5000} and found 2000 optimal for AUC. (4) HVG '
    'count of 2,000 (human/brain): standard Scanpy default; Seurat flavor used for consistency '
    'with the Scanpy ecosystem. (5) Log-base 2 for JS divergence: standard choice giving JS '
    'range [0, 1]; the base does not affect omega = k_f/k_n since it cancels in the ratio. '
    '(6) B = 1,000 for bootstrap: justified by adaptive permutation analysis showing minimum '
    'P = 0.001 is well below BH thresholds for cell-type-level tests (Phase B).'
)

doc.add_page_break()

# ===== SN4: QC and Filtering =====
add_heading('Supplementary Note 4: Dataset Quality Control and Filtering Criteria', 2)

add_para('4.1 Tabula Muris FACS (Mouse)', bold=True)
add_para(
    'Downloaded from GEO (GSE109774). FACS-sorted cells (not droplet-based) were used '
    'to ensure high per-cell gene detection. QC filtering: cells with < 500 detected '
    'genes were removed; cells with > 10% mitochondrial gene expression were removed; '
    'genes detected in < 3 cells were removed. Result: 15,057 cells x 22,308 genes '
    '(post-QC). Cell type annotation: 32 cell type entries with >= 10 cells per group '
    'were retained for pseudobulk construction, spanning 6 organs (Liver, Kidney, '
    'Spleen, Lung, Heart, Bone Marrow).'
)

add_para('4.2 Tabula Sapiens (Human)', bold=True)
add_para(
    'Downloaded from CZ CELLxGENE Discover. QC filtering: cells with < 500 detected '
    'genes were removed; cells with > 20% mitochondrial gene expression were removed. '
    'Result: 108,136 cells retained (6 h5ad files total), with 51,852 genes (filtered from the '
    'original 58,870). Cell type entries: 102 entries across 6 organs (Liver, Kidney, '
    'Heart, Bone Marrow, Spleen, Lung). Cell types included in pairwise omega analysis '
    'were required to have >= 10 cells in at least one donor. Pairwise identity gene '
    'selection (top-200 genes by |Delta expression| ranking) ensures that each comparison '
    'uses the most informative genes for that specific pair. Human HK genes: HRT Atlas '
    'v1.0 reference (1,130 genes; human column, 1,129 matched to data).'
)

add_para('4.3 TCGA Bulk RNA-seq', bold=True)
add_para(
    'Data were obtained from the NCI Genomic Data Commons. Five cancer types were selected: '
    'LUAD (495 tumor + 76 normal), LUSC (567 + 58), LIHC (365 + 57), KIRC (755 + 82), '
    'BRCA (1,032 + 109), totaling n = 3,596 samples. Normalization: TPM values from UCSC Xena, '
    'followed by log2(TPM + 1) transformation. For paired analysis, tumor-normal pairs were '
    'matched by patient barcode (TCGA-XX-XXXX format). Clinical metadata for stratified '
    'analyses were obtained from GDC (via the TCGAbiolinks R package) and the cBioPortal API.'
)

add_para('4.4 Highly Variable Gene (HVG) Selection', bold=True)
add_para(
    'Tabula Muris: Global HVG selection was performed using '
    'scanpy.pp.highly_variable_genes, with parameters flavor="seurat" and '
    'n_top_genes=2,000. The global HVG set was used for all pairwise comparisons '
    'in Phases 3.1-3.2. Tabula Sapiens: Pairwise HVG selection. For each cell type '
    'pair (CT_i, CT_j), the top-200 genes ranked by |mu_i - mu_j| (absolute log1p '
    'expression difference) were selected as identity genes, excluding HK genes. This '
    'avoids the dilution effect of HVG across comparisons involving 102 cell types. '
    'HVG count sensitivity: the parameter sweep (Phase 3.2) tested N_HVG in '
    '{50, 100, 200, 500, 1,000, 2,000}. The global scheme (mouse) achieved peak AUC '
    'at N=2,000, while the pairwise scheme (human) used N=200 to maintain discriminative '
    'power with computational efficiency.'
)

doc.add_page_break()

# ===== Supplementary Tables =====
add_heading('Supplementary Table 1: Parameter Sweep Results', 2)
add_para(
    'Phase 3.2 parameter sweep on Tabula Muris mouse data (n = 703 cell type pairs, '
    '6 organs). The pure identity gene configuration (w1 = 1.0, w2 = 0.0) achieved '
    'optimal cell type discrimination (AUC = 0.847). Data file: '
    'results/phase32_sweep_results.csv. Visualization: results/phase32_sweep_barplot.png.'
)

add_para('')
add_heading('Supplementary Table 2: Cross-Organ Conservation Data', 2)
add_para(
    'Complete dataset of 59 same-cell-type cross-organ pairs in Tabula Sapiens, '
    'including omega, Jensen-Shannon divergence, Spearman distance, Cosine distance, '
    'and Marker Jaccard distance values. Data file: '
    'results/phase35_cross_organ_conservation.csv.'
)

add_para('')
add_heading('Supplementary Table 3: Human Brain Non-neuronal Cell Regional CKI Data', 2)
add_para(
    f'Complete results of CKI brain region analysis for non-neuronal cells from the '
    f'Siletti et al. (2023) human brain atlas, comprising {_br["total_pairs"]:,} pairwise cross-region '
    f'comparisons across 10 cell types (n = {_br["n_nuclei"]:,} nuclei, {_br["n_regions"]} regions, '
    f'{_br["n_genes"]:,} genes). Summary statistics for each cell type (omega mean, '
    f'median, SD, range, k_n and k_f components) are provided in Supplementary Table 3. '
    f'Raw data file: results/brain_siletti_omega_pairs_v3.csv ({_br["total_pairs"]:,} rows). '
    f'Summary file: results/brain_siletti_ct_summary_v3.csv (10-row summary). '
    f'Analysis script: notebooks/07c_brain_siletti_v3.py. '
    f'Figure generation: notebooks/30_nar_figures_fixed_v2.py (Figure 6).'
)

add_para('')
add_heading('Supplementary Table 4: Inter-regional Cell Migration Candidate Data', 2)

# Build dynamic S4 text
n_candidates = len(_candidates)
pct_candidates = n_candidates / len(_mig) * 100
n_strong = int(_br['n_strong'])
n_moderate = int(_br['n_moderate'])
n_weak = int(_br['n_weak'])
pct_strong = float(_br['pct_strong'])
pct_moderate = float(_br['pct_moderate'])
pct_weak = float(_br['pct_weak'])

# Build top-5 list
top5_lines = []
for i, (_, r) in enumerate(_strong.iterrows()):
    top5_lines.append(
        f'{r["cell_type"]} {r["region_a"]} vs. {r["region_b"]} '
        f'(omega = {r["omega"]:.2f}, residual = {r["residual"]:.3f})'
    )

s4_text = (
    f'Results of multiplicative model-based detection of potential inter-regional cell '
    f'migration candidates. Of the {_br["total_pairs"]:,} total pairwise cross-region comparisons, '
    f'{n_candidates:,} pairs ({pct_candidates:.1f}%) were classified as migration candidates '
    f'(residual < 0.75): '
    f'{n_strong} strong signals (residual < {DATA["brain"]["residual_thresholds"]["strong"]}, '
    f'{pct_strong:.2f}%), '
    f'{n_moderate:,} moderate signals (residual < {DATA["brain"]["residual_thresholds"]["moderate"]}, '
    f'{pct_moderate:.2f}%), '
    f'and {n_weak:,} weak signals (residual < {DATA["brain"]["residual_thresholds"]["weak"]}, '
    f'{pct_weak:.2f}%). '
    f'Top-5 strongest migration signals by cell type (ranked by residual): '
    f'1) {top5_lines[0]}, '
    f'2) {top5_lines[1]}, '
    f'3) {top5_lines[2]}, '
    f'4) {top5_lines[3]}, '
    f'5) {top5_lines[4]}. '
    f'Complete candidate dataset: results/brain_siletti_migration_candidates_v3.csv '
    f'({n_candidates:,} rows). '
    f'Analysis script: notebooks/07c_brain_siletti_v3.py.'
)
add_para(s4_text)

add_para('')
add_heading('Supplementary Data 1: Analysis Script Index', 2)
add_para(
    'Complete analysis scripts used in this study are organized in the notebooks/ '
    'directory of the GitHub repository (github.com/zhanglknt/CKI-cell-type-identification). '
    'The package can be installed via: pip install git+https://github.com/zhanglknt/CKI-cell-type-identification.git. '
    'Key scripts include: notebooks/04_phase32_sweep.py (parameter sweep and calibration), '
    'notebooks/06_phase34_v2.py (TCGA pan-cancer analysis), '
    'notebooks/05_phase33_v3_fixed.py (Tabula Sapiens cross-organ analysis), '
    'notebooks/07c_brain_siletti_v3.py (brain regional CKI analysis and migration candidate detection), '
    'and notebooks/30_nar_figures_fixed_v2.py (figure generation).'
)

# ===== Add line numbers (continuous, every line) =====
for sec in doc.sections:
    sect_pr = sec._sectPr
    ln_num = etree.SubElement(sect_pr, qn('w:lnNumType'))
    ln_num.set(qn('w:countBy'), '1')
    ln_num.set(qn('w:start'), '1')
    ln_num.set(qn('w:restart'), 'continuous')

out_path = 'results/CKI_NAR_Supplementary.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
print(f'Paragraphs: {len(doc.paragraphs)}')
