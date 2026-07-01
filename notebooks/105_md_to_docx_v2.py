#!/usr/bin/env python3
"""
Generate CKI_Reproducibility_Guide.docx directly from the CORRECT markdown file.
This ensures the DOCX is ALWAYS in sync with the markdown (single source of truth).
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN
import os

MD = "C:/Users/KnightZ/Desktop/细胞受选择/results/CKI_Reproducibility_Guide.md"
OUT = "C:/Users/KnightZ/Desktop/细胞受选择/results/CKI_Reproducibility_Guide.docx"

FONT = "Arial"
SIZE_BODY = Pt(10)
SIZE_H1 = Pt(16)
SIZE_H2 = Pt(14)
SIZE_H3 = Pt(12)

def parse_markdown(md_path):
    """Parse markdown into a list of (type, content) tuples."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    elements = []
    in_code = False
    code_lines = []
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('```'):
            if in_code:
                elements.append(('code', ''.join(code_lines)))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        
        if in_code:
            code_lines.append(line)
            continue
        
        if stripped.startswith('# '):
            elements.append(('h1', stripped[2:]))
        elif stripped.startswith('## '):
            elements.append(('h2', stripped[3:]))
        elif stripped.startswith('### '):
            elements.append(('h3', stripped[4:]))
        elif stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not all(re.match(r'^-+$', c.strip()) for c in cells if c.strip()):
                table_rows.append(cells)
        elif in_table and not stripped.startswith('|'):
            if table_rows:
                elements.append(('table', table_rows))
            table_rows = []
            in_table = False
            if stripped:
                elements.append(('text', stripped))
        elif stripped:
            elements.append(('text', stripped))
        else:
            elements.append(('blank', ''))
    
    if in_code and code_lines:
        elements.append(('code', ''.join(code_lines)))
    if in_table and table_rows:
        elements.append(('table', table_rows))
    
    return elements

def add_formatted_text(paragraph, text):
    """Add text with basic inline formatting (bold, italic, code)."""
    parts = re.split(r'(\*\*.*?\*\*|__.*?__|`.*?`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = FONT
            run.font.size = SIZE_BODY
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT
            run.font.size = SIZE_BODY
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        else:
            run = paragraph.add_run(part)
            run.font.name = FONT
            run.font.size = SIZE_BODY

def build_docx(elements, out_path):
    """Build DOCX from parsed elements."""
    doc = Document()
    
    for el_type, content in elements:
        if el_type == 'blank':
            doc.add_paragraph()
        elif el_type == 'h1':
            p = doc.add_heading(level=1)
            p.style.font.name = FONT
            run = p.add_run(content)
            run.font.name = FONT
            run.font.size = SIZE_H1
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif el_type == 'h2':
            p = doc.add_heading(level=2)
            p.style.font.name = FONT
            run = p.add_run(content)
            run.font.name = FONT
            run.font.size = SIZE_H2
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif el_type == 'h3':
            p = doc.add_heading(level=3)
            p.style.font.name = FONT
            run = p.add_run(content)
            run.font.name = FONT
            run.font.size = SIZE_H3
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif el_type == 'text':
            p = doc.add_paragraph()
            p.style.font.name = FONT
            add_formatted_text(p, content)
        elif el_type == 'code':
            p = doc.add_paragraph()
            p.style.font.name = "Courier New"
            p.style.font.size = Pt(8)
            run = p.add_run(content.strip())
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        elif el_type == 'table':
            table = doc.add_table(rows=len(content), cols=len(content[0]))
            table.style = 'Table Grid'
            for i, row in enumerate(content):
                for j, cell_text in enumerate(row):
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = FONT
                            run.font.size = SIZE_BODY
    
    doc.save(out_path)
    print(f"DOCX saved to: {out_path}")

if __name__ == '__main__':
    print(f"Reading markdown from: {MD}")
    elements = parse_markdown(MD)
    print(f"Parsed {len(elements)} elements")
    build_docx(elements, OUT)
    print("Done!")
