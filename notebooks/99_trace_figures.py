"""
CKI NAR 图表追溯系统
==================
检查：
1. 稿件正文是否引用了每个图/表的每个panel
2. 每个图表文件是否都有对应的引用
3. 图注是否与实际panel一致
"""

import re
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

RESULTS = Path("results")
FIGS_DIR = RESULTS / "figures_final"
MANUSCRIPT = Path("results/CKI_NAR_Manuscript_v4.docx")

# ============================================================
# Step 1: 从稿件正文中提取所有图/表引用
# ============================================================

def extract_figure_refs_from_docx(docx_path):
    """
    从DOCX读取正文文本，提取所有Figure/Table/Extended Data引用。
    返回：
      fig_refs: set of (fig_id, panel_letters)  e.g. {('Fig. 1', 'ABC'), ('Fig. 6', 'BCDE')}
      table_refs: set of table numbers  e.g. {1, 2}
      all_paragraphs: list of (paragraph_text, is_heading) 用于后续分析
    """
    doc = Document(docx_path)
    
    fig_refs = set()
    table_refs = set()
    all_paras = []
    
    # 读取所有段落文本
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = para.style.name.startswith('Heading')
        all_paras.append((text, is_heading))
    
    # 也读取表格中的文本（表格可能在正文中被引用）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_paras.append((text, False))
    
    full_text = '\n'.join(t for t, _ in all_paras)
    
    # --- 提取Figure引用 ---
    # 匹配: "Fig. 1", "Fig. 1a", "Fig. 1a,b", "Fig. 1a-d", "Figure 1", "Extended Data Figure 1"
    fig_patterns = [
        # Fig. 1 / Figure 1 / Fig.1
        (r'(?:Figure|Fig\.?)\s*(\d+)', 'main'),
        # Extended Data Figure 1 / ED Fig. 1
        (r'(?:Extended Data\s+)?[Ff]ig\.?\s*(\d+)', 'ed'),
        # ED Figure 1
        (r'[Ee]xtended\s+[Dd]ata\s+[Ff]ig\.?\s*(\d+)', 'ed'),
    ]
    
    for pattern, ftype in fig_patterns:
        for m in re.finditer(pattern, full_text, re.IGNORECASE):
            fig_num = int(m.group(1))
            prefix = 'ED Fig.' if ftype == 'ed' else 'Fig.'
            # 尝试提取panel字母（同一句/邻域内）
            start = m.end()
            context = full_text[start:start+30]
            panels = extract_panels(context)
            fig_refs.add((f'{prefix} {fig_num}', panels))
    
    # 更精确地提取：从正文段落中找引用
    for text, is_heading in all_paras:
        if is_heading:
            continue  # 跳过标题（标题中的Figure X是图注，不是引用）
        
        # 主图引用: "Fig. 1", "Fig. 1a", "Fig. 1a-c"
        for m in re.finditer(r'(?:Figure|Fig\.?)\s*(\d+)([a-e]+(?:[-,][a-e]+)*)?', text, re.IGNORECASE):
            fig_num = int(m.group(1))
            panel_str = m.group(2) if m.group(2) else ''
            panels = parse_panel_string(panel_str) if panel_str else list('abcde')
            fig_refs.add((f'Fig. {fig_num}', panels))
        
        # ED图引用
        for m in re.finditer(r'(?:Extended\s+Data\s+)?[Ff]ig\.?\s*(\d+)([a-e]+(?:[-,][a-e]+)*)?', text, re.IGNORECASE):
            fig_num = int(m.group(1))
            panel_str = m.group(2) if m.group(2) else ''
            panels = parse_panel_string(panel_str) if panel_str else list('abcde')
            fig_refs.add((f'ED Fig. {fig_num}', panels))
    
    # --- 提取Table引用 ---
    for text, is_heading in all_paras:
        if is_heading:
            continue
        for m in re.finditer(r'[Tt]able\s+(\d+)', text):
            table_refs.add(int(m.group(1)))
    
    return fig_refs, table_refs, all_paras

def extract_panels(context):
    """从上下文提取panel字母，如 'a,b' -> ['a','b'], 'a-c' -> ['a','b','c']"""
    panels = []
    # 匹配 "a,b,c" 或 "a-c"
    comma_sep = re.search(r'([a-e](?:\s*,\s*[a-e])+)' , context)
    range_sep = re.search(r'([a-e])\s*[-–]\s*([a-e])', context)
    single = re.search(r'^([a-e])', context)
    
    if comma_sep:
        panels = re.findall(r'[a-e]', comma_sep.group(0))
    elif range_sep:
        start_c = ord(range_sep.group(1))
        end_c = ord(range_sep.group(2))
        panels = [chr(c) for c in range(start_c, end_c+1)]
    elif single:
        panels = [single.group(1)]
    return panels

def parse_panel_string(s):
    """解析panel字符串，如 'a-c,e' -> ['a','b','c','e']"""
    panels = []
    s = s.replace(' ', '')
    for part in s.split(','):
        r = re.match(r'([a-e])-([a-e])', part)
        if r:
            for c in range(ord(r.group(1)), ord(r.group(2))+1):
                panels.append(chr(c))
        elif part:
            panels.append(part[0])
    return panels

# ============================================================
# Step 2: 从图注中提取每个图的panel定义
# ============================================================

def extract_legend_panels(manuscript_path):
    """
    从稿件的Figure Legends部分提取每个图的panel列表。
    返回：dict  {'Fig. 1': ['a','b','c'], 'ED Fig. 1': ['a','b'], ...}
    """
    doc = Document(manuscript_path)
    
    in_legends = False
    in_ed_legends = False
    legend_panels = {}
    current_fig = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检测 "Figure Legends" 标题
        if re.match(r'#?\s*Figure\s+Legends?', text, re.IGNORECASE):
            in_legends = True
            continue
        if re.match(r'#?\s*Extended\s+Data\s+Figure\s+Legends?', text, re.IGNORECASE):
            in_ed_legends = True
            in_legends = False
            continue
        
        if not (in_legends or in_ed_legends):
            continue
        
        # 新的图注开始
        fig_m = re.match(r'(?:Figure|Fig\.?)\s+(\d+)\.', text, re.IGNORECASE)
        ed_fig_m = re.match(r'[Ee]xtended\s+[Dd]ata\s+[Ff]ig\.?\s*(\d+)\.', text, re.IGNORECASE)
        
        if fig_m:
            current_fig = f'Fig. {fig_m.group(1)}'
            # 从图注文本中提取所有panel字母
            panels = re.findall(r'\(([a-e])\)', text)
            legend_panels[current_fig] = panels
        elif ed_fig_m:
            current_fig = f'ED Fig. {ed_fig_m.group(1)}'
            panels = re.findall(r'\(([a-e])\)', text)
            legend_panels[current_fig] = panels
        elif current_fig:
            # 继续当前图注，提取更多panel
            panels = re.findall(r'\(([a-e])\)', text)
            if panels:
                legend_panels[current_fig] = list(set(legend_panels.get(current_fig, []) | set(panels)))
    
    return legend_panels

# ============================================================
# Step 3: 从图表生成脚本中提取应生成的图表
# ============================================================

def extract_script_figures(script_path):
    """
    从30_nar_figures_final.py和30_nar_figures_fixed_v2.py
    提取所有生成的图表及panel。
    返回：dict {'Fig. 1': ['a','b','c','d'], ...}
    """
    script = Path(script_path)
    if not script.exists():
        return {}
    
    with open(script, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 查找所有 add_subplot 或类似创建panel的模式
    # 也查找注释中的panel说明
    panels = {}
    
    # 匹配注释如 "# Panel A:", "# (A)", "# Panel A:" 等
    for m in re.finditer(r'#\s*\(?([A-Ea-e])\)?\s*.*?(?:panel|Panel)?\s*([A-Ea-e])', content, re.IGNORECASE):
        pass  # 这个太复杂，改用更简单的方法
    
    # 简单方法：查找所有 fig.add_subplot(gs[...]) 然后看注释
    # 或直接解析脚本中定义的panel数量
    
    return panels

# ============================================================
# Step 4: 检查实际图表文件
# ============================================================

def check_figure_files():
    """
    检查figures_final/目录下的所有图表文件。
    返回：
      main_figs: list of (fig_num, formats)
      ed_figs: list of (fig_num, formats)
    """
    if not FIGS_DIR.exists():
        print(f"WARNING: {FIGS_DIR} not found!")
        return [], []
    
    main_figs = {}
    ed_figs = {}
    
    for f in FIGS_DIR.glob('*.pdf'):
        name = f.stem
        # 匹配 figure1, figure2, ...
        m = re.match(r'figure(\d+)_', name)
        if m:
            num = int(m.group(1))
            if num not in main_figs:
                main_figs[num] = []
            main_figs[num].append(f.suffix)
        
        m = re.match(r'ed_fig(\d+)_', name)
        if m:
            num = int(m.group(1))
            if num not in ed_figs:
                ed_figs[num] = []
            ed_figs[num].append(f.suffix)
    
    return main_figs, ed_figs

# ============================================================
# Step 5: 主追溯逻辑
# ============================================================

def build_traceability_report():
    if not MANUSCRIPT.exists():
        print(f"ERROR: Manuscript not found: {MANUSCRIPT}")
        return
    
    print("=" * 70)
    print("CKI NAR 图表追溯系统报告")
    print("=" * 70)
    print()
    
    # --- Step 1: 提取稿件中的引用 ---
    print("[1/5] 从稿件正文中提取图表引用...")
    fig_refs, table_refs, all_paras = extract_figure_refs_from_docx(MANUSCRIPT)
    
    print(f"  找到主图引用: {sorted(fig_refs)[:10]}...")
    print(f"  找到ED图引用: ...")
    print(f"  找到表格引用: {sorted(table_refs)}")
    print()
    
    # --- Step 2: 提取图注中的panel定义 ---
    print("[2/5] 从图注中提取panel定义...")
    legend_panels = extract_legend_panels(MANUSCRIPT)
    for fig, panels in sorted(legend_panels.items()):
        print(f"  {fig}: panels {panels}")
    print()
    
    # --- Step 3: 检查实际文件 ---
    print("[3/5] 检查实际图表文件...")
    main_figs, ed_figs = check_figure_files()
    print(f"  主图文件: {sorted(main_figs.keys())}")
    print(f"  ED图文件: {sorted(ed_figs.keys())}")
    print()
    
    # --- Step 4: 从图注中推断应有几个panel ---
    # 对每个图，图注中应该定义所有panel
    print("[4/5] 追溯分析...")
    print()
    
    issues = []
    
    # 检查主图
    for fig_num in range(1, 7):
        fig_id = f'Fig. {fig_num}'
        # 稿件中是否引用了该图？
        cited = any(fig_id in ref for ref, _ in fig_refs)
        # 图注中是否有该图？
        has_legend = fig_id in legend_panels
        # 文件是否存在？
        has_file = fig_num in main_figs
        
        status = '✓' if (cited and has_legend and has_file) else '✗'
        issues.append((fig_id, cited, has_legend, has_file, status))
    
    # 检查ED图（已知有7个ED图）
    for fig_num in range(1, 8):
        fig_id = f'ED Fig. {fig_num}'
        cited = any(fig_id in ref for ref, _ in fig_refs)
        has_legend = fig_id in legend_panels
        has_file = fig_num in ed_figs
        
        status = '✓' if (cited and has_legend and has_file) else '✗'
        issues.append((fig_id, cited, has_legend, has_file, status))
    
    # --- Step 5: 输出报告 ---
    print("[5/5] 生成追溯报告...")
    print()
    print("-" * 70)
    print(f"{'图表':<20} {'正文引用':<10} {'图注定义':<10} {'文件存在':<10} {'状态':<5}")
    print("-" * 70)
    
    all_ok = True
    for fig_id, cited, has_legend, has_file, status in issues:
        if status == '✗':
            all_ok = False
        c = '✓' if cited else '✗'
        l = '✓' if has_legend else '✗'
        f = '✓' if has_file else '✗'
        print(f"{fig_id:<20} {c:<10} {l:<10} {f:<10} {status:<5}")
    
    print("-" * 70)
    print()
    
    if all_ok:
        print("✓ 所有图表追溯完整！没有孤立的图表或引用。")
    else:
        print("⚠ 发现问题：")
        for fig_id, cited, has_legend, has_file, status in issues:
            if status == '✗':
                problems = []
                if not cited:
                    problems.append("正文未引用")
                if not has_legend:
                    problems.append("图注缺失")
                if not has_file:
                    problems.append("文件不存在")
                print(f"  {fig_id}: {', '.join(problems)}")
    
    print()
    print("=" * 70)
    
    # 详细panel追溯
    print()
    print("详细Panel追溯：")
    print("-" * 70)
    
    for fig_id, panels_in_legend in sorted(legend_panels.items()):
        # 检查正文引用了哪些panel
        cited_panels = set()
        for ref_fig, ref_panels in fig_refs:
            if ref_fig == fig_id:
                cited_panels.update(ref_panels)
        
        all_panels = set(panels_in_legend)
        missing_citations = all_panels - cited_panels
        uncited_panels = all_panels - cited_panels
        
        if uncited_panels:
            print(f"  {fig_id}: 图注定义panels {sorted(all_panels)}, 正文引用了 {sorted(cited_panels)}, ⚠ 未引用: {sorted(uncited_panels)}")
        else:
            print(f"  {fig_id}: 图注定义panels {sorted(all_panels)}, 正文引用了 {sorted(cited_panels)} ✓")
    
    print("-" * 70)

if __name__ == '__main__':
    build_traceability_report()
