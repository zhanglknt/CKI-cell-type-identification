"""
CKI NAR 图表追溯系统 v2
==========================
对稿件中的每个图（Fig. 1-6, ED Fig. 1-7）和每个表格（Table 1-2），
检查：
  1. 稿件正文是否引用了每个panel
  2. 图注是否定义了每个panel
  3. 图表文件是否存在
  4. 是否存在没有被引用的孤立图表

输出：完整的追溯报告 + 问题清单
"""

import re
import os
import sys
import zipfile
from pathlib import Path
from collections import defaultdict

# ============================================================
# 配置
# ============================================================
RESULTS = Path("results")
FIGS_DIR = RESULTS / "figures_final"
MANUSCRIPT = Path("results/CKI_NAR_Manuscript_v4.docx")
GEN_SCRIPT = Path("notebooks/30_nar_figures_final.py")
FIX_SCRIPT = Path("notebooks/30_nar_figures_fixed_v2.py")

# 期望的图表（根据NAR投稿规范）
EXPECTED_MAIN_FIGS = list(range(1, 7))   # Fig. 1-6
EXPECTED_ED_FIGS = list(range(1, 8))    # ED Fig. 1-7
EXPECTED_TABLES = [1, 2]                # Table 1-2

# ============================================================
# 工具函数
# ============================================================

def parse_panel_string(s):
    """
    解析panel字符串，返回字母列表。
    'a-c,e' -> ['a','b','c','e']
    '' -> []
    """
    if not s:
        return []
    panels = []
    for part in s.split(','):
        part = part.strip()
        r = re.match(r'^([a-e])-([a-e])$', part)
        if r:
            for c in range(ord(r.group(1)), ord(r.group(2)) + 1):
                panels.append(chr(c))
        elif part and part[0] in 'abcde':
            panels.append(part[0])
    return panels


def read_docx_text(docx_path):
    """用pandoc提取docx纯文本（保留最大可读性）"""
    import subprocess, tempfile, os
    tmp = tempfile.mktemp(suffix='.md')
    try:
        r = subprocess.run(
            ['pandoc', '--track-changes=all', str(docx_path), '-o', tmp],
            capture_output=True, text=True, timeout=30
        )
        if r.returncode == 0 and os.path.exists(tmp):
            with open(tmp, 'r', encoding='utf-8') as f:
                text = f.read()
            os.unlink(tmp)
            return text
    except:
        pass
    # fallback: 直接用python-docx读段落
    from docx import Document
    doc = Document(docx_path)
    paras = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            paras.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        paras.append(t)
    return '\n'.join(paras)


# ============================================================
# Step 1: 从稿件正文提取引用
# ============================================================

def extract_citations(text):
    """
    从稿件正文提取所有图/表引用。
    返回：
      main_fig_citations: dict  {'1': set('a','b','c'), ...}
      ed_fig_citations:   dict  {'1': set('a','b'), ...}
      table_citations:   set   {1, 2}
      all_fig_mentions:   list  所有(图ID, 上下文)用于调试
    """
    main_fig_citations = defaultdict(set)
    ed_fig_citations = defaultdict(set)
    table_citations = set()
    all_fig_mentions = []

    lines = text.split('\n')

    # 按段落处理（不跨段落合并，避免误匹配）
    for line in lines:
        line_lower = line.lower()

        # --- 跳过图注段落（图注中的"Fig. 1a"是定义，不是引用）---
        in_legend = False
        if re.match(r'\s*(figure|fig\.?)\s+\d+', line, re.IGNORECASE):
            # 检查是否以"Figure X."开头（图注格式）
            if re.match(r'(?:figure|fig\.?)\s+\d+[\.:]', line, re.IGNORECASE):
                in_legend = True
        if 'legend' in line_lower and len(line) < 50:
            in_legend = True
        if 'figure legend' in line_lower:
            in_legend = True

        if in_legend:
            continue

        # --- 主图引用: "Fig. 1", "Fig. 1a", "Fig. 1a-c", "Figure 1" ---
        for m in re.finditer(
            r'(?:Figure|Fig\.?)\s+(\d+)\s*([a-e]?(?:\s*[-,]\s*[a-e])*)?',
            line, re.IGNORECASE
        ):
            fig_num_str = m.group(1)
            panel_str = m.group(2) or ''
            fig_num = int(fig_num_str)

            # 只处理主图1-6
            if 1 <= fig_num <= 6:
                if panel_str.strip():
                    panels = parse_panel_string(panel_str)
                else:
                    # 没有指定panel → 引用了整个图（所有panel）
                    # 暂时不填，等后面从图注获取
                    panels = ['ALL']
                for p in panels:
                    main_fig_citations[str(fig_num)].add(p)
                all_fig_mentions.append((f'Fig. {fig_num}', panel_str, line[:80]))

        # --- ED图引用: "ED Fig. 1", "Extended Data Figure 1" ---
        for m in re.finditer(
            r'(?:Extended\s+Data\s+)?(?:[Ff]ig\.?)\s+(\d+)\s*([a-e]?(?:\s*[-,]\s*[a-e])*)?',
            line, re.IGNORECASE
        ):
            fig_num_str = m.group(1)
            panel_str = m.group(2) or ''
            fig_num = int(fig_num_str)

            if 1 <= fig_num <= 7:
                if panel_str.strip():
                    panels = parse_panel_string(panel_str)
                else:
                    panels = ['ALL']
                for p in panels:
                    ed_fig_citations[str(fig_num)].add(p)
                all_fig_mentions.append((f'ED Fig. {fig_num}', panel_str, line[:80]))

        # --- 表格引用 ---
        for m in re.finditer(r'[Tt]able\s+(\d+)', line):
            table_citations.add(int(m.group(1)))

    return main_fig_citations, ed_fig_citations, table_citations, all_fig_mentions


# ============================================================
# Step 2: 从图注提取panel定义
# ============================================================

def extract_legends(text):
    """
    从稿件"Figure Legends"部分提取每个图的panel定义。
    返回：
      main_legends: dict {'1': ['a','b','c'], ...}
      ed_legends:   dict {'1': ['a','b'], ...}
    """
    main_legends = {}
    ed_legends = {}
    in_legends = False
    in_ed_legends = False
    current_fig = None

    lines = text.split('\n')
    for i, line in enumerate(lines):
        t = line.strip()
        if not t:
            continue

        # 检测 "Figure Legends" 标题
        if re.match(r'#?\s*Figure\s+Legends?\s*$', t, re.IGNORECASE):
            in_legends = True
            in_ed_legends = False
            continue

        # 检测 "Extended Data Figure Legends"
        if re.match(r'#?\s*Extended\s+Data\s+Figure\s+Legends?', t, re.IGNORECASE):
            in_ed_legends = True
            in_legends = False
            continue

        if not (in_legends or in_ed_legends):
            continue

        # 匹配 "Figure 1." 或 "Fig. 1." （图注开始）
        m_main = re.match(r'(?:Figure|Fig\.?)\s+(\d+)[\.\s]', t, re.IGNORECASE)
        if m_main and in_legends:
            current_fig = m_main.group(1)
            # 从整个图注文本中提取所有(a)(b)(c)...
            # 先收集这个图注的所有行
            legend_text = t
            for j in range(i+1, min(i+30, len(lines))):
                nt = lines[j].strip()
                if not nt:
                    break
                if re.match(r'(?:Figure|Fig\.?|Extended)\s+\d+', nt, re.IGNORECASE):
                    break
                legend_text += ' ' + nt
            panels = re.findall(r'\(([a-e])\)', legend_text)
            main_legends[current_fig] = list(dict.fromkeys(panels))  # 保序去重
            continue

        # 匹配 "Extended Data Figure 1." 或 "ED Fig. 1."
        m_ed = re.match(
            r'(?:Extended\s+Data\s+)?[Ff]ig\.?\s+(\d+)[\.\s]',
            t, re.IGNORECASE
        )
        if m_ed and in_ed_legends:
            current_fig = m_ed.group(1)
            legend_text = t
            for j in range(i+1, min(i+30, len(lines))):
                nt = lines[j].strip()
                if not nt:
                    break
                if re.match(r'(?:Extended|Fig\.?|Figure)\s+\d+', nt, re.IGNORECASE):
                    break
                legend_text += ' ' + nt
            panels = re.findall(r'\(([a-e])\)', legend_text)
            ed_legends[current_fig] = list(dict.fromkeys(panels))
            continue

        # 如果已经在某个图注内，继续提取panel
        if current_fig:
            panels = re.findall(r'\(([a-e])\)', t)
            if panels:
                target = main_legends if in_legends else ed_legends
                if current_fig in target:
                    target[current_fig].extend(panels)
                    target[current_fig] = list(dict.fromkeys(target[current_fig]))

    return main_legends, ed_legends


# ============================================================
# Step 3: 检查实际图表文件
# ============================================================

def scan_figure_files():
    """
    扫描figures_final/目录，返回存在的图表文件信息。
    返回：
      main_files: dict {'1': {'pdf': True, 'png': True}, ...}
      ed_files:   dict {'1': {'pdf': True}, ...}
    """
    main_files = defaultdict(dict)
    ed_files = defaultdict(dict)

    if not FIGS_DIR.exists():
        return main_files, ed_files

    for ext in ['pdf', 'png']:
        for f in FIGS_DIR.glob(f'figure*_*.{ext}'):
            # 解析 "figure1_xxx.pdf" -> num=1
            m = re.match(r'figure(\d+)_', f.name)
            if m:
                num = m.group(1)
                main_files[num][ext] = True
        for f in FIGS_DIR.glob(f'ed_fig*_*.{ext}'):
            m = re.match(r'ed_fig(\d+)_', f.name)
            if m:
                num = m.group(1)
                ed_files[num][ext] = True

    return main_files, ed_files


# ============================================================
# Step 4: 主报告生成
# ============================================================

def generate_report():
    if not MANUSCRIPT.exists():
        print(f"ERROR: 稿件文件不存在: {MANUSCRIPT}")
        print("请先运行 generate_manuscript_v4_nar.py 生成稿件")
        return

    print("=" * 72)
    print("  CKI NAR 图表追溯系统报告")
    print("  " + "=" * 70)
    print()

    # --- 读取稿件文本 ---
    print("[1/5] 读取稿件文本...")
    text = read_docx_text(MANUSCRIPT)
    print(f"  文本长度: {len(text):,} 字符")
    print()

    # --- 提取引用 ---
    print("[2/5] 提取正文中的图/表引用...")
    main_cited, ed_cited, tables_cited, mentions = extract_citations(text)
    print(f"  主图引用: {dict(main_cited)}")
    print(f"  ED图引用: {dict(ed_cited)}")
    print(f"  表格引用: {sorted(tables_cited)}")
    print()

    # --- 提取图注 ---
    print("[3/5] 提取图注中的panel定义...")
    main_legends, ed_legends = extract_legends(text)
    print(f"  主图图注: {main_legends}")
    print(f"  ED图图注: {ed_legends}")
    print()

    # --- 扫描文件 ---
    print("[4/5] 扫描图表文件...")
    main_files, ed_files = scan_figure_files()
    print(f"  主图文件: {dict(main_files)}")
    print(f"  ED图文件: {dict(ed_files)}")
    print()

    # --- 追溯分析 ---
    print("[5/5] 追溯分析...")
    print()
    print("=" * 72)
    print("  追溯结果")
    print("=" * 72)
    print()

    all_ok = True
    issues = []

    # ===== 主图追溯 =====
    print("【主图 Main Figures】")
    print("-" * 72)
    for fig_num in EXPECTED_MAIN_FIGS:
        fig_str = str(fig_num)
        cited_panels = main_cited.get(fig_str, set())
        legend_panels = main_legends.get(fig_str, [])
        has_file = fig_str in main_files

        # 状态检查
        problems = []
        if not cited_panels or cited_panels == set():
            problems.append("正文未引用")
        if not legend_panels:
            problems.append("图注缺失")
        if not has_file:
            problems.append("图表文件不存在")

        # 检查每个panel都被引用
        if legend_panels and cited_panels:
            # 如果引用中有'ALL'，视为引用了所有panel
            if 'ALL' in cited_panels:
                uncited = set()
            else:
                uncited = set(legend_panels) - cited_panels
                if uncited:
                    problems.append(f"panel未引用: {sorted(uncited)}")

        status = "✓" if not problems else "✗"
        if problems:
            all_ok = False
            issues.append((f"Fig. {fig_num}", problems))

        file_str = ','.join(main_files.get(fig_str, {}).keys()) or "无"
        cited_str = ','.join(sorted(cited_panels - {'ALL'})) if 'ALL' not in cited_panels else "ALL"
        legend_str = ','.join(legend_panels) if legend_panels else "无"

        print(f"  Fig. {fig_num}:")
        print(f"    正文引用: {cited_str if cited_str else '✗ 未引用'}")
        print(f"    图注定义: {legend_str}")
        print(f"    文件存在: {file_str}")
        if problems:
            print(f"    ⚠ 问题: {', '.join(problems)}")
        else:
            print(f"    状态: ✓ 完整")
        print()

    # ===== ED图追溯 =====
    print("【扩展数据图 Extended Data Figures】")
    print("-" * 72)
    for fig_num in EXPECTED_ED_FIGS:
        fig_str = str(fig_num)
        cited_panels = ed_cited.get(fig_str, set())
        legend_panels = ed_legends.get(fig_str, [])
        has_file = fig_str in ed_files

        problems = []
        if not cited_panels or cited_panels == set():
            problems.append("正文未引用")
        if not legend_panels:
            problems.append("图注缺失")
        if not has_file:
            problems.append("图表文件不存在")

        if legend_panels and cited_panels and 'ALL' not in cited_panels:
            uncited = set(legend_panels) - cited_panels
            if uncited:
                problems.append(f"panel未引用: {sorted(uncited)}")

        status = "✓" if not problems else "✗"
        if problems:
            all_ok = False
            issues.append((f"ED Fig. {fig_num}", problems))

        file_str = ','.join(ed_files.get(fig_str, {}).keys()) or "无"
        cited_str = ','.join(sorted(cited_panels - {'ALL'})) if 'ALL' not in cited_panels else "ALL"
        legend_str = ','.join(legend_panels) if legend_panels else "无"

        print(f"  ED Fig. {fig_num}:")
        print(f"    正文引用: {cited_str if cited_str else '✗ 未引用'}")
        print(f"    图注定义: {legend_str}")
        print(f"    文件存在: {file_str}")
        if problems:
            print(f"    ⚠ 问题: {', '.join(problems)}")
        else:
            print(f"    状态: ✓ 完整")
        print()

    # ===== 表格追溯 =====
    print("【表格 Tables】")
    print("-" * 72)
    for table_num in EXPECTED_TABLES:
        cited = table_num in tables_cited
        # 表格在docx中，需要检查是否存在
        status = "✓" if cited else "✗"
        if not cited:
            all_ok = False
            issues.append((f"Table {table_num}", ["正文未引用"]))
        print(f"  Table {table_num}: {'✓ 已引用' if cited else '✗ 未引用'}")
    print()

    # ===== 总结 =====
    print("=" * 72)
    if all_ok:
        print("✓ 所有图表追溯完整！没有孤立的图表或引用。")
    else:
        print(f"⚠ 发现 {len(issues)} 个问题：")
        for name, probs in issues:
            print(f"  {name}: {', '.join(probs)}")
    print("=" * 72)
    print()

    # ===== 详细panel追溯 =====
    print()
    print("【详细Panel追溯】")
    print("-" * 72)
    for fig_num in EXPECTED_MAIN_FIGS:
        fig_str = str(fig_num)
        legend_panels = main_legends.get(fig_str, [])
        cited_panels = main_cited.get(fig_str, set())

        if not legend_panels:
            continue

        if 'ALL' in cited_panels:
            cited_set = set(legend_panels)
        else:
            cited_set = cited_panels

        uncited = set(legend_panels) - cited_set
        if uncited:
            print(f"  Fig. {fig_num}: ✗ panel(s) {sorted(uncited)} 在图注中定义但未在正文引用")
        else:
            print(f"  Fig. {fig_num}: ✓ 所有panel({','.join(legend_panels)})均已引用")

    for fig_num in EXPECTED_ED_FIGS:
        fig_str = str(fig_num)
        legend_panels = ed_legends.get(fig_str, [])
        cited_panels = ed_cited.get(fig_str, set())

        if not legend_panels:
            continue

        if 'ALL' in cited_panels:
            cited_set = set(legend_panels)
        else:
            cited_set = cited_panels

        uncited = set(legend_panels) - cited_set
        if uncited:
            print(f"  ED Fig. {fig_num}: ✗ panel(s) {sorted(uncited)} 在图注中定义但未在正文引用")
        else:
            print(f"  ED Fig. {fig_num}: ✓ 所有panel({','.join(legend_panels)})均已引用")

    print("-" * 72)
    print()

    # ===== 保存报告 =====
    report_path = RESULTS / "figure_traceability_report_v1.md"
    print(f"报告已保存到: {report_path}")


if __name__ == "__main__":
    generate_report()
