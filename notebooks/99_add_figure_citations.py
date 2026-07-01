"""
CKI NAR 图表追溯系统 — 最终版
====================================
直接修改 results/CKI_NAR_Manuscript_v4.docx，
在正文中为每个图/表面添加引用，然后验证完整性。

使用方法：
  python notebooks/99_add_figure_citations.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

DOCX_PATH = "results/CKI_NAR_Manuscript_v4.docx"
OUTPUT_PATH = "results/CKI_NAR_Manuscript_v4_cited.docx"

# ============================================================
# 第1步：找到稿件结构边界
# ============================================================

def find_structure_boundaries(docx_path):
    """读取docx，返回关键边界的段落索引。"""
    doc = Document(docx_path)
    paras = []
    for i, para in enumerate(doc.paragraphs):
        paras.append((i, para.text.strip(), para.style.name))
    
    # 找边界
    main_text_end = None     # "Figure Legends" 标题所在段落（正文结束）
    ed_legend_start = None  # "Extended Data Figure Legends" 标题
    reference_start = None   # "References" 标题
    
    for i, text, style in paras:
        if not text:
            continue
        if style.startswith('Heading'):
            lower = text.lower().strip()
            if re.match(r'figure\s+legends?\s*$', lower):
                main_text_end = i
            elif re.match(r'extended\s+data\s+figure\s+legends?', lower):
                ed_legend_start = i
            elif re.match(r'references?\s*$', lower):
                reference_start = i
    
    if main_text_end is None:
        main_text_end = len(paras)
    if ed_legend_start is None:
        ed_legend_start = main_text_end
    if reference_start is None:
        reference_start = len(paras)
    
    print(f"  正文结束于段落 {main_text_end} ('Figure Legends')")
    print(f"  ED图注开始于段落 {ed_legend_start}")
    print(f"  References开始于段落 {reference_start}")
    
    return doc, paras, main_text_end, ed_legend_start, reference_start


# ============================================================
# 第2步：在正文中添加图/表引用
# ============================================================

def add_citations_to_main_text(doc, paras, main_text_end):
    """
    在Results章节的每个子标题后，添加图引用。
    同时在合适位置添加Table引用。
    """
    # 建立标题→图号映射
    heading_fig_map = {
        'Decomposing transcriptomic variation': 'Fig. 1a-c',
        'Calibration confirms neutral behavior at baseline': 'Fig. 2a-d',
        'CKI captures information that standard metrics miss': 'Fig. 3a-d',
        'Cancer analysis reveals unexpected transcriptional convergence': 'Fig. 4a-d',
        'CKI ranks cell types by cross-organ conservation': 'Fig. 5a-c',
        'Brain regional analysis reveals cell-type differentiation gradients': 'Fig. 6a-e',
    }
    
    edits = []
    
    i = 0
    while i < main_text_end:
        text, style = paras[i][1], paras[i][2]
        
        # 检查是否是Results子标题
        if style.startswith('Heading') and style in ('Heading 2', 'Heading 3'):
            heading_text = text.strip()
            if heading_text in heading_fig_map:
                fig_citation = heading_fig_map[heading_text]
                # 在下一个Normal段落开头添加引用
                j = i + 1
                while j < main_text_end:
                    j_text, j_style = paras[j][1], paras[j][2]
                    if j_text and not j_style.startswith('Heading'):
                        # 找到第一个内容段落，在开头插入引用
                        edits.append((j, fig_citation))
                        break
                    j += 1
                i = j + 1
                continue
        
        i += 1
    
    print(f"  找到 {len(edits)} 处需要添加图引用")
    for para_idx, fig_cit in edits:
        print(f"    段落 {para_idx}: 添加 {fig_cit}")
    
    return edits


def apply_edits(doc, paras, edits):
    """将编辑应用到文档。在每个目标段落开头插入图引用。"""
    # 按段落索引分组
    from collections import defaultdict
    edit_map = defaultdict(list)
    for para_idx, fig_cit in edits:
        edit_map[para_idx].append(fig_cit)
    
    # 对doc.paragraphs做同样的索引
    for para_idx, citations in sorted(edit_map.items()):
        if para_idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[para_idx]
        old_text = para.text
        
        # 在开头添加引用
        citation_str = '; '.join(c for c in citations)
        if not old_text.startswith('Fig') and not old_text.startswith('CKI'):
            new_text = old_text
        else:
            # 已经在开头有直接引用，跳过
            continue
        
        # 插入引用到段落开头
        run = para.insert_paragraph_before()
        run.text = f'{citation_str}. '
        # 设置格式
        for r in run.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0, 0, 0)
    
    return doc


# ============================================================
# 第3步：验证追溯完整性
# ============================================================

def verify_traceability(docx_path):
    """验证修改后的稿件中所有图表都被引用。"""
    doc = Document(docx_path)
    paras = [(p.text.strip(), p.style.name) for p in doc.paragraphs]
    
    issues = []
    
    # 检查主图引用
    for fig_num in range(1, 7):
        found = False
        for text, style in paras:
            if not text:
                continue
            if re.search(rf'[Ff]ig\.?\s*{fig_num}[^0-9]', text):
                found = True
                break
        if not found:
            issues.append(f"Fig. {fig_num}: 正文中未找到引用")
    
    # 检查ED图引用（ED图在正文中通常不引用，除非特别指明）
    # 对于NAR，ED图在Extended Data文件中，不必须在正文中引用
    
    # 检查表格引用
    table_found = False
    for text, style in paras:
        if 'Table' in text and re.search(r'Table\s+[12]', text):
            table_found = True
            break
    if not table_found:
        issues.append("Table 1/2: 正文中未找到引用")
    
    return issues


# ============================================================
# 主函数
# ============================================================

def main():
    print("=" * 72)
    print("  CKI NAR 图表追溯系统")
    print("=" * 72)
    print()
    
    # 第1步
    print("[1/3] 读取稿件结构...")
    doc, paras, main_end, ed_start, ref_start = find_structure_boundaries(DOCX_PATH)
    print()
    
    # 第2步
    print("[2/3] 在正文中添加图/表引用...")
    edits = add_citations_to_main_text(doc, paras, main_end)
    
    if edits:
        print(f"  应用 {len(edits)} 处编辑...")
        doc = apply_edits(doc, paras, edits)
        doc.save(OUTPUT_PATH)
        print(f"  已保存到: {OUTPUT_PATH}")
    else:
        print("  无需编辑（可能已有引用）")
        OUTPUT_PATH = DOCX_PATH  # 使用原文件
    print()
    
    # 第3步
    print("[3/3] 验证追溯完整性...")
    issues = verify_traceability(OUTPUT_PATH)
    print()
    
    print("=" * 72)
    if issues:
        print("⚠ 发现问题：")
        for issue in issues:
            print(f"  - {issue}")
    else:
        print("✓ 所有图表追溯完整！")
    print("=" * 72)


if __name__ == '__main__':
    main()
