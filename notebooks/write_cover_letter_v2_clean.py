"""
Rewrite CKI_NAR_Cover_Letter_v2.docx
Front-load the Ka/Ks analogy as the central conceptual innovation.
Use minimal docx API: Document + add_paragraph + add_run only.
"""

import sys
sys.path.insert(0, sys.path[0] + r'\C:\Users\KnightZ\.workbuddy\binaries\python\versions\3.13.12\Lib\site-packages')

from docx import Document

RESULTS = r"C:\Users\KnightZ\Desktop\细胞受选择\results"

doc = Document()

def add(p, text, bold=False, size=12):
    """Add a run to paragraph p with given text and formatting."""
    r = p.add_run(text)
    r.bold = bold
    r.font.size = size * 20  # half-points
    r.font.name = "Times New Roman"
    return r

# --- 0: Date ---
p = doc.add_paragraph("May 24, 2026")
for r in p.runs:
    r.font.name = "Times New Roman"
    r.font.size = 12 * 20

# --- 1: RE ---
p = doc.add_paragraph()
add(p, "RE: \"CKI: A Cell-state Kinetic Index for Quantifying Selective Transcriptomic Remodeling\"", bold=True)

# --- 2: Salutation ---
p = doc.add_paragraph("Dear Editors of Nucleic Acids Research,")
for r in p.runs:
    r.font.name = "Times New Roman"
    r.font.size = 12 * 20

# --- 3: Opening — Ka/Ks analogy (FRONT and CENTER) ---
p = doc.add_paragraph()
add(p,
    "Just as the Ka/Ks framework revolutionized molecular evolution by decomposing sequence divergence into neutral (Ks) and functional (Ka) components, ",
    size=12)
add(p,
    "CKI (Cell-state Kinetic Index) brings this same principled decomposition to single-cell transcriptomics.",
    bold=True, size=12)

p.add_run(
    " CKI partitions transcriptomic divergence into a neutral drift rate (k_n, from housekeeping genes — the transcriptomic Ks) and a functional conversion rate (k_f, from cell-type-identity genes — the transcriptomic Ka), yielding omega = k_f/k_n as a quantitative index of selective transcriptional remodeling."
)
for r in p.runs:
    r.font.name = "Times New Roman"
    r.font.size = 12 * 20

# --- 4: Problem + CKI solution ---
p = doc.add_paragraph()
add(p,
    "Current methods for comparing transcriptomes measure how different two cell populations are, but cannot distinguish functional reprogramming from neutral background noise. ",
    size=12)
add(p,
    "CKI solves this by using housekeeping genes as an internal neutral reference — the transcriptomic equivalent of Ks — and identity genes as the functional readout — the transcriptomic equivalent of Ka.",
    size=12)

# --- 5: Validation overview ---
p = doc.add_paragraph()
add(p,
    "We validated CKI (v0.2.0, available as a pip-installable Python package) across four independent datasets spanning three data modalities: ",
    size=12)
add(p,
    "Tabula Muris (15,057 mouse cells, 38 cell types, 6 organs), Tabula Sapiens (108,136 human cells, 102 cell types, 24 organs), TCGA bulk RNA-seq (3,596 samples, 5 cancer types), and the Siletti et al. human brain atlas (888,263 non-neuronal nuclei, 10 cell classes, 108 regions).",
    size=12)

# --- 6: Key finding 1 — omega is a new dimension ---
p = doc.add_paragraph()
add(p,
    "The Ka/Ks analogy makes a testable prediction: if omega captures a distinct axis of transcriptomic divergence, it should be weakly correlated with standard distance metrics. ",
    size=12)
add(p,
    "Across 5,151 Tabula Sapiens cell-type pairs, CKI omega was weakly correlated with all four standard distance metrics (Spearman r = −0.09 to −0.34, all P < 10^−12), confirming that omega quantifies a previously unmeasured dimension of cell-type identity.",
    size=12)

# --- 7: Key finding 2 — cancer paradox ---
p = doc.add_paragraph()
add(p,
    "Applied to TCGA tumors, CKI revealed a paradox: tumors are transcriptionally more homogeneous than normal tissues (NN/TT omega ratio = 1.76–2.47, all P < 10^−47 across 5 cancer types). ",
    size=12)
add(p,
    "This suggests that cancer progression involves not only selfish transcriptional reprogramming, but also a systematic loss of cell-type-specific identity — a prediction uniquely enabled by the omega = k_f/k_n decomposition.",
    size=12)

# --- 8: Key finding 3 — brain gradient ---
p = doc.add_paragraph()
add(p,
    "In the human brain, CKI revealed a 6.06-fold omega gradient across 10 non-neuronal cell classes (Bergmann glia: ω = 2.37; Astrocytes: ω = 14.36; bootstrap P < 10^−47). ",
    size=12)
add(p,
    "The gradient aligns with known cell biology: vascular cells and fibroblasts (uniform microenvironments) show the lowest omega, while astrocytes (region-dependent functions) show the highest — demonstrating that omega captures position-dependent selective pressure.",
    size=12)

# --- 9: Why NAR ---
p = doc.add_paragraph()
add(p, "Four reasons make this work a strong fit for Nucleic Acids Research: ", bold=True, size=12)
add(p,
    "(1) CKI provides a principled framework for interpreting transcriptomic divergence, directly analogous to the Ka/Ks framework that transformed molecular evolution; "
    "(2) the method is implemented as a fully open-source Python package (pip install git+https://github.com/zhanglknt/CKI-cell-type-identification.git), enabling immediate community adoption; "
    "(3) validation spans three data modalities (scRNA-seq, snRNA-seq, bulk RNA-seq) and four independent datasets; "
    "(4) the cancer and brain findings are of broad interest to NAR readers.",
    size=12)

# --- 10: Competing interests ---
p = doc.add_paragraph()
add(p,
    "All authors have approved the manuscript and declare no competing interests. This manuscript has not been published elsewhere and is not under consideration by any other journal.",
    size=12)

# --- 11: Closing ---
p = doc.add_paragraph("Sincerely,")
for r in p.runs:
    r.font.name = "Times New Roman"
    r.font.size = 12 * 20

# --- 12: Signature ---
p = doc.add_paragraph()
add(p, "Li Zhang (张力)", bold=True, size=12)

# --- 13-15: Affiliations + ORCID + Email ---
for text in [
    "Institute of Blood Transfusion, Chinese Academy of Medical Sciences & Peking Union Medical College, Chengdu, China",
    "Chinese Institute for Brain Research, Beijing, China",
    "ORCID: 0000-0002-0698-0754",
    "Email: knightz@pumc.edu.cn",
]:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = 12 * 20

# Save
out = RESULTS + r"\CKI_NAR_Cover_Letter_v2.docx"
doc.save(out)
print(f"Saved: {out}")
print("Done!")
