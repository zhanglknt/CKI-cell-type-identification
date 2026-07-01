#!/usr/bin/env python3
"""Generate DOCX from markdown - minimal version."""
from docx import Document
from docx.shared import Pt, RGBColor
import os

MD = r"C:\Users\KnightZ\Desktop\细胞受选择\results\CKI_Reproducibility_Guide.md"
OUT = r"C:\Users\KnightZ\Desktop\细胞受选择\results\CKI_Reproducibility_Guide.docx"

FONT = "Arial"

def build():
    doc = Document()
    with open(MD, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code = False
    code_buf = []
    
    for line in lines:
        stripped = line.strip()
        
        # Code block toggle
        if stripped.startswith('```'):
            if in_code:
                # End code block - write buffered lines
                p = doc.add_paragraph()
                p.style.font.name = "Courier New"
                run = p.add_run(''.join(code_buf))
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        
        if in_code:
            code_buf.append(line)
            continue
        
        # Skip empty lines
        if not stripped:
            doc.add_paragraph()
            continue
        
        # Headings
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # Table
        elif stripped.startswith('|'):
            # Simple table handling - just write as monospace text
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        # Normal text
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = FONT
            run.font.size = Pt(10)
    
    # Handle unclosed code block
    if in_code and code_buf:
        p = doc.add_paragraph()
        run = p.add_run(''.join(code_buf))
        run.font.name = "Courier New"
        run.font.size = Pt(8)
    
    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == '__main__':
    build()
