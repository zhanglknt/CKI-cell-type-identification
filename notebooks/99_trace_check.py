"""
CKI NAR 图表追溯检查系统
====================================
验证主稿件中所有图表都在正文中被引用。

使用方法：
  python notebooks/99_trace_check.py
"""

from docx import Document
import re
import json
from collections import defaultdict

DOCX_PATH = "results/CKI_NAR_Manuscript_v4.docx"

def read_docx_paragraphs(docx_path):
    """读取docx，返回段落列表 (text, style_name, is_in_main_text)。"""
    doc = Document(docx_path)
    paragraphs = []
    main_text_end = None
    ed_legend_start = None
    reference_start = None
    
    # 第一遍：找到结构边界
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        
        if not text:
            continue
        
        if style.startswith('Heading'):
            lower = text.lower()
            if re.match(r'figure\s+legends?\s*$', lower):
                main_text_end = i
            elif re.match(r'extended\s+data\s+figure\s+legends?', lower):
                ed_legend_start = i
            elif re.match(r'references?\s*$', lower):
                reference_start = i
    
    if main_text_end is None:
        main_text_end = len(doc.paragraphs)
    if ed_legend_start is None:
        ed_legend_start = len(doc.paragraphs)
    if reference_start is None:
        reference_start = len(doc.paragraphs)
    
    print(f"结构边界：")
    print(f"  正文结束于段落 ~{main_text_end} ('Figure Legends')")
    print(f"  ED图例开始于段落 ~{ed_legend_start}")
    print(f"  References开始于段落 ~{reference_start}")
    print()
    
    # 第二遍：标记每个段落是否在正文中
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        is_main = (i < main_text_end)
        paragraphs.append({
            'index': i,
            'text': text,
            'style': style,
            'is_main_text': is_main
        })
    
    return paragraphs, main_text_end, ed_legend_start, reference_start


def extract_fig_citations_in_main_text(paragraphs):
    """
    从正文中提取所有图表引用。
    只检查正文部分（Figure Legends之前）。
    """
    citations = {
        'figures': defaultdict(set),  # fig_num -> set of panels
        'tables': set(),
        'ed_figures': defaultdict(set),
    }
    
    # 正则：匹配 Fig. 1, Fig. 1a, Fig. 1a-c, Fig. 1a,c,e
    fig_pattern = re.compile(
        r'[Ff]ig\.?\s*(\d+)\s*([a-z]+(?:\s*[-,]\s*[a-z]+)*)?',
        re.IGNORECASE
    )
    # 正则：匹配 Extended Data Fig. 1
    ed_fig_pattern = re.compile(
        r'[Ee]xtended\s+[Dd]ata\s+[Ff]ig\.?\s*(\d+)',
        re.IGNORECASE
    )
    # 正则：匹配 Table 1, Table 1-2
    table_pattern = re.compile(r'[Tt]able\s+(\d+)', re.IGNORECASE)
    
    for para in paragraphs:
        if not para['is_main_text']:
            continue
        text = para['text']
        if not text:
            continue
        
        # 跳过标题行（Heading）
        if para['style'].startswith('Heading'):
            continue
        
        # 查找主图引用
        for m in fig_pattern.finditer(text):
            fig_num = int(m.group(1))
            panel_str = m.group(2)
            
            if panel_str:
                # 解析面板：a-c 或 a,c,e
                panels = set()
                panel_str = panel_str.strip()
                if '-' in panel_str:
                    start, end = re.match(r'([a-z])-([a-z])', panel_str).groups()
                    start_idx = ord(start) - ord('a')
                    end_idx = ord(end) - ord('a')
                    for i in range(start_idx, end_idx + 1):
                        panels.add(chr(ord('a') + i))
                elif ',' in panel_str:
                    for p in panel_str.split(','):
                        p = p.strip()
                        if p:
                            panels.add(p)
                else:
                    panels.add(panel_str)
                citations['figures'][fig_num].update(panels)
            else:
                # 没有指定面板，记录为 "all"
                citations['figures'][fig_num].add('*all*')
        
        # 查找ED图引用
        for m in ed_fig_pattern.finditer(text):
            ed_num = int(m.group(1))
            citations['ed_figures'][ed_num].add('*cited*')
        
        # 查找表格引用
        for m in table_pattern.finditer(text):
            table_num = int(m.group(1))
            citations['tables'].add(table_num)
    
    return citations


def extract_fig_definitions_from_legends(paragraphs, ed_legend_start, reference_start):
    """
    从Figure Legends部分提取所有定义的图表及其面板。
    """
    fig_defs = defaultdict(set)  # fig_num -> set of panels defined
    table_defs = set()
    
    # 只检查 Figure Legends 部分（在 ed_legend_start 之前，或从Heading "Figure Legends" 之后）
    in_fig_legend = False
    current_fig = None
    
    for para in paragraphs:
        idx = para['index']
        text = para['text']
        style = para['style']
        
        # 检测是否进入 Figure Legends 部分
        if style.startswith('Heading'):
            if re.match(r'figure\s+legends?', text, re.IGNORECASE):
                in_fig_legend = True
                continue
            elif re.match(r'extended\s+data|references?', text, re.IGNORECASE):
                in_fig_legend = False
                continue
        
        if not in_fig_legend:
            continue
        
        # 检测 "Figure X." 标题（每个图例的开头）
        # 格式: "Figure 1. The CKI framework. (a) ..."
        m = re.match(r'^Figure\s+(\d+)\.', text)
        if m:
            current_fig = int(m.group(1))
            # 在当前段落中查找面板标签
            panel_pattern = re.compile(r'\(([a-z])\)')
            for pm in panel_pattern.finditer(text):
                fig_defs[current_fig].add(pm.group(1))
            continue
        
        # 在图例段落中，查找面板定义
        if current_fig is not None:
            panel_pattern = re.compile(r'\(([a-z])\)')
            for pm in panel_pattern.finditer(text):
                fig_defs[current_fig].add(pm.group(1))
    
    return fig_defs


def check_traceability(citations, fig_defs):
    """
    检查追溯完整性。
    返回问题列表。
    """
    issues = []
    
    # 检查每个定义的图是否都被引用
    for fig_num, defined_panels in sorted(fig_defs.items()):
        if fig_num not in citations['figures']:
            issues.append(f"Fig. {fig_num}: 定义了但正文中未引用")
            continue
        
        cited_panels = citations['figures'][fig_num]
        
        if '*all*' in cited_panels:
            # 引用了整个图，OK
            continue
        
        # 检查每个定义的面板是否都被引用
        for panel in defined_panels:
            if panel not in cited_panels:
                issues.append(f"Fig. {fig_num}{panel}: 已定义但正文中未引用")
    
    # 检查每个引用的图是否都有定义
    for fig_num, cited_panels in citations['figures'].items():
        if fig_num not in fig_defs:
            issues.append(f"Fig. {fig_num}: 正文中引用但图例中未定义")
    
    # 检查表格
    # （暂时跳过，因为表格可能在单独的文件或正文中）
    
    return issues


def main():
    print("=" * 72)
    print("  CKI NAR 图表追溯检查")
    print("=" * 72)
    print()
    
    # 第1步：读取文档结构
    print("[1/4] 读取文档结构...")
    paragraphs, main_end, ed_start, ref_start = read_docx_paragraphs(DOCX_PATH)
    print(f"  共 {len(paragraphs)} 个段落")
    print()
    
    # 第2步：从正文提取引用
    print("[2/4] 从正文提取图表引用...")
    citations = extract_fig_citations_in_main_text(paragraphs)
    print(f"  主图引用：")
    for fig_num in sorted(citations['figures'].keys()):
        panels = citations['figures'][fig_num]
        print(f"    Fig. {fig_num}: {sorted(panels)}")
    print(f"  表格引用：{sorted(citations['tables']) if citations['tables'] else '无'}")
    print()
    
    # 第3步：从图例提取定义
    print("[3/4] 从图例提取图表定义...")
    fig_defs = extract_fig_definitions_from_legends(paragraphs, ed_start, ref_start)
    print(f"  主图定义：")
    for fig_num in sorted(fig_defs.keys()):
        panels = fig_defs[fig_num]
        print(f"    Fig. {fig_num}: {sorted(panels)}")
    print()
    
    # 第4步：检查追溯完整性
    print("[4/4] 检查追溯完整性...")
    issues = check_traceability(citations, fig_defs)
    print()
    
    print("=" * 72)
    if issues:
        print(f"⚠ 发现 {len(issues)} 个问题：")
        for issue in issues:
            print(f"  - {issue}")
    else:
        print("✓ 所有图表追溯完整！")
    print("=" * 72)
    print()
    
    # 输出详细引用信息供调试
    print("详细引用信息（正文）：")
    for fig_num in sorted(citations['figures'].keys()):
        panels = citations['figures'][fig_num]
        print(f"  Fig. {fig_num} -> panels {sorted(panels)}")
    
    return len(issues) == 0


if __name__ == '__main__':
    success = main()
    exit(0 if success else 1)
