"""
Final audit of MS v6 and submission v11.
Checks: figure citations, file completeness, statistics consistency, placeholder text.
"""
from docx import Document
import re, zipfile, os

doc = Document('results/CKI_NAR_Manuscript_v6.docx')

# ============================================================
# 1. Figure citations
# ============================================================
main_figs = set()
ext_figs = set()

for p in doc.paragraphs:
    t = p.text
    # Extended Data Figure X
    for m in re.finditer(r'Extended\s+Data\s+Figure\s+(\d+)', t, re.IGNORECASE):
        ext_figs.add(int(m.group(1)))
    # Ed Fig X (shorthand)
    for m in re.finditer(r'Ed\s+Fig\.?\s*(\d+)', t, re.IGNORECASE):
        ext_figs.add(int(m.group(1)))
    # Extended Data Fig. X
    for m in re.finditer(r'Extended\s+Data\s+Fig\.?\s*(\d+)', t, re.IGNORECASE):
        ext_figs.add(int(m.group(1)))
    # Main Figure X (only if NOT preceded by 'Extended Data')
    for m in re.finditer(r'Figure\s+(\d+)', t, re.IGNORECASE):
        start = m.start()
        # Check this isn't part of 'Extended Data Figure'
        before = t[max(0, start-30):start].strip()
        if not before.endswith('Extended Data') and 'Extended Data' not in before:
            main_figs.add(int(m.group(1)))

print("=== Figure Citations in MS v6 ===")
print(f"Main figures cited: {sorted(main_figs)}")
print(f"Extended figures cited: {sorted(ext_figs)}")
print(f"Total: main={len(main_figs)}/6, ext={len(ext_figs)}/7")
print()

# ============================================================
# 2. Files in submission v11
# ============================================================
print("=== Submission v11 Contents ===")
with zipfile.ZipFile('results/CKI_NAR_Submission_v11.zip', 'r') as zf:
    names = sorted(zf.namelist())
    for n in names:
        print(f"  {n}")
print()

# ============================================================
# 3. TCGA statistics check
# ============================================================
print("=== TCGA Statistics ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '1.76' in t or '2.47' in t or 'P <' in t:
        print(f"  Para {i}: {t[:250]}")
        print()
        break

# Also check abstract (para 10)
t10 = doc.paragraphs[10].text
print(f"  Abstract (para 10) snippet: {t10[:(t10.find('1.76')+30 if '1.76' in t10 else 200)]}...")
print()

# ============================================================
# 4. Bootstrap CI presence
# ============================================================
print("=== Bootstrap CI Check ===")
human_ok = any('95% CI' in p.text and ('8.65' in p.text or '16.00' in p.text) for p in doc.paragraphs)
brain_ok = any('95% CI' in p.text and ('2.37' in p.text or '14.36' in p.text) for p in doc.paragraphs)
print(f"  Human bootstrap CI in MS: {human_ok}")
print(f"  Brain bootstrap CI in MS: {brain_ok}")
print()

# ============================================================
# 5. Real placeholder check (exclude [IQR], gene brackets)
# ============================================================
print("=== Placeholder Text Check ===")
real_ph = ['TODO', 'PLACEHOLDER', 'FIXME', 'XXXXX', '?????', 'FILL', 'REPLACE']
false_positives = ['[IQR]', '[CI', '95% CI', '[]', '[]', '[]']
found = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    for ph in real_ph:
        if ph in t:
            found.append(f"  Para {i}: {t[:120]}")
if found:
    print("  FOUND placeholders:")
    for f in found[:10]:
        print(f)
else:
    print("  No placeholder text found (clean).")
print()

# ============================================================
# 6. Ka/Ks analogy in CL v2
# ============================================================
print("=== CL v2 Ka/Ks Analogy Check ===")
cl = Document('results/CKI_NAR_Cover_Letter_v2.docx')
found_ka = False
for p in cl.paragraphs:
    if 'Ka/Ks' in p.text or 'Ka/Ks' in p.text:
        print(f"  Found: {p.text[:200]}")
        found_ka = True
        break
if not found_ka:
    print("  WARNING: Ka/Ks not found in CL v2!")
print()

# ============================================================
# 7. Summary
# ============================================================
print("=== AUDIT SUMMARY ===")
issues = []
if len(main_figs) < 6:
    issues.append(f"Missing main figure citations: have {sorted(main_figs)}, need 1-6")
if len(ext_figs) < 7:
    issues.append(f"Missing extended figure citations: have {sorted(ext_figs)}, need 1-7")
if not human_ok:
    issues.append("Human bootstrap CIs not in MS")
if not brain_ok:
    issues.append("Brain bootstrap CIs not in MS")
if issues:
    print("ISSUES FOUND:")
    for iss in issues:
        print(f"  - {iss}")
else:
    print("ALL CHECKS PASSED (0 issues)")
