"""
Generate NAR-compliant Supplementary PDF from supplementary DOCX content.
NAR requires: single PDF file for all supplementary materials.
"""
from docx import Document
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                                Table, TableStyle, Preformatted)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import re

# ── Font setup ───────────────────────────────────────────────────────
try:
    pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
    pdfmetrics.registerFont(TTFont('ArialBd', 'C:/Windows/Fonts/arialbd.ttf'))
    FONT = 'Arial'
    FONT_B = 'ArialBd'
except Exception as e:
    print(f'Font registration failed: {e}')
    FONT = 'Helvetica'
    FONT_B = 'Helvetica-Bold'

# ── Page setup (A4, NAR-compliant margins) ───────────────────────
PAGE_W, PAGE_H = A4
MARGIN_L = 2.5 * cm
MARGIN_R = 2.5 * cm
MARGIN_T = 2.5 * cm
MARGIN_B = 2.5 * cm
CONTENT_W = PAGE_W - MARGIN_L - MARGIN_R

# ── Styles ─────────────────────────────────────────────────────────────
styles = {
    'title': ParagraphStyle(
        'title', fontName=FONT_B, fontSize=16, leading=20,
        spaceAfter=6, spaceBefore=0, textColor=colors.black,
    ),
    'heading1': ParagraphStyle(
        'heading1', fontName=FONT_B, fontSize=12, leading=16,
        spaceBefore=12, spaceAfter=6, textColor=colors.black,
    ),
    'heading2': ParagraphStyle(
        'heading2', fontName=FONT_B, fontSize=11, leading=15,
        spaceBefore=10, spaceAfter=4, textColor=colors.black,
    ),
    'heading3': ParagraphStyle(
        'heading3', fontName=FONT_B, fontSize=10, leading=14,
        spaceBefore=8, spaceAfter=3, textColor=colors.black,
    ),
    'body': ParagraphStyle(
        'body', fontName=FONT, fontSize=10, leading=14,
        spaceAfter=4, textColor=colors.black,
    ),
    'code': ParagraphStyle(
        'code', fontName='Courier', fontSize=8, leading=11,
        spaceAfter=2, leftIndent=12, textColor=colors.black,
    ),
    'caption': ParagraphStyle(
        'caption', fontName=FONT, fontSize=9, leading=12,
        spaceAfter=6, spaceBefore=3, textColor=colors.black,
    ),
}

# ── Helpers ───────────────────────────────────────────────────────────
def clean_text(text: str) -> str:
    """Escape XML special chars for ReportLab."""
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    # Convert common unicode to XML entities
    text = text.replace('\u2018', '&apos;').replace('\u2019', '&apos;')
    text = text.replace('\u201c', '&quot;').replace('\u201d', '&quot;')
    return text.strip()

def para(text: str, style_name: str, **kwargs) -> Paragraph:
    return Paragraph(clean_text(text), styles[style_name], **kwargs)

# ── Main ──────────────────────────────────────────────────────────────
def build_supplementary_pdf():
    docx_path = 'results/NAR_Submission_Final_v2/supplementary/CKI_NAR_Supplementary.docx'
    out_path = 'results/NAR_Submission_Final_v2/supplementary/CKI_NAR_Supplementary.pdf'

    print(f'Reading: {docx_path}')
    doc = Document(docx_path)

    # Collect all paragraphs and tables
    story = []

    # Title page
    story.append(para('Supplementary Information', 'title'))
    story.append(para('CKI: A Cell-state Kinetic Index for Quantifying Selective Transcriptomic Remodeling', 'heading1'))
    story.append(Spacer(1, 0.3 * inch))
    story.append(para('Li Zhang', 'body'))
    story.append(para('Institute of Blood Transfusion, Chinese Academy of Medical Sciences & Peking Union Medical College, Chengdu, China', 'body'))
    story.append(para('Email: knightz@pumc.edu.cn', 'body'))
    story.append(PageBreak())

    # Process DOCX content
    in_refs = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # Skip reference section
        if p.style and 'Heading' in p.style.name and 'Reference' in text:
            in_refs = True
            continue
        if in_refs:
            continue

        # Determine style
        if p.style and 'Heading' in p.style.name:
            lvl = p.style.name.replace('Heading ', '')
            try:
                lvl = int(lvl)
            except:
                lvl = 1
            if lvl == 1:
                story.append(para(text, 'heading1'))
            elif lvl == 2:
                story.append(para(text, 'heading2'))
            elif lvl == 3:
                story.append(para(text, 'heading3'))
            else:
                story.append(para(text, 'heading2'))
        else:
            # Check if looks like code / formula
            if re.match(r'^[A-Za-z0-9_\[\]\(\)\s=+\-*/.,;:]+$', text) and len(text) > 20 and text.count(' ') < 5:
                story.append(Preformatted(text, styles['code']))
            else:
                story.append(para(text, 'body'))

    # Process tables
    for tbl_idx, tbl in enumerate(doc.tables):
        print(f'  Table {tbl_idx}: {len(tbl.rows)} rows x {len(tbl.columns)} cols')
        # Convert to list of lists
        table_data = []
        for row in tbl.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()[:200]  # truncate long cells
                row_data.append(cell_text)
            table_data.append(row_data)

        if table_data:
            t = Table(table_data, repeatRows=1)
            t.setStyle(TableStyle([
                ('FONTNAME', (0,0), (-1,-1), FONT),
                ('FONTSIZE', (0,0), (-1,-1), 8),
                ('BACKGROUND', (0,0), (-1,0), colors.Color(0.9, 0.9, 0.9)),
                ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 4),
                ('RIGHTPADDING', (0,0), (-1,-1), 4),
                ('TOPPADDING', (0,0), (-1,-1), 2),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(Spacer(1, 0.2 * inch))
            story.append(t)
            story.append(Spacer(1, 0.2 * inch))

    # Build PDF
    doc_template = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        leftMargin=MARGIN_L, rightMargin=MARGIN_R,
        topMargin=MARGIN_T, bottomMargin=MARGIN_B,
    )
    doc_template.build(story)
    print(f'\nSaved: {out_path}')
    import os
    sz = os.path.getsize(out_path)
    print(f'Size: {sz:,} bytes ({sz/1024:.1f} KB)')

if __name__ == '__main__':
    build_supplementary_pdf()
