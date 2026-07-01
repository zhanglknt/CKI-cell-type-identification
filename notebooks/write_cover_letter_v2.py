"""
Rewrite CKI_NAR_Cover_Letter_v2.docx
Front-load the Ka/Ks analogy as the central conceptual innovation.
"""

import sys
sys.path.insert(0, sys.path[0] + r'\C:\Users\KnightZ\.workbuddy\binaries\python\versions\3.13.12\Lib\site-packages')

from docx import Document
from docx.text import Paragraph, TextRun
from docx.shared import Pt
from copy import deepcopy

RESULTS = r"C:\Users\KnightZ\Desktop\细胞受选择\results"

# Build new cover letter
doc = Document()

# --- Paragraph 0: Date ---
p_date = doc.add_paragraph("May 24, 2026")
for run in p_date.runs:
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

# --- Paragraph 1: Submission header ---
p_header = doc.add_paragraph()
run = p_header.add_run("RE: \"CKI: A Cell-state Kinetic Index for Quantifying Selective Transcriptomic Remodeling\"")
run.bold = True
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 2: Salutation ---
p_sal = doc.add_paragraph("Dear Editors of Nucleic Acids Research,")
for run in p_sal.runs:
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

# --- Paragraph 3: Opening — Ka/Ks analogy FRONT and CENTER ---
p_open = doc.add_paragraph()
# Sentence 1: The Ka/Ks analogy (FRONT and CENTER)
run = p_open.add_run(
    "Just as the Ka/Ks framework revolutionized molecular evolution by decomposing sequence divergence into neutral (Ks) and functional (Ka) components, "
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

run = p_open.add_run(
    "CKI (Cell-state Kinetic Index) brings this same principled decomposition to single-cell transcriptomics."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run.bold = True  # emphasize the conceptual leap

# Sentence 2: What CKI does
run = p_open.add_run(
    " CKI partitions transcriptomic divergence into a neutral drift rate (k_n, from housekeeping genes) and a functional conversion rate (k_f, from cell-type-identity genes), yielding omega = k_f/k_n as a quantitative index of selective transcriptional remodeling."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 4: The problem + CKI solution ---
p_problem = doc.add_paragraph()
run = p_problem.add_run(
    "Current methods for comparing transcriptomes measure how different two cell populations are, but cannot distinguish functional reprogramming from neutral background noise. "
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run = p_problem.add_run(
    "CKI solves this by using housekeeping genes as an internal neutral reference—the transcriptomic equivalent of Ks—and identity genes as the functional readout—the transcriptomic equivalent of Ka."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 5: Validation overview ---
p_validate = doc.add_paragraph()
run = p_validate.add_run(
    "We validated CKI (v0.2.0, available as a pip-installable Python package) across four independent datasets spanning three data modalities: "
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run = p_validate.add_run(
    "Tabula Muris (15,057 mouse cells, FACS-sorted), Tabula Sapiens (108,136 human cells, 102 cell types, 24 organs), TCGA bulk RNA-seq (3,596 samples, 5 cancer types), and the Siletti human brain atlas (888,263 non-neuronal nuclei, 10 cell classes, 108 regions)."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 6: Key finding 1 — omega is a new dimension ---
p_k1 = doc.add_paragraph()
run = p_k1.add_run(
    "The Ka/Ks analogy makes a testable prediction: if omega captures a distinct axis of transcriptomic divergence, it should be weakly correlated with standard distance metrics. "
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run = p_k1.add_run(
    "Across 5,151 Tabula Sapiens cell-type pairs, CKI omega was weakly correlated with all four standard metrics (Spearman r = −0.09 to −0.34, all P < 0.001), confirming that omega quantifies a previously unmeasured dimension of cell-type identity."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 7: Key finding 2 — cancer paradox ---
p_k2 = doc.add_paragraph()
run = p_k2.add_run(
    "Applied to TCGA tumors, CKI revealed a paradox: tumors are transcriptionally more homogeneous than normal tissues (NN/TT ratio = 1.76–2.47, all P < 10^−16 across 5 cancer types). "
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run = p_k2.add_run(
    "This suggests that cancer progression involves not only selfish transcriptional reprogramming, but also a systematic loss of cell-type-specific identity—a prediction uniquely enabled by the omega = k_f/k_n decomposition."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 8: Key finding 3 — brain gradient ---
p_k3 = doc.add_paragraph()
run = p_k3.add_run(
    "In the human brain, CKI revealed a 6.06-fold omega gradient across 10 non-neuronal cell classes (Bergmann glia: ω = 2.37; Astrocytes: ω = 14.36; all bootstrap P < 10^−47). "
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run = p_k3.add_run(
    "The gradient aligns with known cell biology: vascular cells and fibroblasts (uniform microenvironments) show the lowest omega, while astrocytes (region-dependent functions) show the highest—demonstrating that omega captures position-dependent selective pressure."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 9: Why NAR ---
p_why = doc.add_paragraph()
run = p_why.add_run(
    "Four reasons make this work a strong fit for Nucleic Acids Research: "
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run.bold = True
run = p_why.add_run(
    "(1) CKI provides a principled framework for interpreting transcriptomic divergence, directly analogous to the Ka/Ks framework that transformed molecular evolution; (2) the method is implemented as a fully open-source Python package (pip install cki), enabling immediate community adoption; (3) validation spans three data modalities (scRNA-seq, snRNA-seq, bulk RNA-seq) and four independent datasets; (4) the cancer and brain findings are of broad interest to NAR readers."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 10: Competing interests ---
p_ci = doc.add_paragraph()
run = p_ci.add_run(
    "All authors have approved the manuscript and declare no competing interests. This manuscript has not been published elsewhere and is not under consideration by any other journal."
)
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 11: Closing ---
p_close = doc.add_paragraph()
run = p_close.add_run("Sincerely,")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 12: Signature ---
p_sig = doc.add_paragraph()
run = p_sig.add_run("Li Zhang (张力)")
run.font.size = Pt(12)
run.font.name = "Times New Roman"
run.bold = True

# --- Paragraph 13: Affiliation 1 ---
p_aff1 = doc.add_paragraph()
run = p_aff1.add_run("Institute of Blood Transfusion, Chinese Academy of Medical Sciences & Peking Union Medical College, Chengdu, China")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 14: Affiliation 2 ---
p_aff2 = doc.add_paragraph()
run = p_aff2.add_run("Chinese Institute for Brain Research, Beijing, China")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 15: ORCID ---
p_orcid = doc.add_paragraph()
run = p_orcid.add_run("ORCID: 0000-0002-0698-0754")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# --- Paragraph 16: Email ---
p_email = doc.add_paragraph()
run = p_email.add_run("Email: knightz@pumc.edu.cn")
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# Save
out = RESULTS + r"\CKI_NAR_Cover_Letter_v2.docx"
doc.save(out)
print(f"Saved: {out}")
print("Done!")
