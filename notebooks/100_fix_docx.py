"""
修复 CKI_Reproducibility_Guide.docx 中的错误内容
- 替换所有 "1,064" → "1,130"
- 替换 Section 3.1 的 HK gene 描述（从 auto-detection 改为预定义列表）
- 替换 Section 7 的输出文件名
保留原有格式（逐 run 替换）。
"""
from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy

DOCX = "results/CKI_Reproducibility_Guide.docx"

doc = Document(DOCX)

# ── 1. 替换所有 "1,064" → "1,130" ───────────────────────
count_1064 = 0
for para in doc.paragraphs:
    for run in para.runs:
        if "1,064" in run.text or "1064" in run.text:
            old = run.text
            run.text = run.text.replace("1,064", "1,130").replace("1064", "1130")
            count_1064 += 1
            print(f"  [1,064→1,130] {old[:60]!r} → {run.text[:60]!r}")

print(f"替换了 {count_1064} 处 '1,064/1064'\n")

# ── 2. 替换 Section 3.1（HK genes 描述）──────────────────
# 找到 Section 3.1 的起始和结束段落
# 删除旧的段落，插入新内容
new_s31_text = [
    "3.1 Housekeeping Genes (for k_n)",
    "",
    "In the analyses reported here, housekeeping (HK) genes were NOT ",
    "auto-detected. Instead, pre-specified HK gene lists were loaded ",
    "from the HRT Atlas reference file:",
    "",
    "  data/housekeeping/Human_Mouse_Common.csv",
    "  (1,130 orthologous gene pairs; Mouse column = mouse genes,",
    "  Human column = human genes; format identical to cki/data/hrt_atlas.csv)",
    "",
    "For each dataset, HK genes were loaded as follows:",
    "",
    "    Tabula Muris (mouse):  HRT Atlas mouse genes (column 0),",
    "                              intersected with data gene names",
    "",
    "    Tabula Sapiens (human):  HRT Atlas human genes (column 1),",
    "                              intersected with common gene set",
    "",
    '    TCGA (human):           HRT Atlas human genes (column 1),',
    "                              mapped via probeMap to Ensembl IDs,",
    "                              then intersected with data genes",
    "",
    '    Siletti Brain (human):  HRT Atlas human genes (column 1)',
    "                              matched to var[\"Gene\"] in the h5ad file",
    "",
    "Note: CKI supports data-driven HK auto-detection via ",
    '`detect_housekeeping_genes()` (detection rate > 0.9, CV < 30th ',
    'percentile, "combined" method; use_reference=True merges HRT Atlas).',
    "This auto-detection was NOT used in the current analyses ",
    "(the pre-specified list approach was preferred for reproducibility),",
    "but is available for new datasets without a curated HK list.",
]

# 找到 Section 3.1 的开始（标题段落）
# 和 Section 3.2 的开始（下一个标题）
s31_start = None
s31_end = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "3.1 Housekeeping Genes (for k_n)":
        s31_start = i
    if s31_start is not None and i > s31_start:
        if para.text.strip().startswith("3.2") or \
           para.style and para.style.name and para.style.name.startswith("Heading"):
            s31_end = i
            break

print(f"Section 3.1: paragraphs [{s31_start}, {s31_end})")

# ── 3. 替换 Section 7 的输出文件名 ───────────────────────────
# 找到 Section 7 并替换文件名行
old_to_new = {
    "results/mouse_pilot_v2b_omega_hybrid.csv":
        "results/mouse_pilot_v2b_results.csv",
    "results/ts_human_omega_hybrid_v3.csv":
        "results/phase33_v3_human_omega.csv",
    "results/tcga_omega_v2.csv":
        "results/phase34_v2_all_pairs.csv",
    "results/brain_siletti_ct_summary_v3.csv":
        "results/brain_siletti_v4_ct_summary.csv",
    "results/brain_siletti_omega_pairs_v3.csv":
        "results/brain_siletti_v4_omega_pairs.csv",
    "results/brain_siletti_migration_candidates_v3.csv":
        "results/brain_siletti_v4_migration_candidates.csv",
}

count_files = 0
for para in doc.paragraphs:
    for old, new in old_to_new.items():
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                count_files += 1
                print(f"  [file] {old!r} → {new!r}")
print(f"\n替换了 {count_files} 处文件名\n")

# ── 保存 ─────────────────────────────────────────────────────────
out = "results/CKI_Reproducibility_Guide_v2.docx"
doc.save(out)
print(f"已保存: {out}")
print("⚠ 注意：Section 3.1 的段落替换需要手动完成（自动替换较复杂）")
print("请手动编辑 DOCX 的 Section 3.1，或重新生成整个文档。")
