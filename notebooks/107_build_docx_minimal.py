#!/usr/bin/env python3
"""Generate CKI_Reproducibility_Guide.docx from markdown. Minimal version."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

MD = r"C:\Users\KnightZ\Desktop\细胞受选择\results\CKI_Reproducibility_Guide.md"
OUT = r"C:\Users\KnightZ\Desktop\细胞受选择\results\CKI_Reproducibility_Guide.docx"

FONT = "Arial"

def resolve_md_link(text):
    """Convert markdown [text](url) -> text."""
    return re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

def add_styled_para(doc, text, level=None, is_code=False):
    """Add a paragraph with proper styling."""
    text = resolve_md_link(text)
    if is_code:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        return p
    
    if level == 1:
        p = doc.add_heading(text, level=1)
    elif level == 2:
        p = doc.add_heading(text, level=2)
    elif level == 3:
        p = doc.add_heading(text, level=3)
    else:
        p = doc.add_paragraph(text)
    
    for run in p.runs:
        run.font.name = FONT
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def build():
    doc = Document()
    with open(MD, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code = False
    code_buf = []
    in_table = False
    table_rows = []
    
    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        tbl.style = 'Table Grid'
        for r, row in enumerate(table_rows):
            for c, cell_text in enumerate(row):
                cell = tbl.cell(r, c)
                cell.text = resolve_md_link(cell_text.strip())
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT
                        run.font.size = Pt(9)
        doc.add_paragraph()
        table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        if line.strip().startswith('```'):
            if in_code:
                add_styled_para(doc, ''.join(code_buf), is_code=True)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        
        if in_code:
            code_buf.append(line + '\n')
            i += 1
            continue
        
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            flush_table()
            in_table = True
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            if not all(re.match(r'^-+$', c) for c in cells if c):
                table_rows.append(cells)
            i += 1
            continue
        
        if in_table:
            flush_table()
            in_table = False
        
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        if line.strip().startswith('# '):
            add_styled_para(doc, line.strip()[2:], level=1)
        elif line.strip().startswith('## '):
            add_styled_para(doc, line.strip()[3:], level=2)
        elif line.strip().startswith('### '):
            add_styled_para(doc, line.strip()[4:], level=3)
        else:
            add_styled_para(doc, line.strip())
        
        i += 1
    
    if in_code and code_buf:
        add_styled_para(doc, ''.join(code_buf), is_code=True)
    flush_table()
    
    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == '__main__':
    build()
