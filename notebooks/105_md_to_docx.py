#!/usr/bin/env python3
"""
Generate CKI_Reproducibility_Guide.docx directly from the CORRECT markdown file.
This ensures the DOCX is ALWAYS in sync with the markdown (single source of truth).
"""

import re
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN
import os

MD = "C:/Users/KnightZ/Desktop/细胞受选择/results/CKI_Reproducibility_Guide.md"
OUT = "C:/Users/KnightZ/Desktop/细胞受选择/results/CKI_Reproducibility_Guide.docx"

FONT = "Arial"
SIZE_BODY = Pt(10)
SIZE_H1 = Pt(16)
SIZE_H2 = Pt(14)
SIZE_H3 = Pt(12)

def parse_markdown(md_path):
    """Parse markdown into a list of (type, content) tuples."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    elements = []
    in_code = False
    code_lines = []
    in_table = False
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('```'):
            if in_code:
                elements.append(('code', ''.join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        
        if in_code:
            code_lines.append(line)
            continue
        
        if stripped.startswith('# '):
            elements.append(('h1', stripped[2:]))
        elif stripped.startswith('## '):
            elements.append(('h2', stripped[3:]))
        elif stripped.startswith('### '):
            elements.append(('h3', stripped[4:]))
        elif stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Skip separator rows
            if all(re.match(r'^-+$', c.strip()) for c in cells if c.strip()):
                continue
            table_rows.append(cells)
        elif in_table:
            if stripped.startswith('|'):
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table_rows.append(cells)
            else:
                in_table = False
                if table_rows:
                    elements.append(('table', table_rows))
                    table_rows = []
                # Process current line normally
                if stripped:
                    elements.append(('text', stripped))
        elif stripped.startswith('- ') or re.match(r'^\d+\.\s', stripped):
            elements.append(('list', stripped))
        elif stripped:
            elements.append(('text', stripped))
        else:
            elements.append(('blank', ''))
    
    if table_rows:
        elements.append(('table', table_rows))
    
    return elements


def build_docx(elements, out_path):
    """Build DOCX from parsed elements."""
    doc = Document()
    
    # Remove default paragraph
    if len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    for elem_type, content in elements:
        if elem_type == 'h1':
            p = doc.add_heading(content, level=1)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = SIZE_H1
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif elem_type == 'h2':
            p = doc.add_heading(content, level=2)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = SIZE_H2
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif elem_type == 'h3':
            p = doc.add_heading(content, level=3)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = SIZE_H3
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif elem_type == 'text':
            p = doc.add_paragraph(content)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif elem_type == 'code':
            p = doc.add_paragraph(content)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif elem_type == 'list':
            p = doc.add_paragraph(content, style='List Bullet' if not re.match(r'^\d', content) else 'List Number')
            for run in p.runs:
                run.font.name = FONT
                run.font.size = SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif elem_type == 'table':
            table = doc.add_table(rows=len(content), cols=len(content[0]))
            table.style = 'Table Grid'
            for i, row in enumerate(content):
                for j, cell_text in enumerate(row):
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = FONT
                            run.font.size = SIZE_BODY
                            run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.save(out_path)
    print(f"Generated: {out_path}")
    size = os.path.getsize(out_path)
    print(f"Size: {size / 1024:.1f} KB")


if __name__ == '__main__':
    print("Parsing markdown...")
    elements = parse_markdown(MD)
    print(f"Parsed {len(elements)} elements")
    print("Building DOCX...")
    build_docx(elements, OUT)
    print("Done!")
