"""
Fix NAR submission compliance issues in CKI manuscript.

P0 fixes:
1. Code Availability - add Zenodo DOI placeholder
2. Acknowledgements - add AI disclosure
3. Supplementary material - confirm PDF exists
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

DOCX = 'results/NAR_Submission_Final_v2/manuscript/CKI_NAR_Manuscript_v4.docx'

doc = Document(DOCX)

# ── Fix 1: Code Availability ──────────────────────────────────────
# Find Code Availability paragraph [96]
for i, p in enumerate(doc.paragraphs):
    if 'Code Availability' in p.text and p.style and 'Heading' in p.style.name:
        # Next paragraph with text is [96]
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                target = doc.paragraphs[j]
                old_text = target.text
                new_text = (
                    "The CKI Python package (v0.2.0) is available at "
                    "https://github.com/zhanglknt/CKI-cell-type-identification "
                    "(Zenodo DOI: 10.5281/zenodo.XXXXXXX). "
                    "All analysis notebooks and supplementary scripts are available at the same repository."
                )
                target.text = new_text
                print(f"Fixed Code Availability [{j}]:")
                print(f"  Old: {old_text[:100]}")
                print(f"  New: {new_text[:100]}")
                break
        break

# ── Fix 2: Add AI disclosure to Acknowledgements ───────────────────
# Find Acknowledgements section [98], add AI disclosure after first paragraph
ai_disclosure = (
    "During manuscript preparation, large language models (LLMs) were used to assist with "
    "language polishing and code debugging. All scientific content, analyses, and interpretations "
    "were solely performed by the author. This disclosure complies with NAR's AI tool disclosure policy."
)

for i, p in enumerate(doc.paragraphs):
    if p.style and 'Heading' in p.style.name and 'Acknowledg' in p.text:
        # Find the paragraph after the heading to insert before/after
        # Insert a new paragraph after the heading's first content paragraph
        insert_idx = None
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                insert_idx = j + 1  # insert after first content para
                break
        if insert_idx is not None:
            # Use docx API to insert a new paragraph
            from docx.text.paragraph import Paragraph
            from docx.oxml import OxmlElement
            new_p_element = OxmlElement('w:p')
            r_element = OxmlElement('w:r')
            t_element = OxmlElement('w:t')
            t_element.set(qn('xml:space'), 'preserve')
            t_element.text = ai_disclosure
            r_element.append(t_element)
            new_p_element.append(r_element)
            # Insert after the first content paragraph
            body = doc.element.body
            ref_p = doc.paragraphs[insert_idx - 1]._p
            ref_p.addnext(new_p_element)
            print(f"Added AI disclosure after paragraph [{insert_idx-1}]")
        break

# ── Save ─────────────────────────────────────────────────────────────
OUT = DOCX  # overwrite
doc.save(OUT)
print(f"\nSaved: {OUT}")
