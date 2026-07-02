"""
Add Section 3.3 (Script-to-Manuscript Results Mapping) to
CKI_Reproducibility_Guide_v1.docx.

Inserts the new section before "4. Data Sources & Preprocessing".
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

FONT = "Arial"
BLACK = "000000"
SIZE = Pt(10)  # 10pt

def p(text, bold=False, italic=False, size=None):
    """Create a paragraph with Arial 10pt text."""
    from docx.text.run import Run
    from docx.shared import Pt
    s = size or Pt(10)
    para = docx.text.paragraph.Paragraph.__new__(docx.text.paragraph.Paragraph)
    # Use docx's Paragraph constructor via XML
    # Actually, easier to use add_paragraph and then move
    return None  # will use doc.add_paragraph

# Just use the simple approach: build new paragraphs and insert via XML
from lxml import etree

doc = Document('results/CKI_Reproducibility_Guide_v1.docx')

# ---------------------------------------------------------------------------
# Build the new paragraphs as proper docx Paragraph objects
# We'll insert them by manipulating the document's body XML
# ---------------------------------------------------------------------------
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def make_text_run(text, bold=False, italic=False, size_pt=10):
    """Create a <w:r> element with text."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)
    # Size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size_pt * 2))
    rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def make_para(style=None, runs_text=None, bold_runs=None):
    """Create a <w:p> element."""
    p = OxmlElement('w:p')
    if style:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style)
        pPr.append(pStyle)
        p.append(pPr)
    if runs_text:
        r = make_text_run(runs_text, bold=(bold_runs or False))
        p.append(r)
    return p

def insert_para_before(doc, new_p, before_idx):
    """Insert a new <w:p> element before the paragraph at body index before_idx."""
    body = doc.element.body
    # Find the body child corresponding to doc.paragraphs[before_idx]
    para_els = [c for c in body if c.tag.endswith('}p')]
    if before_idx < len(para_els):
        body.insert(body.index(para_els[before_idx]), new_p)
    else:
        body.append(new_p)

# ---------------------------------------------------------------------------
# Content for Section 3.3
# ---------------------------------------------------------------------------
# We need to insert before doc.paragraphs[74]
# First, let's find the body index of paragraphs[74]

doc = Document('results/CKI_Reproducibility_Guide_v1.docx')
body = doc.element.body
all_p_els = [c for c in body if c.tag.endswith('}p')]

# paragraphs[74] corresponds to all_p_els[74] (assuming no non-para elements in between)
# Actually, doc.paragraphs may not map 1:1 if there are tables etc.
# Let's use a more robust method:

target_text = '4. Data Sources & Preprocessing'
insert_before_elem = None
for i, p in enumerate(all_p_els):
    t_els = p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    text = ''.join(t.text or '' for t in t_els)
    if target_text in text:
        insert_before_elem = p
        print(f'Found insertion point in body (p #{i})')
        break

if insert_before_elem is None:
    print('ERROR: Could not find insertion point in body XML')
    exit(1)

# Now build all new paragraph elements
new_paras = []

# Section 3.3 heading (Heading 2 style)
p_heading = OxmlElement('w:p')
pPr = OxmlElement('w:pPr')
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'Heading 2')
pPr.append(pStyle)
p_heading.append(pPr)
r = make_text_run('3.3 Script-to-Manuscript Results Mapping', bold=True, size_pt=11)
p_heading.append(r)
new_paras.append(p_heading)

# Explanatory paragraph
for txt in [
    "This section clarifies which notebook scripts produce the results "
    "reported in the manuscript, and which scripts are exploratory "
    "(older designs that do not match the manuscript numbers).",
    "",
    "All scripts listed below are in the notebooks/ directory. "
    "Scripts marked \u201cExploratory\u201d are retained for reference but "
    "their output does NOT correspond to manuscript values.",
]:
    p = OxmlElement('w:p')
    r = make_text_run(txt, size_pt=10)
    p.append(r)
    new_paras.append(p)

# Table: Script-to-Results mapping
table_data = [
    # (Dataset, Manuscript Figure, Primary script, Output CSV(s))
    ("Mouse\n(pilot)", "Results 1\u20132,\nFig. 2", "02b_pilot_v2.py\n02c_pilot_v2b.py", "mouse_pilot_v2_results.csv\nmouse_pilot_v2b_results.csv"),
    ("Mouse\n(full\npairwise)", "Fig. 2\n(heatmap)", "03_full_matrix.py", "full_matrix_omega.csv\nfull_matrix_pairs.csv"),
    ("Human", "Fig. 3", "05_phase33_v3_fixed.py", "phase33_v3_human_omega.csv"),
    ("TCGA\n(tumor\nvs normal)", "Fig. 4\n(tumor\nvs normal)", "06_phase34_v2.py", "phase34_v2_summary.csv\nphase34_v2_all_pairs.csv"),
    ("TCGA\n(clinical\nseverity)", "Fig. 4\n(clinical\nseverity)", "07_phase34_clinical.py", "phase34_clinical_severity.csv"),
    ("Brain", "Fig. 5", "07c_brain_siletti_v3.py", "brain_siletti_omega_pairs_v3.csv"),
    ("Method\ncomparison", "Fig. 4\n(AUC)", "13_phase35_method_comparison.py", "phase35_all_metrics_pairs.csv\nphase35_cross_organ_conservation.csv"),
    ("Bootstrap", "Fig. 2\u20133\n(significance)", "08a_tcga_bootstrap.py\n08b_mouse_bootstrap.py", "tcga_bootstrap_results.csv\nmouse_pilot_v2b_bootstrap.csv"),
]

# Build the table
tbl = OxmlElement('w:tbl')
# Table properties (simplified - use the same style as other tables in the doc)
tblPr = OxmlElement('w:tblPr')
# Use "Table Grid" style
tblStyle = OxmlElement('w:tblStyle')
tblStyle.set(qn('w:val'), 'TableGrid')
tblPr.append(tblStyle)
tbl.append(tblPr)
# Grid columns (4 columns)
tblGrid = OxmlElement('w:tblGrid')
for _ in range(4):
    gridCol = OxmlElement('w:gridCol')
    gridCol.set(qn('w:w'), '2880')  # ~2 inches each
    tblGrid.append(gridCol)
tbl.append(tblGrid)

# Header row
hdr_row = OxmlElement('w:tr')
for hdr_txt in ["Dataset", "Manuscript\nFigure", "Primary script", "Output CSV(s)"]:
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    r = make_text_run(hdr_txt, bold=True, size_pt=10)
    p.append(r)
    tc.append(p)
    hdr_row.append(tc)
tbl.append(hdr_row)

# Data rows
for row_data in table_data:
    tr = OxmlElement('w:tr')
    for cell_txt in row_data:
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'top')
        tcPr.append(vAlign)
        tc.append(tcPr)
        p = OxmlElement('w:p')
        r = make_text_run(cell_txt, size_pt=10)
        p.append(r)
        tc.append(p)
        tr.append(tc)
    tbl.append(tr)

new_paras.append(tbl)

# Note about exploratory scripts
for txt in [
    "",
    "Exploratory scripts (do NOT match manuscript numbers):",
    "\u2022 01_pilot_mouse.py: Tissue-level CKI only (not cell-type-level). "
    "Replaced by 02b/02c in the manuscript.",
    "\u2022 02_ct_pilot.py: Early cell-type-level pilot with random-split control. "
    "The manuscript uses mouse-ID-split control (02c_pilot_v2b.py). "
    "[Bug fixed in v0.3.2: fname\u2192fname typo on line 72.]",
    "\u2022 04_phase32_sweep.py: Parameter sweep over w1/w2 weights for multi-component k_f. "
    "The sweep result is NOT exactly reproducible because "
    "gsp.utils.download_library(\u2019H\u2019, \u2019Mouse\u2019) downloads the "
    "latest MSigDB Hallmark definitions at runtime. "
    "To reproduce exactly, use the bundled results/phase32_sweep_results.csv "
    "instead of re-running the sweep.",
    "",
    "Recommended reproduction workflow:",
    "1. Run 02b_pilot_v2.py \u2192 02c_pilot_v2b.py \u2192 03_full_matrix.py "
    "to reproduce all mouse results (Fig. 2, Results 1\u20132).",
    "2. For the parameter sweep, use the bundled phase32_sweep_results.csv "
    "rather than re-running 04_phase32_sweep.py.",
    "3. All other scripts (05 through 08) reproduce manuscript results "
    "exactly when run with the provided data and CKI v0.3.1.",
]:
    p = OxmlElement('w:p')
    r = make_text_run(txt, size_pt=10)
    p.append(r)
    new_paras.append(p)

# ---------------------------------------------------------------------------
# Insert all new paragraphs before the insertion point
# ---------------------------------------------------------------------------
for elem in new_paras:
    body.insert(body.index(insert_before_elem), elem)

# Save
doc.save('results/CKI_Reproducibility_Guide_v1.docx')
print('Saved: results/CKI_Reproducibility_Guide_v1.docx')
print(f'Inserted {len(new_paras)} new elements (Section 3.3)')
