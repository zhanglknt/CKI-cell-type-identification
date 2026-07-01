"""
Rebuild CKI_NAR_Cover_Letter_v3.docx
- Arial font throughout (NAR requirement)
- Add Key Innovations table
- Ka/Ks analogy front-loaded and bolded
- No Unicode encoding issues
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RESULTS = r"C:\Users\KnightZ\Desktop\细胞受选择\results"

doc = Document()

# --- Helper ---
def add(p, text, bold=False, size=12, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Arial"
    r.font.size = Pt(size)
    return r

def set_para_font(p):
    for r in p.runs:
        r.font.name = "Arial"
        if r.font.size is None:
            r.font.size = Pt(12)

# === 0: Date ===
p = doc.add_paragraph("June 9, 2026")
set_para_font(p)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# === 1: RE title ===
p = doc.add_paragraph()
add(p, 'RE: "CKI: A Cell-state Kinetic Index for Quantifying Selective Transcriptomic Remodeling"', bold=True, size=12)
set_para_font(p)

# === 2: Salutation ===
p = doc.add_paragraph("Dear Editors of Nucleic Acids Research,")
set_para_font(p)

# === 3: Ka/Ks analogy (FRONT-LOADED, bold) ===
p = doc.add_paragraph()
add(p, "Just as the Ka/Ks framework revolutionized molecular evolution by decomposing sequence divergence into neutral (Ks) and functional (Ka) components, ", size=12)
add(p, "CKI (Cell-state Kinetic Index) brings this same principled decomposition to single-cell transcriptomics.", bold=True, size=12)
add(p, " CKI partitions transcriptomic divergence into a neutral drift rate (k_n, from housekeeping genes — the transcriptomic Ks) and a functional remodeling rate (k_f, from identity genes — the transcriptomic Ka), yielding omega = k_f/k_n as a quantitative index of selective transcriptional remodeling.", size=12)
set_para_font(p)

# === 4: Table of Key Innovations ===
# NAR cover letters benefit from a clear "highlights" table
tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
headers = ["Innovation", "Description"]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)

innovations = [
    ("1. Principled null model for transcriptomic comparison",
     " omega = k_f/k_n decomposes divergence into neutral (k_n) and functional (k_f) components, analogous to Ka/Ks in molecular evolution. Provides a statistical null for any transcriptomic comparison."),
    ("2. Orthogonal information dimension",
     "Across 5,151 human cell-type pairs, CKI omega is weakly correlated with all standard distance metrics (Spearman r = -0.09 to -0.34, all P < 1e-12), confirming it captures a previously unmeasured axis."),
    ("3. Unexpected cancer transcriptomic convergence",
     "In TCGA (5 cancer types, 3,596 samples), tumors are paradoxically more transcriptionally homogeneous than normal tissues (NN/TT omega ratio = 1.76-2.47, all P < 1e-47), suggesting trans-cancer convergence on shared transcriptional states."),
    ("4. Brain regional cell-type differentiation gradient",
     "In the Siletti et al. human brain atlas (888,263 nuclei, 10 cell classes, 108 regions), CKI revealed a 6.06-fold omega gradient (Bergmann glia: 2.37; Astrocytes: 14.36; bootstrap P < 1e-47), enabling inference of cell migration history from static data."),
]
for ri, (inn, desc) in enumerate(innovations, 1):
    tbl.rows[ri].cells[0].paragraphs[0].add_run(inn).font.name = "Arial"
    tbl.rows[ri].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    tbl.rows[ri].cells[1].paragraphs[0].add_run(desc).font.name = "Arial"
    tbl.rows[ri].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

# Set table font to Arial size 10
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                if r.font.size is None:
                    r.font.size = Pt(10)

# === 5: Problem + CKI solution ===
p = doc.add_paragraph()
add(p, "Current methods for comparing transcriptomes measure how different two cell populations are, but cannot distinguish functional reprogramming from neutral background noise. ", size=12)
add(p, "CKI solves this by using housekeeping genes as an internal neutral reference (the transcriptomic Ks) and cell-type-identity genes as the functional readout (the transcriptomic Ka).", size=12)
set_para_font(p)

# === 6: Validation overview ===
p = doc.add_paragraph()
add(p, "We validated CKI (v0.2.0, pip-installable: ", size=12)
add(p, "pip install cki", italic=True, size=12)
add(p, ") across four independent datasets: Tabula Muris (15,057 mouse cells, 38 cell types, 6 organs), Tabula Sapiens (108,136 human cells, 102 cell types, 6 organs), TCGA bulk RNA-seq (3,596 samples, 5 cancer types), and the Siletti et al. human brain atlas (888,263 non-neuronal nuclei, 10 cell classes, 108 regions).", size=12)
set_para_font(p)

# === 7: Omega = new dimension ===
p = doc.add_paragraph()
add(p, "The Ka/Ks analogy makes a testable prediction: if omega captures a distinct axis of transcriptomic divergence, it should be weakly correlated with standard distance metrics. ", size=12)
add(p, "Across 5,151 Tabula Sapiens cell-type pairs, CKI omega was weakly correlated with all four standard metrics (Spearman |r| = 0.09-0.34, all P < 1e-12), confirming that omega quantifies a previously unmeasured dimension.", size=12)
set_para_font(p)

# === 8: Cancer paradox ===
p = doc.add_paragraph()
add(p, "Applied to TCGA tumors, CKI revealed a paradox: tumors are transcriptionally more homogeneous than normal tissues (NN/TT omega ratio = 1.76-2.47, all P < 1e-47 across 5 cancer types). ", size=12)
add(p, "This suggests that cancer progression involves not only selfish transcriptional reprogramming, but also systematic loss of cell-type-specific identity — a prediction uniquely enabled by the omega = k_f/k_n decomposition.", size=12)
set_para_font(p)

# === 9: Brain gradient ===
p = doc.add_paragraph()
add(p, "In the human brain, CKI revealed a 6.06-fold omega gradient across 10 non-neuronal cell classes (Bergmann glia: omega = 2.37; Astrocytes: omega = 14.36; bootstrap P < 1e-47). ", size=12)
add(p, "The gradient aligns with known cell biology: vascular cells and fibroblasts (uniform microenvironments) show the lowest omega, while astrocytes (region-dependent functions) show the highest — demonstrating that omega captures position-dependent selective pressure.", size=12)
set_para_font(p)

# === 10: Why NAR ===
p = doc.add_paragraph()
add(p, "Four reasons make this work a strong fit for Nucleic Acids Research:", bold=True, size=12)
add(p, " (1) CKI provides a principled framework for interpreting transcriptomic divergence, directly analogous to the Ka/Ks framework that transformed molecular evolution; (2) the method is implemented as a fully open-source Python package (pip install cki), enabling immediate community adoption; (3) validation spans three data modalities (scRNA-seq, snRNA-seq, bulk RNA-seq) and four independent datasets; (4) the cancer and brain findings are of broad interest to NAR readers.", size=12)
set_para_font(p)

# === 11: Competing interests ===
p = doc.add_paragraph()
add(p, "All authors have approved the manuscript and declare no competing interests. This manuscript has not been published elsewhere and is not under consideration by any other journal.", size=12)
set_para_font(p)

# === 12: Closing ===
p = doc.add_paragraph("Sincerely,")
set_para_font(p)

# === 13: Signature ===
p = doc.add_paragraph()
add(p, "Li Zhang (Zhang Li)", bold=True, size=12)
set_para_font(p)

# === 14-17: Affiliations ===
for text in [
    "Institute of Blood Transfusion, Chinese Academy of Medical Sciences & Peking Union Medical College, Chengdu, China",
    "Chinese Institute for Brain Research, Beijing, China",
    "ORCID: 0000-0002-0698-0754",
    "Email: knightz@pumc.edu.cn",
]:
    p = doc.add_paragraph(text)
    set_para_font(p)

# Save
out = RESULTS + r"\CKI_NAR_Cover_Letter_v3.docx"
doc.save(out)
print(f"Saved: {out}")
print("Features: Arial font, Key Innovations table, Ka/Ks analogy front-loaded")
